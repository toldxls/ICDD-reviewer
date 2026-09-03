#!/usr/bin/env python3
"""
bv_check — bond distances and bond-valence sums from a .cif, with a manuscript-table cross-check.

    python3 -m pxrd_review.bv_check <structure.cif> [--table manuscript.docx] [--params gh|bo|ba]
                                    [--ox Fe=2,Mn=3] [--cutoff 3.2] [--no-h] [--word] [--out DIR]
    pxrd bv <structure.cif> [--table manuscript.docx]

What it does
  1. Reads the .cif (cell, symmetry operators, atom sites with occupancies; mixed sites that share
     coordinates are merged), generates every symmetry-equivalent position and finds each cation's
     anion neighbours within a cutoff (3.2 Å; larger for big cations such as K, Ba, Pb, Cs).
  2. Assigns oxidation states — from the type symbol ('Fe3+'), _atom_type_oxidation_number, the
     --ox overrides, or a table of the usual mineral valences (stated in the report as assumed).
  3. Bond valences s = exp((R0 − R)/b) with parameters from I.D. Brown's accumulated table
     (bvparm2020.cif, bundled): Gagné & Hawthorne (2015) for cation–O by default (--params gh),
     or Brese & O'Keeffe (1991) (--params bo) / Brown & Altermatt (1985) (--params ba); U6+–O from
     Burns et al. (1997) unless --u6 params. The report names every parameter it used, and the
     table note cites them the way the journals want ("… from Gagné and Hawthorne (2015); U6+–O
     from Burns et al. (1997)").
  4. Prints (and writes to review_out/<name>_bv.txt):
       • per cation site: each bond with its distance, multiplicity and valence, the mean bond
         length, the bond-valence sum and its deviation from the expected valence;
       • the bond-valence table as journals print it — anion rows × cation columns, '×2↓' where a
         cation receives that bond twice, '×2→' where an anion does, row and column sums;
       • a self-check of the computed distances against the .cif's own _geom_bond loop.
     --word also writes the two tables as a .docx (review_out/<name>_bv.docx) to paste from.
  5. --table manuscript.docx checks the manuscript against the .cif: every 'Cd1–O3² 2.472(3)'
     cell against the computed distance, listed multiplicities, bonds the table omits or that the
     .cif does not have, the arithmetic of '<Cd1–O>' mean rows, and — where the manuscript has a
     bond-valence table — each cell against the computed valence and the arithmetic of its Σ
     column and row.

Conventions: cation sums are unscaled (the sum for the ion present) but each bond is weighted by
the anion's occupancy (a half-occupied O counts half); anion sums weight each contribution by the
cation site's occupancy. Mixed sites count each species by its fraction. Without oxygen in the
structure (sulfides, sulfosalts) the default valences switch to the sulfide ones (As3+, Cu+, Fe2+).
An N with no O within 1.5 Å is ammonium (NH4+) whether or not its H atoms were refined.
Hydrogen bonds (--hbonds oo, the default): their strengths come from the donor–acceptor O···O
distance, s = (d/2.17)^−8.2 + 0.06 (Ferraris & Ivaldi 1988), as mineral descriptions print them.
With H atoms in the .cif the D–H···A pairs are the _geom_hbond loop's (else found from the H
positions); without H the donors are the OH / OW / W sites (or, when nothing is labelled, the O
sites short of bond valence) and the acceptors are chosen from the O···O geometry — each contact
once, no polyhedral edges, H–O–H angles respected, the O with the larger valence deficit
accepting — and the proposal is stated in the report (--donors OW1=2 and --hb OW1>O2 override
it). --hbonds h keeps the older convention: acceptor valences from the H···O distances (Brown
2002) with the donor getting 1 − Σ. Donated valences are listed but not deducted from the
donor's sum, as in the owner's tables.
"""
import os, re, sys, math, argparse
from collections import namedtuple, OrderedDict

# ----------------------------------------------------------------------------- CIF reading

def _cif_tokens(text):
    """CIF tokens: bare words, 'quoted' / "quoted" strings, ;-delimited text fields; # comments dropped."""
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        ln = lines[i]
        if ln.startswith(';'):                          # text field
            buf = [ln[1:]]
            i += 1
            while i < len(lines) and not lines[i].startswith(';'):
                buf.append(lines[i]); i += 1
            yield '\n'.join(buf).strip()
            i += 1
            continue
        pos = 0; n = len(ln)
        while pos < n:
            ch = ln[pos]
            if ch.isspace():
                pos += 1; continue
            if ch == '#':
                break
            if ch in '\'"':
                q = ch; j = pos + 1
                while j < n:
                    if ln[j] == q and (j + 1 == n or ln[j + 1].isspace()):
                        break
                    j += 1
                yield ln[pos + 1:j]; pos = j + 1; continue
            j = pos
            while j < n and not ln[j].isspace():
                j += 1
            yield ln[pos:j]; pos = j
        i += 1

def read_cif(path):
    """[block] — each {'name', 'items': {tag: value}, 'loops': [(tags, rows)]}."""
    with open(path, encoding='utf-8', errors='replace') as f:
        toks = list(_cif_tokens(f.read()))
    blocks = []; cur = None; i = 0
    def new(name):
        b = {'name': name, 'items': {}, 'loops': []}; blocks.append(b); return b
    while i < len(toks):
        t = toks[i]
        if t.lower().startswith('data_'):
            cur = new(t[5:]); i += 1; continue
        if cur is None:
            cur = new('')
        if t.lower() == 'loop_':
            i += 1; tags = []
            while i < len(toks) and toks[i].startswith('_'):
                tags.append(toks[i].lower()); i += 1
            vals = []
            while i < len(toks) and not (toks[i].startswith('_') or toks[i].lower() in ('loop_',)
                                         or toks[i].lower().startswith('data_')):
                vals.append(toks[i]); i += 1
            if tags:
                n = len(tags); rows = [vals[k:k + n] for k in range(0, len(vals) - len(vals) % n, n)]
                cur['loops'].append((tags, rows))
            continue
        if t.startswith('_'):
            if i + 1 < len(toks):
                cur['items'][t.lower()] = toks[i + 1]; i += 2
            else:
                i += 1
            continue
        i += 1                                          # stray value
    return blocks

def _num(s):
    """'1.2345(6)' -> 1.2345; '.' / '?' -> None"""
    if s is None:
        return None
    s = s.strip()
    if s in ('.', '?', ''):
        return None
    m = re.match(r'^([-+]?\d*\.?\d+(?:[eE][-+]?\d+)?)', s)
    return float(m.group(1)) if m else None

def _esd(s):
    """'1.2345(6)' -> 0.0006 (in the value's units); None without one"""
    m = re.match(r'^[-+]?(\d*)\.?(\d*)\((\d+)\)', s or '')
    if not m:
        return None
    return int(m.group(3)) * 10 ** (-len(m.group(2)))

def _loop(block, tag):
    for tags, rows in block['loops']:
        if tag in tags:
            return tags, rows
    return None, None

def _col(tags, rows, tag, default=None):
    if tag not in tags:
        return [default] * len(rows)
    k = tags.index(tag)
    return [r[k] for r in rows]

# ----------------------------------------------------------------------------- symmetry

_TERM = re.compile(r"\s*([+-]?)\s*(?:(\d+(?:\.\d+)?(?:/\d+)?)\s*\*?\s*)?([xyzXYZ])?\s*")

def parse_symop(s):
    """'-y+1/2, x-y, z+0.25' -> (3x3 rotation, translation) as nested lists. No eval()."""
    parts = [p.strip() for p in s.split(',')]
    if len(parts) != 3:
        raise ValueError('bad symmetry operator: %r' % s)
    rot = [[0.0] * 3 for _ in range(3)]; tr = [0.0] * 3
    for i, expr in enumerate(parts):
        pos = 0; expr = expr.replace(' ', '')
        while pos < len(expr):
            m = _TERM.match(expr, pos)
            if not m or m.end() == pos:
                raise ValueError('bad symmetry operator: %r' % s)
            sign, num, var = m.group(1), m.group(2), m.group(3)
            if not num and not var:
                raise ValueError('bad symmetry operator: %r' % s)
            f = -1.0 if sign == '-' else 1.0
            if num:
                f *= (float(num.split('/')[0]) / float(num.split('/')[1])) if '/' in num else float(num)
            if var:
                rot[i]['xyz'.index(var.lower())] += f
            else:
                tr[i] += f
            pos = m.end()
    return rot, tr

def _apply(op, p):
    rot, tr = op
    return [sum(rot[i][j] * p[j] for j in range(3)) + tr[i] for i in range(3)]

def cart_matrix(st):
    """Fractional -> Cartesian (Å) matrix of the cell (a along x, b in the xy plane)."""
    a, b, c, al, be, ga = st.cell
    ca, cb, cg = (math.cos(math.radians(x)) for x in (al, be, ga)); sg = math.sin(math.radians(ga))
    v = st.volume / (a * b * c)
    return [[a, b * cg, c * cb], [0, b * sg, c * (ca - cb * cg) / sg], [0, 0, c * v / sg]]

def cart(st, p):
    M = cart_matrix(st)
    return [sum(M[i][j] * p[j] for j in range(3)) for i in range(3)]

def angle(st, p1, p0, p2):
    """The angle p1–p0–p2 (degrees) between three fractional positions."""
    a, o, b = cart(st, p1), cart(st, p0), cart(st, p2)
    v1 = [a[i] - o[i] for i in range(3)]; v2 = [b[i] - o[i] for i in range(3)]
    n1 = math.sqrt(sum(x * x for x in v1)); n2 = math.sqrt(sum(x * x for x in v2))
    if not n1 or not n2:
        return 0.0
    return math.degrees(math.acos(max(-1.0, min(1.0, sum(v1[i] * v2[i] for i in range(3)) / (n1 * n2)))))

# ----------------------------------------------------------------------------- structure

Species = namedtuple('Species', 'element ox occ')
Site = namedtuple('Site', 'label element frac species positions mult occ_total uiso')
Bond = namedtuple('Bond', 'cation anion dist count vals')   # vals: [(species, s)] per species of the cation site

ANION_ELEMENTS = {'O', 'F', 'Cl', 'Br', 'I', 'S', 'Se', 'Te', 'N', 'H'}
ELEMENTS = set('''H He Li Be B C N O F Ne Na Mg Al Si P S Cl Ar K Ca Sc Ti V Cr Mn Fe Co Ni Cu Zn Ga Ge As Se Br
Kr Rb Sr Y Zr Nb Mo Tc Ru Rh Pd Ag Cd In Sn Sb Te I Xe Cs Ba La Ce Pr Nd Pm Sm Eu Gd Tb Dy Ho Er Tm Yb Lu Hf Ta W
Re Os Ir Pt Au Hg Tl Pb Bi Po At Rn Fr Ra Ac Th Pa U Np Pu Am Cm Bk Cf D'''.split())
# the usual valence of each element in minerals; anything else must come from the .cif or --ox
DEFAULT_OX = {'H': 1, 'Li': 1, 'Be': 2, 'B': 3, 'C': 4, 'N': 5, 'Na': 1, 'Mg': 2, 'Al': 3, 'Si': 4, 'P': 5, 'S': 6,
              'K': 1, 'Ca': 2, 'Sc': 3, 'Ti': 4, 'V': 5, 'Cr': 3, 'Mn': 2, 'Fe': 3, 'Co': 2, 'Ni': 2, 'Cu': 2,
              'Zn': 2, 'Ga': 3, 'Ge': 4, 'As': 5, 'Se': 4, 'Rb': 1, 'Sr': 2, 'Y': 3, 'Zr': 4, 'Nb': 5, 'Mo': 6,
              'Ag': 1, 'Cd': 2, 'In': 3, 'Sn': 4, 'Sb': 3, 'Te': 4, 'Cs': 1, 'Ba': 2, 'La': 3, 'Ce': 3, 'Pr': 3,
              'Nd': 3, 'Sm': 3, 'Eu': 3, 'Gd': 3, 'Tb': 3, 'Dy': 3, 'Ho': 3, 'Er': 3, 'Tm': 3, 'Yb': 3, 'Lu': 3,
              'Hf': 4, 'Ta': 5, 'W': 6, 'Re': 7, 'Hg': 2, 'Tl': 1, 'Pb': 2, 'Bi': 3, 'Th': 4, 'U': 6, 'Np': 5,
              'Pu': 4, 'Au': 3, 'Pt': 4, 'Pd': 2, 'Rh': 3, 'Ru': 4, 'Ir': 4, 'Os': 4, 'Tc': 7, 'Ac': 3, 'Pa': 5,
              'Am': 3, 'Cm': 3, 'Xe': 6, 'D': 1}
ANION_OX = {'O': -2, 'F': -1, 'Cl': -1, 'Br': -1, 'I': -1, 'S': -2, 'Se': -2, 'Te': -2, 'N': -3, 'H': -1}
# without oxygen (sulfides, sulfosalts, halides) the usual valences differ
DEFAULT_OX_NO_O = {'As': 3, 'Sb': 3, 'Bi': 3, 'Cu': 1, 'Fe': 2, 'Ni': 2, 'Co': 2, 'Mn': 2, 'Sn': 2, 'Ge': 2,
                   'Pb': 2, 'Tl': 1, 'Ag': 1, 'Hg': 2, 'Pd': 2, 'Pt': 2, 'Au': 1, 'In': 3, 'Ga': 3, 'Mo': 4,
                   'W': 4, 'V': 3, 'Cr': 3, 'Ti': 4, 'Zn': 2, 'Cd': 2}
# neighbour-search cutoff (Å) by cation element; 3.2 otherwise
CUTOFF = {'K': 3.6, 'Rb': 3.7, 'Cs': 3.9, 'Ba': 3.6, 'Sr': 3.4, 'Pb': 3.6, 'Tl': 3.6, 'NH': 3.4, 'Bi': 3.4,
          'Ag': 3.4, 'Na': 3.3, 'Ca': 3.3, 'La': 3.3, 'Ce': 3.3, 'Pr': 3.3, 'Nd': 3.3, 'Sm': 3.3, 'Eu': 3.3,
          'Gd': 3.3, 'Tb': 3.3, 'Dy': 3.3, 'Ho': 3.3, 'Er': 3.3, 'Tm': 3.3, 'Yb': 3.3, 'Lu': 3.3, 'Y': 3.3,
          'Hg': 3.4, 'Cd': 3.3, 'Th': 3.3, 'U': 3.3, 'H': 2.4}

def _element_of(label, type_symbol, known=None):
    """('Fe', 3) from 'Fe3+'; ('O', None) from 'OH1'/'OW2'/'Ow'; ('Li', None) from 'LiY'.
    `known` (the elements of _chemical_formula_sum) decides a bare 'W1': water unless the
    structure contains tungsten."""
    ox = None
    src = (type_symbol or '').strip()
    # 'Fe3+' (SHELXL), 'Bi+3' / 'S-2' (JANA), 'Fe+' / 'Cl-'
    m = re.match(r'^([A-Z][a-z]?)(?:(\d+)([+-])|([+-])(\d+)|([+-]))?', src)
    el = None
    if m and m.group(1) in ELEMENTS:
        el = m.group(1)
        if m.group(2):
            ox = int(m.group(2)) * (1 if m.group(3) == '+' else -1)
        elif m.group(5):
            ox = int(m.group(5)) * (1 if m.group(4) == '+' else -1)
        elif m.group(6):
            ox = 1 if m.group(6) == '+' else -1
    if el is None:
        lab = label.strip()
        if re.match(r'^(OH|OW|Ow|OA|Wat|Ow)', lab) or lab.upper().startswith('OH') or lab.upper().startswith('OW') \
                or (re.match(r'^W\d', lab) and 'W' not in (known or ())):
            el = 'O'
        else:
            m = re.match(r'^([A-Z][a-z]?)', lab)
            if m and m.group(1) in ELEMENTS:
                el = m.group(1)
            elif m and m.group(1)[0] in ELEMENTS:
                el = m.group(1)[0]
    if el == 'D':
        el = 'H'
    return el, ox

def _cell(block):
    it = block['items']
    vals = [_num(it.get('_cell_length_' + k)) for k in ('a', 'b', 'c')] + \
           [_num(it.get('_cell_angle_' + k)) for k in ('alpha', 'beta', 'gamma')]
    if any(v is None for v in vals):
        raise ValueError('the .cif has no complete unit cell')
    return vals

def _symops(block):
    tags, rows = _loop(block, '_space_group_symop_operation_xyz')
    if rows is None:
        tags, rows = _loop(block, '_symmetry_equiv_pos_as_xyz')
    if rows is not None:
        col = _col(tags, rows, '_space_group_symop_operation_xyz' if '_space_group_symop_operation_xyz' in tags
                   else '_symmetry_equiv_pos_as_xyz')
        return [parse_symop(s) for s in col], len(col)
    # no operator list: only P1 / P-1 can be handled without a space-group table
    sg = (block['items'].get('_space_group_name_h-m_alt') or block['items'].get('_symmetry_space_group_name_h-m')
          or block['items'].get('_space_group_name_h-m') or '').replace(' ', '')
    if sg in ('P1', 'P-1'):
        ops = [parse_symop('x, y, z')] + ([parse_symop('-x, -y, -z')] if sg == 'P-1' else [])
        return ops, len(ops)
    raise ValueError('the .cif lists no symmetry operators (_space_group_symop_operation_xyz) and the space '
                     'group %r cannot be expanded without them — add the operator loop (SHELXL, JANA and '
                     'CrysAlisPro write it)' % sg)

class Structure:
    def __init__(self, path, ox_override=None, include_h=True):
        blocks = [b for b in read_cif(path) if _loop(b, '_atom_site_fract_x')[1]]
        if not blocks:
            raise ValueError('no atom sites (_atom_site_fract_x) in %s' % os.path.basename(path))
        b = blocks[0]; self.block = b; self.path = path
        self.name = (b['items'].get('_chemical_name_mineral') or b['items'].get('_chemical_name_common')
                     or b['name'] or os.path.basename(path))
        self.formula = b['items'].get('_chemical_formula_sum', '')
        self.sg = (b['items'].get('_space_group_name_h-m_alt') or b['items'].get('_symmetry_space_group_name_h-m')
                   or b['items'].get('_space_group_name_h-m') or '?')
        self.cell = _cell(b)
        self.ops, self.n_ops = _symops(b)
        self._metric()
        self.notes = []
        self.include_h = include_h
        self._sites(ox_override or {})

    # -- geometry
    def _metric(self):
        a, b, c, al, be, ga = self.cell
        ca, cb, cg = (math.cos(math.radians(x)) for x in (al, be, ga))
        self.G = [[a * a, a * b * cg, a * c * cb], [a * b * cg, b * b, b * c * ca], [a * c * cb, b * c * ca, c * c]]
        self.volume = a * b * c * math.sqrt(1 - ca * ca - cb * cb - cg * cg + 2 * ca * cb * cg)
        # perpendicular widths of the cell along each axis (for the image range)
        sa, sb, sg = (math.sin(math.radians(x)) for x in (al, be, ga))
        self.widths = [self.volume / (b * c * sa), self.volume / (a * c * sb), self.volume / (a * b * sg)]

    def dist(self, p, q):
        d = [p[i] - q[i] for i in range(3)]
        s = sum(d[i] * self.G[i][j] * d[j] for i in range(3) for j in range(3))
        return math.sqrt(max(s, 0.0))

    # -- sites
    def _sites(self, ox_override):
        b = self.block
        tags, rows = _loop(b, '_atom_site_fract_x')
        labels = _col(tags, rows, '_atom_site_label')
        types = _col(tags, rows, '_atom_site_type_symbol', '')
        xs = [_num(v) for v in _col(tags, rows, '_atom_site_fract_x')]
        ys = [_num(v) for v in _col(tags, rows, '_atom_site_fract_y')]
        zs = [_num(v) for v in _col(tags, rows, '_atom_site_fract_z')]
        occ = [_num(v) if _num(v) is not None else 1.0 for v in _col(tags, rows, '_atom_site_occupancy', '1')]
        uiso = [_num(v) for v in _col(tags, rows, '_atom_site_u_iso_or_equiv', '.')]
        # oxidation numbers from the _atom_type loop, if any
        type_ox = {}
        ttags, trows = _loop(b, '_atom_type_oxidation_number')
        if trows:
            for sym, on in zip(_col(ttags, trows, '_atom_type_symbol'), _col(ttags, trows, '_atom_type_oxidation_number')):
                el, _ = _element_of(sym, sym)
                if el and _num(on) is not None:
                    type_ox[el] = int(round(_num(on)))
        known = set(re.findall(r'[A-Z][a-z]?', b['items'].get('_chemical_formula_sum', ''))) & ELEMENTS
        has_o = any(_element_of(l, t, known)[0] == 'O' for l, t in zip(labels, types))
        raw = []
        for i, lab in enumerate(labels):
            if None in (xs[i], ys[i], zs[i]):
                continue
            el, ox = _element_of(lab, types[i], known)
            if el is None:
                self.notes.append('site %s: element not recognised — skipped' % lab); continue
            if el == 'H' and not self.include_h:
                continue
            # the .cif's own coordinates, NOT wrapped into [0,1): its symmetry codes (n_klm) refer to
            # them, and every distance routine here already searches the neighbouring cells
            raw.append([lab, el, ox, [xs[i], ys[i], zs[i]], occ[i], uiso[i], types[i]])
        # ammonium: an N bonded to H is NH4+
        self.nh4 = set()
        for r in raw:
            if r[1] != 'N':
                continue
            has_h = any(o[1] == 'H' and self._min_dist(r[3], o[3]) < 1.15 for o in raw)
            bonded_o = any(o[1] == 'O' and self._min_dist(r[3], o[3]) < 1.5 for o in raw)
            if has_h or not bonded_o:
                self.nh4.add(r[0])          # ammonium — with or without its H refined
        # merge rows sharing a position into one (mixed) site
        merged = []
        for r in raw:
            for m in merged:
                if self._min_dist(m['frac'], r[3]) < 0.02:
                    m['rows'].append(r); break
            else:
                merged.append({'frac': r[3], 'rows': [r]})
        self.sites = []
        for m in merged:
            species = []
            for lab, el, ox, frac, oc, ui, ts in m['rows']:
                is_nh = lab in self.nh4
                if is_nh:
                    ox_final = 1
                elif ox is not None:
                    ox_final = ox
                elif el in ox_override:
                    ox_final = ox_override[el]
                elif el in type_ox:
                    ox_final = type_ox[el]
                elif el in ANION_ELEMENTS and not (el in ('S', 'Se', 'Te', 'N') and has_o) and el != 'H':
                    ox_final = ANION_OX[el]
                else:
                    ox_final = (DEFAULT_OX_NO_O.get(el) if not has_o else None) or DEFAULT_OX.get(el)
                    if ox_final is None:
                        self.notes.append('%s: no default valence for %s — give --ox %s=N' % (lab, el, el))
                    elif el not in ('O', 'F', 'Cl', 'Br', 'I', 'Si', 'Al', 'Mg', 'Ca', 'Na', 'K', 'H', 'C', 'P', 'B', 'Li', 'Be', 'Sr', 'Ba', 'Rb', 'Cs', 'Zn', 'Zr', 'Y', 'Sc', 'Ga', 'Ge', 'Th'):
                        self.notes.append('%s assumed %s%+d (no oxidation state in the .cif; --ox %s=N to change)' % (lab, el, ox_final, el))
                species.append(Species('NH' if is_nh else el, ox_final, oc))
            label = '/'.join(r[0] for r in m['rows'])
            el0 = species[0].element
            positions = self._equivalents(m['frac'])
            self.sites.append(Site(label, el0, m['frac'], species, positions, len(positions),
                                   sum(s.occ for s in species), m['rows'][0][5]))
        # H atoms of an ammonium group are part of the NH4+ cation, not separate donors
        if self.nh4:
            keep = []
            for s in self.sites:
                if s.element == 'H' and any(self._min_dist(s.frac, n.frac) < 1.15 for n in self.sites if n.label in self.nh4):
                    continue
                keep.append(s)
            self.sites = keep
        # H counts as a cation (donor) only when it has a proper O–H bond
        self.cations = [s for s in self.sites if self.is_cation(s)]
        self.anions = [s for s in self.sites if not self.is_cation(s)]

    def _min_dist(self, p, q):
        return min(self.dist(p, [q[0] + i, q[1] + j, q[2] + k])
                   for i in (-1, 0, 1) for j in (-1, 0, 1) for k in (-1, 0, 1))

    def is_cation(self, s):
        return any(sp.ox is not None and sp.ox > 0 for sp in s.species)

    def _equivalents(self, frac):
        out = []
        for op in self.ops:
            p = [v % 1.0 for v in _apply(op, frac)]
            p = [0.0 if abs(v - 1.0) < 1e-6 else v for v in p]
            if not any(self._same(p, q) for q in out):
                out.append(p)
        return out

    def _same(self, p, q):
        for dx in (-1, 0, 1):
            for dy in (-1, 0, 1):
                for dz in (-1, 0, 1):
                    if self.dist(p, [q[0] + dx, q[1] + dy, q[2] + dz]) < 0.05:
                        return True
        return False

    # -- neighbours
    def neighbours(self, site, cutoff):
        """[(other site, distance, count)] within cutoff of one representative of `site`,
        merged into distinct distances (0.0015 Å) with their multiplicity."""
        p0 = site.positions[0]
        rng = [int(math.ceil(cutoff / w)) + 1 for w in self.widths]
        found = []
        for other in self.sites:
            for q in other.positions:
                for i in range(-rng[0], rng[0] + 1):
                    for j in range(-rng[1], rng[1] + 1):
                        for k in range(-rng[2], rng[2] + 1):
                            d = self.dist(p0, [q[0] + i, q[1] + j, q[2] + k])
                            if 0.3 < d <= cutoff:
                                found.append((other, d))
        found.sort(key=lambda x: (x[0].label, x[1]))
        merged = []
        for other, d in found:
            if merged and merged[-1][0] is other and abs(merged[-1][1] - d) < 0.0015:
                merged[-1][2] += 1
            else:
                merged.append([other, d, 1])
        return [(o, d, n) for o, d, n in merged]

    def images(self, p0, cutoff, sites=None):
        """Every atom image within cutoff of the fractional position p0: [(site, distance,
        position)] by distance — the positions the hydrogen-bond geometry needs."""
        rng = [int(math.ceil(cutoff / w)) + 1 for w in self.widths]
        out = []
        for other in (sites if sites is not None else self.sites):
            for q in other.positions:
                for i in range(-rng[0], rng[0] + 1):
                    for j in range(-rng[1], rng[1] + 1):
                        for k in range(-rng[2], rng[2] + 1):
                            qq = [q[0] + i, q[1] + j, q[2] + k]
                            d = self.dist(p0, qq)
                            if 0.3 < d <= cutoff:
                                out.append((other, d, qq))
        out.sort(key=lambda x: (x[1], x[0].label))
        return out

    def site(self, label):
        """The site holding an atom label ('OH1' finds the merged site 'F1/OH1'); None if absent."""
        for s in self.sites:
            if label == s.label or label in s.label.split('/'):
                return s
        return None

# ----------------------------------------------------------------------------- parameters

PARAM_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'bvparm2020.cif')
# A neighbour within the cutoff counts as a bond only when it contributes at least this much:
# keeps Pb's long bonds (3.4 Å ≈ 0.03 vu) and drops a 3.2 Å P–O (0.015 vu) that no paper lists.
MIN_S = 0.025
PREFER = {'gh': ['bs', 'a', 'b'], 'bo': ['b', 'a', 'bs'], 'ba': ['a', 'b', 'bs']}
# per-cation preferences that override the set: U6+–O from Burns, Ewing & Hawthorne (1997) — the
# parameters every uranyl-mineral description uses (--params still applies to everything else)
PREFER_CATION = {('U', 6): ['r']}
# how the table note cites a reference id (the file's strings carry the journal and page numbers)
SHORT_REF = {'a': 'Brown and Altermatt (1985)', 'b': "Brese and O'Keeffe (1991)", 'bs': 'Gagné and Hawthorne (2015)',
             'r': 'Burns et al. (1997)', 's': 'García-Rodríguez et al. (2000)', 'bc': 'Brown (2002)', 'o': 'Allmann (1975)',
             'p': 'Zachariasen (1978)', 'e': 'Brown (unpublished)', 'q': 'Krivovichev and Brown (2001)',
             'af': 'Locock and Burns (2004)', 'bh': 'Brown (2009)', 'bj': 'Krivovichev (2012)', 'bk': 'Mills and Christy (2013)'}
H_RANGES = [(1.05, 0.907, 0.28), (1.70, 0.569, 0.94), (99.0, 0.990, 0.59)]   # Brown (2002) O–H by distance

class Params:
    def __init__(self, path=PARAM_FILE, prefer='gh', u6='burns'):
        """prefer: 'gh' | 'bo' | 'ba'; u6: 'burns' (U6+–O from Burns et al. 1997, the uranyl
        convention) or 'params' (U6+ from the chosen set like every other cation)."""
        self.table = {}; self.refs = {}
        blocks = read_cif(path)
        for b in blocks:
            tags, rows = _loop(b, '_valence_ref_id')
            if rows:
                for i, r in zip(_col(tags, rows, '_valence_ref_id'), _col(tags, rows, '_valence_ref_reference')):
                    self.refs[i] = r
            tags, rows = _loop(b, '_valence_param_ro')
            if rows:
                for r in rows:
                    d = dict(zip(tags, r))
                    key = (d['_valence_param_atom_1'], int(d['_valence_param_atom_1_valence']),
                           d['_valence_param_atom_2'], int(d['_valence_param_atom_2_valence']))
                    self.table.setdefault(key, []).append((float(d['_valence_param_ro']), float(d['_valence_param_b']),
                                                           d['_valence_param_ref_id'], d.get('_valence_param_details', '')))
        self.prefer = PREFER.get(prefer, PREFER['gh'])
        self.prefer_key = prefer if prefer in PREFER else 'gh'
        self.u6 = u6
        self.used = OrderedDict()

    def short_ref(self, rid):
        """'Gagné and Hawthorne (2015)' for a reference id — the set's own name for a value the
        chosen set reprints (Brese & O'Keeffe 1991 carries every Brown & Altermatt 1985 value)."""
        if rid == 'a' and self.prefer_key == 'bo':
            rid = 'b'
        if rid in SHORT_REF:
            return SHORT_REF[rid]
        txt = self.refs.get(rid, rid)
        m = re.match(r'^\s*(.+?)[,.]?\s*\(?(\d{4})\)?', txt)
        if not m:
            return txt
        authors = [a.strip(' .') for a in re.split(r',| and |&', m.group(1)) if a.strip(' .') and 'et al' not in a]
        if 'et al' in m.group(1) or len(authors) > 2:
            who = authors[0] + ' et al.'
        else:
            who = ' and '.join(authors)
        return '%s (%s)' % (who, m.group(2))

    def note(self):
        """The journal-style attribution of every parameter used: 'Bond-valence parameters from
        Gagné and Hawthorne (2015); U6+–O from Burns et al. (1997); NH4+–O from García-Rodríguez
        et al. (2000)' (hydrogen is cited with the hydrogen bonds)."""
        by_ref = OrderedDict()
        for (cat, cox, an), (r0, b, rid) in self.used.items():
            if cat == 'H':
                continue
            ion = 'NH4+' if cat == 'NH' else cat + ('' if cox == 9 else ('%d+' % cox if cox > 1 else '+'))
            by_ref.setdefault(self.short_ref(rid), []).append('%s–%s' % (ion, an))
        if not by_ref:
            return ''
        main = max(by_ref, key=lambda k: len(by_ref[k]))
        parts = ['Bond-valence parameters from %s' % main]
        for ref, pairs in by_ref.items():
            if ref != main:
                parts.append('%s from %s' % (_join(pairs), ref))
        return '; '.join(parts)

    def get(self, cat, cox, an, aox):
        """(R0, b, ref_id) or None. The preferred reference first; then any; then the
        oxidation-state-unspecified (9) entry."""
        for key in ((cat, cox, an, aox), (cat, 9, an, aox), (cat, cox, an, 9), (cat, 9, an, 9)):
            rows = self.table.get(key)
            if not rows:
                continue
            for ref in (PREFER_CATION.get((cat, cox), []) if self.u6 == 'burns' else []) + self.prefer:
                for r0, b, rid, det in rows:
                    if rid == ref and 'unchecked' not in det:
                        self.used[(cat, cox, an)] = (r0, b, rid); return r0, b, rid
            r0, b, rid, det = rows[0]
            self.used[(cat, cox, an)] = (r0, b, rid); return r0, b, rid
        return None

    def max_length(self, cat, cox, an, aox, min_s=MIN_S):
        """The distance at which the valence drops to min_s (None without parameters)."""
        p = self.get(cat, cox, an, aox)
        return None if p is None else p[0] - p[1] * math.log(min_s)

    def valence(self, cat, cox, an, aox, R):
        if cat == 'H':
            for lim, r0, b in H_RANGES:
                if R < lim:
                    self.used[('H', 1, an)] = (r0, b, 'bc'); return math.exp((r0 - R) / b)
        p = self.get(cat, cox, an, aox)
        if p is None:
            return None
        return math.exp((p[0] - R) / p[1])

def _join(items):
    items = list(items)
    return items[0] if len(items) == 1 else ', '.join(items[:-1]) + ' and ' + items[-1]

# ----------------------------------------------------------------------------- the calculation

def _species_ox(site):
    return sorted({sp.ox for sp in site.species if sp.ox is not None})

def compute(st, params, cutoff=None, hbond='oo', hmax=None, donors=None, force=None):
    """[(cation site, [Bond], bvs, expected, mean_d)], the anion sums, the table cells and the
    hydrogen bonds ([HBond], notes in st.notes).
    hbond: 'oo' — strengths from the donor–acceptor O···O distances (Ferraris & Ivaldi 1988); H
           sites are not cation columns;
           'h'  — H as a cation: acceptor valences from the H···O distances (Brown 2002), the
           donor gets 1 − Σ (the older convention; every H sums to 1 vu);
           'none'.
    Anion sums are cations + accepted hydrogen bonds (the donated O–H valence is reported with
    the hydrogen bonds, neither deducted nor added, as the owner's tables do); in 'h' mode the H
    columns carry it as any cation."""
    result = []
    for c in st.cations:
        if c.element == 'H' and hbond != 'h':
            continue
        cut = cutoff or max(CUTOFF.get(sp.element, 3.2) for sp in c.species)
        if c.element == 'H':
            cut = cutoff or CUTOFF['H']
        bonds = []
        for other, d, n in st.neighbours(c, cut):
            if st.is_cation(other):
                continue
            vals = []
            an_species = [a for a in other.species if a.ox is not None and a.ox < 0] or [Species(other.element, ANION_OX.get(other.element, -2), 1.0)]
            an_tot = sum(a.occ for a in an_species) or 1.0
            for sp in c.species:
                if sp.ox is None or sp.ox <= 0:
                    continue
                s = 0.0; missing = False
                for a in an_species:                    # 'F1/OH1': occupancy-weighted over the anions
                    v = params.valence(sp.element, sp.ox, a.element, a.ox, d)
                    if v is None:
                        missing = True; continue
                    s += (a.occ / an_tot) * v
                vals.append((sp, None if missing and s == 0.0 else s))
            if vals and any(s is not None for _, s in vals):
                if c.element != 'H' and cutoff is None and max((s or 0.0) for _, s in vals) < MIN_S:
                    continue                        # too weak to be a bond in a published table (an
                                                    # explicit --cutoff means: everything within it)
                bonds.append(Bond(c, other, d, n, vals))
        if c.element == 'H' and bonds:
            # X-ray O–H distances are short and unreliable: take the acceptor valences from the
            # H···O distances and give the donor the rest, so every H sums to exactly 1 vu
            donor = min(bonds, key=lambda b: b.dist)
            acc = sum((b.vals[0][1] or 0.0) * b.count for b in bonds if b is not donor)
            if acc > 0.6:
                st.notes.append('%s: acceptor bonds sum to %.2f vu — a very short/symmetric hydrogen bond?' % (c.label, acc))
            bonds = [Bond(b.cation, b.anion, b.dist, b.count, [(b.vals[0][0], max(1.0 - acc, 0.0))]) if b is donor else b
                     for b in bonds]
        tot_occ = sum(sp.occ for sp in c.species if sp.ox and sp.ox > 0) or 1.0
        def site_val(b):                     # occupancy-fraction-weighted valence of one bond,
            aw = min(b.anion.occ_total, 1.0)  # scaled by the anion's occupancy (a half-occupied
            return aw * sum((sp.occ / tot_occ) * (s or 0.0) for sp, s in b.vals)   # O counts half)
        bvs = sum(site_val(b) * b.count for b in bonds)
        expected = sum((sp.occ / tot_occ) * sp.ox for sp in c.species if sp.ox and sp.ox > 0)
        ncoord = sum(b.count for b in bonds)
        mean_d = sum(b.dist * b.count for b in bonds) / ncoord if ncoord else None
        result.append((c, bonds, bvs, expected, mean_d))
    # anion sums: weight by the cation site's occupancy; multiplicity from the site multiplicities
    anion_sum = {a.label: 0.0 for a in st.anions}
    cells = {}                                  # (anion label, cation label) -> [(s, n_down, n_across)]
    for c, bonds, bvs, expected, mean_d in result:
        tot_occ = sum(sp.occ for sp in c.species if sp.ox and sp.ox > 0) or 1.0
        for b in bonds:
            # the anion's own row: the bond as the anion receives it when present (species-weighted
            # over a mixed cation, NOT scaled by the anion's own occupancy — that scaling belongs to
            # the cation's sum, where a half-occupied O is there half the time)
            s_site = sum((sp.occ / tot_occ) * (s or 0.0) for sp, s in b.vals)
            n_across = b.count * c.mult / b.anion.mult
            n_across_r = int(round(n_across)) if abs(n_across - round(n_across)) < 0.02 else n_across
            occ_weight = min(tot_occ, 1.0)
            anion_sum[b.anion.label] += s_site * n_across * occ_weight
            cells.setdefault((b.anion.label, c.label), []).append((s_site, b.count, n_across_r))
    hbonds = estimate_hbonds(st, result, dict(anion_sum), params, hbond, hmax, donors, force)
    if hbond != 'h':
        for hb in hbonds:
            anion_sum[hb.acceptor.label] += hb.s * hb.n_across
    return result, anion_sum, cells, hbonds

# ----------------------------------------------------------------------------- hydrogen bonds

# Ferraris & Ivaldi (1988), Acta Cryst. B44, 341–344: the H···O valence of an O–H···O bond from
# the O···O distance, s = (d/2.17)^−8.2 + 0.06 — the relation mineral descriptions cite
# ("hydrogen-bond strengths based on O–O bond lengths"); reproduces the owner's tables to 0.01 vu.
FI_R0, FI_N, FI_K = 2.17, 8.2, 0.06
HB_MAX = 3.2            # longest O···O counted as a hydrogen bond (the usual limit in the tables)
HB_MIN = 2.45           # shorter than any O–H···O
HB_ANGLE = (70.0, 150.0)   # A···O···A for the two H of a water molecule (H–O–H ≈ 105°)

HBond = namedtuple('HBond', 'donor acceptor d s n_donor n_across via h')
#   donor/acceptor: Site;  d: D···A (Å; H···A in 'h' mode);  s: valence (vu) the acceptor receives;
#   n_donor: bonds per donor atom;  n_across: bonds per acceptor atom;  via: 'loop' | 'H' | 'OO' |
#   'H···A';  h: the H label ('' without one)

def fi_valence(d):
    """Ferraris & Ivaldi (1988): s(H···O) from the O···O distance."""
    return (d / FI_R0) ** -FI_N + FI_K

def label_h(lab):
    """1 for a hydroxyl label, 2 for a water label, 0 otherwise: OH1, Oh2, OW1, Ow3, W4, Wat1,
    O6H, O7W, H2O1 (a merged 'F1/OH1' counts by its OH part)."""
    for part in lab.split('/'):
        if re.match(r'^(OH|Oh)\d*[a-zA-Z]?$', part) or re.match(r'^O\d+[Hh]$', part):
            return 1
        if re.match(r'^(OW|Ow|Wat|W\d|H2O|Hw|OH2)', part) or re.match(r'^O\d+[Ww]$', part):
            return 2
    return 0

def _ratio(a, b):
    r = a / b if b else 1.0
    return int(round(r)) if abs(r - round(r)) < 0.02 else round(r, 2)

def hbond_geometry(st, include_long=True):
    """The D–H⋯A geometry: from the _geom_hbond loop when the .cif has one, else computed from
    the H positions (H⋯A ≤ 2.6 Å, ∠DHA ≥ 110°). [{'D','H','A': Site, 'dh','ha','da','ang': text,
    'da_val','dh_val': float, 'code': str, 'from_loop': bool}] — used by the hydrogen-bond table
    and by the hydrogen-bond valences."""
    out = []
    tags, rows = _loop(st.block, '_geom_hbond_distance_ha')
    if rows:
        g = lambda t, d='': _col(tags, rows, t, d)
        for D, H, A, dh, ha, da, ang, sy in zip(g('_geom_hbond_atom_site_label_d'), g('_geom_hbond_atom_site_label_h'),
                                              g('_geom_hbond_atom_site_label_a'), g('_geom_hbond_distance_dh'),
                                              g('_geom_hbond_distance_ha'), g('_geom_hbond_distance_da'),
                                              g('_geom_hbond_angle_dha'), g('_geom_hbond_site_symmetry_a', '.')):
            sD, sH, sA = st.site(D), st.site(H), st.site(A)
            if sD is None or sA is None:
                continue
            out.append({'D': sD, 'H': sH, 'A': sA, 'dh': dh.strip(), 'ha': ha.strip(), 'da': da.strip(),
                        'ang': re.sub(r'\.000\(0\)$|\.0\(0\)$', '', ang.strip()), 'da_val': _num(da), 'dh_val': _num(dh),
                        'code': sy.strip(), 'from_loop': True, 'labels': (D, H, A)})
        return out
    for h in [x for x in st.sites if x.element == 'H']:
        p_h = h.positions[0]
        near = [(o, d, q) for o, d, q in st.images(p_h, 2.6, [x for x in st.sites if x.element in ('O', 'F', 'N', 'Cl')])]
        if not near or near[0][1] > 1.25:
            continue
        D, d_dh, p_d = near[0]
        for A, d_ha, p_a in near[1:]:
            ang = angle(st, p_d, p_h, p_a)
            if ang < 110:
                continue
            d_da = st.dist(p_d, p_a)
            out.append({'D': D, 'H': h, 'A': A, 'dh': '%.2f' % d_dh, 'ha': '%.2f' % d_ha, 'da': '%.3f' % d_da, 'ang': '%.0f' % ang,
                        'da_val': d_da, 'dh_val': d_dh, 'code': _code_of(st, A, p_a), 'from_loop': False, 'labels': (D.label, h.label, A.label)})
    return out

def _code_of(st, site, pos):
    """The symmetry code 'n_klm' that carries `site`'s listed coordinates to `pos`."""
    for n, op in enumerate(st.ops):
        q = _apply(op, site.frac)
        t = [pos[i] - q[i] for i in range(3)]
        if all(abs(v - round(v)) < 1e-3 for v in t):
            t = [int(round(v)) for v in t]
            return '.' if n == 0 and t == [0, 0, 0] else '%d_%d%d%d' % (n + 1, 5 + t[0], 5 + t[1], 5 + t[2])
    return '.'

def hbond_donors(st, sums, override=None):
    """{anion label: number of H} and how each was decided. Labels first (OH*, OW*, W*, O6H …);
    when nothing is labelled, the O sites short of bond valence (Σcat < 0.75 → water, < 1.5 →
    hydroxyl). `override` ({label: n}) wins."""
    out = OrderedDict(); how = {}
    for a in st.anions:
        if a.element != 'O':
            continue
        if override and a.label in override:
            if override[a.label]:
                out[a.label] = override[a.label]; how[a.label] = 'given'
            continue
        for lab in a.label.split('/'):
            if override and lab in override:
                if override[lab]:
                    out[a.label] = override[lab]; how[a.label] = 'given'
                break
        else:
            n = label_h(a.label)
            if n:
                out[a.label] = n; how[a.label] = 'label'
    if not out:                                     # nothing labelled or given: the valence deficit decides
        for a in st.anions:
            if a.element != 'O' or a.label in out or (override and a.label in override):
                continue
            s = sums.get(a.label, 0.0)
            if s < 0.75:
                out[a.label] = 2; how[a.label] = 'Σcat %.2f' % s
            elif s < 1.5:
                out[a.label] = 1; how[a.label] = 'Σcat %.2f' % s
    return out, how

def estimate_hbonds(st, result, sums, params, mode='oo', hmax=None, donors=None, force=None):
    """[HBond] for the structure. 'oo': from the D–H⋯A geometry when the .cif has H atoms (the
    loop, else the H positions), otherwise proposed from the O···O geometry; strengths from the
    O···O distance (Ferraris & Ivaldi 1988). 'h': the H-as-cation bonds already in `result`.
    Notes go to st.notes."""
    hmax = hmax or HB_MAX
    if mode == 'none':
        return []
    if mode == 'h':
        out = []
        for h, bonds, *_ in result:
            if h.element != 'H' or not bonds:
                continue
            donor = min(bonds, key=lambda b: b.dist)
            for b in bonds:
                if b is donor:
                    continue
                s = (b.vals[0][1] or 0.0) * min(h.occ_total, 1.0)
                out.append(HBond(donor.anion, b.anion, b.dist, s, _ratio(b.count * h.mult, donor.anion.mult),
                                 _ratio(b.count * h.mult, b.anion.mult), 'H···A', h.label))
        return out
    hs = [x for x in st.sites if x.element == 'H']
    out = []
    if hs and not force:
        seen_long = []
        for g in hbond_geometry(st):
            if g['da_val'] is None:
                continue
            if g['da_val'] > hmax:
                seen_long.append('%s–%s⋯%s %.2f Å' % (g['labels'][0], g['labels'][1], g['labels'][2], g['da_val']))
                continue
            if g['dh_val'] is not None and g['dh_val'] > 1.3:
                continue
            H = g['H']
            occ = min(H.occ_total, 1.0) if H is not None else 1.0
            s = fi_valence(g['da_val']) * occ
            m_h = H.mult if H is not None else g['D'].mult
            out.append(HBond(g['D'], g['A'], g['da_val'], s, _ratio(m_h, g['D'].mult), _ratio(m_h, g['A'].mult),
                             'loop' if g['from_loop'] else 'H', H.label if H is not None else ''))
        if seen_long:
            st.notes.append('hydrogen bonds longer than %.2f Å not counted: %s' % (hmax, ', '.join(seen_long)))
        return out
    return _hbonds_blind(st, sums, params, hmax, donors, force)

class _Contact:
    """One class of O···O contacts between two sites at one distance, shared by both ends: k1
    contacts per atom of site 1, k2 per atom of site 2 (k1·mult1 = k2·mult2)."""
    def __init__(self, d, ends):
        self.d = d; self.ends = ends; self.imgs = {}; self.k = {}; self.donor = None; self.m = 0
    def add(self, lab, imgs):
        self.imgs[lab] = imgs; self.k[lab] = len(imgs)
    def other(self, lab):
        a, b = self.ends
        return b if lab == a else a

def _hbonds_blind(st, sums, params, hmax, donors_over, force):
    """No H atoms: propose the hydrogen bonds from the O···O geometry. Donors from the labels /
    valence deficits (hbond_donors); for each, the O within hmax that (1) is no polyhedral edge
    (no cation bonded to both), (2) lies ≥ 80° from every cation bonded to the donor, (3) can
    still accept (Σcat + accepted below its target), taken shortest first, each contact used
    once; when two donors touch, the one with the larger acceptor deficit accepts; a water's two
    acceptors subtend 70–150°. Forced pairs (--hb OW1>O2) are placed first, as given."""
    nH, how = hbond_donors(st, sums, donors_over)
    force = list(force or [])
    for D, A in force:
        sD = st.site(D)
        if sD is None:
            st.notes.append('--hb %s>%s: no site %s' % (D, A, D)); continue
        if sD.label not in nH:
            nH[sD.label] = 0; how[sD.label] = 'given'
        nH[sD.label] = max(nH[sD.label], sum(1 for d2, _ in force if st.site(d2) is sD))
    if not nH:
        return []
    deficit = [l for l, w in how.items() if w.startswith('Σcat')]
    if deficit:
        st.notes.append('no OH/OW labels — donors taken from the bond-valence deficits: %s (--donors to correct)'
                        % ', '.join('%s (%d H, %s)' % (l, nH[l], how[l]) for l in deficit))
    sites = {s.label: s for s in st.sites}
    o_sites = [x for x in st.anions if x.element == 'O']
    # cation environment of each donor atom (first equivalent)
    env = {}
    for lab in nH:
        D = sites[lab]; pD = D.positions[0]; lst = []
        for M, d, q in st.images(pD, 3.9, [c for c in st.cations if c.element != 'H']):
            sp = [x for x in M.species if x.ox and x.ox > 0]
            if not sp:
                continue
            s = max((params.valence(x.element, x.ox, 'O', -2, d) or 0.0) for x in sp)
            if s >= MIN_S:
                lst.append((M, d, q, s, sp))
        env[lab] = lst
    # contact classes per donor
    contacts = {}; classes = {lab: [] for lab in nH}
    for lab in nH:
        D = sites[lab]; pD = D.positions[0]
        groups = OrderedDict()
        for A, d, q in st.images(pD, hmax, o_sites):
            if d < HB_MIN:
                continue
            forced = any(st.site(x) is D and st.site(y) is A for x, y in force)
            if not forced:
                edge = False
                for M, dm, qm, s, sp in env[lab]:
                    dma = st.dist(qm, q)
                    if dma < 3.9 and max((params.valence(x.element, x.ox, 'O', -2, dma) or 0.0) for x in sp) >= MIN_S:
                        edge = True; break
                if edge:
                    continue
                if any(s >= 0.1 and angle(st, qm, pD, q) < 80 for M, dm, qm, s, sp in env[lab]):
                    continue
            groups.setdefault((A.label, round(d, 3)), []).append((A, d, q))
        for (al, dr), imgs in groups.items():
            key = (frozenset([lab, al]), round(dr, 2))
            ct = contacts.get(key)
            if ct is None:
                ct = contacts[key] = _Contact(dr, (lab, al))
            ct.add(lab, imgs)
            classes[lab].append(ct)
    accepted = {}; left = dict(nH); placed = {lab: [] for lab in nH}
    def target(lab):
        return 2.0 - 0.8 * nH.get(lab, 0)
    def need(lab):
        return target(lab) - sums.get(lab, 0.0) - accepted.get(lab, 0.0)
    def ok_angle(lab, q, others):
        return all(HB_ANGLE[0] <= angle(st, q0, sites[lab].positions[0], q) <= HB_ANGLE[1] for _, _, q0 in others)
    def pick_images(lab, ct):
        """The images of a class to place at once: all of them when the donor has H for them
        (symmetry-related H point at symmetry-related acceptors), else one."""
        free = [im for im in ct.imgs[lab] if not any(im[2] is q0 for _, _, q0 in placed[lab])]
        free = [im for im in free if ok_angle(lab, im[2], placed[lab])]
        if not free:
            return []
        if len(free) > 1 and len(free) <= left[lab] and all(ok_angle(lab, a[2], [b]) for i, a in enumerate(free) for b in free[:i]):
            return free
        return free[:1]
    def assign(lab, ct, imgs):
        for A, d, q in imgs:
            ct.donor = lab; ct.m += 1; placed[lab].append((A, d, q)); left[lab] -= 1
            accepted[A.label] = accepted.get(A.label, 0.0) + fi_valence(d) * sites[lab].mult / A.mult * min(sites[lab].occ_total, 1.0)
    # forced pairs first
    for D, A in force:
        sD, sA = st.site(D), st.site(A)
        if sD is None or sA is None:
            st.notes.append('--hb %s>%s: no site %s' % (D, A, A if sD is not None else D)); continue
        cts = [ct for ct in classes[sD.label] if ct.other(sD.label) == sA.label and ct.m < ct.k[sD.label]]
        if not cts:
            st.notes.append('--hb %s>%s: no O···O contact within %.2f Å' % (D, A, hmax)); continue
        ct = min(cts, key=lambda c: c.d)
        assign(sD.label, ct, [ct.imgs[sD.label][ct.m]])
    for relaxed in (False, True):
        while True:
            best = None
            for lab in nH:
                if left[lab] <= 0:
                    continue
                for ct in classes[lab]:
                    if ct.donor not in (None, lab) or ct.m >= ct.k[lab]:
                        continue
                    al = ct.other(lab)
                    if al == lab and ct.m >= ct.k[lab] // 2:
                        continue                          # a site bonded to itself: half donate, half accept
                    if need(al) < -0.15:
                        continue                          # the acceptor is already saturated
                    if not relaxed and al != lab and al in nH and left[al] > 0 and ct.donor is None and need(al) < need(lab) - 0.02:
                        continue                          # the other end has the larger deficit: it accepts
                    imgs = pick_images(lab, ct)
                    if not imgs:
                        continue
                    cand = (ct.d, -left[lab], lab, ct, imgs)
                    if best is None or cand[:2] < best[:2]:
                        best = cand
            if best is None:
                break
            assign(best[2], best[3], best[4])
    for lab in nH:
        if left[lab] > 0:
            st.notes.append('%s: %d of %d H without an acceptor within %.2f Å (no O outside its own polyhedra with room for one)'
                            % (lab, left[lab], nH[lab], hmax))
    out = []
    for lab in nH:
        D = sites[lab]
        by = OrderedDict()
        for A, d, q in placed[lab]:
            by.setdefault((A.label, round(d, 3)), [A, d, 0]); by[(A.label, round(d, 3))][2] += 1
        for (al, dr), (A, d, m) in by.items():
            out.append(HBond(D, A, d, fi_valence(d) * min(D.occ_total, 1.0), m, _ratio(m * D.mult, A.mult), 'OO', ''))
    if any(sites[l].occ_total < 0.999 for l in nH):
        st.notes.append('hydrogen-bond valences of the partly occupied donors (%s) are multiplied by their occupancies'
                        % ', '.join(l for l in nH if sites[l].occ_total < 0.999))
    return out

def hbond_lines(hbonds):
    """The hydrogen bonds as report lines, with the donated valence per donor."""
    L = []
    for hb in hbonds:
        how = {'loop': 'from the .cif hydrogen-bond loop', 'H': 'from the H positions', 'OO': 'O···O geometry, H not located',
               'H···A': 'from the H···A distance'}[hb.via]
        L.append('  %-8s → %-8s %s %.3f Å   s %.2f vu%s%s   (%s)' % (hb.donor.label, hb.acceptor.label,
                 'H⋯A' if hb.via == 'H···A' else 'D⋯A', hb.d, hb.s, ' ×%s per donor' % hb.n_donor if hb.n_donor != 1 else '',
                 ' ×%s→' % hb.n_across if hb.n_across != 1 else '', (hb.h + ', ' if hb.h else '') + how))
    return L

def donated(hbonds):
    """{donor label: Σ of the O–H valences (1 − s) it keeps} — the part the tables do not add."""
    out = OrderedDict()
    for hb in hbonds:
        out[hb.donor.label] = out.get(hb.donor.label, 0.0) + (1.0 - hb.s) * hb.n_donor
    return out

# ----------------------------------------------------------------------------- report

def _fmt_species(site):
    return '/'.join('%s%s' % (sp.element, ('%+d' % sp.ox).replace('+', '+') if sp.ox is not None else '?')
                    for sp in site.species)

def _mark(n_down, n_across):
    m = ''
    if n_down and n_down != 1:
        m += '×%s↓' % n_down
    if n_across and n_across != 1:
        m += '×%s→' % n_across
    return m

def report_text(st, params, result, anion_sum, cells, geom_check=None, hbonds=()):
    L = []
    a, b, c, al, be, ga = st.cell
    L.append('Bond-valence check — %s' % os.path.basename(st.path))
    L.append('  %s   %s   %s' % (st.name, st.formula, st.sg))
    L.append('  a %.4f  b %.4f  c %.4f  α %.3f  β %.3f  γ %.3f  V %.2f Å³   %d symmetry operators'
             % (a, b, c, al, be, ga, st.volume, st.n_ops))
    for n in st.notes:
        L.append('  note: %s' % n)
    L.append('')
    L.append('BOND DISTANCES AND VALENCES (per cation site; s = exp((R0 − R)/b))')
    for site, bonds, bvs, expected, mean_d in result:
        L.append('')
        L.append('  %-12s %-14s mult %-3d occ %.3f' % (site.label, _fmt_species(site), site.mult, site.occ_total))
        for bd in bonds:
            s = sum((sp.occ / (sum(x.occ for x in site.species if x.ox and x.ox > 0) or 1)) * (v or 0.0) for sp, v in bd.vals)
            aw = min(bd.anion.occ_total, 1.0)
            L.append('    –%-10s %-5s %8.3f Å   %6.3f vu%s' % (bd.anion.label, ('×%d' % bd.count) if bd.count > 1 else '', bd.dist, s * aw,
                                                            '  (×%.2f occ)' % aw if aw < 0.999 else ''))
        if mean_d is not None:
            dev = (bvs - expected) / expected * 100 if expected else 0
            flag = '   ◄ check' if abs(dev) > 12 else ''
            L.append('    <%s–X> = %.3f Å (CN %d)      Σ = %.2f vu   expected %.2f  (%+.0f%%)%s'
                     % (site.label, mean_d, sum(bd.count for bd in bonds), bvs, expected, dev, flag))
    L.append('')
    hbonds = list(hbonds)
    if hbonds:
        L.append('HYDROGEN BONDS (s from the O···O distance, Ferraris & Ivaldi 1988: s = (d/2.17)^−8.2 + 0.06)'
                 if hbonds[0].via != 'H···A' else 'HYDROGEN BONDS (acceptor valences from the H···A distances, Brown 2002)')
        L.extend(hbond_lines(hbonds))
        L.append('')
    acc = {}; don = donated(hbonds) if hbonds and hbonds[0].via != 'H···A' else {}
    for hb in hbonds:
        acc.setdefault(hb.acceptor.label, []).append('%.2f%s' % (hb.s, '×%s→' % hb.n_across if hb.n_across != 1 else ''))
    L.append('BOND-VALENCE TABLE (vu; ×n↓ = counted n times in the column sum, ×n→ = n times in the row sum)')
    cats = [r[0] for r in result]
    anions = [an for an in st.anions if any((an.label, ct.label) in cells for ct in cats) or an.label in acc]
    w = max([len(an.label) for an in anions] + [8])
    head = ' ' * (w + 2) + ''.join('%-14s' % ct.label for ct in cats)
    if acc:
        head += '%-14s' % 'H bonds'
    head += '  Σan'
    if don:
        head += '   O–H    Σall'
    L.append('  ' + head)
    for an in anions:
        row = '  %-*s' % (w + 2, an.label)
        for ct in cats:
            vals = cells.get((an.label, ct.label))
            if not vals:
                row += '%-14s' % '–'
            else:
                txt = ', '.join('%.2f%s' % (s, _mark(nd, na)) for s, nd, na in vals)
                row += '%-14s' % txt
        if acc:
            row += '%-14s' % (', '.join(acc[an.label]) if an.label in acc else '–')
        row += '  %.2f' % anion_sum[an.label]
        if don:
            row += '   %-6s %.2f' % ('%.2f' % don[an.label] if an.label in don else '–', anion_sum[an.label] + don.get(an.label, 0.0))
        L.append(row)
    sums = '  %-*s' % (w + 2, 'Σ')
    for site, bonds, bvs, expected, mean_d in result:
        sums += '%-14s' % ('%.2f' % bvs)
    L.append(sums)
    exp = '  %-*s' % (w + 2, 'expected')
    for site, bonds, bvs, expected, mean_d in result:
        exp += '%-14s' % ('%.2f' % expected)
    L.append(exp)
    L.append('')
    if don:
        L.append('  (Σan = cations + accepted hydrogen bonds, as the tables print it; O–H = the donor\'s own Σ(1 − s); Σall = both)')
    L.append('')
    L.append('PARAMETERS USED (R0, b, reference)')
    for (cat, cox, an), (r0, bb, rid) in params.used.items():
        L.append('  %-4s%+d – %-3s  R0 %.3f  b %.3f   %s' % (cat, cox, an, r0, bb, params.refs.get(rid, rid)))
    L.append('  table note: %s.' % params.note())
    if geom_check:
        L.append('')
        L.append(geom_check)
    return '\n'.join(L)

def geom_self_check(st, result):
    """Compare the computed distances with the .cif's own _geom_bond loop."""
    tags, rows = _loop(st.block, '_geom_bond_distance')
    if not rows:
        return 'SELF-CHECK: the .cif has no _geom_bond loop to compare against.'
    l1 = _col(tags, rows, '_geom_bond_atom_site_label_1'); l2 = _col(tags, rows, '_geom_bond_atom_site_label_2')
    dd = [_num(v) for v in _col(tags, rows, '_geom_bond_distance')]
    computed = {}; reach = {}
    for site, bonds, *_ in result:
        cut = max(CUTOFF.get(sp.element, 3.2) for sp in site.species) if site.element != 'H' else CUTOFF['H']
        for lab in site.label.split('/'):
            reach[lab] = cut
        for b in bonds:
            for lab in site.label.split('/'):
                for alab in b.anion.label.split('/'):
                    computed.setdefault((lab, alab), []).append(b.dist)
    n = 0; worst = 0.0; bad = []
    for x, y, d in zip(l1, l2, dd):
        if d is None:
            continue
        cand = computed.get((x, y)) or computed.get((y, x))
        if not cand or d > reach.get(x, reach.get(y, 3.2)) + 1e-6:
            continue
        delta = min(abs(d - c) for c in cand); n += 1; worst = max(worst, delta)
        if delta > (0.01 if 'H' in (x[:1], y[:1]) and (x.startswith('H') or y.startswith('H')) else 0.003):
            bad.append('%s–%s: .cif loop %.4f, computed %.4f' % (x, y, d, min(cand, key=lambda c: abs(d - c))))
    if not n:
        return 'SELF-CHECK: no overlap between the _geom_bond loop and the computed bonds.'
    s = 'SELF-CHECK vs the .cif _geom_bond loop: %d distances compared, largest difference %.4f Å' % (n, worst)
    if bad:
        s += '\n  ' + '\n  '.join(bad[:12])
    else:
        s += ' — consistent.'
    return s

# ----------------------------------------------------------------------------- manuscript tables

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

def _cell_text(tc):
    """Cell text with superscript runs marked ^…^ (the symmetry code of 'O3²')."""
    out = []
    for p in tc.iter(W + 'p'):
        for r in p.iter(W + 'r'):
            t = ''.join(x.text or '' for x in r.iter(W + 't'))
            rpr = r.find(W + 'rPr'); va = rpr.find(W + 'vertAlign') if rpr is not None else None
            if va is not None and va.get(W + 'val') == 'superscript' and t.strip():
                t = '^' + t + '^'
            out.append(t)
        out.append(' ')
    return re.sub(r'\s+', ' ', ''.join(out)).strip()

def read_tables(path):
    from docx import Document
    doc = Document(path)
    out = []
    for ti, t in enumerate(doc.tables):
        out.append([[_cell_text(c._tc) for c in r.cells] for r in t.rows])
    return out

BOND_CELL = re.compile(r"^<?\s*([A-Za-z]{1,3}\d*[A-Za-z]?\d*)\s*[–—−-]\s*([A-Za-z]{1,3}\d*[A-Za-z]*\d*(?:/[A-Za-z]+\d*)?)"
                       r"\s*(?:\^[^^]*\^)?\s*>?\s*(?:[×x]\s*(\d+))?\s*$")
_LABEL_OK = re.compile(r"^(?:OH|OW|Ow|W|[A-Z][a-z]?)\d*[a-zA-Z]?\d*$")

def _atom_like(lab):
    """'Cd1', 'O3', 'OH2', 'OW1', 'F/OH', 'Oyl' — but not 'metal', 'mica'."""
    return all(_LABEL_OK.match(part) and (part[:2] in ('OH', 'OW', 'Ow') or part[:1] == 'W' or
                                          re.match(r'^[A-Z][a-z]?', part).group(0) in ELEMENTS)
               for part in lab.split('/'))
NUM_CELL = re.compile(r"^\s*(\d+\.\d+)\s*(\(\d+\))?\s*(?:[×x]\s*(\d+))?")

def _norm_label(lab):
    return re.sub(r'[()\s]', '', lab).upper()

def manuscript_bonds(tables):
    """[(table idx, row idx, col idx, cation, anion, is_mean, count, d, esd)] from every table."""
    out = []
    for ti, rows in enumerate(tables):
        for ri, row in enumerate(rows):
            for ci in range(len(row) - 1):
                m = BOND_CELL.match(row[ci])
                if not m:
                    continue
                v = NUM_CELL.match(row[ci + 1])
                if not v:
                    continue
                is_mean = row[ci].strip().startswith('<')
                if not (_atom_like(m.group(1)) and _atom_like(m.group(2))) or float(v.group(1)) > 4.5:
                    continue
                count = int(m.group(3) or v.group(3) or 1)
                out.append((ti, ri, ci, m.group(1), m.group(2), is_mean, count, float(v.group(1)),
                            _esd(v.group(1) + (v.group(2) or ''))))
    return out

def check_bond_table(st, result, tables):
    """Findings about the manuscript's bond-distance table(s)."""
    L = []
    mb = manuscript_bonds(tables)
    if not mb:
        return ['no bond-distance table found in the manuscript (cells like "Cd1–O3² | 2.472(3)")']
    comp = {}                                   # (CAT, AN) -> [[dist, count, seen]]
    for site, bonds, *_ in result:
        for b in bonds:
            for lab in site.label.split('/'):
                for alab in b.anion.label.split('/'):
                    comp.setdefault((_norm_label(lab), _norm_label(alab)), []).append([b.dist, b.count, 0])
    # a manuscript anion label like 'F/OH' or 'OW1' vs the .cif's 'F1' / 'Ow1'
    def resolve(cat, an):
        c = _norm_label(cat)
        for a in [_norm_label(x) for x in an.split('/')] + [_norm_label(an)]:
            if (c, a) in comp:
                return (c, a)
            # 'F' / 'OH' for a .cif 'F1' / 'OH1': the only anion with that prefix
            hits = [(cc, aa) for (cc, aa) in comp if cc == c and re.sub(r'\d+$', '', aa) == a]
            if len({aa for _, aa in hits}) == 1:
                return hits[0]
        # 'OH1' written for 'O1'? try without OH/OW prefix variants
        for (cc, aa) in comp:
            if cc == c and (aa.rstrip('0123456789') in ('O', 'OH', 'OW') and _norm_label(an).rstrip('0123456789') in ('O', 'OH', 'OW')
                            and re.sub(r'\D', '', aa) == re.sub(r'\D', '', _norm_label(an))):
                return (cc, aa)
        return None
    n_ok = n_bad = 0
    listed = {}                                 # per table block: for mean checks
    for ti, ri, ci, cat, an, is_mean, count, d, esd in mb:
        if is_mean:
            listed_vals = [x for x in listed.get((ti, ci), []) ]
            if listed_vals:
                def mean_of(vals):
                    return sum(v * n for v, n in vals) / sum(n for v, n in vals)
                mean = mean_of(listed_vals)
                if abs(mean - d) > 0.002:
                    # a qualified mean ('<U–Oyl>', '<Ca–O/F>') may cover a leading or trailing subset
                    subsets = [listed_vals[:k] for k in range(1, len(listed_vals))] + \
                              [listed_vals[k:] for k in range(1, len(listed_vals))]
                    if not any(abs(mean_of(sv) - d) <= 0.002 for sv in subsets):
                        L.append('table %d row %d: <%s–%s> given as %.3f but the %d listed values average %.3f'
                                 % (ti + 1, ri + 1, cat, an, d, sum(n for v, n in listed_vals), mean))
                cif_vals = [(v, n) for (cc, aa), lst in comp.items() if cc == _norm_label(cat) for v, n, _ in lst]
                if cif_vals and abs(mean - d) <= 0.002:      # a plain mean over the whole listed set
                    cm = sum(v * n for v, n in cif_vals) / sum(n for v, n in cif_vals)
                    if abs(cm - d) > 0.004:
                        L.append('table %d row %d: <%s–%s> %.3f vs %.3f from the .cif over %d bonds (different bond set?)'
                                 % (ti + 1, ri + 1, cat, an, d, cm, sum(n for v, n in cif_vals)))
            listed[(ti, ci)] = []
            continue
        listed.setdefault((ti, ci), []).append((d, count))
        if _norm_label(cat).startswith('O') and _norm_label(an).startswith('H'):
            continue                            # O–H donor distances: not a cation bond here
        key = resolve(cat, an)
        if key is None:
            L.append('table %d row %d: %s–%s %.3f — no such bond in the .cif within the cutoff' % (ti + 1, ri + 1, cat, an, d))
            n_bad += 1; continue
        cands = comp[key]
        best = min(cands, key=lambda x: abs(x[0] - d))
        tol = max(0.0015, (esd or 0) * 1.5)
        # a merged site ('F1/OH1') is one bond in the table: mark its other half seen too
        twins = [x for (cc, aa), lst in comp.items() if cc == key[0] and aa != key[1] for x in lst
                 if abs(x[0] - best[0]) < 1e-6 and any(aa in a.label.split('/') and key[1] in a.label.split('/') for a in st.anions)]
        if abs(best[0] - d) <= tol:
            best[2] += count; n_ok += 1
            for x in twins:
                x[2] += count
        elif abs(best[0] - d) <= 0.05:
            L.append('table %d row %d: %s–%s %.3f but the .cif gives %.3f' % (ti + 1, ri + 1, cat, an, d, best[0]))
            best[2] += count; n_bad += 1
        else:
            L.append('table %d row %d: %s–%s %.3f — nearest .cif distance for that pair is %.3f'
                     % (ti + 1, ri + 1, cat, an, d, best[0]))
            n_bad += 1
    # multiplicity and omissions
    for (cc, aa), lst in comp.items():
        for dist, count, seen in lst:
            if seen == 0 and any(s for (c2, a2), l2 in comp.items() if c2 == cc for s in [x[2] for x in l2]):
                L.append('not in the table: %s–%s %.3f%s (the .cif has it within the cutoff)'
                         % (cc, aa, dist, ' ×%d' % count if count > 1 else ''))
            elif seen and seen != count:
                L.append('multiplicity: %s–%s %.3f is listed %d× in the table, the .cif has %d'
                         % (cc, aa, dist, seen, count))
    L.insert(0, '%d bond distances agree with the .cif, %d do not' % (n_ok, n_bad))
    return L

def _bv_cell(txt):
    """Parse one table cell into segments [(value, n_down, n_across)] — one per listed value,
    each with its own marks ('0.70×4↓×2→, 0.64×2↓'). Marks may be '×3↓', '×3 →', or a
    superscript '²↓' (kept as ^2↓^ by the table reader)."""
    segs = []
    for part in re.split(r'\s*[,;]\s*(?=\d)', txt.strip()):
        n_down = n_across = 1
        for mk in re.findall(r'\^([^^]*)\^', part):
            m = re.search(r'(\d+)\s*([↓→]?)', mk)
            if m:
                if m.group(2) == '→':
                    n_across = int(m.group(1))
                else:
                    n_down = int(m.group(1))
        body = re.sub(r'\^[^^]*\^', ' ', part)
        for n, arrow in re.findall(r'[×x]\s*(\d+)\s*([↓→]?)', body):
            if arrow == '→':
                n_across = int(n)
            else:
                n_down = int(n)
        body = re.sub(r'[×x]\s*\d+\s*[↓→]?', ' ', body)
        for v in re.findall(r'\d+\.\d+', body):
            segs.append((float(v), n_down, n_across))
    return segs

def _segs_split(segs):
    """(numbers, n_down, n_across) of a one-value cell — the marks of its first segment."""
    nums = [v for v, _, _ in segs]
    nd = segs[0][1] if segs else 1
    na = segs[0][2] if segs else 1
    return nums, nd, na

def check_bvs_table(st, result, cells, anion_sum, tables, params_label='?'):
    """Findings about a manuscript bond-valence table (anion rows × cation columns).

    Conventions differ between authors, so a cell is accepted when it matches the computed value
    under EITHER reading: one value per bond with a '×n' mark, or the total over the n bonds (and
    a comma list, or the sum, for two distinct distances). Hydrogen columns are not compared —
    donor/acceptor bookkeeping varies too much — but they take part in the row arithmetic: a
    column headed 'D'/'Donor' subtracts, 'A'/'Acceptor'/'H bond' adds."""
    L = []
    cat_labels = {_norm_label(x): r[0].label for r in result for x in r[0].label.split('/')}   # 'Mg' -> 'Mg/Mn'
    cat_labels.update({_norm_label(r[0].label): r[0].label for r in result})
    an_labels = {_norm_label(x): a.label for a in st.anions for x in a.label.split('/')}
    an_labels.update({_norm_label(a.label): a.label for a in st.anions})
    def resolve_anion(lab):
        if lab in an_labels:
            return an_labels[lab]
        parts = lab.split('/')
        hits = set()
        for part in parts:
            cands = {a.label for a in st.anions for x in a.label.split('/') if re.sub(r'\d+$', '', _norm_label(x)) == part}
            hits |= cands
        return hits.pop() if len(hits) == 1 else None
    cat_occ = {x: min(r[0].occ_total, 1.0) for r in result for x in r[0].label.split('/')}
    cat_occ.update({r[0].label: min(r[0].occ_total, 1.0) for r in result})
    bvs_of = {r[0].label: r[2] for r in result}
    h_cols = {r[0].label for r in result if r[0].element == 'H'}
    found = False
    for ti, rows in enumerate(tables):
        if len(rows) < 3:
            continue
        hdr = None
        for ri in range(min(3, len(rows))):
            hits = [ci for ci, x in enumerate(rows[ri]) if _norm_label(x) in cat_labels]
            below = sum(1 for r in rows[ri + 1:ri + 4] if r and _norm_label(r[0]) in an_labels)
            numeric = any(re.search(r'\d\.\d', x) for x in rows[ri])       # a bond-distance table, not a header
            if not numeric and (len(hits) >= 2 or (hits and below >= 1)):
                hdr = ri; break
        if hdr is None:
            continue
        found = True
        header = rows[hdr]
        col_cat = {ci: cat_labels[_norm_label(x)] for ci, x in enumerate(header) if _norm_label(x) in cat_labels}
        sum_col = next((ci for ci, x in enumerate(header) if re.match(r'^\s*(Σ|Sum|Total)', x, re.I)), None)
        col_kind = {}
        for ci, x in enumerate(header):
            if ci in col_cat or ci == sum_col:
                continue
            if re.match(r'^\s*(D|Donor)\b', x.strip(), re.I) or re.search(r'\bdonor\b', x, re.I) and 'vu' not in x.lower():
                col_kind[ci] = 'donor' if re.fullmatch(r'\s*(D|Donor)\s*\*{0,2}', x, re.I) else 'label'
            elif re.match(r'^\s*(A|Acceptor|H[- ]?bonds?|vu)', x.strip(), re.I):
                col_kind[ci] = 'acceptor'
        ncell = nbad = 0
        col_tot = {ci: [0.0, 0.0] for ci in col_cat}       # [per-bond reading, total reading]
        for ri in range(hdr + 1, len(rows)):
            row = rows[hdr + 1:][ri - hdr - 1]
            lab = _norm_label(row[0])
            an = resolve_anion(lab)
            if an is not None:
                row_pb = row_tot = row_alt = 0.0; row_flagged = False; row_h = 0.0
                for ci, cat in col_cat.items():
                    if ci >= len(row):
                        continue
                    segs = _bv_cell(row[ci]); nums, n_down, n_across = _segs_split(segs)
                    if not nums:
                        if (an, cat) in cells and cat not in h_cols:
                            L.append('table %d: %s–%s is blank but the .cif has that bond (%s vu)'
                                     % (ti + 1, an, cat, ', '.join('%.2f' % s for s, _, _ in cells[(an, cat)])))
                        continue
                    calc = cells.get((an, cat))
                    reading = 'perbond'                            # how this cell is written
                    if calc and len(nums) == 1:
                        nd0 = calc[0][1]
                        if nd0 > 1 and abs(nums[0] - calc[0][0] * nd0) <= 0.015 * nd0 + 0.01 \
                                and abs(nums[0] - calc[0][0]) > 0.015:
                            reading = 'total'                      # '1.46' = 3 × 0.49, for the column
                    if cat in h_cols:                                  # H columns: in the row total, not the cation part
                        row_h += sum(nums); col_tot[ci][1] += sum(nums); col_tot[ci][0] += sum(nums)
                        continue
                    if reading == 'total':
                        row_pb += nums[0] / calc[0][1] * n_across; col_tot[ci][0] += nums[0]
                    else:
                        row_pb += sum(v * na for v, _, na in segs); col_tot[ci][0] += sum(v * nd for v, nd, _ in segs)
                    row_tot += sum(nums); col_tot[ci][1] += sum(nums)
                    row_alt += sum(v * max(nd, na) for v, nd, na in segs)   # a '×3↓' some authors also count in the row
                    if cat in h_cols:
                        continue
                    ncell += 1
                    if not calc:
                        L.append('table %d: %s–%s %s in the table but the .cif has no such bond within the cutoff'
                                 % (ti + 1, an, cat, row[ci])); nbad += 1; continue
                    cv = sorted(s for s, _, _ in calc)
                    nd, na = calc[0][1], calc[0][2]
                    ok = False
                    if len(nums) == len(cv) and max(abs(x - y) for x, y in zip(sorted(nums), cv)) <= 0.015:
                        ok = True                                  # per-bond values
                    elif len(nums) == 1:
                        total_down = sum(s * n for s, n, _ in calc)
                        total_across = sum(s * (n2 if isinstance(n2, int) else 1) for s, _, n2 in calc)
                        tol = 0.015 * max(nd, 1) + 0.01
                        if abs(nums[0] - total_down) <= tol or abs(nums[0] - total_across) <= tol \
                                or abs(nums[0] - sum(cv)) <= tol:
                            ok = True                              # the total over the bonds
                    if not ok:
                        if re.search(r'(?<![\d.])0\d\d(?![\d.])', row[ci]):
                            L.append('table %d: %s–%s "%s" — a missing decimal point?' % (ti + 1, an, cat, row[ci].strip()))
                            nbad += 1; row_flagged = True; continue
                        hint = ''
                        if len(cv) > 1:
                            hint = ' (per bond: %s; total %.2f)' % (', '.join('%.2f' % v for v in cv), sum(s * n for s, n, _ in calc))
                        elif nd > 1 or (isinstance(na, int) and na > 1):
                            hint = ' (%.2f per bond, %s)' % (cv[0], _mark(nd, na))
                        L.append('table %d: %s–%s %s vs %.2f computed%s' % (ti + 1, an, cat, row[ci].strip(), cv[0] if len(cv) == 1 else sum(cv), hint))
                        nbad += 1; row_flagged = True
                if sum_col is not None and sum_col < len(row):
                    m = re.search(r'\d+\.\d+', row[sum_col])
                    if m:
                        given = float(m.group(0))
                        extra = 0.0; last = None
                        for ci, x in enumerate(row):
                            if ci in col_cat or ci == sum_col or ci == 0:
                                continue
                            kind = col_kind.get(ci)
                            segs_x = _bv_cell(x); nums = [v for v, _, _ in segs_x]
                            if nums and kind != 'label':
                                v = sum(v_ * max(na_, 1) for v_, _, na_ in segs_x)
                                extra += -v if kind == 'donor' else v
                                last = (ci, v, kind)
                            elif not nums and last and re.search(r'[×x]\s*\d+', x):
                                n = int(re.search(r'[×x]\s*(\d+)', x).group(1))
                                extra += (last[1] * (n - 1)) * (-1 if last[2] == 'donor' else 1)
                        extra += row_h
                        typed = (row_pb + extra, row_tot + extra, row_alt + extra)
                        if row_flagged:
                            pass                                    # a wrong cell already explains the row
                        elif min(abs(t - given) for t in typed) > 0.025:
                            if re.search(r'(?<![\d.])0\d\d(?![\d.])', ' '.join(row)):
                                L.append('table %d: Σ for %s is %.2f but its row adds to %.2f — a cell with a missing decimal point?' % (ti + 1, an, given, typed[0]))
                            else:
                                L.append('table %d: Σ for %s is %.2f but its row adds to %.2f' % (ti + 1, an, given, typed[0]))
                        else:
                            # compare the CATION part of the row with the .cif (hydrogen-bond columns and
                            # H conventions vary; the .cif may have no H at all)
                            cat_only = sum(s * (n2 if isinstance(n2, int) else 1) * cat_occ.get(c2, 1.0)
                                           for (a2, c2), lst in cells.items() if a2 == an and c2 not in h_cols
                                           for s, _, n2 in lst)
                            typed_cat = min(row_pb, row_alt) if abs(row_pb + extra - given) <= abs(row_alt + extra - given) else row_alt
                            typed_cat = row_pb if abs(row_pb + extra - given) <= 0.025 else row_alt
                            if abs(cat_only - typed_cat) > 0.08:
                                L.append('table %d: Σ for %s: the cation part adds to %.2f in the table, %.2f from the .cif (parameters: %s)'
                                         % (ti + 1, an, typed_cat, cat_only, params_label))
            elif re.match(r'^\s*(Σ|Sum|Total)', row[0], re.I):
                for ci, cat in col_cat.items():
                    if ci < len(row):
                        m = re.search(r'\d+\.\d+', row[ci])
                        if m and cat not in h_cols:
                            given = float(m.group(0))
                            if min(abs(t - given) for t in col_tot[ci]) > 0.025:
                                L.append('table %d: Σ for %s is %.2f but its column adds to %.2f' % (ti + 1, cat, given, col_tot[ci][0]))
                            elif abs(bvs_of.get(cat, given) - given) > 0.08:
                                L.append('table %d: Σ for %s %.2f vs %.2f from the .cif (parameters: %s)' % (ti + 1, cat, given, bvs_of.get(cat), params_label))
        L.insert(0, 'bond-valence table %d: %d cells compared, %d disagree (computed with %s; H columns not compared)'
                 % (ti + 1, ncell, nbad, params_label))
    if not found:
        L.append('no bond-valence table found in the manuscript (a header row with the cation labels)')
    return L

# ----------------------------------------------------------------------------- Word output

def write_word(st, result, anion_sum, cells, path, hbonds=()):
    from docx import Document
    from docx.shared import Pt
    from pxrd_review.annotate_review import _save_docx
    doc = Document()
    doc.add_heading('Selected bond distances (Å) for %s' % st.name, level=2)
    t = doc.add_table(rows=0, cols=2)
    for site, bonds, bvs, expected, mean_d in result:
        for b in bonds:
            r = t.add_row().cells
            r[0].text = '%s–%s%s' % (site.label, b.anion.label, ' ×%d' % b.count if b.count > 1 else '')
            r[1].text = '%.3f' % b.dist
        if mean_d is not None:
            r = t.add_row().cells
            r[0].text = '<%s–X>' % site.label; r[1].text = '%.3f' % mean_d
        t.add_row()
    doc.add_paragraph('')
    doc.add_heading('Bond-valence analysis (vu) for %s' % st.name, level=2)
    cats = [r[0] for r in result]
    acc = {}
    for hb in hbonds:
        acc.setdefault(hb.acceptor.label, []).append('%.2f%s' % (hb.s, '×%s→' % hb.n_across if hb.n_across != 1 else ''))
    anions = [an for an in st.anions if any((an.label, ct.label) in cells for ct in cats) or an.label in acc]
    tb = doc.add_table(rows=1, cols=len(cats) + 2 + (1 if acc else 0))
    hdr = tb.rows[0].cells
    hdr[0].text = 'Atom'
    for i, ct in enumerate(cats):
        hdr[i + 1].text = ct.label
    if acc:
        hdr[-2].text = 'H bonds'
    hdr[-1].text = 'Σ'
    for an in anions:
        r = tb.add_row().cells
        r[0].text = an.label
        for i, ct in enumerate(cats):
            vals = cells.get((an.label, ct.label))
            if not vals:
                r[i + 1].text = '–'; continue
            p = r[i + 1].paragraphs[0]
            for k, (s, nd, na) in enumerate(vals):
                p.add_run(('%s' % ', ' if k else '') + '%.2f' % s)
                m = _mark(nd, na)
                if m:
                    run = p.add_run(m); run.font.superscript = True
        if acc:
            r[-2].text = ', '.join(acc.get(an.label, [])) or '–'
        r[-1].text = '%.2f' % anion_sum[an.label]
    r = tb.add_row().cells
    r[0].text = 'Σ'
    for i, (site, bonds, bvs, expected, mean_d) in enumerate(result):
        r[i + 1].text = '%.2f' % bvs
    for row in tb.rows:
        for c in row.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    _save_docx(doc, path)

def write_xlsx(st, params, result, anion_sum, cells, hbonds, path):
    """The bond-valence calculation as a workbook with live formulas, to check a paper's:
    bonds (cation, anion, R, R0, b, s = EXP((R0−R)/b), multiplicities, contributions), the
    cation and anion sums built from those cells, the hydrogen bonds (D···A, s = (d/2.17)^−8.2
    + 0.06), and the parameters used with their sources."""
    import openpyxl
    from openpyxl.styles import Font
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = 'bonds'
    ws.append(['cation', 'species', 'anion', 'R (Å)', 'R0', 'b', 's = exp((R0−R)/b)', '×↓ (per cation)', '×→ (per anion)',
               'anion occ', 'cation occ', 'to the cation sum', 'to the anion sum', 'parameter source'])
    for c in range(1, 15):
        ws.cell(1, c).font = Font(bold=True)
    r = 2; first = {}
    for site, bonds, bvs, expected, mean_d in result:
        tot_occ = sum(sp.occ for sp in site.species if sp.ox and sp.ox > 0) or 1.0
        for b in bonds:
            n_across = b.count * site.mult / b.anion.mult
            for sp, sval in b.vals:
                if sp.ox is None or sp.ox <= 0:
                    continue
                key = (sp.element, sp.ox, b.anion.element)
                p = params.used.get(key)
                if p is None:
                    continue
                r0, bb, rid = p
                ws.append([site.label, '%s%+d' % (sp.element, sp.ox), b.anion.label, round(b.dist, 4), r0, bb,
                           '=EXP((E%d-D%d)/F%d)' % (r, r, r), b.count, round(n_across, 3), min(b.anion.occ_total, 1.0),
                           round(sp.occ / tot_occ, 4), '=G%d*H%d*J%d*K%d' % (r, r, r, r), '=G%d*I%d*K%d*MIN(1,%s)' % (r, r, r, round(min(tot_occ, 1.0), 4)),
                           params.short_ref(rid)])
                first.setdefault(site.label, r); r += 1
    last = r - 1
    for col, w in zip('ABCDEFGHIJKLMN', (10, 8, 10, 9, 8, 7, 18, 12, 12, 9, 10, 16, 16, 30)):
        ws.column_dimensions[col].width = w
    # sums
    wc = wb.create_sheet('cation sums'); wc.append(['cation', 'Σ (vu)', 'expected', 'Σ/expected − 1']); wc.cell(1, 1).font = Font(bold=True)
    for i, (site, bonds, bvs, expected, mean_d) in enumerate(result, 2):
        wc.append([site.label, '=SUMIF(bonds!A:A,A%d,bonds!L:L)' % i, round(expected, 3), '=B%d/C%d-1' % (i, i)])
    wa = wb.create_sheet('anion sums'); wa.append(['anion', 'cations (vu)', 'accepted H bonds (vu)', 'Σan', 'donated O–H (vu)', 'Σall']); wa.cell(1, 1).font = Font(bold=True)
    hb_rows = {}
    for hb in hbonds:
        hb_rows.setdefault(hb.acceptor.label, []).append(hb)
    don = donated(hbonds)
    for i, an in enumerate([a for a in st.anions if any((a.label, r_[0].label) in cells for r_ in result) or a.label in hb_rows], 2):
        wa.append([an.label, '=SUMIF(bonds!C:C,A%d,bonds!M:M)' % i, '=SUMIF(\'H bonds\'!B:B,A%d,\'H bonds\'!F:F)' % i, '=B%d+C%d' % (i, i),
                   round(don.get(an.label, 0.0), 4), '=D%d+E%d' % (i, i)])
    wh = wb.create_sheet('H bonds'); wh.append(['donor', 'acceptor', 'D⋯A (Å)', 's (vu) = (d/2.17)^-8.2 + 0.06', '×→ (per acceptor)', 'to the acceptor sum', 'source'])
    wh.cell(1, 1).font = Font(bold=True)
    for i, hb in enumerate(hbonds, 2):
        formula = '=(C%d/2.17)^(-8.2)+0.06' % i if hb.via != 'H···A' else round(hb.s, 4)
        wh.append([hb.donor.label, hb.acceptor.label, round(hb.d, 4), formula, hb.n_across, '=D%d*E%d' % (i, i),
                   {'loop': 'D–H⋯A from the .cif loop', 'H': 'from the H positions', 'OO': 'PROPOSED from the O⋯O geometry (no H in the .cif)', 'H···A': 'H⋯A distance, Brown 2002'}[hb.via]])
    wp = wb.create_sheet('parameters'); wp.append(['pair', 'R0', 'b', 'reference']); wp.cell(1, 1).font = Font(bold=True)
    for (cat, cox, an), (r0, bb, rid) in params.used.items():
        wp.append(['%s%+d–%s' % (cat, cox, an), r0, bb, params.refs.get(rid, rid)])
    wp.append([]); wp.append(['table note', params.note()])
    wp.column_dimensions['A'].width = 14; wp.column_dimensions['D'].width = 70
    wb.save(path)
    return path

# ----------------------------------------------------------------------------- main

def _parse_ox(s):
    out = {}
    for part in (s or '').split(','):
        if '=' in part:
            el, v = part.split('=', 1)
            out[el.strip().capitalize() if len(el.strip()) > 1 else el.strip().upper()] = int(v)
    return out

PARAM_NAMES = {'gh': 'Gagné & Hawthorne 2015', 'bo': "Brese & O'Keeffe 1991", 'ba': 'Brown & Altermatt 1985'}

def _parse_donors(s):
    """'OW1=2,O5=1' -> {'OW1': 2, 'O5': 1}"""
    out = {}
    for part in (s or '').split(','):
        if '=' in part:
            k, v = part.split('=', 1)
            try:
                out[k.strip()] = int(v)
            except ValueError:
                raise ValueError('--donors: %r is not label=N' % part.strip())
    return out

def _parse_hb(s):
    """'OW1>O2,OW1>O7' -> [('OW1', 'O2'), ('OW1', 'O7')]"""
    out = []
    for part in (s or '').split(','):
        if '>' in part:
            d, a = part.split('>', 1); out.append((d.strip(), a.strip()))
        elif part.strip():
            raise ValueError('--hb: %r is not donor>acceptor' % part.strip())
    return out

def run(cif, table=None, params='gh', ox=None, cutoff=None, include_h=True, word=False, out_dir=None, quiet=False,
        auto_params=True, hbond='oo', hmax=None, donors=None, hb=None, u6='burns', xlsx=False):
    st = Structure(cif, _parse_ox(ox), include_h=include_h)
    tables = read_tables(table) if table else None
    P = Params(prefer=params, u6=u6)
    hb_args = dict(hbond=hbond, hmax=hmax, donors=_parse_donors(donors) if isinstance(donors, str) else donors,
                   force=_parse_hb(hb) if isinstance(hb, str) else hb)
    result, anion_sum, cells, hbonds = compute(st, P, cutoff, **hb_args)
    chosen_note = ''
    if tables and auto_params:
        # the manuscript's bond-valence table tells which parameter set its authors used:
        # score every set and report with the one that agrees best (ties -> the requested set)
        scores = {}
        for key in ('gh', 'bo', 'ba'):
            Pk = Params(prefer=key, u6=u6)
            notes_before = list(st.notes)
            rk = compute(st, Pk, cutoff, **hb_args)
            st.notes[:] = notes_before                      # the trial runs repeat the same notes
            lines = check_bvs_table(st, rk[0], rk[2], rk[1], tables, PARAM_NAMES[key])
            hits = [re.search(r'(\d+) cells compared, (\d+) disagree', ln) for ln in lines]
            hits = [m for m in hits if m]
            if hits and sum(int(m.group(1)) for m in hits):
                scores[key] = (sum(int(m.group(2)) for m in hits), 0 if key == params else 1, Pk, rk)
        if scores:
            best = min(scores, key=lambda k: scores[k][:2])
            if best != params:
                P, (result, anion_sum, cells, hbonds) = scores[best][2], scores[best][3]
                chosen_note = ('  the manuscript table agrees best with %s parameters (%d cells disagree, vs %d with %s) '
                               '— the report below uses them; --params %s forces a set\n' %
                               (PARAM_NAMES[best], scores[best][0], scores[params][0] if params in scores else -1,
                                PARAM_NAMES[params], params))
                params = best
    text = report_text(st, P, result, anion_sum, cells, geom_self_check(st, result), hbonds)
    if table:
        text += '\n\nMANUSCRIPT TABLE CHECK — %s\n' % os.path.basename(table) + chosen_note + '  '
        text += '\n  '.join(check_bond_table(st, result, tables))
        text += '\n  ' + '\n  '.join(check_bvs_table(st, result, cells, anion_sum, tables, PARAM_NAMES[params]))
    if not quiet:
        print(text)
    out_dir = out_dir or os.path.join(os.path.dirname(os.path.abspath(cif)), 'review_out')
    os.makedirs(out_dir, exist_ok=True)
    stem = os.path.splitext(os.path.basename(cif))[0]
    rep = os.path.join(out_dir, stem + '_bv.txt')
    with open(rep, 'w', encoding='utf-8') as f:
        f.write(text + '\n')
    if not quiet:
        print('  report → %s' % rep)
    if word:
        wp = os.path.join(out_dir, stem + '_bv.docx')
        write_word(st, result, anion_sum, cells, wp, hbonds)
        if not quiet:
            print('  tables → %s' % wp)
    if xlsx:
        xp = os.path.join(out_dir, stem + '_bv.xlsx')
        write_xlsx(st, P, result, anion_sum, cells, hbonds, xp)
        if not quiet:
            print('  workbook → %s' % xp)
    return st, result, anion_sum, cells, text

def main(argv=None):
    ap = argparse.ArgumentParser(prog='pxrd bv', description=__doc__.split('\n\n')[1],
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument('cif')
    ap.add_argument('--table', help='manuscript .docx whose bond-distance / bond-valence tables to check')
    ap.add_argument('--params', default='gh', choices=['gh', 'bo', 'ba'],
                    help='gh = Gagné & Hawthorne 2015 (default), bo = Brese & O\'Keeffe 1991, ba = Brown & Altermatt 1985')
    ap.add_argument('--ox', help='oxidation states, e.g. Fe=2,Mn=3 (override the .cif / defaults)')
    ap.add_argument('--cutoff', type=float, help='neighbour cutoff in Å for every cation (default 3.2, larger for big cations)')
    ap.add_argument('--no-h', action='store_true', help='ignore hydrogen atoms')
    ap.add_argument('--hbonds', default='oo', choices=['oo', 'h', 'none'],
                    help='hydrogen-bond strengths: oo = from the O···O distances (Ferraris & Ivaldi 1988, default); '
                         'h = from the H···O distances (Brown 2002, H as a cation column); none')
    ap.add_argument('--hmax', type=float, help='longest O···O counted as a hydrogen bond (default %.2f Å)' % HB_MAX)
    ap.add_argument('--donors', help='hydrogen count per O site when the labels do not say, e.g. OW1=2,O5=1 (0 = not a donor)')
    ap.add_argument('--hb', help='hydrogen bonds to place as given, e.g. OW1>O2,OW1>O7 (donor>acceptor; the rest are still proposed)')
    ap.add_argument('--u6', default='burns', choices=['burns', 'params'],
                    help='U6+–O parameters: burns = Burns et al. (1997), the uranyl convention (default); params = the chosen set')
    ap.add_argument('--word', action='store_true', help='also write the tables as a .docx')
    ap.add_argument('--xlsx', action='store_true', help='also write the calculation as a workbook with live formulas (review_out/<name>_bv.xlsx)')
    ap.add_argument('--out', help='output folder (default <cif dir>/review_out)')
    a = ap.parse_args(argv)
    try:
        run(a.cif, a.table, a.params, a.ox, a.cutoff, not a.no_h, a.word, a.out,
            hbond=a.hbonds, hmax=a.hmax, donors=a.donors, hb=a.hb, u6=a.u6, xlsx=a.xlsx)
    except ValueError as e:
        raise SystemExit('bv_check: %s' % e)
    return 0

if __name__ == '__main__':
    sys.exit(main())
