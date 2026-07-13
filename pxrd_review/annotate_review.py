#!/usr/bin/env python3
"""
Write the cell/wavelength review back INTO each entry .docx as Word comments and
highlights, so the reviewer sees the findings in context instead of in a console
report.

  * Highlights (yellow) the Author's-Cell value cells the checker flags.
  * Adds a comment on each flagged value explaining the discrepancy
    (value / significant-figures / esd), and a comment on the Radiation cell
    when the anode/λ is flagged.
  * Adds one summary comment on the 'Author's Cell' label cell: the overall
    verdict, the matched PDF cell, and the λ status.

Comparison logic is reused verbatim from cell_lambda_check (single source of
truth); this module only renders the result into the document.

Flagged entries are written as '<name>_edited.docx' so the ones the tool
commented on stand out in the folder listing; clean entries keep the source name.

Usage:
    python3 -m pxrd_review.annotate_review "/path/to/Part 1"                 # -> <folder>/review_out
    python3 -m pxrd_review.annotate_review "/path/to/Part 1" --id I003416
    python3 -m pxrd_review.annotate_review "/path/to/Part 1" --out DIR
    python3 -m pxrd_review.annotate_review "/path/to/Part 1" --inplace       # edit originals (asks nothing)
"""
import sys, os, re, glob, shutil, argparse, textwrap, zipfile, io, datetime, json, hashlib, time, copy

from pxrd_review import cell_lambda_check as C
from pxrd_review import extra_checks as X
from pxrd_review import errors as E
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from lxml import etree

# Hardened XML parser for docx XML: a crafted .docx could carry a DTD with external/local-file
# entities (XXE) or an entity bomb. docx OOXML never legitimately uses a DTD, so disable DTD
# loading + entity resolution + any network access. The 5 predefined entities (&amp; etc.) still
# work, so real documents are unaffected. Use this everywhere instead of bare etree.fromstring().
_SAFE_XML = etree.XMLParser(resolve_entities=False, no_network=True, load_dtd=False,
                            dtd_validation=False, huge_tree=False)
def _xml(data):
    return etree.fromstring(data, _SAFE_XML)

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
AUTHOR, INITIALS = C.TOOL_AUTHOR, 'PXRD'               # 'PXRD Review Tool' — single source in C
def _q(tag): return W + tag
PARAM_COL = {'a': 1, 'b': 2, 'c': 3, 'α': 4, 'β': 5, 'γ': 6,
             'SG': 7, 'Z': 8}                                  # Author's-Cell column per parameter
KIND_LABEL = {'value': 'VALUE MISMATCH', 'precision': 'significant figures differ',
              'esd': 'uncertainty (esd) differs'}

def _tidy(s):
    """Cosmetic clean-up of comment text: collapse the double spaces inherited
    from the report formatter, and render anode names with the α subscript
    (docx writes 'MoKa'/'CuKa1', the literature 'MoKα')."""
    s = re.sub(r'(?<=[A-Za-z])Ka', 'Kα', s)
    return re.sub(r' {2,}', ' ', s).strip()

NO_MATCH = 'No matching .pdf cell found.'

def _has_lam_flag(res):
    return res['lam'] is not None and res['lam'][0] == 'flag'

def _save_docx(doc, path):
    """Save to a temp file in the destination directory, then os.replace() it into
    place. A crash (or a full disk) mid-write can then never truncate the target —
    which under --inplace, or when refreshing onto a hand-edited output, IS the
    reviewer's only copy. python-docx streams a zip, so a direct doc.save(path)
    leaves an unopenable file if it dies part-way through."""
    d = os.path.dirname(path) or '.'
    os.makedirs(d, exist_ok=True)
    tmp = os.path.join(d, '~save_' + os.path.basename(path))
    try:
        doc.save(tmp)
        os.replace(tmp, path)
    finally:
        if os.path.exists(tmp):
            os.remove(tmp)

class _RunLock:
    """Advisory single-writer lock over one output folder.

    Two annotate_review processes on the same folder — the classic case is clicking 'Rerun all'
    in the GUI while a terminal run is still going — write the same output docx and the same
    fixed-name temp files. Interleaved, that can leave a corrupt .docx, and it corrupts it
    SILENTLY. The writers are the only thing that needs serialising, so an O_EXCL lockfile in
    the output folder is enough. A lock whose owner is gone (crash) is taken over, so a stale
    file can never wedge the tool permanently."""
    STALE_S = 3 * 3600

    def __init__(self, out_dir):
        self.path = os.path.join(out_dir, '.annotate.lock')
        self.held = False

    def _owner_alive(self):
        try:
            with open(self.path, encoding='utf-8') as f:
                pid = int((f.read().split('\n', 1)[0] or '0').strip())
        except Exception:
            return False                     # unreadable/garbled -> treat as abandoned
        if pid <= 0:
            return False
        try:
            os.kill(pid, 0)                  # signal 0: existence check, no signal delivered
            return True
        except OSError:
            return False                     # no such process -> abandoned by a crashed run

    def acquire(self):
        os.makedirs(os.path.dirname(self.path) or '.', exist_ok=True)
        for _ in range(2):
            try:
                fd = os.open(self.path, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
                with os.fdopen(fd, 'w') as f:
                    f.write('%d\n%s\n' % (os.getpid(), datetime.datetime.now().isoformat()))
                self.held = True
                return True
            except FileExistsError:
                stale = (not self._owner_alive() or
                         time.time() - os.path.getmtime(self.path) > self.STALE_S)
                if not stale:
                    return False
                try:                          # abandoned by a dead run — reclaim it
                    os.remove(self.path)
                except OSError:
                    return False
        return False

    def release(self):
        if self.held:
            try:
                os.remove(self.path)
            except OSError:
                pass
            self.held = False

    def __enter__(self):
        return self

    def __exit__(self, *exc):
        self.release()

def _writable_extras(res):
    """Which extra-check findings get written into the docx: every hard FLAG, plus
    the authoritative Mindat group statement (informational but routinely added by
    the reviewer). Soft 'note'/other 'info' findings stay console-only."""
    return [f for f in res.get('extra', [])
            if f.sev == 'flag' or (f.code == 'classification' and f.anchor == 'name')]

# Comment banner reflects the finding's domain so it reads accurately: composition/
# formula findings say 'Chemistry check'; everything else (cell, indexing, instrument,
# radiation, lines, symmetry) stays 'PXRD check'.
_CHEM_CODES = {'analysis', 'ideal_formula', 'name_formula', 'mindat_chem'}
def _check_banner(code):
    return 'Chemistry check' if code in _CHEM_CODES else 'PXRD check'

def _is_clean(res):
    """Nothing to report: the cell matched exactly, with no parameter
    discrepancies, no radiation hard-flag, and no writable extra findings. (Soft
    λ 'verify'/'unrec' notes are not errors and are intentionally not reported.)"""
    return (res['cell'][0] == 'match' and not res['params']
            and not _has_lam_flag(res) and not _writable_extras(res))

# ----------------------------------------------------------------------------- analysis
def analyze(docx_path, pdf_path, cif_path=None, dft_path=None):
    """Return a structured verdict mirroring cell_lambda_check.report()."""
    d = C.parse_docx(docx_path)
    res = {'docx': d, 'cell': ('nopdf',), 'params': {}, 'param_labels': {}, 'lam': None, 'extra': []}
    text = C.pdf_text(pdf_path) if pdf_path else None
    # a scanned .pdf has NO extractable text layer — pdf_text comes back blank; treat
    # it as no text at all (the checks already tolerate None, the no-pdf path), so the
    # radiation check can't emit a false 'anode NOT found in .pdf' flag on blank text.
    if text is not None and not text.strip():
        text = None
    cif_data = X.parse_cif(cif_path) if cif_path else {}
    dft_data = X.parse_dft(dft_path) if dft_path else {}
    # extra checks run regardless of a PDF (symmetry/indexing/calculated are
    # docx-internal). Never let them break the core cell/λ verdict.
    entry = None
    try:
        entry = X.parse_entry(docx_path)
        res['extra'] = X.run_all(entry, text, cif_data, dft_data)
    except Exception as ex:
        # Do NOT fail silently. If parse_entry dies, ALL 22 extra checks produce nothing —
        # and an entry with no findings is indistinguishable from a clean one, so a parse
        # failure used to read as a PASS. Record it as a note (console/log/GUI, never written
        # into the docx — severity discipline) and stamp res['error'] so the sweep counts it.
        res['extra'] = [X.Finding('parse_error', 'note',
                                  'the extra checks could not run on this entry — %s'
                                  % E.explain(ex, docx_path), None, None)]
        res['error'] = E.explain(ex, docx_path)
        print('  !! %s: extra checks SKIPPED — %s'
              % (os.path.basename(docx_path), E.explain(ex, docx_path)))
    # a CALCULATED powder pattern uses the modelling wavelength (calc software
    # such as PLATON commonly defaults to CuKα), not the experimental radiation
    is_calc = bool(entry and (entry.instr.get('spacing_instr') or '').strip().lower() == 'calculated')
    if not pdf_path:
        return res
    if text is None:
        # .pdf present but no extractable text layer (scanned images): nothing to
        # compare against — a distinct verdict so the GUI/console can say so instead
        # of reporting a bogus no-match cell / missing anode.
        res['cell'] = ('notext',)
        return res
    cands = C.find_cells(text)
    if d.authors_cell and any(d.authors_cell[:3]):
        cd, nmatch, ncomp, dev, mode = C.best_match(d.authors_cell[:3], cands, C.entry_name(docx_path))
        if cd is None:
            res['cell'] = ('nocell',)
        elif ncomp >= 2 and nmatch == ncomp:
            res['cell'] = ('match', cd, nmatch, ncomp, dev, mode)
            keys = ['a', 'b', 'c', 'α', 'β', 'γ']
            docx_p = dict(zip(keys, d.authors_cell[:6]))
            pdf_p = dict(zip(keys, [cd.a, cd.b, cd.c, cd.al, cd.be, cd.ga]))
            # collapse symmetry-equivalent axes (cubic a=b=c, uniaxial a=b) so the same
            # sig-fig/esd error flags ONCE, on the representative axis, under a combined label
            for label, rep, kind, note in C.grouped_axis_issues(docx_p, pdf_p, keys):
                res['params'].setdefault(rep, []).append((kind, note))
                res['param_labels'][rep] = label
            # Z (formula units): a/b/c match, so this IS the same cell — a different Z is a
            # real inconsistency. Compare the docx Z to an EXPLICIT 'Z = N' in the matched
            # cell's own sentence (the '=' form, so a digit in the space-group symbol, e.g.
            # P2_12_12_1, isn't misread as Z). The .pdf is ground truth (grokhovskyite: docx
            # Z=2, but the paper states 'Z = 3').
            zm = re.search(r'\bZ\s*=\s*(\d+)', cd.snippet or '')
            dz = re.sub(r'\D', '', str(d.authors_cell[7])) if len(d.authors_cell) > 7 else ''
            if zm and dz and zm.group(1) != dz:
                res['params']['Z'] = [('value', 'docx Z=%s but the .pdf states Z=%s — verify'
                                       % (dz, zm.group(1)))]
            # cell SOURCE: ICDD entries use the powder cell — flag/note an SCXRD cell
            src = C.cell_source_finding(d.authors_cell, cd, cands)
            if src:
                res.setdefault('extra', []).append(
                    X.Finding('cell_source', src[0], src[1], src[2], 'cell:a'))
        else:
            res['cell'] = ('investigate', cd, nmatch, ncomp, dev, mode)
            res['cell_diffs'] = C.cell_axis_deltas(d.authors_cell, cd)
    # radiation
    rads = C.find_radiation(text)
    dk = C.anode_key(d.radiation)
    powder = [r for r in rads if r[2] == 'powder']
    # Choose the powder radiation from the PAPER'S OWN evidence, not the docx. A radiation
    # named in an explicit powder-collection statement / beside a powder-only geometry
    # (r[3] = powder_collected_near) is a CONFIRMED powder radiation, ranked above one that
    # merely landed in a 'powder' context by section proximity — tarutinoite: the SC MoKα sits
    # just after a 'powder diffraction patterns' results line, so _section_mode tags it powder,
    # but only the CoKα sits in 'Powder XRD … collected … Debye-Scherrer … CoKα'. Among the
    # CONFIRMED radiations a docx-anode match breaks ties, so a multi-sample paper that
    # collected powder with several anodes (I002787: CuKα for the K member, CoKα for the Cs
    # member) attributes the right one to this phase. This is validation, NOT circular trust:
    # the docx tiebreak only orders GENUINE powder radiations (a misclassified single-crystal
    # radiation has no collection evidence, so it can never win), and the verdict still FLAGS
    # whenever the paper's powder-collection evidence does not include the docx anode.
    powder.sort(key=lambda r: (not r[3], C.anode_key(r[0]) != dk, r[1] is None))
    pk = powder[0] if powder else None
    any_match = any(C.anode_key(r[0]) == dk for r in rads)
    if dk is None:
        if C.is_sync_anode(d.radiation):
            # synchrotron: λ is beamline-tunable, so it legitimately matches no
            # characteristic tube line — 'Sync' is the correct designator, not an error.
            conf = '; .pdf confirms synchrotron radiation' if C.pdf_mentions_synchrotron(text) else ''
            res['lam'] = ('ok', 'anode %s — synchrotron radiation (λ is beamline-specific, '
                          'not a characteristic tube line)%s' % (d.radiation, conf))
        else:
            res['lam'] = ('unrec', 'docx anode not recognised (%s)' % d.radiation)
    elif pk is not None:
        if C.anode_key(pk[0]) == dk:
            via = 'a .pdf powder-collection statement' if pk[3] else '.pdf powder radiation'
            res['lam'] = ('ok', 'anode %s matches %s' % (d.radiation, via))
        elif pk[3]:
            res['lam'] = ('flag', 'docx anode %s but the .pdf collected the powder pattern with %sKα'
                          % (d.radiation, pk[0].capitalize()))
        else:
            res['lam'] = ('flag', 'docx anode %s but .pdf POWDER radiation is %sKα'
                          % (d.radiation, pk[0].capitalize()))
        if pk[1] and d.lam and not C.close(float(pk[1]), C.num_val(d.lam), abstol=0.003):
            res['lam'] = (res['lam'][0], res['lam'][1] + ' (λ docx=%s pdf=%s)' % (d.lam, pk[1]))
    elif any_match:
        # When the paper names a SINGLE anode and it matches the docx, it is
        # unambiguously THE radiation — the powder pattern shares the source even if
        # the only mention sits in a single-crystal sentence (Gandolfi / crystal-
        # rotation powder on a single-crystal instrument). Don't ask to verify.
        distinct = {C.anode_key(r[0]) for r in rads}; distinct.discard(None)
        if distinct == {dk}:
            res['lam'] = ('ok', 'anode %s matches the .pdf radiation (single source — powder shares it, '
                          'e.g. Gandolfi/crystal-rotation on a single-crystal instrument)' % d.radiation)
        else:
            res['lam'] = ('verify', 'anode %s appears in .pdf but no clear powder-context radiation found' % d.radiation)
    else:
        anodes = sorted(set(r[0].capitalize() + 'Kα' for r in rads)) or ['(none found)']
        res['lam'] = ('flag', 'docx anode %s NOT found in .pdf; .pdf mentions: %s'
                      % (d.radiation, ', '.join(anodes)))
    # For a CALCULATED pattern the docx anode/λ is the modelling wavelength and
    # legitimately differs from the experimental radiation in the article
    # (which is usually the single-crystal collection). Don't flag that mismatch.
    if is_calc and res['lam'] and res['lam'][0] == 'flag':
        res['lam'] = ('calc', res['lam'][1] +
                      ' — but pattern is CALCULATED, so this is the modelling '
                      'wavelength, not the experimental radiation (no action needed)')
    # Structured evidence for the GUI '? look', so it never parses the message prose (which
    # silently broke the zoom twice): the .pdf radiation the verdict concerns — the matched /
    # conflicting POWDER radiation when one was identified (anode element + its λ), else the
    # docx anode (verify/unrec: where it does or does not appear). Anode is the lowercase
    # element ('cu'/'mo'/…) or None when unrecognised (e.g. synchrotron).
    if pk is not None:
        res['lam_evidence'] = {'anode': pk[0], 'lam': pk[1]}
    else:
        res['lam_evidence'] = {'anode': C.anode_key(d.radiation), 'lam': None}
    return res

# ----------------------------------------------------------------------------- docx writing
def _cell_runs(cell):
    runs = []
    for p in cell.paragraphs:
        runs.extend(p.runs)
    return runs

def _highlight(cell, color='yellow'):
    for r in cell._tc.findall('.//' + W + 'r'):
        rPr = r.find(W + 'rPr')
        if rPr is None:
            rPr = r.makeelement(W + 'rPr', {}); r.insert(0, rPr)
        hl = rPr.find(W + 'highlight')
        if hl is None:
            hl = rPr.makeelement(W + 'highlight', {}); rPr.append(hl)
        hl.set(W + 'val', color)

def _cell_text(cell):
    """The cell's current visible text (w:t only — a w:delText is struck through, i.e. gone)."""
    return ''.join(t.text or '' for t in cell._tc.iter(_q('t')))

def _cell_human_marks(cell):
    """True when a PERSON has already worked on this cell — any tracked insertion/deletion whose
    author is not the tool. Their correction stands: the tool must not overwrite it, nor re-apply
    a change they already made (or deliberately made differently)."""
    for tag in ('ins', 'del'):
        for el in cell._tc.iter(_q(tag)):
            if (el.get(_q('author')) or '') != AUTHOR:
                return True
    return False

_FIX_ID = [90000]                     # tracked-change ids, kept clear of Word's own numbering

def _apply_tracked_fix(cell, new_text):
    """Write `new_text` into the cell as a Word TRACKED CHANGE authored by the tool: the old runs
    become a <w:del>, the new text a <w:ins>.

    Tracked, not a silent overwrite. The reviewer opens the docx and sees exactly what the tool
    changed, in Word's own review UI, and can Accept or Reject it — a field the tool rewrote
    invisibly would be unreviewable, and that is the thing this tool has always refused to do.
    The write lands in the review_out COPY; the source docx is never touched.

    Returns True when it wrote something."""
    ps = list(cell.paragraphs)
    if not ps or not new_text:
        return False
    # ONE text-bearing paragraph only. With more, every paragraph's runs would be swept into a
    # single <w:del> in the FIRST paragraph — so REJECTING the tool's change would restore the
    # words but not the paragraph break, silently collapsing 'Old Title.' / 'Author, A.' into one
    # line. No reference cell in the 466-entry corpus is laid out that way, so refusing costs
    # nothing and removes the only way this writer can lose a reviewer's formatting. Fall back to
    # comment-only (the caller reports it) rather than guess.
    with_text = [p for p in ps if p.text.strip()]
    if len(with_text) > 1:
        return False
    runs = [r for para in ps for r in para.runs]
    if not runs or not _cell_text(cell).strip():
        return False
    now = datetime.datetime.now().replace(microsecond=0).isoformat()
    _FIX_ID[0] += 1
    dele = ps[0]._p.makeelement(_q('del'), {_q('id'): str(_FIX_ID[0]),
                                            _q('author'): AUTHOR, _q('date'): now})
    _FIX_ID[0] += 1
    ins = ps[0]._p.makeelement(_q('ins'), {_q('id'): str(_FIX_ID[0]),
                                           _q('author'): AUTHOR, _q('date'): now})
    first = runs[0]._r
    parent = first.getparent()
    parent.insert(list(parent).index(first), dele)
    # move every existing run into the deletion, rewriting w:t -> w:delText (Word requires it)
    for r in runs:
        el = r._r
        if el.getparent() is not None:
            el.getparent().remove(el)
        for t in el.findall(_q('t')):
            t.tag = _q('delText')
        dele.append(el)
    # the replacement run, carrying the first run's formatting so the citation keeps its look
    newr = dele.makeelement(_q('r'), {})
    rPr = runs[0]._r.find(_q('rPr'))
    if rPr is not None:
        newr.append(copy.deepcopy(rPr))
    t = newr.makeelement(_q('t'), {'{http://www.w3.org/XML/1998/namespace}space': 'preserve'})
    t.text = new_text
    newr.append(t)
    ins.append(newr)
    dparent = dele.getparent()
    dparent.insert(list(dparent).index(dele) + 1, ins)
    return True

def _rows(doc):
    ac = rad = None
    for t in doc.tables:
        for row in t.rows:
            h = row.cells[0].text.strip()
            if h.startswith('Author') and ac is None:
                ac = row
            if h.startswith('Radiation') and rad is None:
                rad = row
    return ac, rad

# --- extra-check anchoring: map a Finding.anchor to the docx cell to comment on
def _find_cell(doc, pred):
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                if pred(c.text):
                    return c
    return None

def _find_value(doc, pred):
    """The value cell following a label cell matching pred; falls back to the
    label cell when the value is blank, so there are always runs to anchor to."""
    for t in doc.tables:
        for row in t.rows:
            cells = row.cells
            for i, c in enumerate(cells):
                if pred(c.text):
                    nxt = cells[i + 1] if i + 1 < len(cells) else None
                    return nxt if (nxt is not None and nxt.text.strip()) else c
    return None

def _find_field_value(doc, pred):
    """Like _find_value, but skips the horizontal-merge duplicates of the label cell
    (a merged label repeats its text across row.cells, e.g. 'Spacing Instr. :' at
    [5][6] then 'Other' at [7]) so the highlight/comment lands on the VALUE. Used only
    for anchors whose value is known non-blank, so it never runs past into the next
    field's label."""
    for t in doc.tables:
        for row in t.rows:
            cells = row.cells
            for i, c in enumerate(cells):
                if pred(c.text):
                    label = c.text.strip()
                    j = i + 1
                    while j < len(cells) and cells[j].text.strip() == label:
                        j += 1                      # skip merged-span duplicates
                    nxt = cells[j] if j < len(cells) else None
                    return nxt if (nxt is not None and nxt.text.strip()) else c
    return None

def _anchor_cell(doc, ac_row, anchor):
    if anchor and anchor.startswith('cell:'):
        col = PARAM_COL.get(anchor.split(':', 1)[1])
        return ac_row.cells[col] if (col is not None and col < len(ac_row.cells)) else None
    if anchor == 'instr':
        return (_find_value(doc, lambda t: 'spacing instr' in t.lower())
                or _find_value(doc, lambda t: t.strip().lower().startswith('radiation')))
    # the instr_class flags highlight the field VALUE ('Other') — use the merge-aware
    # resolver so it lands on the value, not the (merged) label cell
    if anchor == 'spacing_instr':
        return (_find_field_value(doc, lambda t: 'spacing instr' in t.lower())
                or _find_value(doc, lambda t: 'spacing instr' in t.lower()))
    if anchor == 'intensity_instr':
        return (_find_field_value(doc, lambda t: 'intensity instr' in t.lower())
                or _find_field_value(doc, lambda t: 'spacing instr' in t.lower()))
    # An Intensity Type finding is about the Intensity Type field — it used to share the generic
    # 'instr' anchor, which highlighted the Spacing Instr. cell (and sent the GUI's '? look' to
    # the Radiation row). Land it on its own field.
    if anchor == 'intensity_type':
        return (_find_field_value(doc, lambda t: 'intensity type' in t.lower())
                or _find_field_value(doc, lambda t: 'intensity instr' in t.lower())
                or _find_field_value(doc, lambda t: 'spacing instr' in t.lower()))
    if anchor == 'refl':
        return _find_cell(doc, lambda t: t.strip() in ('d(A)', 'd(Å)'))
    # the citation itself (References section) — highlight the value cell, not the label
    if anchor == 'reference':
        return (_find_value(doc, lambda t: t.strip().lower().startswith('primary reference'))
                or _find_cell(doc, lambda t: t.strip() == 'References'))
    if anchor == 'ima':
        # IMA numbers are reported in the Comments section. When this entry has no
        # 'IMA Number' row, anchor to the Comments-section header rather than letting
        # it fall through to the Author's Cell label (the generic fallback).
        return (_find_value(doc, lambda t: t.strip() == 'IMA Number')
                or _find_cell(doc, lambda t: t.strip() == 'Comments'))
    if anchor == 'formula':
        return (_find_value(doc, lambda t: t.strip() == 'Empirical')
                or _find_value(doc, lambda t: t.strip() == 'Chemical'))
    if anchor == 'analysis':
        # the 'Analysis' comment cell (where the microprobe wt.% data is given); falls back to
        # the 'Analysis' label cell when empty, and — when the entry has no Analysis row at all
        # — to the Comments-section header (the Analysis field lives there), rather than letting
        # it drop to the Author's-Cell fallback.
        return (_find_value(doc, lambda t: t.strip() == 'Analysis')
                or _find_cell(doc, lambda t: t.strip() == 'Comments'))
    if anchor == 'radiation':
        return _find_value(doc, lambda t: t.strip().lower().startswith('radiation'))
    if anchor == 'filter':
        return _find_value(doc, lambda t: t.strip().lower().startswith('filter')
                           and 'type' not in t.strip().lower())
    if anchor == 'optical':
        return _find_value(doc, lambda t: t.strip() == 'Optical Data')
    if anchor == 'name':
        return _find_value(doc, lambda t: t.strip() == 'Mineral')
    if anchor == 'primary':
        return _find_value(doc, lambda t: t.strip() == 'Primary')
    return None

# --- optional triage feedback (review_gui sidecar) ------------------------
# A rerun may pass the reviewer's per-finding verdicts so the regenerated docx
# reflects them — strictly COMMENT-ONLY: a 'dismiss' suppresses that comment/
# highlight, a 'confirm'/'look' note is folded into the comment text, and the
# Accept decision can be overridden. The tool still NEVER rewrites a field cell.
# All of this is gated on a non-empty `triage`; with triage=None the output is
# byte-for-byte the same as before (regression depends on this).
def finding_key(f):
    """Content-stable triage key for an extra finding (positional keys drift when the list changes)."""
    return 'f:' + hashlib.sha1(('%s|%s|%s' % (f.code, f.anchor or '', (f.msg or '')[:120]))
                               .encode('utf-8')).hexdigest()[:10]

def finding_keys(findings):
    """Keys for a findings list, '#2'/'#3' suffixes de-duplicating identical findings in order."""
    keys, seen = [], {}
    for f in findings:
        k = finding_key(f)
        seen[k] = seen.get(k, 0) + 1
        keys.append(k if seen[k] == 1 else '%s#%d' % (k, seen[k]))
    return keys

_OLD_FKEY = re.compile(r'^f\d+$')

def _legacy_fkey(triage, f, label_counts):
    """Old positional triage key ('f0'/'f1'/…) for this finding, recovered via the
    GUI's saved label ('code: msg[:80]') when the content-stable key misses. Only
    honoured when the label is unambiguous on BOTH sides — one saved record, one
    current finding — because a lost dismissal merely re-surfaces a comment, while
    a wrong match would suppress the wrong one. Returns the old key or None."""
    label = '%s: %s' % (f.code, (f.msg or '')[:80])
    if label_counts.get(label, 0) != 1:
        return None
    hits = [k for k, v in (triage.get('findings') or {}).items()
            if _OLD_FKEY.match(k) and isinstance(v, dict) and v.get('label') == label]
    return hits[0] if len(hits) == 1 else None

def _t_dismissed(triage, fkey):
    return bool(triage and (triage.get('findings') or {}).get(fkey, {}).get('verdict') == 'dismiss')

def _t_note(triage, fkey):
    v = (triage.get('findings') or {}).get(fkey, {}) if triage else {}
    return v.get('note') if (v.get('verdict') in ('confirm', 'look') and v.get('note')) else None

def _with_note(text, triage, fkey):
    note = _t_note(triage, fkey)
    return _tidy(text + '  · reviewer: ' + note) if note else text

def _write_extras(doc, ac_row, res, rec, triage=None):
    """Write the writable extra-check findings as comments (flags also highlight
    their anchor cell). Authored as the same 'PXRD Review Tool' so they filter
    apart from human reviewers'. `triage` (optional) suppresses dismissed
    findings and folds reviewer notes into the comment text."""
    writable_ids = {id(f) for f in _writable_extras(res)}
    extras = res.get('extra', [])
    fkeys = finding_keys(extras)                       # content-stable (positional keys drift)
    label_counts = {}                                  # for the old-key fallback's ambiguity guard
    for f in extras:
        lab = '%s: %s' % (f.code, (f.msg or '')[:80])
        label_counts[lab] = label_counts.get(lab, 0) + 1
    for i, f in enumerate(extras):
        if id(f) not in writable_ids:
            continue
        fkey = fkeys[i]
        # backward compat: a triage.json saved before content-stable keys still uses
        # positional 'f0'/'f1' keys — recover the record via its saved label.
        if triage and fkey not in (triage.get('findings') or {}):
            fkey = _legacy_fkey(triage, f, label_counts) or fkey
        if _t_dismissed(triage, fkey):
            rec.setdefault('suppressed', []).append(f.code)
            continue
        cell = _anchor_cell(doc, ac_row, f.anchor) or ac_row.cells[0]
        runs = _cell_runs(cell) or _cell_runs(ac_row.cells[0])
        if not runs:
            continue
        if f.sev == 'flag':
            _highlight(cell)
            rec['highlights'].append(f.anchor or f.code)
        text = _with_note(_tidy('%s [%s] — %s' % (_check_banner(f.code), f.code, f.msg)), triage, fkey)
        doc.add_comment(runs, text=text, author=AUTHOR, initials=INITIALS)
        rec['comments'].append(text)
        # A finding carrying a `fix` gets it WRITTEN into the cell, as a tracked change — but only
        # where no person has been there first. If the reviewer already edited this cell, their
        # version stands and the tool stays a comment. Every other check has no fix and remains
        # comment-only, unchanged.
        if f.fix and f.sev == 'flag':
            if _cell_human_marks(cell):
                rec.setdefault('fix_skipped', []).append(f.code)
            elif _cell_text(cell).strip() == f.fix.strip():
                pass                                   # already reads as the correction
            elif _apply_tracked_fix(cell, f.fix):
                rec.setdefault('fixes', []).append((f.code, f.fix))
            else:
                # the writer declined (a layout it will not risk, e.g. a multi-paragraph cell) —
                # the comment still carries the correction, so nothing is lost but the automation
                rec.setdefault('fix_skipped', []).append(f.code)

def _mark_accept(doc):
    """Put an 'x' in the checkbox cell after the 'Accept' label — the reviewer's
    convention (verified against past reviews). Skips if a human already marked
    any of Accept/Reject/Replace. Returns True if Accept was (or already is) set."""
    for t in doc.tables:
        for row in t.rows:
            cells = row.cells
            idx = {c.text.strip().lower(): i for i, c in enumerate(cells)}
            if not ({'accept', 'reject', 'replace'} <= set(idx)):
                continue
            boxes = [cells[idx[k] + 1] for k in ('accept', 'reject', 'replace')
                     if idx[k] + 1 < len(cells)]
            if any(b.text.strip() for b in boxes):
                return True                       # already decided — don't touch
            acc = cells[idx['accept'] + 1]
            p = acc.paragraphs[0] if acc.paragraphs else acc.add_paragraph()
            if p.runs:
                p.runs[0].text = 'x'
            else:
                p.add_run('x')
            # centre the 'x' in the checkbox cell (it otherwise sits left, hard against
            # the 'Accept' label); also vertically centre within the cell.
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                acc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            except Exception:
                pass
            return True
    return False

def _clear_accept(doc):
    """Empty the Accept checkbox cell, but ONLY when it holds exactly the tool's own
    'x' (case/whitespace-insensitive). A triage 'disagree' must actively remove the
    mark on the refresh path — the base there is a previous output where an earlier
    run already stamped Accept, so merely skipping _mark_accept leaves it standing.
    Anything else in the box is the reviewer's decision and is never touched."""
    for t in doc.tables:
        for row in t.rows:
            cells = row.cells
            idx = {c.text.strip().lower(): i for i, c in enumerate(cells)}
            if not ({'accept', 'reject', 'replace'} <= set(idx)):
                continue
            if idx['accept'] + 1 >= len(cells):
                return False
            acc = cells[idx['accept'] + 1]
            if acc.text.strip().lower() != 'x':
                return False                      # not the tool's mark — don't touch
            for p in acc.paragraphs:
                for r in p.runs:
                    r.text = ''
            return True
    return False

def _is_severe(res):
    """Withhold auto-Accept ONLY when the cell is fundamentally wrong ('way off').
    Per the reviewer, essentially every new mineral is an Accept — minor cell /
    significant-figure / esd discrepancies and Z mismatches are NOT severe (the
    editorial team fixes those). Only a completely wrong cell (several axes off,
    or a single axis grossly off) is left blank for a manual decision."""
    if res.get('cell', ('',))[0] != 'investigate':
        return False
    out = [t for t in (res.get('cell_diffs') or []) if not t[4]]   # axes off tol
    if len(out) >= 2:                                              # multiple axes wrong
        return True
    for _lab, dv, nv, _d, _m in out:                              # one axis grossly off
        x, y = C.num_val(dv), C.num_val(nv)
        if x and y and abs(x - y) / y > 0.02:
            return True
    return False

def _body_signature(doc):
    """Concatenated body text (paragraphs + table cells), used to tell a human-edited
    output from the untouched tool output. ONLY the Accept value cell is excluded — the
    tool auto-stamps its 'x' there, so including it would make every clean output look
    hand-edited. Reject and Replace are NEVER written by the tool, so a mark in either is
    a reviewer's decision that MUST register as a hand edit; otherwise a rerun rebuilds
    the entry from source and silently re-stamps Accept over the reviewer's choice."""
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            cells = row.cells
            idx = {c.text.strip().lower(): i for i, c in enumerate(cells)}
            skip = set()
            if {'accept', 'reject', 'replace'} <= set(idx) and idx['accept'] + 1 < len(cells):
                skip.add(idx['accept'] + 1)         # the tool's own auto-Accept 'x' only
            for i, c in enumerate(cells):
                if i not in skip:
                    parts.append(c.text)
    return '\n'.join(parts)

def _format_signature(doc):
    """Per-run FORMATTING signature over the same scope as _body_signature (body
    paragraphs + top-level table cells, excluding the Accept checkbox cell): a hash
    of (text, highlight, bold, italic, underline, strike, font colour) for every
    run, read through the python-docx API so an absent rPr normalises to None and
    round-trip serialisation differences can't false-positive. Catches formatting-
    only reviewer edits (e.g. a manual highlight) that leave the body TEXT alone."""
    h = hashlib.sha1()
    def add(paragraphs):
        for p in paragraphs:
            for r in p.runs:
                f = r.font
                color = f.color.rgb if (f.color is not None and f.color.type is not None) else None
                h.update(repr((r.text, f.highlight_color, f.bold, f.italic,
                               f.underline, f.strike, color)).encode('utf-8'))
    add(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            cells = row.cells
            idx = {c.text.strip().lower(): i for i, c in enumerate(cells)}
            skip = set()
            if {'accept', 'reject', 'replace'} <= set(idx) and idx['accept'] + 1 < len(cells):
                skip.add(idx['accept'] + 1)         # the tool's own auto-Accept 'x' only
            for i, c in enumerate(cells):
                if i not in skip:
                    add(c.paragraphs)
    return h.hexdigest()

def _has_foreign_comment(path):
    """True if the docx carries a comment authored by someone other than the tool."""
    try:
        z = zipfile.ZipFile(path)
        if 'word/comments.xml' not in z.namelist():
            return False
        cx = z.read('word/comments.xml').decode('utf-8', 'replace')
        return any(a != AUTHOR for a in re.findall(r'w:author="([^"]*)"', cx))
    except Exception:
        return False

def _has_tool_comment(path):
    """True if the docx carries a comment authored by the tool itself — i.e. a
    previous run already annotated it (an --inplace rerun must strip those first,
    or every comment/highlight would be added a second time)."""
    try:
        z = zipfile.ZipFile(path)
        if 'word/comments.xml' not in z.namelist():
            return False
        cx = z.read('word/comments.xml').decode('utf-8', 'replace')
        return AUTHOR in re.findall(r'w:author="([^"]*)"', cx)
    except Exception:
        return False

def _has_tracked_changes(path):
    """True if the docx contains a HUMAN's tracked changes (w:ins / w:del).

    The tool now makes revisions of its own — it writes an applied fix as a tracked change so the
    reviewer can Accept/Reject it — so 'any w:ins/w:del' no longer means 'a person edited this'.
    Counting the tool's own change as a human edit would make every rerun think the output was
    hand-edited: it would take a backup each time and report the tool's own work back to the
    reviewer as theirs. Author is what separates them."""
    try:
        root = _xml(zipfile.ZipFile(path).read('word/document.xml'))
        for tag in ('ins', 'del'):
            for el in root.iter(_q(tag)):
                if (el.get(_q('author')) or '') != AUTHOR:
                    return True
        return False
    except Exception:
        return False

def _extract_reviewer_edits(path):
    """Human-readable summary of the reviewer's OWN marks in a hand-edited docx —
    tracked changes (w:ins/w:del) and comments NOT authored by the tool. For the
    annotation log only (so the reviewer's edits are visible alongside the tool's
    findings); never used to modify a docx. Best-effort: returns [] on any error."""
    out = []
    try:
        z = zipfile.ZipFile(path)
        names = z.namelist()
        droot = _xml(z.read('word/document.xml'))
        # tracked changes, in document order; pair an adjacent del+ins (same author,
        # either order) into a single 'old -> new' replacement.
        revs = []
        for el in droot.iter():
            # the tool's OWN revisions (its applied fixes) are not "the reviewer's edits" —
            # they are reported separately, under 'fixes written', and must not be echoed back
            # to the reviewer as though they had made them.
            if el.tag in (_q('ins'), _q('del')) and (el.get(_q('author')) or '') == AUTHOR:
                continue
            if el.tag == _q('ins'):
                revs.append(('ins', el.get(_q('author')) or '?',
                             ''.join(t.text or '' for t in el.iter(_q('t')))))
            elif el.tag == _q('del'):
                revs.append(('del', el.get(_q('author')) or '?',
                             ''.join(t.text or '' for t in el.iter(_q('delText')))))
        i = 0
        while i < len(revs):
            kind, auth, txt = revs[i]
            nxt = revs[i + 1] if i + 1 < len(revs) else None
            if nxt and nxt[1] == auth and {kind, nxt[0]} == {'ins', 'del'}:
                old, new = (txt, nxt[2]) if kind == 'del' else (nxt[2], txt)
                out.append("tracked change: '%s' -> '%s'  (%s)" % (old.strip(), new.strip(), auth))
                i += 2
            elif kind == 'ins':
                out.append("inserted '%s'  (%s)" % (txt.strip(), auth)); i += 1
            else:
                out.append("deleted '%s'  (%s)" % (txt.strip(), auth)); i += 1
        # comments authored by someone other than the tool
        if 'word/comments.xml' in names:
            croot = _xml(z.read('word/comments.xml'))
            for c in croot:
                if c.tag == _q('comment') and c.get(_q('author')) != AUTHOR:
                    body = ''.join(t.text or '' for t in c.iter(_q('t'))).strip()
                    out.append("comment (%s): %s" % (c.get(_q('author')) or '?', body))
        z.close()
    except Exception:
        return out
    return out

def output_hand_edited(out_path, src_path):
    """The output docx was edited by a human since the tool wrote it — tracked
    changes, a non-tool comment, body text differing from the source (beyond the
    Accept checkbox), or a formatting-only edit (e.g. a manual highlight) the text
    compare can't see. Such edits must be preserved across reruns, so on ANY doubt
    (an unreadable/corrupt output included) err toward hand-edited: that path backs
    up + refreshes, whereas a False here rebuilds from source and destroys work."""
    try:
        if _has_tracked_changes(out_path) or _has_foreign_comment(out_path):
            return True
        # Compare against the source only AFTER stripping the tool's own marks off a TEMP copy.
        # Stripping reverts its applied fixes too, so what remains is the source plus whatever a
        # PERSON did. Comparing the un-stripped output would read the tool's own corrected title
        # as a body-text change and call every fixed entry 'hand-edited' — a backup on every
        # rerun, and the tool's work reported back to the reviewer as theirs.
        bdir = os.path.join(os.path.dirname(out_path), '.edit_backup')
        os.makedirs(bdir, exist_ok=True)
        tmp = os.path.join(bdir, '~sig_' + os.path.basename(out_path))
        try:
            shutil.copy(out_path, tmp)
            _strip_tool_annotations(tmp)
            stripped = Document(tmp)
            src = Document(src_path)
            if _body_signature(stripped) != _body_signature(src):
                return True
            # text-identical: a reviewer's manual highlight/bold with no text change must still
            # register as a hand edit, so compare run-level FORMATTING as well.
            return _format_signature(stripped) != _format_signature(src)
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)
    except Exception:
        return True

def _strip_tool_annotations(path):
    """Remove the tool's OWN comments (author == AUTHOR) and its yellow highlights
    from a docx in place, leaving everything human: tracked changes, body edits,
    and any non-tool comments. Lets a rerun refresh the tool's findings without
    disturbing the reviewer's work."""
    z = zipfile.ZipFile(path)
    parts = {n: z.read(n) for n in z.namelist()}
    z.close()
    tool_ids = set()
    if 'word/comments.xml' in parts:
        croot = _xml(parts['word/comments.xml'])
        for c in list(croot):
            if c.tag == _q('comment') and c.get(_q('author')) == AUTHOR:
                tool_ids.add(c.get(_q('id')))
                croot.remove(c)
        parts['word/comments.xml'] = etree.tostring(croot, xml_declaration=True,
                                                    encoding='UTF-8', standalone=True)
    droot = _xml(parts['word/document.xml'])
    # Identify the tool's OWN yellow highlights BEFORE the comment ranges are removed.
    # The annotator highlights a cell and comments the SAME runs, so every tool
    # highlight sits inside that finding's comment range; a yellow highlight that is
    # NOT inside any tool comment range is the reviewer's own and must be preserved.
    # Walk in document order (preorder), tracking how many tool ranges are open.
    tool_highlights, open_ranges = [], 0
    for el in droot.iter():
        tag = el.tag
        if tag == _q('commentRangeStart') and el.get(_q('id')) in tool_ids:
            open_ranges += 1
        elif tag == _q('commentRangeEnd') and el.get(_q('id')) in tool_ids:
            open_ranges = max(0, open_ranges - 1)
        elif open_ranges and tag == _q('highlight') and el.get(_q('val')) == 'yellow':
            tool_highlights.append(el)
    # drop the comment anchors/refs for tool comments
    for tag in ('commentRangeStart', 'commentRangeEnd'):
        for el in list(droot.iter(_q(tag))):
            if el.get(_q('id')) in tool_ids:
                el.getparent().remove(el)
    for ref in list(droot.iter(_q('commentReference'))):
        if ref.get(_q('id')) in tool_ids:
            run = ref.getparent()                      # the <w:r> wrapping the ref
            if run is not None and run.getparent() is not None:
                run.getparent().remove(run)
    # drop ONLY the tool's own yellow highlights (re-applied fresh on re-annotation);
    # a reviewer's manual yellow highlight, outside any tool comment range, is kept.
    for hl in tool_highlights:
        parent = hl.getparent()
        if parent is not None:
            parent.remove(hl)
    # REVERT the tool's own tracked changes (its applied fixes) so a rerun re-derives them from
    # scratch instead of stacking a second delete/insert on the first. This is a REJECT of the
    # tool's edit: drop what it inserted, restore what it struck through. A PERSON's tracked
    # changes (any other author) are never touched — including one that accepted the tool's fix
    # by retyping it, which then simply reads as their own edit.
    for ins in list(droot.iter(_q('ins'))):
        if (ins.get(_q('author')) or '') == AUTHOR and ins.getparent() is not None:
            ins.getparent().remove(ins)
    for dl in list(droot.iter(_q('del'))):
        if (dl.get(_q('author')) or '') != AUTHOR:
            continue
        parent = dl.getparent()
        if parent is None:
            continue
        at = list(parent).index(dl)
        for run in list(dl):                            # put the original runs back, in place
            for t in run.findall(_q('delText')):
                t.tag = _q('t')
            dl.remove(run)
            parent.insert(at, run)
            at += 1
        parent.remove(dl)
    parts['word/document.xml'] = etree.tostring(droot, xml_declaration=True,
                                                encoding='UTF-8', standalone=True)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zo:
        for n, data in parts.items():
            zo.writestr(n, data)
    with open(path, 'wb') as f:
        f.write(buf.getvalue())
    return tool_ids

def _backup_edited(out):
    """Timestamped copy of a hand-edited output into .edit_backup/ — so ANY rerun, including
    --force, can never lose the reviewer's work irrecoverably. Returns the backup path or None."""
    try:
        bdir = os.path.join(os.path.dirname(out), '.edit_backup')
        os.makedirs(bdir, exist_ok=True)
        stem, ext = os.path.splitext(os.path.basename(out))
        ts = datetime.datetime.now().strftime('%Y-%m-%dT%H%M%S')
        dst = os.path.join(bdir, '%s.%s%s' % (stem, ts, ext))
        shutil.copy(out, dst)
        return dst
    except Exception:
        return None

def annotate(docx_path, res, out_path, inplace=False, base_path=None, triage=None):
    """Write comments/highlights for `res` into the docx — ERRORS ONLY. Entries
    whose cell matches the PDF with no discrepancies get no annotation (the
    unmodified file is copied through, preserving the original bytes). Returns
    the number of comments added.

    Reported (each a separate comment):
      * per-parameter value / significant-figure / esd mismatches (highlighted);
      * a radiation anode hard-flag (highlighted), only when the cell matched;
      * a single brief 'No matching PDF cell found.' when no cell could be
        matched (no cell parsed, or no exact match — both reduce to this);
      * extra-check FLAGS — symmetry/precision (on the offending cell parameter),
        Calculated pattern (Instrumentation), hkl-indexing (Reflection List),
        missing IMA number, ideal-formula — each highlighted on its anchor cell;
      * the authoritative Mindat group statement (informational, on the mineral
        name; not highlighted).

    Returns a record: {'status', 'highlights': [cell labels], 'comments':
    [texts], 'clean': bool} for the run log."""
    rec = {'status': res['cell'][0], 'highlights': [], 'comments': [], 'clean': False,
           'accept': False}
    # `base_path` is the document to annotate ONTO — the source for a normal run,
    # or the reviewer's edited output (already stripped of old tool comments) when
    # refreshing in place so manual edits are preserved.
    base = base_path or docx_path
    if _is_clean(res):
        rec['clean'] = True
        # clean entries are never severe → auto-Accept. (We now open+save to set
        # the Accept 'x', so they are no longer byte-identical copies.)
        out = docx_path if inplace else out_path
        doc = Document(base)
        # clean entry is never severe → auto-Accept, unless the reviewer disagreed
        # (then also UNDO a previous run's 'x' — refresh/inplace bases carry it).
        if not (triage and triage.get('accept') == 'disagree'):
            rec['accept'] = _mark_accept(doc)
        else:
            _clear_accept(doc)
        _save_docx(doc, out)
        return rec

    doc = Document(base)
    ac_row, rad_row = _rows(doc)
    if ac_row is None:
        raise RuntimeError('no Author\'s Cell row found')
    status = res['cell'][0]

    if status == 'match':
        # per-parameter flags: highlight the value cell + explain
        for k, issues in res['params'].items():
            col = PARAM_COL.get(k)
            if col is None or col >= len(ac_row.cells):
                continue
            if _t_dismissed(triage, 'param:%s' % k):
                rec.setdefault('suppressed', []).append('param:%s' % k); continue
            cell = ac_row.cells[col]
            _highlight(cell)
            runs = _cell_runs(cell) or _cell_runs(ac_row.cells[0])
            body = '; '.join('%s — %s' % (KIND_LABEL[kind], note) for kind, note in issues)
            text = _with_note(_tidy('PXRD check — %s: %s' % (res.get('param_labels', {}).get(k, k), body)), triage, 'param:%s' % k)
            doc.add_comment(runs, text=text, author=AUTHOR, initials=INITIALS)
            rec['highlights'].append("Author's Cell:%s" % k)
            rec['comments'].append(text)
        # radiation hard-flag: highlight anode cell + explain
        if (_has_lam_flag(res) and not _t_dismissed(triage, 'lam')
                and rad_row is not None and len(rad_row.cells) > 1):
            cell = rad_row.cells[1]
            _highlight(cell)
            runs = _cell_runs(cell) or _cell_runs(rad_row.cells[0])
            text = _with_note(_tidy('PXRD check — radiation: ' + res['lam'][1]), triage, 'lam')
            doc.add_comment(runs, text=text, author=AUTHOR, initials=INITIALS)
            rec['highlights'].append('Radiation:anode')
            rec['comments'].append(text)
        elif _has_lam_flag(res) and _t_dismissed(triage, 'lam'):
            rec.setdefault('suppressed', []).append('lam')
    elif _t_dismissed(triage, 'cell'):
        # reviewer dismissed the cell-level comment (investigate / no-match) entirely.
        rec.setdefault('suppressed', []).append('cell')
    elif status == 'investigate' and res.get('cell_diffs'):
        # closest PDF cell exists but isn't an exact match — pinpoint the axis.
        diffs = res['cell_diffs']
        out_axes = [t for t in diffs if not t[4]]
        in_axes = [t for t in diffs if t[4]]
        if len(out_axes) == 1 and in_axes:
            lab, dv, nv, dd, _ = out_axes[0]
            col = PARAM_COL.get(lab)
            cell = ac_row.cells[col] if (col is not None and col < len(ac_row.cells)) else ac_row.cells[0]
            _highlight(cell)
            text = _with_note(_tidy('PXRD check — %s: likely transcription error — docx=%s but .pdf cell gives %s '
                         '(Δ=%.4f Å); the other axes match exactly.' % (lab, dv, nv, dd)), triage, 'cell')
            doc.add_comment(_cell_runs(cell) or _cell_runs(ac_row.cells[0]),
                            text=text, author=AUTHOR, initials=INITIALS)
            rec['highlights'].append("Author's Cell:%s" % lab)
            rec['comments'].append(text)
        else:
            text = _with_note('No exact cell match — closest .pdf cell differs on %s; it may be a different '
                    'phase/cell in a multi-cell .pdf (the matching cell may be unparsed).'
                    % ', '.join(t[0] for t in out_axes), triage, 'cell')
            doc.add_comment(_cell_runs(ac_row.cells[0]), text=text, author=AUTHOR, initials=INITIALS)
            rec['comments'].append(text)
    else:
        # no cell parsed at all (nocell / nopdf / investigate w/o diffs): one brief flag
        doc.add_comment(_cell_runs(ac_row.cells[0]),
                        text=_with_note(NO_MATCH, triage, 'cell'), author=AUTHOR, initials=INITIALS)
        rec['comments'].append(NO_MATCH)

    # the 10 extra reviewer-comment checks (symmetry, calculated, indexing, IMA,
    # ideal-formula, and the authoritative Mindat group statement)
    _write_extras(doc, ac_row, res, rec, triage)

    # auto-Accept unless the entry is severe (then leave blank); the reviewer can
    # override the tool's Accept decision via triage ('disagree' flips it).
    _tool_severe = _is_severe(res)
    _severe = (not _tool_severe) if (triage and triage.get('accept') == 'disagree') else _tool_severe
    if not _severe:
        rec['accept'] = _mark_accept(doc)
    elif triage and triage.get('accept') == 'disagree':
        _clear_accept(doc)     # refresh/inplace base carries a previous run's 'x' — undo it

    _save_docx(doc, out_path)
    return rec

# ----------------------------------------------------------------------------- main
def output_name(src_basename, edited):
    """Output filename for a source docx: flagged entries (the tool added a
    comment/highlight) get an '_edited' suffix so they stand out in the folder
    listing; clean entries keep the source name."""
    stem, ext = os.path.splitext(src_basename)
    return stem + ('_edited' if edited else '') + ext

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('folder')
    ap.add_argument('--id', help='only this entry id, e.g. I003416')
    ap.add_argument('--out', help='output folder (default <folder>/review_out)')
    ap.add_argument('--inplace', action='store_true', help='edit the original .docx in place')
    ap.add_argument('--limit', type=int, help='only the first N entries')
    ap.add_argument('--log', help='log file path (default <out>/annotation_log.txt)')
    ap.add_argument('--force', action='store_true',
                    help='regenerate every output even if it has manual edits (default: preserve hand-edited outputs)')
    ap.add_argument('--triage', help='review_gui triage.json — suppress dismissed findings, fold reviewer '
                    'notes into the comments, override Accept (COMMENT-ONLY; never rewrites a field)')
    ap.add_argument('--no-logs', action='store_true',
                    help='skip the folder-level annotation_log.txt / mindat_discrepancies.txt writes '
                         '(used by the GUI single-entry rerun so it does not clobber the batch logs)')
    args = ap.parse_args()

    # Refuse to run ON the tool's own output tree. discover() can now index a review_out folder
    # (the GUI opens one to view/resume a review), so without this guard a mistyped path — e.g.
    # tab-completing one level too deep — would re-annotate the outputs: duplicate comments, and
    # with --inplace written straight INTO the reviewer's hand-edited copies.
    if os.path.basename(os.path.abspath(args.folder)) in ('review_out', '.edit_backup'):
        sys.exit("refusing to annotate the tool's own output folder (%s) — run on its parent instead"
                 % args.folder)

    # Refuse to write the outputs INTO the source folder. The outputs are named after the
    # source ('<name>.docx' when clean, '<name>_edited.docx' when flagged), so an --out that
    # resolves to the source folder makes the tool's own stale-twin bookkeeping treat a SOURCE
    # docx as a previous run's output: it would delete or rename it. Copies only — never the
    # source. (--inplace is the explicit, documented way to edit originals.)
    if args.out and not args.inplace and \
       os.path.realpath(args.out) == os.path.realpath(args.folder):
        sys.exit("refusing --out '%s': that is the source folder, and the tool would overwrite or "
                 "delete the source .docx files. Omit --out (writes to <folder>/review_out), point "
                 "it elsewhere, or use --inplace to edit the originals deliberately." % args.out)

    triage = {}
    if args.triage and os.path.exists(args.triage):
        try:
            with open(args.triage, encoding='utf-8') as _tf:
                triage = json.load(_tf)
        except Exception as ex:
            print('  !! could not read triage %s (%s) — ignoring' % (args.triage, ex))

    idx = C.pdf_index(args.folder)
    cif_idx = C.cif_index(args.folder)
    dft_idx = C.dft_index(args.folder)
    try:                                    # keep the Mindat structural cache current (cross-source check)
        from pxrd_review import mindat
        mindat.refresh_struct_if_stale()
        # State the cache's age up front. The Mindat-backed checks go SILENT when the cache is
        # absent — they just find nothing — so without this line a missing cache is indistinguishable
        # from a clean corpus.
        mindat.print_cache_banner()
    except Exception:
        pass
    docs = sorted(C.discover(args.folder).values())   # recursive: docx at any depth under the root
    if args.id:
        # Match the parsed entry id exactly: the corpus mixes 5- and 6-digit ids, so a plain
        # substring test on the filename makes '--id I10126' also select 'I101261…' and rerun
        # (and rewrite the output of) a neighbouring entry. Fall back to the substring form
        # only when nothing matches exactly, so a partial id still works interactively.
        want = args.id.strip().upper()
        exact = [d for d in docs if (C.entry_id(d) or '').upper() == want]
        docs = exact or [d for d in docs if want in os.path.basename(d).upper()]
    if args.limit:
        docs = docs[:args.limit]
    out_dir = args.out or os.path.join(args.folder, 'review_out')

    if not docs:
        print('no entry .docx found under %s%s'
              % (args.folder, (' matching --id %s' % args.id) if args.id else ''))
        return

    # Single writer per output folder (see _RunLock): the GUI's 'Rerun all' and a terminal run
    # on the same folder would otherwise interleave writes to the same docx and corrupt it
    # without saying so. --inplace writes to the SOURCE folder, so lock that instead.
    lock = _RunLock(os.path.dirname(docs[0]) if args.inplace else out_dir)
    if not lock.acquire():
        sys.exit("another annotate_review run is already writing to this folder (lock: %s).\n"
                 "Wait for it to finish, or — if no run is active — delete that lock file."
                 % lock.path)
    try:
        _run(args, docs, out_dir, idx, cif_idx, dft_idx, triage)
    finally:
        lock.release()

def _run(args, docs, out_dir, idx, cif_idx, dft_idx, triage):
    records = []           # (filename, rec) for the log
    skipped = []
    for dp in docs:
        eid = C.entry_id(dp)
        pdf = idx.get(eid)
        cif = cif_idx.get(eid)
        dft = dft_idx.get(eid)
        # A file that cannot even be OPENED (an HTML error page named .docx, a truncated .pdf,
        # a docx held open by Word) must cost one entry, not the batch: analyze() -> parse_docx()
        # raises before any check runs, and unguarded that aborted the run and left every
        # remaining entry unreviewed — silently, since the ones already written looked fine.
        try:
            res = analyze(dp, pdf, cif, dft)
        except Exception as e:
            print('  !! %s: SKIPPED — %s' % (os.path.basename(dp), E.explain(e, dp)))
            skipped.append(os.path.basename(dp))
            continue
        # Name the output by what the tool found: flagged entries -> '<name>_edited.docx',
        # clean ones keep the source name. Cleanliness is deterministic for a given
        # source, so the path is stable across reruns (refresh below still works).
        if args.inplace:
            out = dp
            # An --inplace RERUN would stack a second copy of every comment/highlight on
            # top of the first: if the source already carries the tool's own comments,
            # strip them first. Strip a TEMP copy and swap it in atomically (os.replace)
            # so a crash mid-strip can never leave a half-written source.
            if _has_tool_comment(dp):
                bdir = os.path.join(os.path.dirname(dp), '.edit_backup')
                tmp = os.path.join(bdir, '~strip_' + os.path.basename(dp))
                try:
                    os.makedirs(bdir, exist_ok=True)
                    shutil.copy(dp, tmp)
                    _strip_tool_annotations(tmp)
                    os.replace(tmp, dp)
                except Exception as e:
                    print('  !! %s: could not strip the previous run\'s annotations — left '
                          'untouched.\n     %s' % (os.path.basename(dp), E.explain(e, dp)))
                    continue
                finally:
                    if os.path.exists(tmp):
                        os.remove(tmp)
        else:
            edited = not _is_clean(res)
            out = os.path.join(out_dir, output_name(os.path.basename(dp), edited))
            # A prior run may have written the opposite-named twin (clean<->edited can
            # flip if the source/PDF changed, or it predates this naming). Keep one
            # output per entry without ever losing the reviewer's work:
            #   * a pristine tool output under the old name -> drop it;
            #   * a hand-edited old name with no new-name output yet -> RENAME it onto
            #     the new name, so refresh-in-place (below) carries the edits forward;
            #   * both names already present and the old one edited -> leave it, warn.
            stale = os.path.join(out_dir, output_name(os.path.basename(dp), not edited))
            if os.path.exists(stale):
                if not output_hand_edited(stale, dp):
                    os.remove(stale)
                elif not os.path.exists(out):
                    os.rename(stale, out)
                    print('  -- %s: migrated hand-edited %s -> %s (edits preserved)'
                          % (os.path.basename(dp), os.path.basename(stale), os.path.basename(out)))
                else:
                    print('  !! %s: hand-edited %s sits alongside %s — reconcile manually'
                          % (os.path.basename(dp), os.path.basename(stale), os.path.basename(out)))
        # If the reviewer hand-edited this output, ALWAYS preserve a timestamped backup before
        # touching it — even under --force — so their work can never be lost irrecoverably.
        # Then either REFRESH in place (strip only the tool's OWN comments/highlights and
        # re-annotate ONTO the edited file, keeping the human edits), or — under --force —
        # rebuild from source, with the backup as the safety net.
        hand_edited = (not args.inplace and os.path.exists(out) and output_hand_edited(out, dp))
        # capture the reviewer's OWN marks (tracked changes / non-tool comments) from the
        # still-hand-edited output, before annotate() refreshes it — for the log only.
        user_edits = _extract_reviewer_edits(out) if hand_edited else []
        if hand_edited:
            # No backup means NO safety net — do not touch the entry at all (neither the
            # refresh nor a --force rebuild), or the reviewer's edits could be lost.
            if _backup_edited(out) is None:
                print('  !! %s: BACKUP FAILED — could not copy the hand-edited %s to '
                      '.edit_backup/; left untouched' % (os.path.basename(dp), os.path.basename(out)))
                records.append((os.path.basename(dp),
                                {'status': 'preserved', 'highlights': [], 'comments': [],
                                 'clean': False, 'accept': None, 'refreshed': True,
                                 'user_edits': user_edits}))
                continue
            if args.force:
                print('  !! %s: --force rebuilt a HAND-EDITED output from source; your edits were '
                      'saved to .edit_backup/ first' % os.path.basename(dp))
        refresh = hand_edited and not args.force
        base = None
        if refresh:
            # strip a TEMP copy (never the real output) so a later annotate() failure can't
            # leave it half-done; the timestamped backup above is the durable safeguard.
            bdir = os.path.join(os.path.dirname(out), '.edit_backup')
            os.makedirs(bdir, exist_ok=True)
            try:
                base = os.path.join(bdir, '~strip_' + os.path.basename(out))
                shutil.copy(out, base)
                _strip_tool_annotations(base)    # strip the temp, not `out`
            except Exception as e:
                print('  !! %s: strip failed — left untouched.\n     %s'
                      % (os.path.basename(dp), E.explain(e, out)))
                records.append((os.path.basename(dp),
                                {'status': 'preserved', 'highlights': [], 'comments': [],
                                 'clean': False, 'accept': None, 'refreshed': True,
                                 'user_edits': user_edits}))
                continue
        try:
            tkey = os.path.splitext(os.path.basename(dp))[0]    # matches the GUI's triage key
            rec = annotate(dp, res, out, inplace=args.inplace, base_path=base,
                           triage=triage.get(tkey))
        except Exception as e:
            # `out` is untouched on failure (annotate read the temp, not `out`), and _save_docx
            # writes via a temp + os.replace, so a failure mid-save cannot truncate it either.
            print('  !! %s: SKIPPED — %s' % (os.path.basename(dp), E.explain(e, dp)))
            continue
        finally:
            if base and os.path.exists(base):
                os.remove(base)
        if refresh:
            rec['refreshed'] = True
        rec['user_edits'] = user_edits
        rec['outfile'] = os.path.basename(out)
        # .dft cross-check notes are console/log only (the .dft is a co-equal proxy);
        # collect them for the log's 'verify' section — they are NOT written into the docx.
        rec['dft'] = [f.msg for f in res.get('extra', []) if f.code == 'dft']
        # Mindat cross-check notes (cell / chemistry / status) — a co-equal proxy,
        # so these go to a SEPARATE feedback log, never into the docx.
        rec['mindat'] = [(f.code, f.msg) for f in res.get('extra', [])
                         if f.code in ('mindat_fix', 'mindat_chem', 'ima_status')]
        records.append((os.path.basename(dp), rec))
        n = len(rec['comments'])
        tag = ' [REFRESHED: manual edits kept]' if rec.get('refreshed') else ''
        renamed = '' if (args.inplace or rec['clean']) else ' -> %s' % rec['outfile']
        dtag = '  {%d .dft-verify}' % len(rec['dft']) if rec['dft'] else ''
        # An applied fix EDITS the copy (as a tracked change), so it must never be silent.
        ftag = ''
        if rec.get('fixes'):
            ftag = '  [%d fix%s written: %s]' % (len(rec['fixes']),
                                                 '' if len(rec['fixes']) == 1 else 'es',
                                                 ', '.join(c for c, _ in rec['fixes']))
        if rec.get('fix_skipped'):
            ftag += '  [fix left to the reviewer (already hand-edited): %s]' % ', '.join(rec['fix_skipped'])
        print('%-44s cell=%-12s %-26s (%d comment%s)%s%s%s%s'
              % (os.path.basename(dp), rec['status'],
                 'clean' if rec['clean'] else ('highlights: ' + ', '.join(rec['highlights']) or 'flag'),
                 n, '' if n == 1 else 's', renamed, dtag, tag, ftag))

    # Say plainly that some entries were NOT reviewed. An unreadable entry produces no
    # output docx, so without this line it just quietly isn't there.
    if skipped:
        print('\n!! %d of %d entries NOT reviewed (unreadable — see the !! lines above for the '
              'cause and the fix): %s' % (len(skipped), len(docs), ', '.join(skipped)))

    # The single-entry GUI rerun passes --no-logs so it regenerates one docx
    # without rewriting the batch-wide annotation_log.txt / mindat_discrepancies.txt
    # (which, scoped to one entry, would otherwise clobber the full-folder files).
    if args.no_logs:
        return

    # A filtered run (--id / --limit) covers only a subset: rewriting the batch logs
    # would scope annotation_log.txt to that subset and delete a full-folder
    # mindat_discrepancies.txt whose entries simply weren't rerun. Leave both alone;
    # an EXPLICIT --log path is still honoured (scoped to the subset).
    filtered = bool(args.id or args.limit)
    if filtered:
        print('(filtered run — batch logs left untouched%s)'
              % ('; explicit --log written, scoped to this subset' if args.log else ''))
        if not args.log:
            return

    # --- write the run log -------------------------------------------------
    log_path = args.log or os.path.join(out_dir if not args.inplace else args.folder, 'annotation_log.txt')
    refreshed = [(f, r) for f, r in records if r.get('refreshed')]
    edited = [(f, r) for f, r in records if not r['clean']]
    os.makedirs(os.path.dirname(log_path) or '.', exist_ok=True)
    with open(log_path, 'w', encoding='utf-8') as fh:
        fh.write('PXRD review — annotation log\n')
        fh.write('source folder : %s\n' % os.path.abspath(args.folder))
        fh.write('output        : %s\n' % ('in place' if args.inplace else os.path.abspath(out_dir)))
        fh.write('entries       : %d total | %d edited | %d clean (untouched)\n'
                 % (len(records), len(edited), len(records) - len(edited)))
        fh.write('total comments: %d | total highlights: %d\n'
                 % (sum(len(r['comments']) for _, r in records),
                    sum(len(r['highlights']) for _, r in records)))
        # Applied fixes CHANGE the copy's text, so the log states them explicitly — the reviewer
        # must be able to see, without opening a single docx, exactly what the tool rewrote.
        fixed = [(f, r) for f, r in records if r.get('fixes')]
        skipped_fix = [(f, r) for f, r in records if r.get('fix_skipped')]
        if fixed or skipped_fix:
            fh.write('fixes written : %d entr%s — the correction is a Word TRACKED CHANGE in the copy '
                     '(review it in Word: Accept or Reject)\n'
                     % (len(fixed), 'y' if len(fixed) == 1 else 'ies'))
            for f, r in fixed:
                for code, newtext in r['fixes']:
                    fh.write('  %s [%s]\n' % (C.entry_id(f) or f, code))
                    fh.write(textwrap.fill(newtext, width=76, initial_indent='      -> ',
                                           subsequent_indent=' ' * 9) + '\n')
            if skipped_fix:
                fh.write('  left to the reviewer (the cell was already hand-edited): %s\n'
                         % ', '.join(C.entry_id(f) or f for f, _ in skipped_fix))
        accepted = [f for f, r in records if r.get('accept') is True]
        withheld = [f for f, r in records if r.get('accept') is False]
        fh.write('Accept marked : %d auto-accepted | %d left blank (severe — decide manually)\n'
                 % (len(accepted), len(withheld)))
        if withheld:
            fh.write('  left blank  : %s\n' % ', '.join(C.entry_id(f) or f for f in withheld))
        if refreshed:
            fh.write('refreshed     : %s  (hand-edited — manual edits kept, tool comments refreshed; --force to rebuild from source)\n'
                     % ', '.join(C.entry_id(f) or f for f, r in refreshed))
        fh.write('=' * 78 + '\n\nEDITED ENTRIES (highlights / comments)\n')
        for f, r in edited:
            hl = ', '.join(r['highlights']) if r['highlights'] else '(none)'
            name = (C.entry_name(f) or C.entry_id(f) or f).upper()
            eid = C.entry_id(f)
            fh.write('\n' + '=' * 78 + '\n')
            fh.write('  %s   (%s)\n' % (name, eid or '?'))
            fh.write('-' * 78 + '\n')
            fh.write('  highlights: %s\n' % hl)
            for c in r['comments']:
                c = re.sub(r'^(?:PXRD|Chemistry) check (?:\[[\w+\-]+\] — (?:Mindat: )?|— )', '', c)
                fh.write(textwrap.fill(c, width=78,
                                       initial_indent='  comment   : ',
                                       subsequent_indent=' ' * 14) + '\n')
        fh.write('\n' + '=' * 78 + '\n\n')
        # Reviewer's OWN marks (tracked changes / non-tool comments): preserved across
        # reruns but never written by the tool — surfaced here so they show in the log.
        rev_recs = [(f, r) for f, r in records if r.get('user_edits')]
        if rev_recs:
            fh.write('REVIEWER EDITS (tracked changes / non-tool comments — preserved, NOT tool-written)\n'
                     + '-' * 78 + '\n')
            for f, r in rev_recs:
                fh.write('\n  %s   (%s)\n' % ((C.entry_name(f) or C.entry_id(f) or f).upper(), C.entry_id(f) or '?'))
                for e in r['user_edits']:
                    fh.write(textwrap.fill(e, width=78, initial_indent='    - ', subsequent_indent=' ' * 6) + '\n')
            fh.write('\n' + '=' * 78 + '\n\n')
        fh.write('CLEAN ENTRIES (no edits — copied unchanged)\n' + '-' * 78 + '\n')
        for f, r in records:
            if r['clean']:
                fh.write('  %-32s (%s)\n' % (C.entry_name(f).upper(), C.entry_id(f) or '?'))
        # .dft cross-check: a co-equal ICDD proxy (NOT ground truth) — these are
        # surfaced for the reviewer to verify against the paper; they are NOT
        # written into any docx.
        dft_recs = [(f, r) for f, r in records if r.get('dft')]
        if dft_recs:
            fh.write('\n' + '=' * 78 + '\n\n')
            fh.write('ICDD .dft CROSS-CHECK (verify against the paper — NOT written to docx)\n'
                     + '-' * 78 + '\n')
            for f, r in dft_recs:
                fh.write('\n  %s   (%s)\n' % ((C.entry_name(f) or C.entry_id(f) or f).upper(), C.entry_id(f) or '?'))
                for c in r['dft']:
                    c = re.sub(r'^verify vs ICDD \.dft — ', '', c)
                    fh.write(textwrap.fill(c, width=78, initial_indent='    - ',
                                           subsequent_indent=' ' * 6) + '\n')
    print('\nlog written -> %s' % log_path)

    if filtered:
        return          # never rewrite/remove mindat_discrepancies.txt for a subset

    # --- separate Mindat cross-check feedback log -------------------------
    # Mindat is a co-equal proxy (its data is transcribed from delayed CNMNC
    # newsletters), so these verify items are NOT written into any docx — a list to
    # check against the paper and follow up (either side may be the one to fix).
    md_recs = [(f, r) for f, r in records if r.get('mindat')]
    md_path = os.path.join(out_dir if not args.inplace else args.folder, 'mindat_discrepancies.txt')
    if not md_recs and os.path.exists(md_path):
        os.remove(md_path)                       # no discrepancies now — drop a stale file
    if md_recs:
        LABEL = {'mindat_fix': 'CELL (docx & .cif disagree with Mindat — verify which is correct)',
                 'mindat_chem': 'CHEMISTRY (ideal formula differs from Mindat)',
                 'ima_status': 'IMA STATUS'}
        with open(md_path, 'w', encoding='utf-8') as fh:
            fh.write('Mindat cross-check — verify against the paper and follow up; either side may be '
                     'the one to fix (Mindat data can lag the CNMNC newsletter). Not written into any docx.\n')
            fh.write('%d entr%s.\n' % (len(md_recs), 'y' if len(md_recs) == 1 else 'ies') + '=' * 78 + '\n')
            for f, r in md_recs:
                fh.write('\n  %s   (%s)\n' % ((C.entry_name(f) or C.entry_id(f) or f).upper(), C.entry_id(f) or '?'))
                for code, msg in r['mindat']:
                    fh.write('    [%s]\n' % LABEL.get(code, code))
                    fh.write(textwrap.fill(msg, width=78, initial_indent='      - ', subsequent_indent=' ' * 8) + '\n')
        print('mindat cross-check -> %s  (%d)' % (md_path, len(md_recs)))

    _next_steps(args, out_dir, len(md_recs))

def _next_steps(args, out_dir, n_mindat):
    """Close a run by naming the commands worth running next. The tool has several entry
    points and a reviewer should not have to re-read the README to remember which one does
    what — nor to discover that the docx it just wrote are meant to be triaged in the GUI."""
    folder = os.path.abspath(args.folder)
    out = os.path.abspath(out_dir)
    print('\n' + '-' * 78)
    print('NEXT')
    print('  Triage the findings side-by-side with the paper (recommended):')
    print('      pxrd gui "%s"' % folder)
    print("  The tool's own change log (what it commented, and why):")
    print('      %s' % os.path.join(out, 'annotation_log.txt'))
    if n_mindat:
        print('  %d entr%s disagree with Mindat. Verify against the paper — neither side is'
              % (n_mindat, 'y' if n_mindat == 1 else 'ies'))
        print('  automatically right (Mindat lags the CNMNC newsletter):')
        print('      %s' % os.path.join(out, 'mindat_discrepancies.txt'))
    print('  Reviewed copies: %s'
          % ('the SOURCE folder (--inplace was used)' if args.inplace else out))
    if not args.inplace:
        print('  The source .docx were not modified.')

if __name__ == '__main__':
    main()
