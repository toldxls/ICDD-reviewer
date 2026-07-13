#!/usr/bin/env python3
"""Regression check — locks in the false-positive fixes and genuine catches that
were established by hand-auditing one private review batch (see README "Behavioral
contract"). Run after ANY change to the checks; every case must still PASS.

    python3 -m pxrd_review.regression_check "/path/to/fixtures"
    PXRD_REGRESSION_DIR="/path/to/fixtures" python3 -m pxrd_review.regression_check

The fixtures are a private ICDD review batch (not shipped with this repo);
point the script at your local copy via the argument or $PXRD_REGRESSION_DIR.

Read-only: it runs analyze() and inspects the verdict — it does not write docx.
Exit code 0 = all pass, 1 = a regression.
"""
import sys, os, glob, re

from pxrd_review import cell_lambda_check as C
from pxrd_review import annotate_review as A
from pxrd_review import extra_checks as X
from pxrd_review import mindat as M
from pxrd_review import candidate_groups as G

def _stub_entry(z, formula):   # 'testite' name so _cif_name_ok passes and the Z logic runs
    return type('S', (), {'cell': {'Z': str(z)}, 'formulas': {'Empirical': formula},
                          'name': 'testite', 'comments': {}})()

def _stub(**kw):               # minimal parse_entry stand-in for single-check unit cases
    base = {'instr': {}, 'name': 'testite', 'primary': '', 'comments': {}, 'cell': {}, 'formulas': {}}
    base.update(kw)
    return type('S', (), base)()

# Fixtures are private ICDD data, resolved at RUN time (in main), NOT at import: importing
# this module must be side-effect-free and must not consume sys.argv — so it can be
# introspected/unit-tested, and so a future importer is not silently handed our argv.
FOLDER = ''
_idx, _cif, _dft = {}, {}, {}
_cache = {}

def _init_fixtures(folder=None):
    """Resolve the private review-batch fixtures (explicit arg > $PXRD_REGRESSION_DIR) and
    build the docx/cif/dft indices. Called by main() before the suite runs; exits with a
    clear message when the fixtures are absent (they are not shipped in this repo)."""
    global FOLDER, _idx, _cif, _dft, _cache
    FOLDER = folder or (sys.argv[1] if len(sys.argv) > 1 else os.environ.get('PXRD_REGRESSION_DIR', ''))
    if not FOLDER:
        sys.exit("Regression fixtures not found. Pass the private review-batch folder as an "
                 "argument or set $PXRD_REGRESSION_DIR — it is private ICDD data and is "
                 "not included in this repository.")
    _idx, _cif, _dft = C.pdf_index(FOLDER), C.cif_index(FOLDER), C.dft_index(FOLDER)
    _cache = {}
    return FOLDER
class FixtureMissing(Exception):
    """A case references an entry the fixture batch doesn't contain."""

def _find_docx(eid):
    """Locate the fixture docx for an entry id: a root-level match is preferred;
    otherwise search recursively, skipping the tool's own outputs (review_out/,
    .edit_backup/) so an _edited copy is never analyzed as the fixture."""
    hits = sorted(glob.glob(os.path.join(FOLDER, eid + '*.docx')))
    if hits:
        return hits[0]
    for root, dirs, files in os.walk(FOLDER):
        dirs[:] = sorted(d for d in dirs if d not in ('review_out', '.edit_backup'))
        for f in sorted(files):
            if f.startswith(eid) and f.endswith('.docx'):
                return os.path.join(root, f)
    return None

def res_for(eid):
    if eid not in _cache:
        dp = _find_docx(eid)
        if not dp:
            # a missing fixture must be LOUD: caching None here would turn every
            # negative assertion ("must NOT flag") for the entry into a vacuous PASS
            raise FixtureMissing('fixture missing: %s' % eid)
        _cache[eid] = A.analyze(dp, _idx.get(eid), _cif.get(eid), _dft.get(eid))
    return _cache[eid]

def _extras_of(eid):
    """Every extra finding for an entry (or None when the fixture isn't present)."""
    r = res_for(eid)
    return (r['extra'] if r else None)

def _ref_entry(ref):
    """A minimal entry whose only row is the Primary Reference — the fixture for check 26."""
    return type('S', (), {'raw_rows': [['References'], ['Primary Reference', ref]]})()

def _ref_cell(cell_xml):
    """A python-docx cell built from raw XML — the fixture for the tracked-change writer, which
    has to cope with whatever run/paragraph layout a real docx throws at it."""
    import tempfile, zipfile as _zf, io as _io
    from docx import Document as _D
    doc = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
           '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
           '<w:body><w:tbl><w:tr>'
           '<w:tc><w:p><w:r><w:t>Primary Reference</w:t></w:r></w:p></w:tc>'
           '<w:tc>' + cell_xml + '</w:tc>'
           '</w:tr></w:tbl></w:body></w:document>')
    fd, path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    _D().save(path)                                   # a valid docx skeleton (all the other parts)
    z = _zf.ZipFile(path)
    parts = {n: z.read(n) for n in z.namelist()}
    z.close()
    parts['word/document.xml'] = doc.encode()
    buf = _io.BytesIO()
    with _zf.ZipFile(buf, 'w', _zf.ZIP_DEFLATED) as zo:
        for n, d in parts.items():
            zo.writestr(n, d)
    with open(path, 'wb') as f:
        f.write(buf.getvalue())
    return _D(path).tables[0].rows[0].cells[1]

def _docx_with_tracked_change(author):
    """A one-cell docx carrying a single tracked insertion by `author` — the fixture for
    'is a tracked change by the TOOL a reviewer edit?' (it must not be: the tool now writes its
    applied fixes as tracked changes, and counting those as human edits would make every rerun
    back the output up and hand the tool's own work back to the reviewer as theirs)."""
    import tempfile, zipfile as _zf, io as _io
    from lxml import etree as _et
    W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    doc = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
           '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
           '<w:body><w:p>'
           '<w:ins w:id="1" w:author="%s" w:date="2026-01-01T00:00:00">'
           '<w:r><w:t>added</w:t></w:r></w:ins>'
           '</w:p></w:body></w:document>' % author)
    fd, path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    with _zf.ZipFile(path, 'w') as z:
        z.writestr('word/document.xml', doc)
    return path

def extras(eid, code=None, sev=None, substr=None):
    r = res_for(eid)
    out = []
    for f in (r['extra'] if r else []):
        if code and f.code != code: continue
        if sev and f.sev != sev: continue
        if substr and substr.lower() not in (f.msg or '').lower(): continue
        out.append(f)
    return out

def cell_verdict(eid):
    r = res_for(eid); return r['cell'][0] if r else None
def lam_flag(eid):
    r = res_for(eid); return bool(r and r['lam'] and r['lam'][0] == 'flag')
def lam_verdict(eid):
    r = res_for(eid); return (r['lam'][0] if (r and r['lam']) else None)
def params(eid):
    r = res_for(eid); return set((r.get('params') or {})) if r else set()
def _anchor_text(eid, anchor):
    """Text of the docx cell a finding with the given anchor would comment on."""
    from docx import Document
    dp = _find_docx(eid)
    if not dp:
        raise FixtureMissing('fixture missing: %s' % eid)
    doc = Document(dp)
    cell = A._anchor_cell(doc, doc.tables[0].rows[0], anchor)   # ac_row unused for ima/analysis
    return cell.text.strip() if cell else None

def _xxe_blocked():
    """Security invariant: the docx XML parsers must not be XXE-able. The stdlib-ET parsers
    reject a DTD/DOCTYPE outright; the lxml parser refuses to resolve an external entity."""
    dt = b'<!DOCTYPE a [<!ENTITY e SYSTEM "file:///etc/hostname">]><a>&e;</a>'
    for mod in (C, X):                       # cell_lambda_check + extra_checks (stdlib ET)
        try:
            mod._xml_fromstring(dt)
            return False                     # should have raised on the DTD
        except ValueError:
            pass
    return ''.join(A._xml(dt).itertext()) == '&e;'   # lxml: entity left unresolved, no file read

def _discover_recursive_ok():
    """discover() finds entry docx at ANY depth under the root, skips the tool's own review_out/
    output, and prefers the '(MineralName)' transcription over a supplementary docx of the same id."""
    import tempfile
    with tempfile.TemporaryDirectory() as td:
        sub = os.path.join(td, 'sub'); os.makedirs(sub)
        ro = os.path.join(td, 'review_out'); os.makedirs(ro)
        open(os.path.join(sub, 'I999001(Test).docx'), 'w').close()        # nested real entry
        open(os.path.join(sub, 'I999001_Supp.docx'), 'w').close()         # supp of same id -> not chosen
        open(os.path.join(ro, 'I999002(Out)_edited.docx'), 'w').close()   # tool output -> skipped
        got = C.discover(td)
        return ('I999001' in got and got['I999001'].endswith('(Test).docx')
                and 'I999002' not in got)

def _docx_with(td, rel, body):
    """Write a minimal valid .docx (zip) at td/rel whose word/document.xml contains `body`.
    Enough for discover()/_review_activity to read it (they only need the zip + document.xml)."""
    import zipfile
    p = os.path.join(td, rel); os.makedirs(os.path.dirname(p), exist_ok=True)
    with zipfile.ZipFile(p, 'w') as z:
        z.writestr('word/document.xml',
                   '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                   '<w:body>%s</w:body></w:document>' % body)
    return p

def _discover_prefers_reviewed_ok():
    """When an id has several parallel '(MineralName)' transcriptions (the multi-reviewer training
    tree), discover() picks the MOST-REVIEWED copy (tracked changes/comments), not the alphabetically
    first RAW copy. Guards the raw-copy tiebreak bug (the 'primary_name 9/9 FP' false alarm)."""
    import tempfile
    with tempfile.TemporaryDirectory() as td:
        # 'A_raw' sorts before 'Z_reviewed' alphabetically, so the old tiebreak picked A_raw.
        _docx_with(td, 'A_raw/I999004(Test).docx', '<w:p><w:r><w:t>clean</w:t></w:r></w:p>')
        _docx_with(td, 'Z_reviewed/I999004(Test).docx',
                   '<w:ins w:id="1" w:author="R"><w:r><w:t>fix</w:t></w:r></w:ins>'
                   '<w:del w:id="2" w:author="R"><w:r><w:delText>old</w:delText></w:r></w:del>')
        got = C.discover(td)
        return 'I999004' in got and 'Z_reviewed' in got['I999004']

def _no_errored_checks():
    """Corpus-wide: no check may ERROR on any fixture entry (an errored check silently
    skips its assertions; run_all surfaces it as a finding whose msg says so). Matched
    on the MESSAGE substring, not the finding code, so it holds however the errored
    finding is filed."""
    bad = []
    for eid, dp in sorted(C.discover(FOLDER).items()):
        try:
            r = _cache.get(eid) or A.analyze(dp, _idx.get(eid), _cif.get(eid), _dft.get(eid))
        except Exception as ex:
            bad.append('%s (analyze crashed: %s)' % (eid, ex)); continue
        for f in r.get('extra', []):
            m = (f.msg or '').lower()
            if 'check errored' in m or 'errored:' in m:
                bad.append('%s [%s] %s' % (eid, f.code, (f.msg or '')[:80]))
    if bad:
        print('      errored findings: ' + '; '.join(bad[:8]) + (' …' if len(bad) > 8 else ''))
    return not bad

def _errored_check_filed_under_real_code():
    """run_all files a CRASHED check under the code its findings actually carry
    (check19_intensity_detector -> 'intensity_type'), not the raw function name —
    so sweep's fire table and code-filtered lookups can see a broken check."""
    def check19_intensity_detector(e, text):
        raise NameError('regression probe')
    saved = X.CHECKS
    try:
        X.CHECKS = [check19_intensity_detector]
        fs = X.run_all(_stub(), '')
    finally:
        X.CHECKS = saved
    return any(f.code == 'intensity_type' and 'check errored:' in (f.msg or '') for f in fs)

def _pdf_index_uppercase_ok():
    """pdf_index pairs publisher-style '.PDF' (uppercase) files like cif/dft_index do."""
    import tempfile
    with tempfile.TemporaryDirectory() as td:
        open(os.path.join(td, 'I999003(Test).PDF'), 'w').close()
        return 'I999003' in C.pdf_index(td)

def _no_two_species_chimera():
    """The inline grab window must truncate at a second 'a =': a two-mineral sentence
    (I003403 ferriphoxite / carboferriphoxite class) must not dress the first species'
    cell in the neighbour's α/γ."""
    t = ('Ferriphoxite is triclinic with a = 10.318(3), b = 7.395(5), c = 8.842(2), '
         'while carboferriphoxite has a = 10.402(4), alpha = 95.950(7), gamma = 100.525(8).')
    ours = [c for c in C.find_cells(t) if c.a == '10.318(3)']
    return bool(ours) and all(c.al is None and c.ga is None for c in ours)

def _vertical_tables_stay_separate():
    """Distant vertical-table label blocks yield one candidate EACH — never a single
    document-wide chimera mixing axes from different tables."""
    t = ('a (Å)\n5.100\nc (Å)\n10.200\n' + 'filler line\n' * 15
         + 'a (Å)\n7.300\nb (Å)\n8.400\nc (Å)\n9.500\n')
    got = {(c.a, c.b, c.c) for c in C.find_cells(t)}
    return (('7.300', '8.400', '9.500') in got
            and not any(c.a == '5.100' and c.b == '8.400' for c in C.find_cells(t)))

# (description, predicate) — predicate returns True when behaviour is correct
CASES = [
 # --- calculated-pattern false positives (measured/comparison, not calculated) ---
 ("I003448 not flagged 'calculated'",        lambda: not extras('I003448', 'calculated', 'flag')),
 ("I003563 not 'calculated'",      lambda: not extras('I003563', 'calculated', 'flag')),
 ("I003815 not 'calculated'",          lambda: not extras('I003815', 'calculated', 'flag')),
 # 'dcalc from a Rietveld unit-cell refinement' / measured Iobs/Imeas table = MEASURED pattern
 ("calculated: Rietveld unit-cell dcalc not flagged", lambda: X.check4_calculated(
     type('S', (), {'instr': {}, 'name': 'testite', 'primary': ''}),
     'Calculated from PXRD Rietveld unit cell refinement with a = 9.00, b = 9.01 from the structure.') == []),
 ("calculated: real structure-simulated pattern still flags", lambda: [f for f in X.check4_calculated(
     type('S', (), {'instr': {}, 'name': 'testite', 'primary': ''}),
     'The powder X-ray diffraction pattern was calculated from the single-crystal structure model.') if f.sev == 'flag'] != []),
 # an explicit "powder data were NOT collected" means the pattern is calculated/theoretical -> a
 # docx marked as a measured Diffractometer is wrong (grokhovskyite I003744, structure from EBSD).
 ("I003744 calculated flag (EBSD; powder not collected, docx wrongly Diffractometer)",
  lambda: bool(extras('I003744', 'calculated', 'flag'))),
 # but the akasakaite group IS correctly marked Calculated (also 'powder … not collected'); the
 # not-collected relaxation must NOT turn their console note into a docx flag.
 ("I003810 no calculated FLAG (already Calculated; not-collected relax stays a note)",
  lambda: not extras('I003810', 'calculated', 'flag')),
 ("calculated: 'not collected' relax still needs a calc-powder sentence", lambda: X.check4_calculated(
     type('S', (), {'instr': {}, 'name': 'testite', 'primary': ''}),
     'Powder diffraction data were not collected owing to the tiny crystal size.') == []),
 # --- radiation ---
 ("O002127 radiation OK (CuK found)",  lambda: not lam_flag('O002127')),
 ("I003815 radiation NOT flagged",     lambda: not lam_flag('I003815')),
 ("I003699 radiation IS flagged (real)",   lambda: lam_flag('I003699')),
 # single anode matching the docx = OK, not 'verify' (powder shares the source via
 # Gandolfi/crystal-rotation on a single-crystal instrument); two anodes stay 'verify'
 ("I003562 lam OK (single MoKα source)",     lambda: lam_verdict('I003562') == 'ok'),
 # a radiation is attached to a collection verb whose SUBJECT sets the mode: 'powder
 # XRD data were collected … using CuKα' is the powder radiation even when a single-
 # crystal instrument/comparison sits nearer the token. docx CuKα matches -> OK.
 ("I003521 lam OK (powder collected with CuKα, matches docx)", lambda: lam_verdict('I003521') == 'ok'),
 ("I003823 lam OK (powder XRD with CuKα despite SC-instrument mention)", lambda: lam_verdict('I003823') == 'ok'),
 # 'Sync' is a valid synchrotron designator (λ is beamline-tunable, matches no tube
 # line) — OK, not 'unrec'/'docx anode not recognised'
 ("I003599 lam OK (Sync synchrotron anode)", lambda: lam_verdict('I003599') == 'ok'),
 ("I003554 lam OK (Sync synchrotron anode)", lambda: lam_verdict('I003554') == 'ok'),
 # a single-crystal radiation can be misclassified into a powder context (tarutinoite: the SC
 # MoKα sits nearer a 'powder diffraction patterns' phrase than 'single crystal data', so
 # _section_mode tags it powder). The powder radiation is ranked by the PAPER's evidence — only
 # CoKα sits in 'Powder XRD … collected … Debye-Scherrer … CoKα' — NOT by the docx anode.
 ("I003636 lam OK (CoKα has the powder-collection binding; MoKα is the SC source)", lambda: lam_verdict('I003636') == 'ok'),
 # the ranking must NOT lean on the docx: a token's powder-collection evidence is what counts
 ("radiation: CoKα powder-collection binding detected", lambda: C.powder_collected_near(
     'Powder X-ray diffraction data were collected using Debye-Scherrer geometry, CoKa', 78)),
 ("radiation: SC mounting sentence has NO powder binding", lambda: not C.powder_collected_near(
     'the crystal was mounted on a glass fibre and examined through a Supernova diffractometer, MoKa', 92)),
 # --- 1. geometry: a 'neutron'/ToF mention in a planned/future-work sentence must NOT
 #     masquerade as the collection method; the real X-ray geometry is still found ---
 ("I003599 no geometry 'neutron' grab (future-work)", lambda: not extras('I003599', 'geometry', substr='neutron')),
 # geometry method is descriptive metadata: when Spacing Instr. is already a valid
 # designator (Diffractometer), don't nag to annotate the specific method (I003510 is
 # Diffractometer and its DC comment already documents the Debye-Scherrer geometry)
 ("I003510 no geometry nag (Spacing Instr. already Diffractometer)", lambda: not extras('I003510', 'geometry')),
 # --- instr_class: a diffractometer named in a POWDER sentence -> 'Other'/blank
 #     Spacing/Intensity Instr. should be 'Diffractometer' (FLAG). Requires the powder
 #     tie-in (a single-crystal-only instrument is not assumed) and skips Calculated. ---
 ("I003414 Spacing Instr. -> Diffractometer flag",  lambda: bool(extras('I003414', 'instr_class', 'flag', 'Spacing Instr.'))),
 ("I003414 Intensity Instr. -> Diffractometer flag", lambda: bool(extras('I003414', 'instr_class', 'flag', 'Intensity Instr.'))),
 ("I003528 instr_class flag (powder on R-AXIS)",    lambda: bool(extras('I003528', 'instr_class', 'flag'))),
 # anaphoric 'the same diffractometer' (a pseudo-Gandolfi crystal-rotation setup on a
 # single-crystal area detector): the model is named only in the SC sentence, but the
 # literal word 'diffractometer' in the powder-COLLECTION sentence still pins Spacing/
 # Intensity Instr. -> Diffractometer, and the crystal-rotation method (= area detector)
 # -> Intensity Type Integrated. The weak 'no named geometry' note must not co-fire.
 ("I003548 Spacing Instr. -> Diffractometer flag (bare 'diffractometer')", lambda: bool(extras('I003548', 'instr_class', 'flag', 'Spacing Instr.'))),
 ("I003548 Intensity Instr. -> Diffractometer flag (bare 'diffractometer')", lambda: bool(extras('I003548', 'instr_class', 'flag', 'Intensity Instr.'))),
 ("I003548 Intensity Type Peak->Integrated (crystal rotation = pseudo-Gandolfi area detector)", lambda: bool(extras('I003548', 'intensity_type', 'flag'))),
 ("I003548 no weak geometry note (superseded by instr_class flag)", lambda: not extras('I003548', 'geometry')),
 # a .pdf line-break hyphenation ('R-\nAxis' -> 'R- Axis', yellowcatite/I003561) must still bind
 # the R-AXIS Rapid as the powder diffractometer (regex tolerates hyphen+whitespace, \b-anchored);
 # docx Other -> Diffractometer flag. \b stops 'polar axis rapid' matching. Unit-level (synthetic).
 ("instr_class: 'R- Axis Rapid' line-break hyphen recognised (Other -> Diffractometer)",
  lambda: bool([f for f in X.check1_geometry(
      type('S', (), {'instr': {'spacing_instr': 'Other', 'intensity_instr': 'Other'}}),
      'Powder X-ray diffraction data were collected on a Rigaku R- Axis Rapid II microdiffractometer.')
      if f.code == 'instr_class'])),
 ("instr_class: '...polar axis rapid...' must NOT match R-AXIS (\\b guard)",
  lambda: not [f for f in X.check1_geometry(
      type('S', (), {'instr': {'spacing_instr': 'Other', 'intensity_instr': 'Other'}}),
      'Powder data were collected with the sample on the polar axis rapidly spinning.')
      if f.code == 'instr_class']),
 ("I003745 no instr_class flag (D8 single-crystal only)", lambda: not extras('I003745', 'instr_class')),
 ("I003807 no instr_class flag (no powder instrument; 'expert' != Empyrean)", lambda: not extras('I003807', 'instr_class')),
 ("I003246 no instr_class flag (already Diffractometer)", lambda: not extras('I003246', 'instr_class')),
 # --- 2. cell provenance: powder-refined cells must NOT be flagged 'from single-crystal' ---
 ("I003562 no cell_provenance (powder-refined)", lambda: not extras('I003562', 'cell_provenance')),
 ("I003599 no cell_provenance (ADP, not cell)",  lambda: not extras('I003599', 'cell_provenance')),
 ("I003750 no cell_provenance (Dcalc boilerplate)", lambda: not extras('I003750', 'cell_provenance')),
 # --- cell SOURCE: docx using the SCXRD cell when a same-phase powder cell exists = FLAG;
 #     an SCXRD-looking cell with no powder cell = soft NOTE; a powder cell = nothing ---
 ("I003632 cell_source FLAG (SC cell + powder cell exists)", lambda: bool(extras('I003632', 'cell_source', 'flag'))),
 # the same-phase powder conflict drives the FLAG even without a definitive SC cue in the
 # matched cell's sentence (spaltiite: docx used the SCXRD cell, PXRD cell 15.821 refined)
 ("I003807 cell_source FLAG (SCXRD used, PXRD cell refined)", lambda: bool(extras('I003807', 'cell_source', 'flag'))),
 # julgoldite-(Fe2+) is a POWDER-ONLY study (no single-crystal anywhere in the .pdf), so its
 # cell — reported in the abstract / a summary table, far from the powder-methods text — is the
 # powder cell, not SCXRD. The document-level guard settles it; no cell_source SCXRD note fires.
 ("I003657 no cell_source (powder-only paper, no single-crystal mention)", lambda: not extras('I003657', 'cell_source')),
 # guard mechanics (unit-level): a paper with NO single-crystal mention classifies as powder;
 # the SC-doc detector is generous (any 'single[-/ws]crystal' / SCXRD) and quiet on powder-only.
 ("guard: powder-only text -> cell context powder", lambda: C.classify_context('x-ray powder diffraction data; a = 9.05 A', 30) == 'powder'),
 ("guard: SC-doc detector matches single-crystal/SCXRD, not powder-only",
  lambda: bool(C._SINGLE_CRYSTAL_DOC.search('single crystal')) and bool(C._SINGLE_CRYSTAL_DOC.search('SCXRD')) and not C._SINGLE_CRYSTAL_DOC.search('powder-only study, a = 9.0')),
 ("I003510 no cell_source flag (powder cell, not SC)", lambda: not extras('I003510', 'cell_source', 'flag')),
 ("I003566 no cell_source flag (GSAS powder cell)",  lambda: not extras('I003566', 'cell_source', 'flag')),
 ("I003636 no cell_source flag (UnitCell powder cell)", lambda: not extras('I003636', 'cell_source', 'flag')),
 ("I003562 no cell_source (powder cell)",            lambda: not extras('I003562', 'cell_source')),
 # the CHECKCELL powder cell sits under a 'Powder XRD data were collected …' subsection whose
 # opener uses the XRD abbreviation; _section_mode must read that as powder (not inherit the
 # earlier single-crystal subsection) so the matched cell isn't mislabelled SCXRD (plumbojohntomaite).
 ("I003563 no cell_source flag (powder cell; 'Powder XRD' subsection)", lambda: not extras('I003563', 'cell_source')),
 ("section: POWDER_SEC matches the 'Powder XRD' abbreviation", lambda: bool(C.POWDER_SEC.search('powder xrd data were collected')) and bool(C.POWDER_SEC.search('x-ray powder diffraction'))),
 # --- instrument/geometry lexicon (classify_context fallback when no bare powder/
 #     single word sits near the cell). Derived from DC fields; mode-determining terms
 #     only. Unit-level (no fixture needed) so the contract is locked regardless of
 #     which private entries happen to exercise it. ---
 ("lexicon: Gandolfi geometry -> powder",   lambda: C._instr_mode('gandolfi-like geometry, recording') == 'powder'),
 ("lexicon: pseudo-Gandolfi motion -> powder", lambda: C._instr_mode('collected with pseudo-gandolfi motion') == 'powder'),
 # 'crystal rotation' is the same pseudo-Gandolfi technique under a different name -> powder
 ("lexicon: crystal-rotation method -> powder (pseudo-Gandolfi synonym)", lambda: C._instr_mode('obtained using the crystal rotation method') == 'powder'),
 ("lexicon: Debye-Scherrer -> powder",      lambda: C._instr_mode('in debye-scherrer geometry') == 'powder'),
 ("lexicon: kappa four-circle -> single",   lambda: C._instr_mode('bruker kappa four-circle goniometer') == 'single'),
 ("lexicon: three-circle goniometer -> single", lambda: C._instr_mode('a three-circle diffractometer') == 'single'),
 # X'Celerator (PANalytical 1D powder detector); the X'-prefix keeps it off 'accelerator'
 ("lexicon: X'Celerator -> powder", lambda: C._instr_mode("equipped with an X'celerator silicon-strip detector") == 'powder'),
 ("lexicon: 'accelerator' is not a powder cue", lambda: C._instr_mode('at the synchrotron accelerator facility') is None),
 # D8 Advance/Discover = dedicated powder Bragg-Brentano; D8 Venture (Photon area detector)
 # is dual-use -> must NOT classify as single from the model name alone
 ("lexicon: D8 Advance=powder, D8 Venture undetermined",
                                            lambda: C._instr_mode('bruker d8 advance') == 'powder' and C._instr_mode('bruker d8 venture diffractometer') is None),
 # dual-use area-detector single-crystal instruments run pseudo-Gandolfi/Gandolfi-like to
 # collect POWDER -> their model names must NOT force 'single'; the geometry/motion word wins
 ("lexicon: pseudo-Gandolfi on XtaLAB/Photon -> powder (dual-use models)",
                                            lambda: C._instr_mode('rigaku xtalab synergy, pseudo-gandolfi motion') == 'powder'
                                                    and C._instr_mode('bruker photon iii detector, gandolfi-like scan') == 'powder'),
 # bare single-crystal MODEL with no geometry word -> undetermined (dual-use), not 'single'
 ("lexicon: bare XtaLAB/Apex model -> undetermined", lambda: C._instr_mode('rigaku xtalab synergy diffractometer') is None and C._instr_mode('bruker apex ii detector') is None),
 # dual-use Rigaku R-AXIS Rapid II -> must NOT classify (guards against re-adding it;
 # it mislabelled a single-crystal cell as powder in I002535)
 ("lexicon: dual-use R-AXIS Rapid -> unknown", lambda: C._instr_mode('rigaku r-axis rapid ii curved imaging plate microdiffractometer') is None),
 # both modes named -> one-sided rule returns None rather than guessing
 ("lexicon: both modes present -> None",    lambda: C._instr_mode('gandolfi motion on a kappa four-circle goniometer') is None),
 # refinement-software / method cues (Rietveld = powder by definition; SHELX/OLEX refine a
 # single-crystal STRUCTURE). Derived+validated against the paired PDFs.
 ("lexicon: Rietveld/GSAS software -> powder", lambda: C._instr_mode('refined by the rietveld method using gsas-ii') == 'powder'),
 ("lexicon: TOPAS/FullProf -> powder",      lambda: C._instr_mode('whole-pattern refinement in topas') == 'powder' and C._instr_mode('fullprof suite') == 'powder'),
 ("lexicon: SHELX/OLEX -> single",          lambda: C._instr_mode('structure refined with shelxl-2018') == 'single' and C._instr_mode('olex2 was used') == 'single'),
 # CrysAlisPro + Bruker APEX software suites = single-crystal workup (the version digit
 # keeps 'apex4' off the dual-use 'Apex II' instrument; 'crysalis' is not in 'chrysalis')
 ("lexicon: CrysAlisPro/APEX4 -> single", lambda: C._instr_mode('data reduced with crysalispro') == 'single' and C._instr_mode('processed in the bruker apex4 suite') == 'single'),
 ("lexicon: APEX software + Gandolfi -> None (pseudo-Gandolfi protected)", lambda: C._instr_mode('apex4 then gandolfi powder extraction') is None),
 ("lexicon: 'chrysalis' is not a cue", lambda: C._instr_mode('emerged from its chrysalis') is None),
 # dual-use software (does BOTH) must NOT classify — both Jana2006 and Jana2020
 ("lexicon: JANA (does both) -> undetermined", lambda: C._instr_mode('refined using jana2006') is None and C._instr_mode('refined using jana2020') is None),
 # canon-substring collision guards: these common phrases/minerals must NOT read as a cue
 ("lexicon: 'unit cell' phrase is not a powder cue", lambda: C._instr_mode('the unit cell parameters were') is None),
 ("lexicon: 'jadeite' mineral is not a powder cue",  lambda: C._instr_mode('jadeite and omphacite inclusions') is None),
 # collision-safe rescues of versioned/suffixed names (snippet the collider can't contain)
 ("lexicon: JADE 2010 -> powder (not jadeite)", lambda: C._instr_mode('processed with MDI Jade 2010') == 'powder' and C._instr_mode('jadeite-bearing eclogite') is None),
 ("lexicon: SIR2011 -> single (not bare 'sir')", lambda: C._instr_mode('solved by direct methods in sir2011') == 'single'),
 # full-corpus (438-pair) re-scan additions: profile/whole-pattern fitting + 1D powder
 # detectors (LynxEye/MYTHEN are NOT dual-use area detectors)
 ("lexicon: whole-pattern/profile fit -> powder", lambda: C._instr_mode('whole-pattern profile fit refinement') == 'powder'),
 ("lexicon: LynxEye 1D detector -> powder",  lambda: C._instr_mode('d8 with a lynxeye detector') == 'powder'),
 # 'xpert' was NOT added: its canon form collides with 'expert' -> must stay undetermined
 ("lexicon: 'expert' is not a powder cue (xpert collision avoided)", lambda: C._instr_mode('in the expert opinion of the authors') is None),
 # --- entry-id pairing: both 5-digit (ICDD Task Group) and 6-digit ids; underscore-
 #     prefixed pdf names; a longer digit run must NOT be bitten into a false id ---
 ("id: 6-digit docx",         lambda: C.entry_id('I003559(Selenolaurite).docx') == 'I003559'),
 ("id: 5-digit docx",         lambda: C.entry_id('I10636(Yarzhemskiite).docx') == 'I10636'),
 ("id: underscore-prefixed 5-digit pdf", lambda: C.entry_id('77539_I11149.pdf') == 'I11149'),
 ("id: 7-digit run is not an id", lambda: C.entry_id('scan_I1234567.tif') is None),
 ("id: lowercase prefix normalised", lambda: C.entry_id('i11397(Grimmite).docx') == 'I11397'),
 ("id: lowercase i in a word is not an id", lambda: C.entry_id('doi10665.pdf') is None),
 # pdf range pairing: a hyphen 'A-B' expands ONLY its adjacent pair, never ids[0]..ids[-1]
 # (filenames carry unrelated comma/underscore ids) — guards the 550-id overshoot bug
 ("pdf range: A-B plus extra id does not overshoot",
   lambda: set(C._id_keys('2226_I001146-I001147_I001695.pdf')) == {'I001146','I001147','I001695'}),
 ("pdf range: comma id plus range",
   lambda: set(C._id_keys('76264_I10126,I10499-I10500.pdf')) == {'I10126','I10499','I10500'}),
 ("pdf range: a true consecutive range still expands",
   lambda: set(C._id_keys('-60814_I000738-I000748.pdf')) == {'I%06d' % n for n in range(738, 749)}),
 # --- cif Z-check: reconcile across formula-unit conventions (docx vs CIF formula differ by
 #     a multiple, so Z differs while the cell CONTENTS match) — was an ~89% FP storm ---
 ("formula atoms keep fractional occupancy", lambda: X._formula_atoms('Pb1.50 O4.50') == {'Pb': 1.5, 'O': 4.5}),
 ("cif Z: MoO3·H2O Z=8 reconciles with Mo2O8 Z=4 (no flag)",
   lambda: not X.check_cif(_stub_entry(8, 'H2 Mo O4'), {'Z': 4, 'formula': 'Mo2 O8', 'mineral_name': 'testite'})),
 ("cif Z: fractional-occupancy CIF reconciles (Pb1.5 anchor, no flag)",
   lambda: not X.check_cif(_stub_entry(24, 'O3 Pb Te'), {'Z': 16, 'formula': 'O4.50 Pb1.50 Te1.50', 'mineral_name': 'testite'})),
 # a CIF lacking the mineral's dominant cation is a mis-filed/garbage CIF -> NOTE, not a
 # misleading 'Z mismatch' flag (e.g. I002904's Cd/Se CIF vs the As/Sc mineral)
 ("cif Z: chemistry-mismatched CIF -> note, not flag",
   lambda: [f.sev for f in X.check_cif(_stub_entry(4, 'As H4 O6 Sc'), {'Z': 8, 'formula': 'Cd H4 O6 Se', 'mineral_name': 'testite'})] == ['note']),
 # _cif_name_ok must normalize accents, else a legitimately-paired CIF is silently skipped
 ("cif name: accents normalized (Ríosecoite ~ riosecoite)", lambda: X._cif_name_ok(
     {'mineral_name': 'riosecoite'}, type('S', (), {'name': 'Ríosecoite', 'primary': ''})) is True),
 ("cif name: genuinely different mineral still rejected", lambda: X._cif_name_ok(
     {'mineral_name': 'zincostottite'}, type('S', (), {'name': 'Anorthoyttrialite-(Y)', 'primary': ''})) is False),
 # --- indexing: overlapped reflections are packed per column (h='10',k='01',l='44' =
 #     (1,0,4)+(0,1,4)); split so a matching sub-line isn't a false 'disagrees' flag ---
 ("indexing: overlapped hkl splits per column",
   lambda: {(1, 0, 4), (0, 1, 4)} <= set(X._candidate_hkls('10', '01', '44'))),
 ("indexing: SIGNED overlapped hkl splits ('-41'/'03'/'10' = (-4,0,1)+(1,3,0))",
   lambda: {(-4, 0, 1), (1, 3, 0)} <= set(X._candidate_hkls('-41', '03', '10'))),
 # --- temperature: a non-ambient docx temp the paper doesn't state = likely transcription
 #     slip (PXRD is normally room T); a paper-confirmed non-ambient value is real ---
 ("temperature: non-ambient docx vs room-T pdf -> slip flag",
   lambda: bool(X.check13_temperature(type('S', (), {'comments': {'Temperature': '273 K', 'DC': ''}})(), 'Temperature (K) 293 collected reflections'))),
 ("temperature: paper-confirmed non-ambient -> no flag",
   lambda: not X.check13_temperature(type('S', (), {'comments': {'Temperature': '100 K', 'DC': ''}})(), 'data collected at 100 K')),
 ("temperature: ambient docx -> no flag",
   lambda: not X.check13_temperature(type('S', (), {'comments': {'Temperature': '293 K', 'DC': ''}})(), 'x')),
 ("indexing: genuine multi-digit index not split",
   lambda: X._candidate_hkls('1', '0', '14') == [(1, 0, 14)]),
 ("indexing: signed index falls through to literal",
   lambda: X._candidate_hkls('-2', '2', '4') == [(-2, 2, 4)]),
 # --- _norm_pdf font/glyph fixes (validated on the second reference set) ---
 ("norm: þ -> + (charge mojibake)",   lambda: C._norm_pdf('Fe3þ and [4þ1]') == 'Fe3+ and [4+1]'),
 ("norm: spaced angstrom A ˚ -> Å",   lambda: 'Å' in C._norm_pdf('a = 5.93 A˚')),
 # Osc2tab/Osc2xrd (Britvin) generate a POWDER pattern from single-crystal frames
 ("lexicon: Osc2tab/Osc2xrd -> powder", lambda: C._instr_mode('pattern produced with osc2tab') == 'powder' and C._instr_mode('using osc2xrd software') == 'powder'),
 ("lexicon: STOE WinXPOW -> powder", lambda: C._instr_mode('refined with stoe winxpow software') == 'powder'),
 # find_radiation: a parenthesised source with kV/mA tube settings is the X-ray SOURCE,
 # not a microprobe standard (recovers I002189's 'rotating anode (CoKα, 40 kV, 15 mA)')
 ("radiation: parenthesised (CoKα, 40 kV) is found",
   lambda: any(a == 'co' for a, *_ in C.find_radiation('rotating anode (CoKα, 40 kV, 15 mA), imaging plate'))),
 ("radiation: microprobe standard (FeKα) still skipped",
   lambda: not any(a == 'fe' for a, *_ in C.find_radiation('analysed against albite, diopside (FeKα) and apatite standards'))),
 ("norm: thin/nbsp spaces -> normal", lambda: C._norm_pdf('Cu\u2009K\u00a0radiation') == 'Cu K radiation'),
 # a 'calculated [X-ray] powder' pattern is simulated from the single-crystal structure,
 # so it must NOT make a single-crystal cell read 'powder' (I003398: synchrotron crystallite
 # refinement whose only nearby 'powder' is 'Calculated X-ray powder diffraction data')
 ("calc-powder phrase is NOT powder context",
   lambda: (lambda s: C.classify_context(s, s.find('a =')))('refinement gave a = 3.09(1). calculated X-ray powder diffraction data are listed in Table S1.') != 'powder'),
 ("real 'refined from powder data' IS powder context",
   lambda: (lambda s: C.classify_context(s, s.find('a =')))('the unit-cell parameters refined from powder data are a = 5.24(1) and so on.') == 'powder'),
 # preceding powder-software cue (JADE Pro) outranks a bare 'single-crystal' that opens the
 # NEXT sentence — the I002373 case (figure caption pushed the bare 'powder' out of range)
 ("classify_context: before powder-software cue beats after 'single-crystal'",
   lambda: (lambda s: C.classify_context(s, s.find('a =')))('observed d values by profile fitting using jade pro software. refined unit-cell parameters are a = 5.04 b = 6.0. single-crystal diffraction data were collected at the synchrotron') == 'powder'),
 # a Le Bail / Rietveld / Pawley refined cell is powder even when 'single-crystal' is the
 # NEARER word (it's the starting-point/comparison reference) — the I002960 case
 ("classify_context: Le Bail refinement beats a nearer 'single-crystal' starting ref",
   lambda: (lambda s: C.classify_context(s, s.find('a =')))('powder x-ray diffraction data were collected. unit-cell parameters refined using the le bail profile-fitting method starting from parameters determined from single-crystal techniques are a = 9.20 and b = 12.4') == 'powder'),
 # a 'calculated PXRD' pattern (e.g. exported by a visualizer like Mercury/PLATON) is NOT a
 # powder experiment — same as 'calculated powder' (I003398), now also covering the 'pxrd' word
 ("calc-PXRD (Mercury export) is NOT powder context",
   lambda: (lambda s: C.classify_context(s, s.find('a =')))('refined structure a = 7.21(2). A calculated PXRD pattern was generated in Mercury.') != 'powder'),
 # 'simulated' patterns (CrystalDiffract/Mercury) are calculated too, not measured powder
 ("simulated PXRD (CrystalDiffract) is NOT powder context",
   lambda: (lambda s: C.classify_context(s, s.find('a =')))('single-crystal structure a = 7.21(2). A simulated powder pattern from CrystalDiffract.') != 'powder'),
 # but 'calcined powder' is a REAL powder sample — must NOT be swallowed by the calc guard
 ("'calcined powder' IS powder context",
   lambda: (lambda s: C.classify_context(s, s.find('a =')))('cell refined from calcined powder data: a = 5.24(1) and so on.') == 'powder'),
 # pure visualizers are NOT mode cues (VESTA/Diamond/CrystalMaker/Mercury)
 ("lexicon: visualizers are not mode cues", lambda: C._instr_mode('structure drawn in vesta and mercury') is None),
 # --- SG <-> crystal-system consistency (docx-internal, blind-spot check #1) ---
 ("sg_system: classifier maps HM symbols", lambda: all(X._sg_system(s) == exp for s, exp in
     [('Pbca','o'), ('P21/c','m'), ('P121/c1','m'), ('Fd-3m','c'), ('P213','c'), ('F432','c'),
      ('P63/mmc','h'), ('P-3m1','h'), ('R-3m','r'), ('I4/mmm','t'), ('P-1','a'), ('P212121','o')])),
 # A 4-fold SCREW and a cubic body-diagonal triad are indistinguishable once the spaces are gone:
 # 'P4132' is 4₁·3·2 (cubic), 'P43212' is 4₃·2₁·2 (TETRAGONAL). Testing '3' before the 4-fold made
 # every P4₃… "cubic" — a false disagreement WRITTEN into the docx, and it contradicted itself
 # (P4₁2₁2 tetragonal, its enantiomorph P4₃2₁2 cubic). The docx keeps the spaced form, which still
 # carries the grouping; where it doesn't, abstain rather than guess.
 ("sg_system: a 4-fold screw is not cubic (spaced form disambiguates)", lambda: all(
     X._sg_system(s) == exp for s, exp in
     [('P 43 21 2', 't'), ('P 41 3 2', 'c'), ('P 42/n n m', 't')])),
 # The 36 cubic groups are a CLOSED SET, so membership is exact — no heuristic can mistake a 4₃
 # screw subscript for a body-diagonal triad. This is the pair that broke it: P4₃2₁2 (tetragonal,
 # #96) and P4₁3₂ (cubic, #213) look identical once the spaces are gone.
 ("sg_system: a 4-screw is tetragonal, a body-diagonal 3 is cubic (compact form)", lambda: all(
     X._sg_system(s) == exp for s, exp in
     [('P43212', 't'), ('P4322', 't'), ('P43', 't'), ('P41', 't'),
      ('P4132', 'c'), ('P4332', 'c'), ('F432', 'c'), ('P432', 'c'),
      ('I-43d', 'c'), ('Ia-3d', 'c'), ('Fd-3mZ', 'c')])),
 ("sg_system: mismatch flags (monoclinic system, ortho SG)", lambda: [f for f in X.check23_sg_system(
     type('S', (), {'crystal_system': 'm', 'space_group': 'Pbca', 'cell': {}})) if f.sev=='flag'] != []),
 ("sg_system: agreement no flag (ortho system, ortho SG)", lambda: X.check23_sg_system(
     type('S', (), {'crystal_system': 'o', 'space_group': 'Pbca', 'cell': {}})) == []),
 ("sg_system: rhombohedral system + hex-setting SG compatible", lambda: X.check23_sg_system(
     type('S', (), {'crystal_system': 'h', 'space_group': 'R3m', 'cell': {}})) == []),
 ("sg_system: corpus clean (no SG/system disagreement)", lambda: not extras('I003510', 'sg_system')
     and not extras('I003807', 'sg_system') and not extras('I003632', 'sg_system')),
 # A BLANK Crystal System must abstain, not flag. The guard used to be `cs not in 'amothrc'`
 # on a possibly-empty string — and '' is a substring of everything, so it never returned and
 # wrote a bogus "Crystal System () disagrees with…" comment INTO the docx.
 ("sg_system: blank Crystal System abstains (does not flag)", lambda: X.check23_sg_system(
     type('S', (), {'crystal_system': '', 'space_group': 'P21/c', 'cell': {}})) == []),
 ("sg_system: unknown Crystal System word abstains", lambda: X.check23_sg_system(
     type('S', (), {'crystal_system': 'Wurtzitic', 'space_group': 'P21/c', 'cell': {}})) == []),
 # 'Triclinic'/'Trigonal'/'Tetragonal' all begin with 't': a first-letter classifier read the
 # first two as TETRAGONAL, which both false-flagged here and imposed a=b, α=β=γ=90 in
 # _constraints (a false-positive storm out of check8). Whole words only.
 ("sg_system: IUCr 'Triclinic' is not read as tetragonal", lambda: X.check23_sg_system(
     type('S', (), {'crystal_system': 'Triclinic', 'space_group': 'P-1', 'cell': {}})) == []),
 ("sg_system: IUCr 'Trigonal' is not read as tetragonal", lambda: X.check23_sg_system(
     type('S', (), {'crystal_system': 'Trigonal', 'space_group': 'R-3m', 'cell': {}})) == []),
 ("sg_system: ICDD 'Anorthic' == triclinic", lambda: X.check23_sg_system(
     type('S', (), {'crystal_system': 'Anorthic', 'space_group': 'P-1', 'cell': {}})) == []),
 ("_constraints: triclinic imposes NO equal-axis / fixed-angle constraint", lambda:
     X._constraints(type('S', (), {'crystal_system': 'Triclinic', 'space_group': 'P-1',
                                   'cell': {}})) == ([], [], {})),
 # a near-cubic R cell (α=β=γ≈90, a=b=c) is rhombohedral AXES, not the hexagonal setting —
 # the old |γ-90|>1 test threw it into the hex branch and demanded γ=120.
 ("_constraints: near-90 rhombohedral cell is not forced to γ=120", lambda:
     X._constraints(type('S', (), {'crystal_system': 'r', 'space_group': 'R-3',
         'cell': {'a': '9.1', 'b': '9.1', 'c': '9.1', 'α': '90.5', 'β': '90.5', 'γ': '90.5'}}))
     == ([('a', 'b', 'c')], [('α', 'β', 'γ')], {})),
 # --- anode detection: the SYMBOL as a token, never a bare substring ---
 # 'CoKα (Fe filter)' read as an IRON anode (the filter metal); 'Copper' contains 'co'.
 ("anode_key: filter metal does not win over the K-line anode", lambda:
     C.anode_key('CoKa (Fe filter)') == 'co' and C.anode_key('Fe-filtered CoKα') == 'co'),
 ("anode_key: 'Copper' is copper, not cobalt", lambda:
     C.anode_key('Copper Kα') == 'cu' and C.anode_key('Cobalt') == 'co'),
 ("anode_key: the ordinary forms still resolve", lambda: all(C.anode_key(s) == k for s, k in
     [('CuKa','cu'), ('Cu Ka1','cu'), ('CuKα','cu'), ('Cu-Kα','cu'), ('MoKa','mo'),
      ('FeKα','fe'), ('CrKα','cr'), ('AgKα','ag'), ('Cu Kα (Ni filter)','cu')])),
 ("anode_key: a non-anode string stays None", lambda:
     C.anode_key('Sync') is None and C.anode_key('graphite monochromator') is None),
 # --- rounding at the common precision: half-UP on the decimal, not half-even on the float ---
 # pdf 2.675 vs a correctly-rounded docx 2.68: round(2.675, 2) == 2.67, which reported a
 # correct transcription as a VALUE MISMATCH instead of the intended sig-figs note.
 ("axis_issues: a correctly-rounded .5 boundary is precision, not a value mismatch", lambda:
     [k for k, _ in C.axis_issues('2.68', '2.675')] == ['precision']),
 ("axis_issues: a genuine value difference still flags", lambda:
     [k for k, _ in C.axis_issues('10.13', '10.123')] == ['value']),
 # --- reference title case (check26) -----------------------------------------------------
 # Mined from 165 reviewer-corrected reference cells: 53 were case-only title fixes and ALL 53
 # went Title Case -> sentence case. The suggestion must NEVER alter letters, and must never
 # lowercase chemistry, a Levinson suffix, a Roman numeral, or an unverifiable proper noun.
 ("ref_title: Title Case is rewritten to sentence case", lambda:
     X._title_sentence_case('Amurselite, A New Uranyl Selenite From the Burro Mine, Colorado')[1]
     == 'Amurselite, a new uranyl selenite from the Burro mine, Colorado'),
 ("ref_title: 'mine' is lowercase, the place name is not", lambda:
     'Burro mine' in X._title_sentence_case('A New Mineral From the Burro Mine, Colorado')[1]),
 # A formal place name keeps its head-noun ('Elba Island', 'Ingul Gold Placer') even though the
 # paper writes the bare word lowercase elsewhere — the papers disagree with each other here, so
 # the reviewer's rule wins. 'mine' is the stated exception. A species is normally lowercase but
 # is part of the name when it sits inside one ('Rose Quartz Quarry' vs 'isotypic with jamesite').
 ("ref_title: a place name keeps its head-noun ('Elba Island')", lambda:
     'Elba Island' in X._title_sentence_case(
         'Elbaite, the Neotype Material From the Rosina Pegmatite, Elba Island, Italy',
         'Tourmaline from Elba Island. The Elba Island locality is famous; Elba Island '
         'material is the neotype. The Rosina Pegmatite is on Elba Island.')[1]),
 ("ref_title: 'Ingul Gold Placer' keeps all three words", lambda:
     'Ingul Gold Placer' in X._title_sentence_case(
         'Selenolaurite, a New Mineral From the Ingul Gold Placer, South Urals',
         'Found in the Ingul Gold Placer. The Ingul Gold Placer lies in the South Urals; '
         'the Ingul Gold Placer is a new mineral locality. Gold placer deposits occur.')[1]),
 ("ref_title: a species inside a place name is kept ('Rose Quartz Quarry')", lambda:
     'Rose Quartz Quarry' in X._title_sentence_case(
         'Occurrences of Bazzite in Canada – the Quadeville Rose Quartz Quarry, Ontario',
         'the specimens were collected at the Quadeville Rose Quartz Quarry in eastern\n'
         'ontario, where beryl is abundant. material from the Quadeville Rose Quartz Quarry\n'
         'has been described before, and the Quadeville Rose Quartz Quarry remains the only\n'
         'canadian locality. the quarry is worked for quartz, and bazzite is rare there.\n'
         'bazzite from this quarry is associated with quartz and albite.')[1]),
 ("ref_title: a species with an ordinary neighbour still lowercases", lambda: (lambda s:
     'isotypic with jamesite' in s and 'dongchuanite group' in s)(X._title_sentence_case(
         'Desorite, a New Mineral Isotypic with Jamesite of the Dongchuanite Group',
         'Desorite is isotypic with jamesite. The dongchuanite group includes jamesite; '
         'this new mineral is isotypic with jamesite of the dongchuanite group.')[1])),
 # an UNATTESTED word is left alone, but must not go on to vouch for its neighbours
 ("ref_title: an unattested word does not anchor a name ('a new layered …')", lambda:
     'a new layered' in X._title_sentence_case(
         'Kodamaite, A New Layered Alkali Fluorosilicate',
         'we describe a new layered fluorosilicate from the massif. the new mineral is a\n'
         'layered alkali fluorosilicate, and its layered structure is unusual. a new\n'
         'alkali fluorosilicate of this kind has not been reported.')[1]),
 ("ref_title: Levinson suffix survives (-(Ce) never -(ce))", lambda:
     '-(Ce)' in X._title_sentence_case('Vielleaureite-(Ce), a New Epidote of the Group')[1]),
 ("ref_title: Levinson suffix survives a trailing comma", lambda:
     '-(Y),' in X._title_sentence_case('Two Minerals: Nacareniobsite-(Y), a New Mineral Here')[1]),
 ("ref_title: an element-prefixed compound is untouched (Al-bearing)", lambda:
     'Al-bearing' in X._title_sentence_case('Petermegawite, a new Al-bearing Selenite Mineral')[1]),
 ("ref_title: a Roman numeral is not sentence-cased (IV. not Iv.)", lambda:
     'IV.' in X._title_sentence_case('New Minerals From the Redmond Mine: IV. Haywoodite Here')[1]),
 ("ref_title: a site variable is not the article ('(A = K, Rb, Cs)')", lambda:
     '(A = K, Rb, Cs)' in X._title_sentence_case(
         'Crystal Structures and Powder Data for AAlGe2O6 Synthetic Analogs (A = K, Rb, Cs)')[1]),
 ("ref_title: chemistry is never re-cased", lambda:
     'Pb2(Fe3+6Zn)O2(PO4)4(OH)8' in X._title_sentence_case(
         'Desorite, Pb2(Fe3+6Zn)O2(PO4)4(OH)8, a New Phosphate Mineral Isotypic with Jamesite.')[1]),
 # the paper is the oracle: a word IT writes lowercase mid-sentence is an ordinary word; one it
 # capitalizes is a name; one it writes in caps is an acronym to restore ('Usa' -> 'USA').
 # The paper decides the ORDINARY words ('new', 'mineral', 'fumarole', 'deposits'). The head-noun
 # of a place name is kept regardless — 'Tolbachik Volcano' is the volcano's name, even though the
 # paper writes the bare word 'volcano' lowercase elsewhere.
 ("ref_title: the paper's own usage decides the ordinary words", lambda:
     X._title_sentence_case(
         'Natromolybdite, a New Mineral From Fumarole Deposits of the Tolbachik Volcano',
         'samples came from the Tolbachik volcano. a new mineral was found at the Tolbachik\n'
         'volcano in fumarole deposits. the new mineral is a fumarole product, and the\n'
         'volcano hosts many fumarole deposits of this kind.')[1]
     == 'Natromolybdite, a new mineral from fumarole deposits of the Tolbachik Volcano'),
 ("ref_title: an acronym mangled by the title-caser is restored (Usa -> USA)", lambda:
     'USA' in X._title_sentence_case(
         'Amurselite, A New Selenite From the Burro Mine, Colorado, Usa.',
         'Collected in the USA. The USA locality is well known. Specimens from the USA were '
         'studied. This USA occurrence is new.')[1]),
 # the one harmful mistake this check could make: lowercasing a name it cannot verify
 ("ref_title: an unverifiable capitalized word is LEFT ALONE, never lowercased", lambda:
     'Koštálov' in X._title_sentence_case(
         'Julgoldite-(Fe2+) and Fe-rich Prehnite From Pectolite Veins at Koštálov Quarry')[1]),
 ("ref_title: never alters letters — only case", lambda: all(
     re.sub(r'\s+', ' ', X._title_sentence_case(t)[1]).lower() == re.sub(r'\s+', ' ', t).lower()
     for t in ['Amurselite, A New Uranyl Selenite From the Burro Mine, San Miguel County, Usa.',
               'Vielleaureite-(Ce), a New Epidote of the Dollaseite Group, France',
               'New Minerals From the Redmond Mine, North Carolina, USA: IV. Haywoodite Here',
               'Desorite, Pb2(Fe3+6Zn)O2(PO4)4(OH)8, a New Phosphate Mineral with Jamesite.'])),
 # conservative: one arguable word is not enough to fire, and a correct title stays silent
 ("ref_title: a sentence-case title does NOT flag", lambda: not X.check26_reference_title_case(
     type('S', (), {'raw_rows': [['References'], ['Primary Reference',
        'Okruginite, Cu2SnSe3, a new mineral from the Ozernovskoe deposit, Kamchatka, Russia. '
        'Vymazalova, A., Kozlov, V. V. Mineral. Mag.. 2024. 88(1) 31 - 39.']]})(), None)),
 ("ref_title: a machine-Title-Cased title DOES flag, anchored on the reference cell", lambda: [
     (f.sev, f.anchor) for f in X.check26_reference_title_case(
         type('S', (), {'raw_rows': [['References'], ['Primary Reference',
            'Amurselite, A New Uranyl Selenite From the Burro Mine, San Miguel County, Colorado. '
            'Kampf, A. R., Olds, T. A. Can. J. Mineral. Petrol.. 2025. 63(3) 305-315.']]})(), None)]
     == [('flag', 'reference')]),
 ("ref_title: no Primary Reference row -> abstains", lambda: X.check26_reference_title_case(
     type('S', (), {'raw_rows': [['Comments'], ['Desc.', 'text']]})(), None) == []),
 # --- what the check is allowed to WRITE (it is the only one that edits the docx) --------------
 # The PAPER is the oracle for what is a name. With no .pdf paired (25/405 corpus entries, plus any
 # scanned paper with no text layer) every decision falls back to a word list, which cannot know
 # 'Mexico' from 'Mineral' — and duly wrote 'new Mexico' into a reviewer's citation. Suggest, never
 # write.
 ("ref_title: with NO .pdf the check suggests but NEVER writes", lambda: (lambda fs:
     bool(fs) and fs[0].fix is None)(X.check26_reference_title_case(_ref_entry(
        'Raydemarkite, the Natural Analogue of Synthetic MoO3, from Cookes Peak, New Mexico. '
        'Yang, H., Gu, X. Can. J. Mineral. 2023. 61(1) 203.'), None))),
 ("ref_title: with the .pdf it does write", lambda: (lambda fs:
     bool(fs) and bool(fs[0].fix))(X.check26_reference_title_case(_ref_entry(
        'Raydemarkite, the Natural Analogue of Synthetic MoO3, from Cookes Peak, New Mexico. '
        'Yang, H., Gu, X. Can. J. Mineral. 2023. 61(1) 203.'),
        'the specimens come from new mexico. a new analogue of synthetic MoO3 occurs in New\n'
        'Mexico. material from New Mexico was studied; the New Mexico locality is natural and\n'
        'synthetic analogues are known. cookes peak lies in New Mexico.'))),
 # a citation with no author pattern: we cannot tell where the title ends, and treating the WHOLE
 # string as the title case-rewrote the journal name and series. Abstain.
 ("ref_title: a citation with no authors abstains (never re-cases the journal)", lambda:
     X.check26_reference_title_case(_ref_entry('Physics And Chemistry Of Minerals 45, 1-12 (2018)'),
                                    'physics and chemistry of minerals is a journal. minerals are '
                                    'studied. chemistry and physics of minerals.') == []),
 # a paper that writes a name inconsistently (OCR, typos) still means it is a NAME: the 70% bar
 # demoted 'Utah' (capitalised in only half its sightings) to an ordinary word and lowercased a US
 # state into the citation.
 ("ref_title: a name the paper capitalises only half the time is still a name", lambda:
     X._paper_word_forms('found in Utah. the Utah deposit is large. Utah hosts it.\n'
                         'in utah the ore is rich. utah samples were taken. utah is arid.'
                         ).get('utah') == 'Utah'),
 # 'Ca-(OH)' / 'Fe-(III)' are chemistry, not Levinson-suffixed species
 ("ref_title: an element with a parenthesised group is not a species (Ca-(OH))", lambda: (lambda s:
     'Ca-(OH)' in s and 'Fe-(III)' in s)(
     X._title_sentence_case('A New Mineral with Ca-(OH) and Fe-(III) Groups')[1])),
 # --- the APPLIED fix (the one check that writes into the docx, as a tracked change) ----------
 # The fix must be the whole citation with ONLY the title's case changed: the authors, journal,
 # year and pages stay byte-identical, and no letter anywhere may change. Anything else means
 # the tool would be rewriting text it cannot account for, and it must fall back to a comment.
 ("ref_title fix: rewrites the title, leaves the rest of the citation byte-identical", lambda: (
     lambda f: f.fix == ('Desorite, a new phosphate mineral isotypic with jamesite. '
                         'Kampf, A. R., Olds, T. A. Am. Mineral.. 2025. 110(7) 1105-1111.'))(
     X.check26_reference_title_case(_ref_entry(
        'Desorite, a New Phosphate Mineral Isotypic with Jamesite. '
        'Kampf, A. R., Olds, T. A. Am. Mineral.. 2025. 110(7) 1105-1111.'), """desorite is a new phosphate mineral. the new mineral is isotypic with jamesite.\na new phosphate mineral of this kind is rare; jamesite and desorite are isotypic.""")[0])),
 ("ref_title fix: differs from the docx in CASE only — never a letter", lambda: (
     lambda f, src: f.fix.lower() == src.lower())(
     X.check26_reference_title_case(_ref_entry(
        'Amurselite, A New Uranyl Selenite From the Burro Mine, Colorado. '
        'Kampf, A. R., Olds, T. A. Can. J. Mineral. Petrol.. 2025. 63(3) 305-315.'),
        "amurselite is a new uranyl selenite from the Burro mine in Colorado. the new mineral\n"
        "occurs at the Burro mine; a new selenite from Colorado. uranyl selenite minerals are rare\n"
        "and the Burro mine in Colorado is their type locality.")[0],
     'Amurselite, A New Uranyl Selenite From the Burro Mine, Colorado. '
     'Kampf, A. R., Olds, T. A. Can. J. Mineral. Petrol.. 2025. 63(3) 305-315.')),
 # every OTHER check stays comment-only: a fix is opt-in per finding, not a new default
 ("only the reference-title check carries a fix (everything else is comment-only)", lambda: all(
     f.fix is None for e in ('I003414', 'I003416', 'I003448', 'I003510')
     for f in (_extras_of(e) or []) if f.code != 'ref_title_case')),
 # the tool's own tracked change is NOT a human edit — otherwise every rerun backs the output up
 # and reports the tool's own work back to the reviewer as theirs
 ("a tool-authored tracked change does not count as a reviewer edit", lambda:
     not A._has_tracked_changes(_docx_with_tracked_change(A.AUTHOR))),
 ("a human tracked change DOES count as a reviewer edit", lambda:
     A._has_tracked_changes(_docx_with_tracked_change('Tony Kampf'))),
 # The writer sweeps a cell's runs into ONE <w:del>. Across several paragraphs that would restore
 # the words but not the paragraph BREAK when the reviewer hits Reject — so it declines instead
 # and stays a comment. (No reference cell in the corpus is laid out that way; refusing is free.)
 ("applied fix: a single-paragraph cell is rewritten", lambda:
     A._apply_tracked_fix(_ref_cell('<w:p><w:r><w:t>Old Title. Author, A.</w:t></w:r></w:p>'),
                          'Old title. Author, A.') is True),
 # the rewrite is CASE-ONLY by construction: anything that would change a letter is refused
 ("applied fix: a rewrite that changes a LETTER is refused", lambda:
     A._apply_tracked_fix(_ref_cell('<w:p><w:r><w:t>Old Title. Author, A.</w:t></w:r></w:p>'),
                          'New title. Author, A.') is False),
 ("applied fix: a MULTI-paragraph cell is declined (never collapses a line break)", lambda:
     A._apply_tracked_fix(_ref_cell('<w:p><w:r><w:t>Old Title.</w:t></w:r></w:p>'
                                    '<w:p><w:r><w:t>Author, A.</w:t></w:r></w:p>'),
                          'Old title. Author, A.') is False),
 ("applied fix: an empty cell is declined", lambda:
     A._apply_tracked_fix(_ref_cell('<w:p/>'), 'New title.') is False),
 # --- the write path's three reproduced criticals (all found in the final review) --------------
 # python-docx's cell.text does NOT see runs nested in w:ins/w:del, so a citation the reviewer
 # replaced with track-changes ON reads as an EMPTY cell -> _find_value falls through to the row's
 # LABEL cell. Without a guard the tool struck out "Primary Reference" and pasted the citation in.
 ("applied fix: a cell whose text is not the fix (bar case) is NEVER written", lambda:
     A._apply_tracked_fix(_ref_cell('<w:p><w:r><w:t>Primary Reference</w:t></w:r></w:p>'),
                          'Desorite, a new phosphate mineral. Kampf, A. R.') is False),
 ("_cell_text sees text nested in a tracked change (cell.text does not)", lambda:
     A._cell_text(_ref_cell('<w:p><w:ins w:id="9" w:author="Tony" w:date="2026-01-01T00:00:00">'
                            '<w:r><w:t>hidden</w:t></w:r></w:ins></w:p>')) == 'hidden'),
 # the rewrite is case-only, so it is sliced back onto the EXISTING runs — italic title, bold year
 # and plain authors all survive. Collapsing them into one run turned the whole citation italic.
 ("applied fix: per-run formatting survives (italic title, bold year)", lambda: (lambda c: (
     A._apply_tracked_fix(c, 'Desorite, a new phosphate mineral. Kampf. 2025.') and
     [(('i' if r.find(A._q('rPr')) is not None and r.find(A._q('rPr')).find(A._q('i')) is not None
        else 'b' if r.find(A._q('rPr')) is not None and r.find(A._q('rPr')).find(A._q('b')) is not None
        else 'p'),
       ''.join(t.text or '' for t in r.findall(A._q('t'))))
      for ins in c._tc.iter(A._q('ins')) for r in ins.findall(A._q('r'))]
     == [('i', 'Desorite, a new phosphate mineral. '), ('p', 'Kampf. '), ('b', '2025.')]))(
     _ref_cell('<w:p>'
               '<w:r><w:rPr><w:i/></w:rPr><w:t xml:space="preserve">Desorite, a New Phosphate Mineral. </w:t></w:r>'
               '<w:r><w:t xml:space="preserve">Kampf. </w:t></w:r>'
               '<w:r><w:rPr><w:b/></w:rPr><w:t>2025.</w:t></w:r></w:p>'))),
 # a cell carrying ANY tracked change is the reviewer's: after the strip, the tool's own are gone,
 # so whatever survives is a person's (or a tool insertion a person has since edited)
 ("applied fix: a cell with a tracked change is left to the reviewer", lambda:
     A._cell_has_revisions(_ref_cell(
         '<w:p><w:ins w:id="9" w:author="Tony" w:date="2026-01-01T00:00:00">'
         '<w:r><w:t>Desorite, a New Phosphate Mineral.</w:t></w:r></w:ins></w:p>')) is True),
 # --- the three bugs the hand-check of all 26 fires caught (all would corrupt a title) ------
 # 1. An ordinary word INSIDE a multi-word place name must survive: 'New Mexico' is not
 #    'new Mexico'. The word is only kept because a proper noun follows it — 'a New Mineral'
 #    (ordinary word following) must still lowercase.
 ("ref_title: 'New Mexico' survives ('new' before a name is part of the name)", lambda:
     'New Mexico' in X._title_sentence_case(
         'Raydemarkite, the Natural Analogue of Synthetic MoO3, from Cookes Peak, New Mexico',
         'The natural analogue was found. A new analogue of synthetic MoO3 occurs in New Mexico. '
         'Specimens from New Mexico were studied; the New Mexico locality is natural.')[1]),
 ("ref_title: 'a New Mineral' still lowercases (ordinary word, ordinary neighbour)", lambda:
     'a new mineral' in X._title_sentence_case(
         'Desorite, a New Mineral Isotypic with Jamesite',
         'This is a new mineral. The new mineral is isotypic. Another new mineral occurs.')[1]),
 ("ref_title: 'Vanadium Queen mine, La Sal' survives", lambda: (lambda s:
     'Vanadium Queen' in s and 'La Sal' in s)(X._title_sentence_case(
         'Lasalite, a new decavanadate Mineral Species from the Vanadium Queen mine, La Sal District',
         'Found at the Vanadium Queen mine. The Vanadium Queen mine lies in the La Sal District. '
         'The La Sal district hosts a new mineral species; vanadium is abundant.')[1])),
 # 2. The paper's own Title-Cased title/running head must not teach the oracle that ordinary
 #    words are proper nouns — it repeats on every page and would block the whole correction.
 # ('A' opens the sentence after the colon, so it stays capital — only the ordinary words move)
 ("ref_title: a Title-Cased running head does not poison the evidence", lambda:
     'natural compound with a layered architecture' in X._title_sentence_case(
         'Kanatzidisite: A Natural Compound with a Layered Architecture',
         'Kanatzidisite: A Natural Compound With A Layered Architecture\n'
         'Kanatzidisite: A Natural Compound With A Layered Architecture\n'
         'We describe a natural compound. The natural compound is layered. '
         'Its architecture is layered and the compound is natural. The architecture '
         'of this natural compound is layered throughout.')[1]),
 # 3. Case is only ever LOWERED, never raised: all 53 corpus corrections went one way.
 ("ref_title: a lowercase word after a colon is not up-cased", lambda:
     ': occurrence' in X._title_sentence_case(
         'Stibiotantalite from the Nanyangshan LCT Pegmatite, Central China: occurrence and Data',
         'Stibiotantalite occurs there. The occurrence is in Central China. Central China hosts '
         'the pegmatite; data and occurrence are reported.')[1]),
 # --- optical 2V / sign vs refractive indices (docx-internal, blind-spot check #3) ---
 # I002959 (sulfatoredmondite): β=1.850 sits next to γ=1.860 (far from α=1.780) -> optically
 # NEGATIVE, and 2Vcalc≈40 matches the docx 2V=40, but the docx says Sign=+ (real sign error)
 ("optical_2v: sign contradicting the indices flags", lambda: [f for f in X.check24_optical_2v(
     type('S', (), {'comments': {'Optical Data': 'A=1.780(5), B=1.850(calc), Q=1.860(calc), Sign=+, 2V=40(3)'}}))
     if f.sev=='flag'] != []),
 ("optical_2v: gross 2V gap flags", lambda: [f for f in X.check24_optical_2v(
     type('S', (), {'comments': {'Optical Data': 'A=1.540, B=1.545, Q=1.600, Sign=+, 2V=80'}})) if f.sev=='flag'] != []),
 ("optical_2v: near-uniaxial (β≈γ) stays silent", lambda: X.check24_optical_2v(
     type('S', (), {'comments': {'Optical Data': 'A=1.795, B=1.820, Q=1.820, Sign=+, 2V=80'}})) == []),
 ("optical_2v: consistent sign+2V no flag", lambda: X.check24_optical_2v(
     type('S', (), {'comments': {'Optical Data': 'A=1.540, B=1.545, Q=1.600, Sign=+, 2V=35'}})) == []),
 # --- reflection geometry vs wavelength (docx-internal, blind-spot check #2) ---
 # I002983: d=0.70 Å is impossible under CuKα (λ/2=0.7703); I003406: d=0.7719 -> 2θ≈173° (a
 # synchrotron λ=0.708 mislabelled as CuKα). A normal d-range stays quiet.
 ("reflection_geom: d < λ/2 is impossible -> flag", lambda: [f for f in X.check25_reflection_geometry(
     type('S', (), {'instr': {'lam': '1.54056', 'anode': 'CuKa1'},
                    'refl': [('0.7000','11','4','1','11'), ('2.5','100','1','1','1')]})) if f.sev=='flag'] != []),
 ("reflection_geom: implausibly high 2θ -> flag", lambda: [f for f in X.check25_reflection_geometry(
     type('S', (), {'instr': {'lam': '1.54056', 'anode': 'CuKa1'},
                    'refl': [('0.7719','14','','',''), ('3.0','100','1','1','1')]})) if f.sev=='flag'] != []),
 ("reflection_geom: ordinary d-range stays quiet", lambda: X.check25_reflection_geometry(
     type('S', (), {'instr': {'lam': '1.54056', 'anode': 'CuKa1'},
                    'refl': [('1.5','100','1','1','1'), ('3.0','50','2','2','2')]})) == []),
 ("reflection_geom: no wavelength -> no flag", lambda: X.check25_reflection_geometry(
     type('S', (), {'instr': {}, 'refl': [('0.7','11','1','1','1')]})) == []),
 # a CALCULATED pattern is exempt — presentation λ (CuKα) + geometric d's + high 2θ are all fine
 ("reflection_geom: calculated pattern exempt", lambda: X.check25_reflection_geometry(
     type('S', (), {'instr': {'lam': '1.54056', 'anode': 'CuKa1', 'spacing_instr': 'Calculated'},
                    'refl': [('0.7000','11','4','1','11'), ('2.5','100','1','1','1')]})) == []),
 # --- document-structure (section) tier: the cell belongs to its subsection's experiment ---
 ("section: powder subsection -> powder", lambda: (lambda s: C._section_mode(s, s.find('a =')))(
     'x-ray powder diffraction data were collected (table 2). the refined unit cell is a = 5.0 and b = 6.0') == 'powder'),
 ("section: single-crystal subsection -> single", lambda: (lambda s: C._section_mode(s, s.find('a =')))(
     'single-crystal x-ray diffraction was carried out. the cell is a = 5.0') == 'single'),
 ("section: umbrella heading abstains (None)", lambda: (lambda s: C._section_mode(s, s.find('a =')))(
     'crystal structure determination was performed. the cell is a = 5.0') is None),
 ("section: incidental 'single-crystal X' doesn't beat a powder subsection", lambda: (lambda s: C._section_mode(s, s.find('a =')))(
     'single-crystal fragments were unavailable. x-ray powder diffraction gave a = 5.0') == 'powder'),
 ("section tier wired into classify_context", lambda: (lambda s: C.classify_context(s, s.find('a =')))(
     'single-crystal fragments were unavailable. x-ray powder diffraction gave a = 5.0') == 'powder'),
 # --- cell parsing (cubic / uniaxial / rounded-table) ---
 ("I003634 cubic cell matches",         lambda: cell_verdict('I003634') == 'match'),
 ("I003751 uniaxial: only c flagged",        lambda: params('I003751') == {'c'}),
 # symmetry-equivalent axes (cubic a=b=c, uniaxial a=b) collapse so an identical sig-fig/esd
 # error flags ONCE on the representative axis, not per axis (jonlarsenite I003697 is cubic).
 ("I003697 cubic precision collapses to one axis", lambda: params('I003697') == {'a'}),
 ("grouped_axis_issues: cubic a=b=c collapses each issue to one label",
  lambda: [l for l, r, k, n in C.grouped_axis_issues(
      {'a': '8.701(10)', 'b': '8.701(10)', 'c': '8.701(10)', 'α': '90', 'β': '90', 'γ': '90'},
      {'a': '8.70(1)', 'b': '8.70(1)', 'c': '8.70(1)', 'α': '90', 'β': '90', 'γ': '90'})] == ['a=b=c', 'a=b=c']),
 ("grouped_axis_issues: distinct axes stay separate",
  lambda: sorted(set(l for l, r, k, n in C.grouped_axis_issues(
      {'a': '5.00(1)', 'b': '6.00(1)', 'c': '7.00(2)', 'α': '90', 'β': '90', 'γ': '90'},
      {'a': '5.0(1)', 'b': '6.0(1)', 'c': '7.00(1)', 'α': '90', 'β': '90', 'γ': '90'}))) == ['a', 'b', 'c']),
 # Z (formula units): a/b/c match but the .pdf explicitly states a different 'Z = N'
 # (grokhovskyite: docx Z=2, paper 'Z = 3'); a docx-Z == paper-Z entry must NOT flag
 ("I003744 Z mismatch flagged (docx 2 vs .pdf Z=3)", lambda: 'Z' in params('I003744')),
 ("I003634 no Z flag (docx Z matches .pdf)",        lambda: 'Z' not in params('I003634')),
 ("I003747 clean (no rounding noise)",  lambda: not params('I003747')),
 ("I003698 clean (full-precision cell)", lambda: not params('I003698')),
 ("I003562 indexing flagged (>3%)",     lambda: bool(extras('I003562', 'indexing', 'flag'))),
 # weak reflections (I≤35) only flag at >5%: the I=16 line 4.1440 [3.2%] is dropped,
 # the strong I=80 line 5.0790 [3.1%] is kept
 ("I003562 indexing keeps strong 5.0790",  lambda: bool(extras('I003562', 'indexing', substr='5.0790'))),
 ("I003562 indexing drops weak 4.1440",     lambda: not extras('I003562', 'indexing', substr='4.1440')),
 # --- optical sign (control-char glyphs must NOT be guessed) ---
 ("I003528 no optical-sign flag",       lambda: not extras('I003528', 'optical')),
 ("I003633 no optical-sign flag",       lambda: not extras('I003633', 'optical')),
 # multi-mineral comparison table (several distinct signs) -> can't attribute -> no flag
 ("optical: multi-mineral table doesn't flag", lambda: X.check10_optical(
     type('S', (), {'comments': {'Optical Data': 'Sign=-'}}),
     'Optical class biaxial (+) biaxial (+) biaxial (-) biaxial (+) biaxial (-) biaxial (-)') == []),
 ("optical: single consistent sign still flags", lambda: X.check10_optical(
     type('S', (), {'comments': {'Optical Data': 'Sign=+'}}),
     'the mineral is optically biaxial (-), 2V = 50') != []),
 # strongest line: use the entry's prose list, not a wrong column of a comparison table
 ("strongest_lines: prose I=100 beats comparison-table column", lambda: X.check15_strongest_lines(
     type('S', (), {'refl': [('2.583', '100', '2', '0', '0'), ('3.20', '20', '1', '1', '0')]}),
     'table dmeas(I) 3.20(20) 3.174(100) 3.00(34). The strongest lines in the powder '
     'diffraction pattern are (d A, I %, hkl): 4.49, 31, (110); 2.583, 100, (200).') == []),
 ("strongest_lines: genuinely missing I=100 line still flags", lambda: X.check15_strongest_lines(
     type('S', (), {'refl': [('3.20', '20', '1', '1', '0')]}),
     'The strongest lines (d A, I %, hkl): 4.49, 31, (110); 2.583, 100, (200).') != []),
 # --- IMA number (new mineral vs reinvestigation/reference) ---
 ("I003633 IMA flag (new mineral)",     lambda: bool(extras('I003633', 'ima'))),
 ("I003688 IMA flag (new mineral)",    lambda: bool(extras('I003688', 'ima'))),
 # the IMA-missing comment must anchor in the Comments section (where IMA numbers live),
 # not fall through to the Author's Cell label; entries WITH an 'IMA Number' row keep
 # anchoring to its value cell. (moxuanxueite I003633 has no IMA Number row.)
 ("I003633 IMA comment anchors to Comments section (no IMA Number row)", lambda: _anchor_text('I003633', 'ima') == 'Comments'),
 ("I003246 IMA anchor stays on the value cell when an IMA Number row exists", lambda: _anchor_text('I003246', 'ima') not in (None, 'Comments')),
 # same for the Analysis field: when an entry has no Analysis row, the note anchors in the
 # Comments section (where Analysis lives), not on the Author's Cell (esdanaite-(Ce) I003699
 # has the analytical data misplaced in 'Absolute Configuration', no Analysis row).
 ("I003699 analysis comment anchors to Comments section (no Analysis row)", lambda: _anchor_text('I003699', 'analysis') == 'Comments'),
 ("I003548 analysis anchor stays on the Analysis value cell when present", lambda: _anchor_text('I003548', 'analysis') not in (None, 'Comments')),
 ("I003687 NO IMA flag (reinvest.)",     lambda: not extras('I003687', 'ima')),
 ("I003511 NO IMA flag (reinvest.)",    lambda: not extras('I003511', 'ima')),
 # --- classification (docx already names the group) ---
 ("I003657 no classification comment",    lambda: not extras('I003657', 'classification', 'info')),
 # --- 3. structure relation: don't restate a docx Structure comment; DO check when the
 #     .pdf asserts a relation and the docx has none ---
 ("I003416 no Structure-comment noise (docx already has it)", lambda: not extras('I003416', 'classification', substr='structure')),
 # fluormacraeite is already classified in the Paulkerrite Group, so 'isostructural with
 # paulkerrite' is redundant (group membership implies it) -> suppress the note
 ("I003246 isostructural-with-own-group note suppressed", lambda: not extras('I003246', 'classification', substr='no Structure')),
 # but an OUT-of-group structural relation must still be surfaced
 ("structure-relation: out-of-group isostructural still noted", lambda: [f for f in X.check3_classification(
     type('S', (), {'comments': {}, 'raw_rows': [['IMA Classifications: Some Group Otherite']], 'name': 'zzznotarealmineral', 'primary': ''}),
     'the new phase is isostructural with stranskiite and shares its topology.') if 'no Structure' in f.msg] != []),
 ("structure-relation: isostructural-with-own-group-namesake suppressed", lambda: [f for f in X.check3_classification(
     type('S', (), {'comments': {}, 'raw_rows': [['IMA Classifications: Paulkerrite Group Paulkerrite']], 'name': 'zzznotarealmineral', 'primary': ''}),
     'the new phase is isostructural with paulkerrite, the group prototype.') if 'no Structure' in f.msg] == []),
 # a mineral can't be 'related to itself' (paper names THIS species as the anchor for others)
 ("structure-relation: self-reference suppressed", lambda: [f for f in X.check3_classification(
     type('S', (), {'comments': {}, 'raw_rows': [], 'name': 'Keutschite', 'primary': 'Keut'}),
     'minerals chemically or structurally related to keutschite, including agmantinite and stannite.') if 'no Structure' in f.msg] == []),
 # --- precision check names the esd to add ---
 ("I003566 names esd to add",      lambda: bool(extras('I003566', 'precision', substr='.pdf gives'))),
 # --- analysis field ---
 ("I003510 misplaced-analysis flag",  lambda: bool(extras('I003510', 'analysis', substr='moved'))),
 # the analysis-count flag carries the .pdf wording ('average of N …') as evidence so the GUI
 # '? look' lands on that phrase, not a generic spot (gorbunovite I003704 -> 'Average of 10 analyses')
 ("I003704 analysis-count flag carries 'average of N' as look evidence",
  lambda: any(f.code == 'analysis' and (f.evidence or '').lower().startswith('average of 10')
              for f in (res_for('I003704') or {}).get('extra', []))),
 # multi-species paper: each species' count comes from the 'average of N' beside ITS OWN name
 # ("auropolybasite (average of 11 analyses) and auropearceite (average of 17 analyses)"), not
 # the first global match — so auropearceite is 17, not auropolybasite's 11.
 ("I003823 auropearceite analysis count = 17 (species-specific)", lambda: bool(extras('I003823', 'analysis', substr='n=17'))),
 ("I003822 auropolybasite analysis count = 11 (species-specific)", lambda: bool(extras('I003822', 'analysis', substr='n=11'))),
 # a spelled-out count tied to grains ("(five analyses on one grain)", keutschite I003806) is
 # recognised; spelled-out numbers are allowed ONLY in this 'on … grain(s)' form (too ambiguous
 # elsewhere — 'two grains', 'Table 3 analyses nos.').
 ("I003806 keutschite analysis count = 5 ('five analyses on one grain')", lambda: bool(extras('I003806', 'analysis', substr='n=5'))),
 ("analysis: _count_val parses spelled-out numbers", lambda: X._count_val('five') == 5 and X._count_val('12') == 12 and X._count_val('xyz') is None),
 # --- density check is deregistered (never fires) ---
 ("No density flags anywhere",               lambda: not any(extras(e, 'density')
                                                   for e in ('I003704','I003448','I003562'))),
 # --- 16. instrumentation designators (corpus-curated) ---
 ("I003246 no geometry nag (Spacing Instr. already Diffractometer)", lambda: not extras('I003246', 'geometry')),
 ("I003747 no instr_vocab flag (clean)",     lambda: not extras('I003747', 'instr_vocab', 'flag')),
 ("I003698 no instr_vocab flag (clean)",     lambda: not extras('I003698', 'instr_vocab', 'flag')),
 # 'Debye-Scherrer' is a geometry, not an instrument class -> normalize to 'Diffractometer'
 # (reviewers did this 15x; the value never survives review). Diffractometer itself stays clean.
 ("instr_vocab: Debye-Scherrer -> Diffractometer", lambda: any(
     f.code == 'instr_vocab' and 'Diffractometer' in (f.msg or '')
     for f in X.check16_instr_vocab(type('S', (), {'instr': {'intensity_instr': 'Debye-Scherrer'}})(), None))),
 ("instr_vocab: Diffractometer is clean", lambda: not [
     f for f in X.check16_instr_vocab(type('S', (), {'instr': {'intensity_instr': 'Diffractometer'}})(), None)
     if f.sev == 'flag']),
 # --- 16b. measured-data completeness (blank anode/intensity-type) ---
 ("I003600 blank anode -> derive Sync (synch λ)", lambda: bool(extras('I003600', 'instr_vocab', substr='set the anode to Sync'))),
 ("I003246 no blank-field flag",             lambda: not extras('I003246', 'instr_vocab', substr='blank')),
 # --- 17. PDF monochromator/β-filter -> fill blank Filter (powder-context only) ---
 ("I003747 .pdf graphite mono (powder)",      lambda: bool(extras('I003747', 'instr_filter', substr='graphite'))),
 ("I003745 SC mono NOT attributed",          lambda: not extras('I003745', 'instr_filter')),
 ("I003566 SC mono NOT attributed",          lambda: not extras('I003566', 'instr_filter')),
 # --- 19. Intensity Type follows the detector (area→Integrated, BB→Peak) ---
 ("I003657 Bragg-Brentano -> Peak flag",     lambda: bool(extras('I003657', 'intensity_type'))),
 ("I003563 Gandolfi/area -> Integrated flag", lambda: bool(extras('I003563', 'intensity_type'))),
 # a powder pattern whose OBSERVED intensities are ALL multiples of 10 was visually estimated
 # from film -> Intensity Type Visual. This OVERRIDES the Gandolfi/area rule (a Gandolfi FILM
 # camera is not a digital detector): spaltiite (I003807, all-×10 Irel) -> Visual not Peak. The
 # measured film cameras (intensities not all-×10) and modern Gandolfi (I003563) are unaffected.
 ("I003807 visual: all-×10 intensities -> Intensity Type Visual, not Peak",
  lambda: bool(extras('I003807', 'intensity_type', 'flag', 'Visual'))),
 # a measured (diffractometer) pattern can still carry VISUALLY-ESTIMATED intensities; the .pdf
 # saying so is its own signal (keutschite: Rigaku Rapid II d-spacings, Table-5 vs/s/ms/w Irel).
 ("I003806 keutschite: '.pdf visually estimated' -> Intensity Type Visual, not Peak",
  lambda: bool(extras('I003806', 'intensity_type', 'flag', 'Visual'))),
 # the Visual flag must carry '?look' evidence (the 'visually estimated' phrase, at the PXRD-table
 # legend) so the button lands on the table instead of going nowhere.
 ("I003806 visual flag carries '?look' evidence",
  lambda: any(f.code == 'intensity_type' and (f.evidence or '').lower().startswith('visual')
              for f in (res_for('I003806') or {}).get('extra', []))),
 ("I003563 Gandolfi->Integrated not overridden (intensities not all-×10)",
  lambda: bool(extras('I003563', 'intensity_type', 'flag', 'Integrated')) and not extras('I003563', 'intensity_type', 'flag', 'Visual')),
 ("I003822 calc pattern -> no intensity_type", lambda: not extras('I003822', 'intensity_type')),
 # --- 7. synthetic: the 'Synthetic Rock-forming minerals' subfile is a CATEGORY (a
 #     natural mineral is also filed there) — not a 'this sample is synthetic' flag ---
 ("I003512 no synthetic note (Synthetic Rock-forming category, not a synthetic sample)",
                                              lambda: not extras('I003512', 'synthetic')),
 # --- 20. calculated pattern, λ not stated in the paper ---
 # only a NON-standard calc wavelength is worth flagging; I003511 is calculated at the
 # default CuKα (1.54056), legitimately absent from a paper using another source -> no flag
 ("I003511 no calc_wavelength flag (standard CuKα calc default)", lambda: not extras('I003511', 'calc_wavelength')),
 ("calc_wavelength: non-standard λ still flags",
   lambda: bool(X.check20_calc_wavelength(type('S', (), {'instr': {'spacing_instr': 'Calculated', 'lam': '2.28970', 'anode': 'CrKa'}})(),
                'the powder pattern was calculated from the single-crystal structure'))),
 # --- 21. Primary name normalization: corrected names must NOT flag ---
 ("I003523 primary name OK (no FP)",         lambda: not extras('I003523', 'primary_name')),
 # mindat name matching folds diacritics: Mindat keeps the accented species name
 # ('Åsgruvanite-(Ce)', 'Désorite') but the docx uses the ASCII form — they must match,
 # else a real species reads as 'not found among IMA species'. (Cache-independent unit check.)
 ("mindat _norm folds diacritics (Å/é/č/ě/š == ASCII)",
  lambda: M._norm('Åsgruvanite-(Ce)') == M._norm('Asgruvanite-(Ce)')
          and M._norm('Désorite') == M._norm('Desorite')
          and M._norm('Quartz') == 'quartz'),
 # security: docx XML parsing is not XXE-able (DTD rejected / external entity unresolved)
 ("docx XML parsers block XXE (DTD/external entity)", _xxe_blocked),
 # discovery is recursive (docx at any depth), skips review_out/, prefers the transcription
 ("discover() finds docx recursively, skips review_out/ + supp", _discover_recursive_ok),
 ("discover() picks the most-reviewed copy, not the raw transcription", _discover_prefers_reviewed_ok),
 # --- 22. cross-source: synthetic skips Mindat (natural≠synthetic); agreeing entry clean ---
 ("I003599 synthetic -> Mindat cell skipped",  lambda: not extras('I003599', 'mindat_fix')),
 ("I003246 no cross-source flag (agrees)",   lambda: not extras('I003246', 'cell_cif') and not extras('I003246', 'mindat_fix')),
 # --- .dft (DataQuacker) soft cross-check: console/log only, never written ---
 ("I003599 .dft Z verify note",              lambda: bool(extras('I003599', 'dft', substr='Z:'))),
 ("I003599 .dft notes are sev=note (not flag)", lambda: not [f for f in (res_for('I003599')['extra'] if res_for('I003599') else []) if f.code == 'dft' and f.sev == 'flag']),
 ("I003246 no .dft note (agrees)",           lambda: not extras('I003246', 'dft')),
 # --- 18. name vs ideal formula: correctly-named REE species must NOT flag ---
 ("I003511 Allanite-(Y) name OK",            lambda: not extras('I003511', 'name_formula', 'flag')),
 ("I003523 Parisite-(Nd) name OK",           lambda: not extras('I003523', 'name_formula', 'flag')),
 ("I003521 Marsaalamite-(Y) name OK",        lambda: not extras('I003521', 'name_formula', 'flag')),
 # --- 2026-07 fix batch: errored-check filing, parser hardening, misfire guards ---
 ("run_all files an errored check under its real code", _errored_check_filed_under_real_code),
 ("pdf_index pairs uppercase '.PDF'", _pdf_index_uppercase_ok),
 ("esd rejoin: a reference '(3)' line-opener never fuses across the newline",
  lambda: C._norm_pdf('value 112\n(3) Smith ref') == 'value 112\n(3) Smith ref'),
 ("esd rejoin: line-wrapped esd with a unit still fuses (I002066 class)",
  lambda: '5.2783(3)' in C._norm_pdf('a = 5.2783\n(3) Å')),
 ("find_radiation: nm-quoted λ converted to Å",
  lambda: any(r[1] == '1.54178' for r in C.find_radiation(
      'The powder pattern was collected with CuKα radiation, λ = 0.154178 nm.'))),
 ("find_radiation: a nearby cell parameter is not read as λ",
  lambda: all(r[1] is None for r in C.find_radiation(
      'Powder data were collected with CuKα radiation. Unit-cell parameter a = 5.4602 Å.'))),
 ("best_match: a partial candidate is never zip-mispaired (mode stays direct)",
  lambda: C.best_match(['5.0', '7.0', '10.0'],
      [C.CellCand('7.0', None, '10.0', None, None, None, None, None, 'powder', 0, '', None)])[4] == 'direct'),
 ("parse_docx pads a short (merged) Author's-Cell row to 8",
  lambda: C._pad8(['1', '2']) == ['1', '2', '', '', '', '', '', '']),
 ("fixed-angle skip: docx γ=90 vs .pdf 120 now compares",
  lambda: any(k == 'value' for _, _, k, _ in C.grouped_axis_issues({'γ': '90'}, {'γ': '120'}, ['γ']))),
 ("fixed-angle skip: docx 120 with a silent .pdf still skipped",
  lambda: C.grouped_axis_issues({'γ': '120'}, {'γ': None}, ['γ']) == []),
 ("find_cells: two-species sentence never chimeras α/γ (I003403 class)", _no_two_species_chimera),
 ("find_cells: distant vertical tables stay separate candidates", _vertical_tables_stay_separate),
 ("named rows: EPMA 'Point Ca …' is not a cell; a mineral row with a 2-char SG survives",
  lambda: C.find_cells('Point Ca 3.45 2.98 4.01') == []
          and any(c.phase == 'Whatzite' for c in C.find_cells('Whatzite Cc 5.20 9.00 10.40'))),
 ("I003523 polytype-table paper still cell-matches",
  lambda: cell_verdict('I003523') == 'match'),
 ("check4: another species' not-collected/calculated does not flag a measured entry",
  lambda: X.check4_calculated(_stub(name='alphaite', primary='Alphaite',
                                    instr={'spacing_instr': 'Diffractometer'}),
      'Powder X-ray diffraction data for betaite could not be collected owing to the paucity of '
      'material. The powder X-ray diffraction pattern of betaite was calculated from the '
      'single-crystal structure model.') == []),
 ("check4: the not-collected species itself still flags",
  lambda: [f for f in X.check4_calculated(_stub(name='betaite', primary='Betaite',
                                                instr={'spacing_instr': 'Diffractometer'}),
      'Powder X-ray diffraction data for betaite could not be collected. The powder X-ray '
      'diffraction pattern of betaite was calculated from the single-crystal structure model.')
      if f.sev == 'flag'] != []),
 ("check4: English -ite words are not species (stoplist)",
  lambda: [f for f in X.check4_calculated(_stub(instr={'spacing_instr': 'Diffractometer'}),
      'Despite the composite satellite reflections, the powder pattern was calculated from the '
      'structure model.') if f.sev == 'flag'] != []),
 ("_ree_coeffs: a flattened charge digit is not a coefficient",
  lambda: X._ree_coeffs('(La0.6Ce3+0.4)PO4') == {'La': 0.6, 'Ce': 0.4}
          and X._ree_coeffs('Ce3+') == {'Ce': 1.0}),
 ("check18: charge-digit formula no longer suggests a wrong rename",
  lambda: X.check18_name_formula(_stub(primary='Testite-(La)', name='Testite-(La)',
      formulas={'Empirical': '(La0.6Ce3+0.4)PO4'}), '') == []),
 ("check18: a genuine wrong Levinson suffix still flags",
  lambda: bool(X.check18_name_formula(_stub(primary='Testite-(Ce)', name='Testite-(Ce)',
      formulas={'Empirical': '(La0.6Ce3+0.4)PO4'}), ''))),
 ("_pdf_esd_for: no suffix match inside a longer number",
  lambda: X._pdf_esd_for('12.219', 'beta = 112.219(5)') is None
          and X._pdf_esd_for('12.219', 'c = 12.219(3)') == '3'),
 ("check10: a prose dash after 'biaxial' is not an optic sign",
  lambda: X.check10_optical(_stub(comments={'Optical Data': 'Sign=+'}),
      'The mineral is biaxial – the sign could not be measured.') == []),
 ("check10: a parenthesized sign still reads ('biaxial (–)' flags vs docx '+')",
  lambda: bool(X.check10_optical(_stub(comments={'Optical Data': 'Sign=+'}),
      'Optically it is biaxial (–), 2V = 60 degrees.'))),
 ("candidate_groups: setting-swapped cell reads similar, not sub/super-cell",
  lambda: G.cell_relation({'a': 5, 'b': 7, 'c': 10}, {'a': 10, 'b': 7, 'c': 5}, 0.04)[0] == 'similar'
          and G.cell_relation({'a': 5, 'b': 7, 'c': 10}, {'a': 10, 'b': 14, 'c': 20}, 0.04)[0] == 'sub/super-cell'),
 ("candidate_groups: volume carries the angle term; γ=0 (Mindat uniaxial) reads as 90°",
  lambda: G._vol({'a': 10, 'b': 10, 'c': 10, 'be': 125}) < G._vol({'a': 9.8, 'b': 9.8, 'c': 9.8, 'be': 91})
          and G._vol({'a': 5, 'b': 5, 'c': 5, 'ga': 0}) == G._vol({'a': 5, 'b': 5, 'c': 5})),
 # --- corpus-wide: no check may error on any fixture entry (kept LAST so it reuses
 #     the analyses the pointwise cases already cached) ---
 ("No errored checks across the fixture corpus", _no_errored_checks),
]

def main():
    _init_fixtures()
    if not glob.glob(os.path.join(FOLDER, 'I*.docx')) and not C.discover(FOLDER):
        print("!! batch not found at %r — pass the path as an argument" % FOLDER)
        return 2
    fails = 0
    for desc, pred in CASES:
        try:
            ok = bool(pred())
        except Exception as e:
            ok = False; desc += "  (threw: %s)" % e
        print("%s  %s" % ("PASS" if ok else "FAIL", desc))
        fails += not ok
    # policy: no entry in this batch should withhold auto-Accept (none are 'severe')
    severe = [C.entry_id(d) for d in sorted(glob.glob(os.path.join(FOLDER, '*.docx')))
              if '~$' not in d and A._is_severe(A.analyze(d, _idx.get(C.entry_id(d)), _cif.get(C.entry_id(d))))]
    ok = not severe
    print("%s  No entry withheld from auto-Accept (severe=%s)" % ("PASS" if ok else "FAIL", severe))
    fails += not ok
    print("\n%d case(s) FAILED" % fails if fails else "\nAll regression cases pass.")
    return 1 if fails else 0

if __name__ == '__main__':
    sys.exit(main())
