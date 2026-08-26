#!/usr/bin/env python3
"""
tables — publishable structure tables from a .cif, formatted the way the mineralogy journals (and
the corpus manuscripts) print them:

  1. Atom coordinates and displacement parameters (Å²): Atoms | s.o. | x | y | z | Ueq/Uiso |
     U11 U22 U33 U12 U13 U23 — site occupancy as element + subscript fraction ('Ca1.00',
     'F0.80O0.20' for a mixed site), fixed special coordinates as ⅓ ⅔ ½ ¼ ¾, isotropic atoms with
     '-' in the U columns, values verbatim from the .cif (esds kept).
  2. Selected bond distances (Å): one block per cation in three label–distance column pairs,
     symmetry-equivalent atoms marked with a superscript code, '<M–O>' mean rows (uranyl gets
     <U–Oyl> and <U–Oeq>), a 'Symmetry codes' note. Distances and esds come from the .cif's own
     _geom_bond loop when it has one, else they are computed (no esd).
  3. Hydrogen bonds: D–H⋯A | D–H | H⋯A | D⋯A | ∠DHA from the _geom_hbond loop, else computed
     from the H positions (H⋯A ≤ 2.6 Å, angle ≥ 110°).
  4. Bond-valence analysis (vu): anion rows × cation columns with ×n↓ / ×n→ marks, hydrogen-bond
     Donor | vu columns on the acceptor rows, Σan (cations + accepted H bonds) and a Σcat row —
     from bv_check, same parameters and conventions (--params, --ox, --cutoff, --no-h).

    python3 -m pxrd_review.tables <structure.cif> [--word] [--journal ammin|minmag|cjmp|ejm] [--params gh|bo|ba] [--ox Fe=2] [--out DIR]
    pxrd tables <structure.cif> --word --journal cjmp

--journal picks a journal's conventions (caption form, header labels, sum labels, rules, notes,
font) from the JOURNALS registry below — each rule notes whether it comes from the journal's
author instructions or from what its published papers do.

Console shows the tables as text; --word writes review_out/<name>_tables.docx with captions and
notes (italic x y z U, superscript indices and codes, subscript occupancies) to paste from.
"""
import os, re, sys, math, argparse
from collections import OrderedDict

from pxrd_review import bv_check as B

# ----------------------------------------------------------------------------- journal styles
# Each journal's table conventions. Sources: the journal's own author instructions where they say
# something ('doc'), otherwise what its published papers in the corpus do ('corpus', ~300 papers
# surveyed) — the owner's notes may correct the latter. Every table the tool writes follows the
# chosen journal: caption form, header labels, sum labels, rules, notes, font.
JOURNALS = {
    'ammin': {
        'name': 'American Mineralogist',
        'caption': 'Table {n}.',            # doc: brief titles; printed 'Table 3.' (corpus: 32 of 36 captions)
        'title_case': 'sentence',
        'font': None,                       # not prescribed
        'rules': 'three',                   # doc: no vertical/diagonal rules, no shading (top/header/bottom kept)
        'notes_prefix': 'Notes: ',          # corpus: 'Note:'/'Notes:' lead the note line; doc: footnotes end with periods
        'foot_marks': ['*', '†', '‡', '§'],
        'atom_head': 'Atom', 'coord_heads': ('x', 'y', 'z'), 'u_head': 'Ueq/Uiso',
        'sum_labels': ('Σcat', 'Σan'),      # corpus: Σ used (147); Σcat/Σan as the owner's manuscripts
        'symcodes': 'Symmetry codes:',
        'bonds_caption': 'Selected bond distances (Å) for {name}',
        'bvs_caption': 'Bond-valence analysis (in valence units) for {name}',
        'coords_caption': 'Atom coordinates and displacement parameters (Å²) for {name}',
        'hbonds_caption': 'Hydrogen-bond geometry (Å, °) for {name}',
    },
    'minmag': {
        'name': 'Mineralogical Magazine',
        'caption': 'Table {n}.',            # doc: 'Table 5. In bold, full stop, space after full stop'
        'title_case': 'sentence',
        'font': None,
        'rules': 'three',                   # doc: 'remove internal borders and merged cells'
        'notes_prefix': '',                 # doc: abbreviations defined in captions or footnotes ('LA = …'); corpus: '*' / 'Note:'
        'foot_marks': ['*', '†', '‡', '§'],
        'atom_head': 'Site', 'coord_heads': ('x/a', 'y/b', 'z/c'), 'u_head': 'Ueq',   # corpus: 'Site' 39, 'x/a' 25
        'sum_labels': ('Σcat', 'Σan'),
        'symcodes': 'Symmetry operators:',  # corpus (2 of 4)
        'bonds_caption': 'Selected bond distances (Å) for {name}',
        'bvs_caption': 'Bond-valence sums (in valence units) for {name}',
        'coords_caption': 'Atom coordinates, site occupancies and displacement parameters (Å²) for {name}',
        'hbonds_caption': 'Hydrogen-bond geometry (Å, °) for {name}',
    },
    'cjmp': {
        'name': 'Canadian Journal of Mineralogy and Petrology',
        'caption': 'TABLE {n}.',            # doc (template): all caps title, period after the number, none at the end
        'title_case': 'caps',               # doc: caps except element symbols / case-sensitive abbreviations
        'font': 'Arial',                    # doc: Arial, including symbols, title and notes
        'rules': 'three',                   # doc: one rule across the top, one below the header row, one at the bottom
        'notes_prefix': '',                 # doc: notes below the table, outside the body ('* Notes go here.')
        'foot_marks': ['*', '§', '†', '‡'], # corpus: '*' 63, '§' 43
        'atom_head': 'Atom', 'coord_heads': ('x', 'y', 'z'), 'u_head': 'Ueq',
        'sum_labels': ('Sum', 'Sum'),       # corpus: 'BVS'/'Sum', no Σ
        'symcodes': 'Symmetry operators:',
        'bonds_caption': 'Selected interatomic distances (Å) in {name}',
        'bvs_caption': 'Bond-valence analysis for {name}',
        'coords_caption': 'Atom coordinates and displacement parameters (Å²) for {name}',
        'hbonds_caption': 'Hydrogen-bond geometry (Å, °) in {name}',
    },
    'ejm': {
        'name': 'European Journal of Mineralogy',
        'caption': 'Table {n}.',            # corpus: 37 of 37 sentence case
        'title_case': 'sentence',
        'font': None,
        'rules': 'three',
        'notes_prefix': '',
        'foot_marks': ['*', '†', '‡'],
        'atom_head': 'Site', 'coord_heads': ('x', 'y', 'z'), 'u_head': 'Ueq',   # corpus: 'Site' 32
        'sum_labels': ('Σcat', 'Σan'),
        'symcodes': 'Symmetry codes:',
        'bonds_caption': 'Selected bond distances (Å) for {name}',
        'bvs_caption': 'Bond-valence analysis (valence units) for {name}',
        'coords_caption': 'Fractional atomic coordinates and displacement parameters (Å²) for {name}',
        'hbonds_caption': 'Hydrogen-bond geometry (Å, °) for {name}',
    },
}
JOURNALS['manuscript'] = dict(JOURNALS['ammin'], name='the owner\'s manuscript style (Am Min)', atom_head='Atoms')
DEFAULT_JOURNAL = 'manuscript'

def journal(key):
    return JOURNALS.get(key or DEFAULT_JOURNAL, JOURNALS[DEFAULT_JOURNAL])

def _title(J, text):
    """The caption text in the journal's case; element symbols and labels survive a caps title."""
    if J['title_case'] != 'caps':
        return text
    def keep(m):
        return m.group(0)
    out = text.upper()
    # put back mixed-case tokens that are element symbols / formulas / mineral names the journal keeps
    for tok in re.findall(r"[A-Z][a-z](?:\d|\b)|\([A-Za-z0-9]+\)|Å²|Å", text):
        out = out.replace(tok.upper(), tok, 1)
    return out

# ----------------------------------------------------------------------------- rich cells
# A cell is a list of (text, style) runs; style is a set of 'i' (italic), 'sup', 'sub', 'b'.
def R(text, *styles):
    return [(str(text), set(styles))]

def C(*parts):
    """Concatenate runs / plain strings into one cell."""
    out = []
    for p in parts:
        if p is None:
            continue
        if isinstance(p, str):
            out.append((p, set()))
        else:
            out.extend(p)
    return out

def T_plain(cell):
    return ''.join(t for t, _ in cell)

def plain(cell):
    """Plain text with Unicode super/subscripts where they exist (console / clipboard)."""
    SUP = str.maketrans('0123456789+-', '⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻')
    SUB = str.maketrans('0123456789.', '₀₁₂₃₄₅₆₇₈₉.')
    out = []
    for t, st in cell:
        if 'sup' in st and re.fullmatch(r'[\d+\-]+', t):
            out.append(t.translate(SUP))
        elif 'sub' in st and re.fullmatch(r'[\d.]+', t):
            out.append(t.translate(SUB))
        else:
            out.append(t)
    return ''.join(out)

def html(cell):
    import html as _h
    out = []
    for t, st in cell:
        s = _h.escape(t)
        if 'sup' in st: s = '<sup>%s</sup>' % s
        if 'sub' in st: s = '<sub>%s</sub>' % s
        if 'i' in st: s = '<i>%s</i>' % s
        if 'b' in st: s = '<b>%s</b>' % s
        out.append(s)
    return ''.join(out)

# ----------------------------------------------------------------------------- formatting helpers
FRACTIONS = [(1/3, '⅓'), (2/3, '⅔'), (0.5, '½'), (0.25, '¼'), (0.75, '¾'), (1/6, '⅙'), (5/6, '⅚'), (1/8, '⅛'), (3/8, '⅜'), (5/8, '⅝'), (7/8, '⅞')]

def coord(s):
    """A fractional coordinate as printed: fixed special values as fractions, esds kept."""
    s = (s or '').strip()
    if '(' in s:
        return s
    v = B._num(s)
    if v is None:
        return s
    if abs(v) < 1e-6:
        return '0'
    for f, sym in FRACTIONS:
        if abs(abs(v) - f) < 2e-4:
            return ('-' if v < 0 else '') + sym
    if abs(v - round(v)) < 1e-6:
        return str(int(round(v)))
    return s

def _u(s):
    s = (s or '').strip()
    return '0' if re.fullmatch(r'-?0\.0+', s) else s

def _elem_occ(species):
    """'Ca1.00' / 'F0.80O0.20' as runs: element, subscript occupancy."""
    cell = []
    for sp in species:
        cell += R(sp.element) + R('%.2f' % sp.occ, 'sub')
    return cell

def _const(v):
    """A symmetry-operator constant as a fraction string ('½', '1', '−⅓')."""
    sign = '−' if v < 0 else ''
    a = abs(v)
    if abs(a - round(a)) < 1e-6:
        return sign + str(int(round(a))) if round(a) else ''
    for f, sym in FRACTIONS:
        if abs(a - f) < 1e-6:
            return sign + sym
    for den in (2, 3, 4, 6, 8):                     # an improper fraction: 4/3, 3/2 …
        if abs(a * den - round(a * den)) < 1e-6:
            return '%s%d/%d' % (sign, int(round(a * den)), den)
    return sign + ('%.4f' % a).rstrip('0')

def symop_string(rot, tr):
    """(rot, tr) -> 'x−y+1, x, −z+½' in the journals' style."""
    parts = []
    for i in range(3):
        s = ''
        for j, v in enumerate('xyz'):
            c = rot[i][j]
            if abs(c) < 1e-6:
                continue
            s += ('−' if c < 0 else ('+' if s else '')) + (v if abs(abs(c) - 1) < 1e-6 else '%g%s' % (abs(c), v))
        k = _const(tr[i])
        if k:
            s += (k if (k.startswith('−') or not s) else '+' + k)
        parts.append(s or '0')
    return ', '.join(parts)

class SymCodes:
    """Superscript numbering of symmetry codes, shared by the bond and hydrogen-bond tables:
    a CIF code '17_544' (operator 17, translation 0,−1,−1) -> (n, 'x−y, x, −z−1')."""
    def __init__(self, st):
        self.st = st; self.codes = OrderedDict(); self.used = set()
    def begin_table(self):
        self.used = set()
    def index(self, code):
        code = (code or '.').strip()
        if code in ('.', '', '?', '1_555'):
            return None
        self.used.add(code)
        if code not in self.codes:
            m = re.match(r'^(\d+)(?:_(\d)(\d)(\d))?$', code)
            if not m:
                self.codes[code] = code; return len(self.codes)
            n = int(m.group(1)); t = [int(m.group(k)) - 5 for k in (2, 3, 4)] if m.group(2) else [0, 0, 0]
            if 1 <= n <= len(self.st.ops):
                rot, tr = self.st.ops[n - 1]
                self.codes[code] = symop_string(rot, [tr[i] + t[i] for i in range(3)])
            else:
                self.codes[code] = code
        return list(self.codes).index(code) + 1
    def note(self, phrase='Symmetry codes:'):
        """The codes THIS table used (numbering shared across tables)."""
        items = [(i + 1, s) for i, (c, s) in enumerate(self.codes.items()) if c in self.used]
        if not items:
            return ''
        return phrase + ' ' + '; '.join('(%d) %s' % (i, s) for i, s in items) + '.'

def _label_cell(label, code_idx):
    return C(label, R(str(code_idx), 'sup') if code_idx else None)

def _bond_cell(cat, an, code_idx):
    return C(cat, '–', an, R(str(code_idx), 'sup') if code_idx else None)

# ----------------------------------------------------------------------------- the tables

def coordinates_table(st, J=None):
    """Rows for 'Atom coordinates and displacement parameters'."""
    J = J or journal(None)
    b = st.block
    tags, rows = B._loop(b, '_atom_site_fract_x')
    labels = B._col(tags, rows, '_atom_site_label')
    xs, ys, zs = (B._col(tags, rows, '_atom_site_fract_' + k) for k in 'xyz')
    uiso = B._col(tags, rows, '_atom_site_u_iso_or_equiv', '')
    adp = B._col(tags, rows, '_atom_site_adp_type', '')
    atags, arows = B._loop(b, '_atom_site_aniso_u_11')
    aniso = {}
    if arows:
        al = B._col(atags, arows, '_atom_site_aniso_label')
        for i, lab in enumerate(al):
            aniso[lab] = {k: B._col(atags, arows, '_atom_site_aniso_u_' + k, '')[i] for k in ('11', '22', '33', '12', '13', '23')}
    have_aniso = bool(aniso)
    by_label = {}
    for s in st.sites:
        for lab in s.label.split('/'):
            by_label[lab] = s
    out = []; done = set()
    for i, lab in enumerate(labels):
        site = by_label.get(lab)
        if site is None or site.label in done:
            continue
        done.add(site.label)
        row = [C(site.label), _elem_occ(site.species), C(coord(xs[i])), C(coord(ys[i])), C(coord(zs[i])), C(uiso[i].strip() or '-')]
        if have_aniso:
            a = aniso.get(lab)
            if a and adp[i].strip().lower() != 'uiso':
                row += [C(_u(a[k])) for k in ('11', '22', '33', '12', '13', '23')]
            else:
                row += [C('-')] * 6
        out.append(row)
    cx, cy, cz = J['coord_heads']
    uh = C(R('U', 'i'), R('eq', 'sub'), '/', R('U', 'i'), R('iso', 'sub')) if J['u_head'] == 'Ueq/Uiso' else C(R('U', 'i'), R('eq', 'sub'))
    head = [C(J['atom_head']), C('s.o.'), R(cx, 'i'), R(cy, 'i'), R(cz, 'i'), uh]
    if have_aniso:
        head += [C(R('U', 'i'), R(k, 'sup')) for k in ('11', '22', '33', '12', '13', '23')]
    note = 's.o. = site occupancy.'
    if have_aniso:
        note += ' Isotropic atoms show Uiso.' if J['u_head'] == 'Ueq/Uiso' else ' * = Uiso (isotropic atom).'
        if J['u_head'] != 'Ueq/Uiso':
            for row, lab in zip(out, [r[0] for r in out]):
                if T_plain(row[6]) == '-':
                    row[5] = C(T_plain(row[5]), '*')
    return {'caption': _title(J, J['coords_caption'].format(name=st.name)), 'head': head, 'rows': out, 'note': note}

def _bond_source(st, P):
    """[(cation label, anion label, distance string, code)] from the _geom_bond loop (esds) or
    computed (no esd), cation–anion pairs only, H excluded."""
    cat = {lab: s for s in st.cations if s.element != 'H' for lab in s.label.split('/')}
    an = {lab: s for s in st.anions for lab in s.label.split('/')}
    an.update({s.label: s for s in st.anions})
    tags, rows = B._loop(st.block, '_geom_bond_distance')
    out = []
    if rows:
        l1 = B._col(tags, rows, '_geom_bond_atom_site_label_1'); l2 = B._col(tags, rows, '_geom_bond_atom_site_label_2')
        dd = B._col(tags, rows, '_geom_bond_distance'); sy = B._col(tags, rows, '_geom_bond_site_symmetry_2', '.')
        for a, b_, d, s in zip(l1, l2, dd, sy):
            if a in cat and b_ in an:
                out.append((a, b_, d.strip(), s.strip()))
            elif b_ in cat and a in an:
                # the loop lists the anion first, so its symmetry code moves the CATION; the table
                # puts the code on the anion, which needs the inverse operator
                out.append((b_, a, d.strip(), _inverse_code(st, s.strip())))
        # only bonds a paper would list: within the cutoff AND worth >= MIN_S vu (SHELXL's BOND
        # list can reach further)
        keep = []
        for a, b_, d, s in out:
            site = cat[a]; cut = max(B.CUTOFF.get(sp.element, 3.2) for sp in site.species)
            dv = B._num(d) or 0
            if dv > cut + 1e-6:
                continue
            asite = an[b_]
            lim = [P.max_length(sp.element, sp.ox, asite.element, B.ANION_OX.get(asite.element, -2))
                   for sp in site.species if sp.ox and sp.ox > 0]
            lim = [x for x in lim if x is not None]
            if lim and dv > max(lim) + 1e-6:
                continue
            keep.append((a, b_, d, s))
        return keep, True
    # computed: keep every symmetry-equivalent bond with its own code
    for site in st.cations:
        if site.element == 'H':
            continue
        cut = max(B.CUTOFF.get(sp.element, 3.2) for sp in site.species)
        for other, d, code in _neighbours_with_codes(st, site, cut):
            if st.is_cation(other):
                continue
            lim = [P.max_length(sp.element, sp.ox, other.element, B.ANION_OX.get(other.element, -2))
                   for sp in site.species if sp.ox and sp.ox > 0]
            lim = [x for x in lim if x is not None]
            if lim and d > max(lim) + 1e-6:
                continue
            out.append((site.label, other.label, '%.3f' % d, code))
    return out, False

def _inverse_code(st, code):
    """The symmetry code that moves the OTHER atom of a bond: for a code 'n_abc' (operator n, lattice
    translation abc) return the code of the inverse operation, so 'O1 Ca1 d n_abc' can be printed
    as Ca1–O1^(inverse). Identity stays identity; an unparseable code is returned as is."""
    code = (code or '.').strip()
    if code in ('.', '', '?', '1_555'):
        return '.'
    m = re.match(r'^(\d+)(?:_(\d)(\d)(\d))?$', code)
    if not m or not (1 <= int(m.group(1)) <= len(st.ops)):
        return code
    n = int(m.group(1)); t = [int(m.group(k)) - 5 for k in (2, 3, 4)] if m.group(2) else [0, 0, 0]
    rot, tr = st.ops[n - 1]
    tr = [tr[i] + t[i] for i in range(3)]
    # inverse rotation via the adjugate (det = ±1 for a symmetry operation)
    a = rot
    det = (a[0][0] * (a[1][1] * a[2][2] - a[1][2] * a[2][1]) - a[0][1] * (a[1][0] * a[2][2] - a[1][2] * a[2][0])
           + a[0][2] * (a[1][0] * a[2][1] - a[1][1] * a[2][0]))
    if abs(abs(det) - 1) > 1e-6:
        return code
    cof = [[(a[(j + 1) % 3][(i + 1) % 3] * a[(j + 2) % 3][(i + 2) % 3] - a[(j + 1) % 3][(i + 2) % 3] * a[(j + 2) % 3][(i + 1) % 3]) / det
            for j in range(3)] for i in range(3)]
    inv = [[round(cof[i][j]) for j in range(3)] for i in range(3)]
    tr_inv = [-sum(inv[i][j] * tr[j] for j in range(3)) for i in range(3)]
    # which listed operator has this rotation? its translation differs from tr_inv by a lattice vector
    for k, (r2, t2) in enumerate(st.ops):
        if all(abs(r2[i][j] - inv[i][j]) < 1e-6 for i in range(3) for j in range(3)):
            lat = [tr_inv[i] - t2[i] for i in range(3)]
            if all(abs(v - round(v)) < 1e-4 for v in lat):
                lat = [int(round(v)) for v in lat]
                if k == 0 and lat == [0, 0, 0]:
                    return '.'
                return '%d_%d%d%d' % (k + 1, 5 + lat[0], 5 + lat[1], 5 + lat[2])
    return code

def _neighbours_with_codes(st, site, cutoff):
    """[(other site, distance, code)] — every neighbour with the symmetry code (op_translation)
    that generates it from the other site's first equivalent."""
    p0 = site.positions[0]
    rng = [int(math.ceil(cutoff / w)) + 1 for w in st.widths]
    found = []
    for other in st.sites:
        for n, op in enumerate(st.ops):
            q = B._apply(op, other.frac)
            for i in range(-rng[0], rng[0] + 1):
                for j in range(-rng[1], rng[1] + 1):
                    for k in range(-rng[2], rng[2] + 1):
                        qq = [q[0] + i, q[1] + j, q[2] + k]
                        d = st.dist(p0, qq)
                        if 0.3 < d <= cutoff:
                            # translation relative to the operator's own image: the code digits
                            t = [int(math.floor(qq[m] - (q[m] % 1.0) + 1e-9)) + 0 for m in range(3)]
                            tt = [int(round(qq[m] - (B._apply(op, other.frac)[m]))) for m in range(3)]
                            code = '%d_%d%d%d' % (n + 1, 5 + tt[0], 5 + tt[1], 5 + tt[2])
                            found.append((other, d, code, tuple(round(v, 4) for v in qq)))
    # dedupe identical positions reached by different operators (special positions)
    seen = set(); out = []
    for other, d, code, pos in sorted(found, key=lambda x: (x[0].label, x[1], x[2])):
        key = (other.label, pos)
        if key in seen:
            continue
        seen.add(key); out.append((other, d, code))
    return out

def bond_table(st, codes, P, J=None):
    J = J or journal(None)
    codes.begin_table()
    src, from_loop = _bond_source(st, P)
    blocks = []
    for site in st.cations:
        if site.element == 'H':
            continue
        mine = [x for x in src if x[0] == site.label or x[0] in site.label.split('/')]
        if not mine:
            continue
        rows = []; seen = set()
        for cat, an, d, code in mine:
            key = (an, d, code)                  # a mixed site (Ca1/Na1) lists each bond once, not per species
            if key in seen:
                continue
            seen.add(key)
            rows.append((_bond_cell(site.label, an, codes.index(code)), C(d), B._num(d), an))
        # means: uranyl gets <U–Oyl> (the two shortest) and <U–Oeq>
        anion_elems = sorted({st.sites[[s.label for s in st.sites].index(an)].element if an in [s.label for s in st.sites]
                              else re.sub(r'\d.*$', '', an) for _, _, _, an in rows}, key=lambda e: (e != 'O', e))
        xlab = '/'.join(anion_elems) if anion_elems else 'X'
        means = []
        vals = [v for _, _, v, _ in rows if v is not None]
        # '<Ca–O>' when Ca has one site, '<Ca1–O>' when it has several
        el_sites = [s for s in st.cations if s.element == site.element and s.element != 'H']
        catlab = site.label if len(el_sites) > 1 or '/' in site.label else re.sub(r'\d+$', '', site.label)
        if site.element == 'U' and any(sp.ox == 6 for sp in site.species) and len(vals) >= 4:
            sv = sorted(vals)
            yl, eq = sv[:2], sv[2:]
            means.append((C('<', R(catlab + '–O', 'i'), R('yl', 'i', 'sub'), '>'), C('%.3f' % (sum(yl) / len(yl)))))
            means.append((C('<', R(catlab + '–O', 'i'), R('eq', 'i', 'sub'), '>'), C('%.3f' % (sum(eq) / len(eq)))))
        elif vals:
            means.append((C('<', R('%s–%s' % (catlab, xlab), 'i'), '>'), C('%.3f' % (sum(vals) / len(vals)))))
        blocks.append({'cation': site.label, 'rows': [(a, b_) for a, b_, _, _ in rows], 'means': means})
    # lay the blocks out in three columns: each block goes to the shortest column, in order
    cols = [[], [], []]
    for blk in blocks:
        col = min(range(3), key=lambda k: len(cols[k]))
        if cols[col]:
            cols[col].append(None)                       # a blank row between blocks
        cols[col].extend(blk['rows'] + blk['means'])
    n = max((len(c) for c in cols), default=0)
    rows = []
    for i in range(n):
        row = []
        for c in cols:
            cell = c[i] if i < len(c) else None
            row += [cell[0], cell[1]] if cell else [C(''), C('')]
        rows.append(row)
    return {'caption': _title(J, J['bonds_caption'].format(name=st.name)), 'head': None, 'rows': rows,
            'note': (codes.note(J['symcodes']) + ('' if from_loop else ' Distances computed from the .cif coordinates (no esds: the .cif has no _geom_bond loop).')).strip(),
            'blocks': blocks}

def _cart_matrix(st):
    a, b, c, al, be, ga = st.cell
    ca, cb, cg = (math.cos(math.radians(x)) for x in (al, be, ga)); sg = math.sin(math.radians(ga))
    v = st.volume / (a * b * c)
    return [[a, b * cg, c * cb], [0, b * sg, c * (ca - cb * cg) / sg], [0, 0, c * v / sg]]

def _angle(st, p_d, p_h, p_a):
    M = _cart_matrix(st)
    def cart(p): return [sum(M[i][j] * p[j] for j in range(3)) for i in range(3)]
    d, h, a = cart(p_d), cart(p_h), cart(p_a)
    v1 = [d[i] - h[i] for i in range(3)]; v2 = [a[i] - h[i] for i in range(3)]
    n1 = math.sqrt(sum(x * x for x in v1)); n2 = math.sqrt(sum(x * x for x in v2))
    if not n1 or not n2:
        return 0.0
    return math.degrees(math.acos(max(-1.0, min(1.0, sum(v1[i] * v2[i] for i in range(3)) / (n1 * n2)))))

def hbond_table(st, codes, J=None):
    J = J or journal(None)
    codes.begin_table()
    tags, rows = B._loop(st.block, '_geom_hbond_distance_ha')
    out = []
    if rows:
        g = lambda t, d='': B._col(tags, rows, t, d)
        for D, H, A, dh, ha, da, ang, sy in zip(g('_geom_hbond_atom_site_label_d'), g('_geom_hbond_atom_site_label_h'), g('_geom_hbond_atom_site_label_a'),
                                              g('_geom_hbond_distance_dh'), g('_geom_hbond_distance_ha'), g('_geom_hbond_distance_da'),
                                              g('_geom_hbond_angle_dha'), g('_geom_hbond_site_symmetry_a', '.')):
            out.append([C(D, '–', H, '⋯', A, R(str(codes.index(sy)), 'sup') if codes.index(sy) else None),
                        C(dh.strip()), C(ha.strip()), C(da.strip()), C(re.sub(r'\.000\(0\)$|\.0\(0\)$', '', ang.strip()))])
        note = ''
    else:
        # computed from the H positions
        hs = [s for s in st.sites if s.element == 'H']
        donors = [s for s in st.sites if s.element in ('O', 'F', 'N')]
        for h in hs:
            near = [(o, d, code) for o, d, code in _neighbours_with_codes(st, h, 2.6) if o.element in ('O', 'F', 'N', 'Cl')]
            if not near:
                continue
            near.sort(key=lambda x: x[1])
            D, d_dh, _ = near[0]
            if d_dh > 1.25:
                continue
            p_h = h.positions[0]
            # the donor's actual position next to this H
            for o, dist, code in near[1:]:
                p_a = _pos_from_code(st, o, code)
                p_d = _pos_from_code(st, D, near[0][2])
                ang = _angle(st, p_d, p_h, p_a)
                if ang < 110:
                    continue
                da = st.dist(p_d, p_a)
                out.append([C(D.label, '–', h.label, '⋯', o.label, R(str(codes.index(code)), 'sup') if codes.index(code) else None),
                            C('%.2f' % d_dh), C('%.2f' % dist), C('%.3f' % da), C('%.0f' % ang)])
        note = 'Computed from the refined H positions (no esds: the .cif has no _geom_hbond loop); H⋯A ≤ 2.6 Å, ∠DHA ≥ 110°.' if out else ''
    if not out:
        return None
    head = [C(R('D', 'i'), '–', R('H', 'i'), '⋯', R('A', 'i')), C(R('D', 'i'), '–', R('H', 'i')), C(R('H', 'i'), '⋯', R('A', 'i')), C(R('D', 'i'), '⋯', R('A', 'i')), C('∠', R('DHA', 'i'))]
    return {'caption': _title(J, J['hbonds_caption'].format(name=st.name)), 'head': head, 'rows': out, 'note': (codes.note(J['symcodes']) + ' ' + note).strip()}

def _pos_from_code(st, site, code):
    """The site's position under a symmetry code: 'n_klm', a bare 'n', or '.' (identity)."""
    code = (code or '.').strip()
    m = re.match(r'^(\d+)(?:_(\d)(\d)(\d))?$', code)
    if not m or not (1 <= int(m.group(1)) <= len(st.ops)):
        return list(site.frac)
    n = int(m.group(1)); t = [int(m.group(k)) - 5 for k in (2, 3, 4)] if m.group(2) else [0, 0, 0]
    q = B._apply(st.ops[n - 1], site.frac)
    return [q[i] + t[i] for i in range(3)]

def bvs_table(st, params, result, anion_sum, cells, J=None):
    J = J or journal(None)
    cats = [r[0] for r in result if r[0].element != 'H']
    hres = [r for r in result if r[0].element == 'H']
    bvs_of = {r[0].label: r[2] for r in result}
    # hydrogen bonds accepted per anion: donor label, valence, count across
    accepted = {}
    for h, bonds, *_ in hres:
        if not bonds:
            continue
        donor = min(bonds, key=lambda b: b.dist)
        for b in bonds:
            if b is donor:
                continue
            s = b.vals[0][1] or 0.0
            n_across = b.count * h.mult / b.anion.mult
            n_across = int(round(n_across)) if abs(n_across - round(n_across)) < 0.02 else 1
            accepted.setdefault(b.anion.label, []).append((donor.anion.label, s * min(h.occ_total, 1.0), n_across))
    has_h = bool(hres)
    head = [C(J['atom_head'] if J['atom_head'] != 'Atoms' else 'Atom')] + [C(ct.label) for ct in cats]
    if has_h:
        head += [C('Donor'), C('vu'), C('H bond')]
    sum_cat, sum_an = J['sum_labels']
    def sumcell(lab):
        return C('Σ', R(lab[1:], 'sub')) if lab.startswith('Σ') else C(lab)
    head += [sumcell(sum_an)]
    rows = []
    for an in st.anions:
        if not any((an.label, ct.label) in cells for ct in cats) and an.label not in accepted:
            continue
        row = [C(an.label)]; total = 0.0
        for ct in cats:
            vals = cells.get((an.label, ct.label))
            if not vals:
                row.append(C('-')); continue
            occ = min(ct.occ_total, 1.0)
            parts = []
            for s, nd, na in vals:
                parts.append('%.2f%s' % (s, B._mark(nd, na)))
                total += s * float(na) * occ         # the same weighting as bv_check's anion sums
            row.append(C(', '.join(parts)))
        if has_h:
            acc = accepted.get(an.label, [])
            if acc:
                row += [C(', '.join(d for d, _, _ in acc)), C(', '.join('%.2f' % s for _, s, _ in acc)),
                        C(', '.join(('×%d →' % n) if n > 1 else '-' for _, _, n in acc))]
                total += sum(s * n for _, s, n in acc)
            else:
                row += [C('-'), C('-'), C('-')]
        row.append(C('%.2f' % total))
        rows.append(row)
    last = [sumcell(sum_cat)] + [C('%.2f' % bvs_of[ct.label]) for ct in cats]
    if has_h:
        last += [C(''), C(''), C('')]
    last.append(C(''))
    rows.append(last)
    refs = OrderedDict()
    for (cat, cox, an), (r0, b, rid) in params.used.items():
        if cat != 'H':
            refs[params.refs.get(rid, rid)] = True
    note = 'Bond-valence parameters from %s' % '; '.join(refs) if refs else ''
    if has_h:
        note += '; hydrogen-bond valences from the H⋯A distances (Brown 2002), Σan includes the accepted hydrogen bonds'
    return {'caption': _title(J, J['bvs_caption'].format(name=st.name)), 'head': head, 'rows': rows, 'note': note + '.' if note else ''}

# ----------------------------------------------------------------------------- build + render

def build(cif, params='gh', ox=None, cutoff=None, include_h=True, journal_key=None):
    st = B.Structure(cif, B._parse_ox(ox), include_h=include_h)
    P = B.Params(prefer=params)
    J = journal(journal_key)
    result, anion_sum, cells = B.compute(st, P, cutoff)
    codes = SymCodes(st)
    tabs = [('coords', coordinates_table(st, J)), ('bonds', bond_table(st, codes, P, J))]
    hb = hbond_table(st, codes, J) if include_h else None
    if hb:
        tabs.append(('hbonds', hb))
    tabs.append(('bvs', bvs_table(st, P, result, anion_sum, cells, J)))
    for i, (k, t) in enumerate(tabs, 1):
        t['n'] = i
        t['label'] = J['caption'].format(n=i)
        if t.get('note') and J['notes_prefix'] and not t['note'].startswith(J['notes_prefix']):
            t['note'] = J['notes_prefix'] + t['note']
        t['journal'] = J
    return st, tabs

def render_text(tabs):
    L = []
    for k, t in tabs:
        L.append('%s %s' % (t.get('label', 'Table %d.' % t['n']), t['caption']))
        rows = ([t['head']] if t.get('head') else []) + t['rows']
        txt = [[plain(c) for c in r] for r in rows]
        if not txt:
            L.append('  (nothing to list)')
        else:
            w = [max(len(r[i]) if i < len(r) else 0 for r in txt) for i in range(max(len(r) for r in txt))]
            for r in txt:
                L.append('  ' + '  '.join(x.ljust(w[i]) for i, x in enumerate(r)).rstrip())
        if t.get('note'):
            L.append('  ' + t['note'])
        L.append('')
    return '\n'.join(L)

def render_html(tabs):
    out = []
    for k, t in tabs:
        J = t.get('journal') or journal(None)
        out.append('<div class="pubtable" data-kind="%s"%s><div class="pubcap"><b>%s</b> %s</div><table>'
                   % (k, ' style="font-family:%s"' % J['font'] if J.get('font') else '', html(C(t.get('label', 'Table %d.' % t['n']))), html(C(t['caption']))))
        if t.get('head'):
            out.append('<thead><tr>' + ''.join('<th>%s</th>' % html(c) for c in t['head']) + '</tr></thead>')
        out.append('<tbody>')
        for r in t['rows']:
            out.append('<tr>' + ''.join('<td>%s</td>' % html(c) for c in r) + '</tr>')
        out.append('</tbody></table>')
        if t.get('note'):
            out.append('<div class="pubnote">%s</div>' % html(C(t['note'])))
        out.append('</div>')
    return '\n'.join(out)

def write_word(st, tabs, path):
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from pxrd_review.annotate_review import _save_docx
    doc = Document()
    J = (tabs[0][1].get('journal') if tabs else None) or journal(None)
    style = doc.styles['Normal']; style.font.size = Pt(10)
    if J.get('font'):
        style.font.name = J['font']
    def runs(par, cell, size=9):
        for text, st_ in cell:
            r = par.add_run(text); r.font.size = Pt(size)
            if J.get('font'): r.font.name = J['font']
            if 'i' in st_: r.font.italic = True
            if 'b' in st_: r.font.bold = True
            if 'sup' in st_: r.font.superscript = True
            if 'sub' in st_: r.font.subscript = True
    def border(cell, side):
        tcPr = cell._tc.get_or_add_tcPr()
        b = tcPr.find(qn('w:tcBorders'))
        if b is None:
            b = OxmlElement('w:tcBorders'); tcPr.append(b)
        e = OxmlElement('w:' + side); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '6'); e.set(qn('w:color'), '000000'); b.append(e)
    for k, t in tabs:
        cap = doc.add_paragraph(); r = cap.add_run(t.get('label', 'Table %d.' % t['n']) + ' '); r.font.bold = True
        r2 = cap.add_run(t['caption'])
        if J.get('font'):
            r.font.name = r2.font.name = J['font']
        rows = ([t['head']] if t.get('head') else []) + t['rows']
        if not rows:
            doc.add_paragraph('(nothing to list)'); continue
        ncol = max(len(r_) for r_ in rows)
        tb = doc.add_table(rows=len(rows), cols=ncol)
        for i, r_ in enumerate(rows):
            for j in range(ncol):
                cell = tb.cell(i, j); par = cell.paragraphs[0]
                runs(par, r_[j] if j < len(r_) else C(''))
                if i == 0: border(cell, 'top')
                if i == 0 and t.get('head'): border(cell, 'bottom')
                if i == len(rows) - 1: border(cell, 'bottom')
        if t.get('note'):
            p = doc.add_paragraph(); r = p.add_run(t['note']); r.font.size = Pt(8)
        doc.add_paragraph('')
    _save_docx(doc, path)

def run(cif, word=False, params='gh', ox=None, cutoff=None, include_h=True, out_dir=None, quiet=False, journal_key=None):
    st, tabs = build(cif, params, ox, cutoff, include_h, journal_key)
    text = render_text(tabs)
    if not quiet:
        print(text)
    out_dir = out_dir or os.path.join(os.path.dirname(os.path.abspath(cif)), 'review_out')
    os.makedirs(out_dir, exist_ok=True)
    stem = os.path.splitext(os.path.basename(cif))[0]
    with open(os.path.join(out_dir, stem + '_tables.txt'), 'w', encoding='utf-8') as f:
        f.write(text)
    if word:
        wp = os.path.join(out_dir, stem + '_tables.docx')
        write_word(st, tabs, wp)
        if not quiet:
            print('  tables → %s' % wp)
    return st, tabs, text

def main(argv=None):
    ap = argparse.ArgumentParser(prog='pxrd tables', description=__doc__.split('\n\n')[1], formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument('cif')
    ap.add_argument('--word', action='store_true', help='write review_out/<name>_tables.docx')
    ap.add_argument('--params', default='gh', choices=['gh', 'bo', 'ba'])
    ap.add_argument('--ox'); ap.add_argument('--cutoff', type=float); ap.add_argument('--no-h', action='store_true'); ap.add_argument('--out')
    ap.add_argument('--journal', default=DEFAULT_JOURNAL, choices=sorted(JOURNALS),
                    help='table style: ammin, minmag, cjmp (Canadian Journal of Mineralogy and Petrology), ejm, or manuscript (default)')
    a = ap.parse_args(argv)
    try:
        run(a.cif, a.word, a.params, a.ox, a.cutoff, not a.no_h, a.out, journal_key=a.journal)
    except ValueError as e:
        raise SystemExit('tables: %s' % e)
    return 0

if __name__ == '__main__':
    sys.exit(main())
