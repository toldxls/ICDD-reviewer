"""Unit tests for pxrd_review.pxrd_table — observed / calculated line matching and the table.

    python3 -m unittest tests.test_pxrd_table -v
"""
import os, shutil, tempfile, unittest

from pxrd_review import pxrd_table as P

OBS = """spanoite-like
#\tAngle\td(Å)\tI%(f)\t( h k l)\tθ(°)
1\t6.15\t6.609\t22.5\t( 1 0 0)\t3.08
2\t6.28\t6.479\t5.0\t( 0 1 1)\t3.14
3\t7.82\t5.199\t0.1\t( 1 1 0)\t3.91
4\t12.5\t3.304\t86.9\t( 2 0 0)\t6.2
5\t13.0\t3.150\t43.8\t(-1 1 3)\t6.5
6\t14.0\t2.732\t4.0\t( 1 2 2)\t7.0
7\t15.0\t2.600\t4.0\t\t7.5
8\t16.0\t2.172\t4.0\t( 0 2 4)\t8.0
9\t16.0\t2.172\t4.0\t(-3 1 2)\t8.0
"""
CALC = """spanoite-like
#\t( h k l)\tAngle\td(Å)\tI%(f)\t| F |\tm
1\t( 1 0 0)\t6.30\t6.648\t27.22\t170\t2
2\t( 0 1 1)\t6.41\t6.534\t5.90\t61\t4
3\t( 1 1 0)\t7.95\t5.234\t0.88\t27\t4
4\t(-1 1 1)\t8.13\t5.117\t6.67\t75\t4
5\t( 1 1 2)\t12.4\t3.325\t53.0\t200\t4
6\t( 2 0 0)\t12.4\t3.324\t50.0\t200\t2
7\t(-1 1 3)\t13.0\t3.179\t71.1\t250\t4
8\t( 1 2 2)\t14.0\t2.752\t2.2\t30\t4
9\t( 9 9 9)\t15.0\t2.605\t9.0\t80\t4
10\t( 0 2 4)\t16.0\t2.190\t1.0\t20\t4
11\t(-3 1 2)\t16.0\t2.180\t2.0\t25\t4
12\t( 5 5 5)\t30.0\t1.300\t6.0\t70\t4
"""


class Matching(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp(prefix='pxrd_')
        self.o = os.path.join(self.tmp, 'obs.txt'); self.c = os.path.join(self.tmp, 'calc.txt')
        open(self.o, 'w').write(OBS); open(self.c, 'w').write(CALC)
        self.obs = P.load_lines(self.o); self.calc = P.load_lines(self.c)

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)

    def test_readers(self):
        self.assertEqual(len(self.obs), 9); self.assertEqual(self.obs[0].hkl, (1, 0, 0)); self.assertIsNone(self.obs[6].hkl)
        self.assertEqual(self.calc[6].hkl, (-1, 1, 3)); self.assertAlmostEqual(self.calc[6].d, 3.179)
        bare = os.path.join(self.tmp, 'bare.txt'); open(bare, 'w').write('6.609 22.5\n3.304 86.9\n')
        b = P.load_lines(bare); self.assertEqual(len(b), 2); self.assertAlmostEqual(b[1].I, 86.9)
        tt = os.path.join(self.tmp, 'tt.txt'); open(tt, 'w').write('2theta I\n26.64 100\n')
        t = P.load_lines(tt, wavelength=1.5406); self.assertAlmostEqual(t[0].d, 3.343, places=3)

    def test_rules(self):
        rows = P.match(self.obs, self.calc, tol_pct=1.2, min_i=3.5)
        by = {}
        for r in rows:
            by.setdefault(r.hkl, []).append(r)
        # hkl match with the observed intensity carried
        self.assertEqual((round(by[(1, 0, 0)][0].Iobs), by[(1, 0, 0)][0].Icalc), (22, 27.22))
        # an unobserved reflection inside an observed peak is attached (Iobs/dobs repeated) — 1 1 2 under 3.304
        r112 = by[(1, 1, 2)][0]; self.assertAlmostEqual(r112.dobs, 3.304); self.assertAlmostEqual(r112.Iobs, 86.9)
        # a calculated line far from any peak stands alone; below threshold it is dropped
        self.assertIsNone(by[(-1, 1, 1)][0].Iobs); self.assertNotIn((1, 1, 0), by)
        # an observed line without hkl takes the nearest calculated line within tol
        self.assertEqual(by[(9, 9, 9)][0].Iobs, 4.0)
        # a peak whose only reflection is weak keeps the line with blank calc
        r122 = by[(1, 2, 2)][0]; self.assertIsNone(r122.dcalc); self.assertAlmostEqual(r122.dobs, 2.732)
        # a peak with several weak reflections keeps ONE row (the strongest)
        self.assertIn((-3, 1, 2), by); self.assertNotIn((0, 2, 4), by)
        # dmin cuts the tail
        self.assertNotIn((5, 5, 5), {r.hkl for r in P.match(self.obs, self.calc, dmin=1.45)})
        # order: decreasing d, a peak's rows together
        ds = [(r.dobs if r.dobs is not None else r.dcalc) for r in rows]
        self.assertEqual(ds, sorted(ds, reverse=True))

    def test_table_bold_blocks_and_outputs(self):
        rows = P.match(self.obs, self.calc)
        top = P.strongest(rows, 2)
        self.assertEqual({round(r.Iobs) for r in rows if r.obs_id in top}, {87, 44})
        t = P.build_table(rows, 'test', blocks=2, bold_n=2)
        self.assertEqual(len(t['head']), 11)                                  # 5 + spacer + 5
        from pxrd_review import tables as T
        text = T.render_text([('pxrd', t)])
        self.assertIn('Powder X-ray diffraction data', text)
        bold = [c for r in t['rows'] for c in r for txt, st in c if 'b' in st]
        # Iobs + dobs bold on EVERY row of the 2 strongest peaks (the 87 peak carries two reflections)
        self.assertEqual(len(bold), 2 * len([r for r in rows if r.obs_id in top])); self.assertEqual(len(bold), 6)
        out = os.path.join(self.tmp, 'o.docx'); T.write_word(None, [('pxrd', t)], out)
        from docx import Document
        d = Document(out); self.assertEqual(len(d.tables), 1)
        self.assertTrue(any(run.font.bold for row in d.tables[0].rows for c in row.cells for p in c.paragraphs for run in p.runs))
        x = P.write_xlsx(self.obs, self.calc, rows, os.path.join(self.tmp, 'o.xlsx'))
        import openpyxl
        self.assertEqual(openpyxl.load_workbook(x).sheetnames, ['matched', 'obs', 'calc'])


if __name__ == '__main__':
    unittest.main()
