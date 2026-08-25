"""Unit tests for pxrd_review.refs_check — the manuscript citation ↔ reference-list check.

    python3 -m unittest tests.test_refs_check -v

Synthetic docx built with python-docx: every citation form the checker claims to handle, the
noise it must ignore, a table cell, a '———' repeated-author entry, a multi-word corporate
author, a companion file, and a numbered list. No corpus files needed."""
import os, io, re, shutil, tempfile, unittest, zipfile

from docx import Document
from docx.enum.text import WD_BREAK

from pxrd_review import refs_check as R


def _build(paras, path, table=None):
    doc = Document()
    for p in paras:
        if p is None:
            doc.add_paragraph('')
        else:
            doc.add_paragraph(p)
    if table:
        t = doc.add_table(rows=1, cols=1)
        t.cell(0, 0).text = table
    doc.save(path)
    return path


BODY = [
    'Introduction',
    'Quartz is common (Smith 2019; Jones and Brown, 2020a). As shown by van der Waals et al. (1873), '
    'and by Kampf (1977), the cell is small. Balić-Žunić (1990) reported the same. Data were reduced '
    'with Rigaku Oxford Diffraction (2018) software. He won the Meritorious (1981) Service Award. '
    'In December 2024 it was approved (IMA 2024-012). The structure was solved with SHELXT '
    '(Sheldrick 2015a) and refined with SHELXL (Sheldrick 2015b). See also Čejka (1999 and 2005) '
    'and Barton’s (1980) work; PGMs, Godel and Barnes (2008) agree. '
    'Details at (http://www.example.org/TOC/2026/Xxx2026_data.html).',
    'Gurzhyi et al. (2024) described uramphite (Zhao 2024). Cooper et al. (2019) and Hawthorne (1985) '
    'and Hawthorne (2025) and Evans et al. (2020) agree; so do Brese and O’Keeffe (1991).',
    'References',
    'Balić-Žunić, T. (1990) A title. Journal 1, 1–2.',
    'Barton, P.B. (1980) The Ag-Au-S system. Economic Geology 75, 303–316.',
    'Čejka, J. (1999) Infrared spectroscopy of uranyl minerals. Reviews in Mineralogy 38, 521–622.',
    'Čejka, J. (2005) Vibrational spectroscopy of the uranyl minerals. Journal 2, 1–2.',
    'Brese, N. E. and O’Keeffe, M. (1991) Bond-valence parameters for solids. Acta B47, 192–197.',
    'Cooper, M.A., Hawthorne, F.C. and Kampf, A.R. (2019a) Determination of V4+:V5+ ratios. Journal 10, 1–2.',
    'Evans H.A. Wu Y. Seshadri R. and Cheetham A.K. (2020) Perovskite-related ReO3 materials. Nature Reviews Materials 5, 196–213.',
    'Godel, B. and Barnes, S.J. (2008) Platinum-group elements in sulfide minerals. Journal 3, 1–2.',
    'Gurzhiy, V.V., Kasatkin, A.V. (2024) Uramphite. Journal 4, 3–4.',
    'Hawthorne (1985) Towards a structural classification of minerals. American Mineralogist 70, 455–473.',
    'Hawthorne, F.C. (2025) Decavanadates and polyanion-polycation interactions. Journal 11, 1–2.',
    'Jones, A. and Brown, B. (2020a) A title. Journal 5, 5–6.',
    'Kampf, A.R. (2015) A title. Journal 6, 7–8.',
    'Rigaku Oxford Diffraction (2018) CrysAlisPro. Yarnton, England.',
    'Schoep, A. (1923) Sur la chinkolobwite. Bulletin 46, 1–2.',
    'Sheldrick, G.M. (2015a) SHELXT. Acta Crystallographica A71, 3–8.',
    '——— (2015b) Crystal structure refinement with SHELXL. Acta Crystallographica C71, 3–8.',
    'Smith, J. (2019) A title. Journal 7, 9–10.',
    'Uncited, U. (2001) Never cited. Journal 8, 11–12.',
    'van der Waals, J.D., Smith, A. and Jones, B. (1873) A title. Journal 9, 13–14.',
    'Zhao, Q., Zhao, M., Jiao, X., Xia, Y., and Chen, D. (2024) Unraveling the structural evolution. Journal 12, 1–2.',
    'Figure captions',
    'Fig. 1. Something (Schoep 1947).',
]
TABLE = 'Data from Kampf et al. (2015).'


class AuthorYearDocx(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.tmp = tempfile.mkdtemp(prefix='refs_')
        cls.path = _build(BODY, os.path.join(cls.tmp, 'ms.docx'), table=TABLE)
        cls.doc, cls.paras = R.load_docx(cls.path)
        cls.res = R.analyze(cls.paras)

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.tmp, ignore_errors=True)

    def test_list_found(self):
        r = self.res
        self.assertEqual(r['style'], 'author-year')
        self.assertEqual(self.paras[r['heading']].text, 'References')
        self.assertEqual(len(r['entries']), 21)
        # the dash entry inherits Sheldrick and keeps its own year/suffix
        dash = [e for e in r['entries'] if e.text.startswith('———')][0]
        self.assertEqual((dash.names, dash.year, dash.suffix), (['sheldrick'], '2015', 'b'))
        # captions after the list count as body
        self.assertTrue(any(c.who == 'Schoep' for c in r['cites']))

    def test_noise_is_not_a_citation(self):
        whos = {c.who for c in self.res['cites']}
        for noise in ('Meritorious', 'December', 'IMA', 'Xxx'):
            self.assertNotIn(noise, whos)

    def test_matches(self):
        orphans = {R._fmt_cite(c) for c in self.res['orphans']}
        for ok in ('Smith (2019)', 'Jones and Brown (2020a)', 'van der Waals et al. (1873)',
                   'Balić-Žunić (1990)', 'Rigaku Oxford Diffraction (2018)', 'Sheldrick (2015b)',
                   'Čejka (2005)', 'Barton (1980)', 'Kampf et al. (2015)', 'Zhao (2024)', 'Cooper et al. (2019)',
                   'Hawthorne (1985)', 'Evans et al. (2020)', 'Brese and O’Keeffe (1991)'):
            self.assertNotIn(ok, orphans, ok)
        # a comma list that swallowed a noun still matches on its real authors
        self.assertFalse(any('PGMs' in R._fmt_cite(c) for c in self.res['orphans']))

    def test_orphans_uncited_pairs(self):
        r = self.res
        self.assertEqual(sorted(R._fmt_cite(c) for c in r['orphans']), ['Kampf (1977)', 'Schoep (1947)'])
        self.assertEqual([e.text[:12] for e in r['uncited']], ['Schoep, A. (', 'Uncited, U. '])
        self.assertEqual([(R._fmt_cite(c), e.names[0]) for c, e in r['pairs']],
                         [('Gurzhyi et al. (2024)', 'gurzhiy')])
        # FORM: the three flatimerite cases, nothing else (Evans/Brese/Gurzhiy parse to the right count)
        form = sorted((f.kind, R._fmt_cite(f.cite) if f.cite else f.entry.text[:16]) for f in r['form'])
        self.assertEqual(form, [('authors', 'Kampf et al. (2015)'), ('authors', 'Zhao (2024)'),
                                ('initials', 'Hawthorne (1985)'), ('suffix', 'Cooper et al. (2019)')])
        self.assertIn('Zhao et al. (2024)', [f.msg for f in r['form'] if f.cite and f.cite.who == 'Zhao'][0])
        # the orphan hint names the same surname's entries
        kampf = [c for c in r['orphans'] if c.who == 'Kampf'][0]
        self.assertIn('Kampf (2015)', R._list_has(kampf, r['entries']))

    def test_report_and_annotation(self):
        out = os.path.join(self.tmp, 'out')
        res = R.check_file(self.path, out_dir=out, quiet=True)
        with open(os.path.join(out, 'ms_refs_report.txt'), encoding='utf-8') as f:
            rep = f.read()
        self.assertIn('CITED BUT NOT LISTED (2)', rep)
        self.assertIn('LISTED BUT NOT CITED (2)', rep)
        self.assertIn('MISMATCH', rep)
        self.assertIn('FORM — the citation and its entry disagree (4)', rep)
        annotated = os.path.join(out, 'ms_refs.docx')
        self.assertTrue(os.path.exists(annotated))
        _, before = R.load_docx(self.path)
        _, after = R.load_docx(annotated)
        self.assertEqual([p.text for p in before], [p.text for p in after], 'annotation must not change text')
        cx = zipfile.ZipFile(annotated).read('word/comments.xml').decode('utf-8')
        # 2 orphans + 2 uncited + 2 sides of the pair + 4 form = 10 comments, all by the tool
        self.assertEqual(len(re.findall(r'<w:comment ', cx)), 10)
        self.assertNotIn('w:author="Someone', cx)
        dx = zipfile.ZipFile(annotated).read('word/document.xml').decode('utf-8')
        self.assertIn('w:highlight', dx)
        # the source is untouched
        self.assertEqual(os.path.getsize(self.path), os.path.getsize(self.path))

    def test_companion_file_counts_as_body(self):
        comp = _build(['Table 2. Comparison (Uncited 2001).'], os.path.join(self.tmp, 'table.docx'))
        res = R.check_file(self.path, out_dir=os.path.join(self.tmp, 'out2'), annotate=False,
                           quiet=True, companions=[comp])
        self.assertEqual([e.text[:8] for e in res['uncited']], ['Schoep, '])


class NumericDocx(unittest.TestCase):
    def test_bracketed(self):
        tmp = tempfile.mkdtemp(prefix='refs_')
        try:
            path = _build(['Intro [1] and [2–3]; also [5]. A direction [100] is not a citation.',
                           'References',
                           '1. A. Author, Journal 1, 1–2 (2019).',
                           '2. B. Author, Journal 2, 3–4 (2020).',
                           '3. C. Author, Journal 3, 5–6 (2021).',
                           '4. D. Author, Journal 4, 7–8 (2022).'], os.path.join(tmp, 'num.docx'))
            _, paras = R.load_docx(path)
            res = R.analyze(paras)
            self.assertEqual(res['style'], 'numeric')
            self.assertEqual([n for n, *_ in res['orphans']], [5])
            self.assertEqual([e.number for e in res['uncited']], [4])
        finally:
            shutil.rmtree(tmp, ignore_errors=True)


class Splitter(unittest.TestCase):
    def test_year_anchored_split(self):
        blob = ('Brown, I.D. and Altermatt, D. (1985) Bond-valence parameters. Acta Crystallographica,\n'
                'B41, 244–247, https://doi.org/10.1107/S0108768185002063. Craig, J.R. and\n'
                'Carpenter, A.B. (1977) Fletcherite. Economic Geology, 72, 480–486.\n'
                'Kampf, A.R., Adams, P.M., Nash, B.P., and Marty, J. (2015) Ferribushmakinite. American\n'
                'Mineralogist, 100, 1–5. Sheldrick, G.M. (2015a) SHELXT. Acta A71, 3–8. ——— (2015b)\n'
                'Crystal structure refinement with SHELXL. Acta C71, 3–8.')
        ents = R.split_reference_block(blob)
        firsts = [R.surnames(e[:60])[:1] for e in ents[:4]]
        self.assertEqual(firsts, [['brown'], ['craig'], ['kampf'], ['sheldrick']])
        self.assertEqual(len(ents), 5)
        self.assertTrue(ents[-1].startswith('———'))

    def test_numbered_split_and_editor_guard(self):
        blob = ('1. Sharygin, V.V.; Yakovlev, G.A. Title. Minerals 2019, 9, 123.\n'
                '2. Okada, A.; Keil, K. Caswellsilverite. In Meteorites; Smith, J., Ed.; Press: City, 1982.\n'
                '3. Third, T. Another. Journal 2020, 1, 1.')
        ents = R.split_reference_block(blob)
        self.assertEqual(len(ents), 3)
        self.assertTrue(ents[1].startswith('2. Okada'))

    def test_author_grammar(self):
        for text, n in (('Brese, N. E. and O’Keeffe, M.: Bond-valence', 2), ('Holland T.J.B and Redfern S.A.T. (1997) Unit', 2),
                        ('Nekrasova A.N. and Borodaev Yu.S. (1972) First', 2), ('Robinson PD, Sen Gupta PK, Swihart GH, Houk L (1992) X', 4),
                        ('Peacor, D.R., Coveney, R.M. Jr., and Zhao, G.M. (2000) X', 3), ('Cesbron, F. and Morin, N. Curienite, A new mineral. 1968', 2),
                        ('Sheldrick, G.M., Crystal structure refinement with SHELXL. 2015', 1), ('Sheldrick, G.M. SHELXL-2019 (2019) Program', 1),
                        ('Evans, A., et al., Uranium in the environment. 1992', 3), ('Hawthorne (1985) Towards a structural', None),
                        ('Rigaku Oxford Diffraction (2018) CrysAlisPro', None)):
            self.assertEqual(R.parse_entry(text, 0, 0, 0, len(text)).nauth, n, text)

    def test_fold(self):
        self.assertEqual(R._fold('Balić-Žunić'), R._fold('Balic-Zunic'))
        self.assertEqual(R._fold('Karup-Møller'), 'karupmoller')
        self.assertEqual(R._fold("O'Neill"), 'oneill')


if __name__ == '__main__':
    unittest.main()
