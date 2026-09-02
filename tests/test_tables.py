"""Unit tests for pxrd_review.tables — publishable tables from a .cif (rutile, no corpus needed).

    python3 -m unittest tests.test_tables -v
"""
import os, re, shutil, tempfile, unittest

from pxrd_review import tables as T
from pxrd_review import bv_check as B
from tests.test_bv_check import RUTILE, HYDRATE, HYDRATE_H, _write


class Cells(unittest.TestCase):
    def test_rich_cells(self):
        c = T.C('Ca1–O3', T.R('3', 'sup'))
        self.assertEqual(T.plain(c), 'Ca1–O3³')
        self.assertEqual(T.html(c), 'Ca1–O3<sup>3</sup>')
        self.assertEqual(T.plain(T._elem_occ([B.Species('F', -1, 0.8), B.Species('O', -2, 0.2)])), 'F₀.₈₀O₀.₂₀')
        self.assertEqual(T.html(T.C(T.R('U', 'i'), T.R('eq', 'sub'))), '<i>U</i><sub>eq</sub>')

    def test_coord_and_symop_formatting(self):
        for s, want in (('0.666667', '⅔'), ('0.3333', '⅓'), ('0.5', '½'), ('0.25', '¼'), ('0', '0'), ('0.30479(5)', '0.30479(5)'), ('0.5000(2)', '0.5000(2)')):
            self.assertEqual(T.coord(s), want, s)
        rot, tr = B.parse_symop('-y+1/2, x+1/2, z+1/2')
        self.assertEqual(T.symop_string(rot, tr), '−y+½, x+½, z+½')
        rot, tr = B.parse_symop('x-y, x, -z')
        self.assertEqual(T.symop_string(rot, [tr[0] + 1, tr[1] - 1, tr[2] + 4 / 3]), 'x−y+1, x−1, −z+4/3')


class Rutile(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.tmp = tempfile.mkdtemp(prefix='tbl_')
        cls.cif = _write(cls.tmp, 'rutile.cif', RUTILE)
        cls.st, cls.tabs = T.build(cls.cif, params='bo')

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.tmp, ignore_errors=True)

    def test_tables_built(self):
        kinds = [k for k, _ in self.tabs]
        self.assertEqual(kinds, ['coords', 'bonds', 'bvs'])         # no H: no hydrogen-bond table
        coords = dict(self.tabs)['coords']
        self.assertEqual([T.plain(c) for c in coords['head']], ['Atoms', 's.o.', 'x', 'y', 'z', 'Ueq/Uiso'])  # no aniso loop: no U columns
        self.assertEqual([T.plain(c) for c in coords['rows'][0]], ['Ti1', 'Ti₁.₀₀', '0', '0', '0', '-'])
        self.assertEqual([T.plain(c) for c in coords['rows'][1]][2:4], ['0.30479', '0.30479'])
        self.assertIn('rutile', coords['caption'])

    def test_bond_table_from_loop(self):
        bonds = dict(self.tabs)['bonds']
        text = T.render_text([('bonds', bonds)])
        self.assertIn('Ti1–O1', text)
        self.assertIn('1.9485(5)', text)                              # esds from the .cif loop
        self.assertIn('<Ti–O>', text)                                 # one Ti site: mean by element
        self.assertIn('Symmetry codes: (1) −y+½, x+½, z+½', text)
        self.assertNotIn('no _geom_bond loop', bonds['note'])

    def test_bvs_table(self):
        bvs = dict(self.tabs)['bvs']
        rows = [[T.plain(c) for c in r] for r in bvs['rows']]
        self.assertEqual([T.plain(c) for c in bvs['head']], ['Atom', 'Ti1', 'Σan'])
        self.assertEqual(rows[0][0], 'O1')
        self.assertIn('0.70×4↓×2→', rows[0][1])
        self.assertEqual(rows[-1][:2], ['Σcat', '4.07'])
        self.assertIn("Brese and O'Keeffe (1991)", bvs['note'])       # the set asked for, not the file's ref id
        self.assertIn('Multiplicity is indicated by', bvs['note'])
        self.assertNotIn('Brown and Altermatt', bvs['note'])

    def test_render_and_word(self):
        h = T.render_html(self.tabs)
        self.assertIn('<sup>1</sup>', h); self.assertIn('data-kind="bvs"', h)
        self.assertNotIn('<script', h)
        out = os.path.join(self.tmp, 'out')
        st, tabs, text = T.run(self.cif, word=True, params='bo', out_dir=out, quiet=True)
        self.assertTrue(os.path.exists(os.path.join(out, 'rutile_tables.txt')))
        from docx import Document
        d = Document(os.path.join(out, 'rutile_tables.docx'))
        self.assertEqual(len(d.tables), 3)
        sup = [r for row in d.tables[1].rows for p in row.cells[0].paragraphs for r in p.runs if r.font.superscript]
        self.assertTrue(sup, 'the symmetry code is a superscript run')


class Journals(unittest.TestCase):
    def test_journal_styles(self):
        tmp = tempfile.mkdtemp(prefix='tbl_')
        try:
            cif = _write(tmp, 'rutile.cif', RUTILE)
            st, tabs = T.build(cif, params='bo', journal_key='cjmp')
            coords, bonds, bvs = dict(tabs)['coords'], dict(tabs)['bonds'], dict(tabs)['bvs']
            self.assertEqual(coords['label'], 'TABLE 1.')
            self.assertEqual(coords['caption'], 'ATOM COORDINATES AND DISPLACEMENT PARAMETERS (Å²) FOR RUTILE')
            self.assertTrue(bonds['note'].startswith('Symmetry operators:'))
            self.assertEqual([T.plain(c) for c in bvs['rows'][-1]][:1], ['Sum'])
            self.assertIn('font-family:Arial', T.render_html(tabs))
            st, tabs = T.build(cif, params='bo', journal_key='minmag')
            coords = dict(tabs)['coords']
            self.assertEqual([T.plain(c) for c in coords['head']][:4], ['Site', 's.o.', 'x/a', 'y/b'])
            self.assertEqual(coords['label'], 'Table 1.')
            st, tabs = T.build(cif, params='bo')                     # the default: the owner's style
            self.assertEqual([T.plain(c) for c in dict(tabs)['coords']['head']][:1], ['Atoms'])
            self.assertEqual(T.build(cif, params='bo', journal_key='nope')[1][0][1]['label'], 'Table 1.')   # unknown -> default
        finally:
            shutil.rmtree(tmp, ignore_errors=True)


class Hydrate(unittest.TestCase):
    def test_hbond_columns_and_proposed_contacts(self):
        tmp = tempfile.mkdtemp(prefix='tbl_')
        try:
            cif = _write(tmp, 'h.cif', HYDRATE)
            st, tabs = T.build(cif)
            d = dict(tabs)
            self.assertEqual([k for k, _ in tabs], ['coords', 'bonds', 'hbonds', 'bvs'])
            self.assertIn('proposed', d['hbonds']['caption'].lower())
            self.assertIn('PROPOSED', d['hbonds']['note'])
            head = [T.plain(c) for c in d['bvs']['head']]
            self.assertEqual(head, ['Atom', 'Ca1', 'Donor', 'vu', 'H bond', 'Σan'])
            rows = {T.plain(r[0]): [T.plain(c) for c in r] for r in d['bvs']['rows']}
            s = B.fi_valence(2.8)
            self.assertEqual(rows['O2'][2:5], ['OW1', '%.2f' % s, '-'])
            self.assertEqual(rows['O2'][-1], '%.2f' % s)                          # Σan = accepted only (no cation)
            self.assertEqual(rows['OW1'][2:5], ['-', '-', '-'])                   # donors: nothing deducted
            self.assertIn('Ferraris and Ivaldi (1988)', d['bvs']['note'])
            self.assertIn('not located', d['bvs']['note'])
            # H located: the geometry table, and no 'proposed' wording
            st, tabs = T.build(_write(tmp, 'hh.cif', HYDRATE_H))
            d = dict(tabs)
            self.assertIn('OW1–H1⋯O2', T.render_text([('hbonds', d['hbonds'])]))
            self.assertNotIn('not located', d['bvs']['note'])
            # no H at all
            st, tabs = T.build(cif, include_h=False)
            self.assertEqual([k for k, _ in tabs], ['coords', 'bonds', 'bvs'])
            st, tabs = T.build(cif, hbond='none')
            self.assertNotIn('Donor', [T.plain(c) for c in dict(tabs)['bvs']['head']])
        finally:
            shutil.rmtree(tmp, ignore_errors=True)


class Computed(unittest.TestCase):
    def test_bonds_without_loop(self):
        tmp = tempfile.mkdtemp(prefix='tbl_')
        try:
            cif = _write(tmp, 'r.cif', RUTILE.split('loop_\n_geom_bond')[0])   # drop the bond loop
            st, tabs = T.build(cif, params='bo')
            bonds = dict(tabs)['bonds']
            text = T.render_text([('bonds', bonds)])
            self.assertIn('no _geom_bond loop', bonds['note'])
            self.assertIn('1.949', text); self.assertIn('1.980', text)
            self.assertEqual(text.count('Ti1–O1'), 6)                 # every symmetry-equivalent bond, each with its code
        finally:
            shutil.rmtree(tmp, ignore_errors=True)


if __name__ == '__main__':
    unittest.main()
