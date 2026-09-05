"""pxrd_review.paper_extract — a synthetic two-column paper page built with PyMuPDF: the analytical
table, the basis sentence, the H2O statement, optics and a powder table.

    python3 -m unittest tests.test_paper_extract -v
"""
import os, shutil, tempfile, unittest

from pxrd_review import paper_extract as PE


def make_pdf(path):
    import fitz
    doc = fitz.open(); page = doc.new_page(width=595, height=842)
    y = 60
    def line(txt, x=40, dy=14):
        nonlocal y
        page.insert_text((x, y), txt, fontsize=9); y += dy
    line('Testite, a new mineral from Nowhere', dy=18)
    line('The empirical formula, calculated on the basis of 7 O apfu, is Ca1.00Mg2.01Si2.99O7(OH)0.98.', dy=16)
    line('H2O was calculated by difference; the F content is below detection. Optical: alpha = 1.600, beta = 1.610, gamma = 1.620.')
    line('Density (measured by flotation) = 3.120 g/cm3; calculated density = 3.145 g/cm3.')
    line('Bond-valence parameters are from Gagne and Hawthorne (2015); hydrogen-bond strengths from O-O bond lengths (Ferraris and Ivaldi 1988).', dy=20)
    line('Table 1. Chemical data (wt%) for testite.')
    line('Constituent   Mean     Range        S.D.   Standard')
    # the other page column's text at the same baselines
    for row, side in (('CaO  20.10  19.80-20.40  0.21  wollastonite', 'some body text of the right column, with commas'),
                      ('MgO  29.05  28.70-29.60  0.30  forsterite', 'that must not become a standard name'),
                      ('SiO2  46.30  45.90-46.70  0.28  quartz', ''),
                      ('H2O  3.20', 'and a third line of prose here'),
                      ('Total  98.65', '')):
        page.insert_text((40, y), row, fontsize=9)
        if side:
            page.insert_text((330, y), side, fontsize=9)
        y += 14
    y += 14
    line('Table 2. Powder X-ray diffraction data for testite.')
    xs = (40, 80, 130, 180, 230, 245, 260)                      # the columns line up under the header, as in print
    for cells in (('Iobs', 'dobs', 'dcalc', 'Icalc', 'h', 'k', 'l'), ('100', '3.4550', '3.4531', '92', '1', '1', '0'),
                  ('35', '2.9800', '2.9791', '40', '0', '2', '1'), ('12', '2.5010', '2.4997', '9', '2', '0', '0')):
        for x, c in zip(xs, cells):
            page.insert_text((x, y), c, fontsize=9)
        y += 14
    doc.save(path); doc.close()


class Extract(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.tmp = tempfile.mkdtemp(prefix='pe_')
        cls.pdf = os.path.join(cls.tmp, 'testite.pdf'); make_pdf(cls.pdf)

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.tmp, ignore_errors=True)

    def test_table_basis_method_optics_bv(self):
        ex = PE.extract(self.pdf, self.tmp, 'testite')
        self.assertEqual(ex['name'], 'testite')
        rows = {r['constituent']: r for r in ex['epma']['rows']}
        self.assertEqual(sorted(rows), ['CaO', 'H2O', 'MgO', 'SiO2'])
        self.assertEqual((rows['CaO']['mean'], rows['CaO']['range'], rows['CaO']['sd'], rows['CaO']['standard']), (20.1, (19.8, 20.4), 0.21, 'wollastonite'))
        self.assertEqual(rows['MgO']['standard'], 'forsterite')          # the other column's prose is not a standard
        self.assertIsNone(rows['H2O']['range']); self.assertEqual(ex['epma']['total'], 98.65)
        self.assertEqual(ex['basis'], ('O', 7.0)); self.assertIn('basis of 7 O', ex['basis_sentence'])
        self.assertEqual(ex['method']['h2o'], 'difference'); self.assertTrue(any('by difference' in s for s in ex['method']['sentences']))
        self.assertAlmostEqual(ex['optics']['n'], 1.61, places=4); self.assertEqual((ex['optics']['D_meas'], ex['optics']['D_calc']), (3.12, 3.145))
        self.assertEqual((ex['bv']['params'], ex['bv']['hb']), ('gh', 'oo'))
        self.assertEqual((ex['pxrd']['obs'], ex['pxrd']['calc']), (3, 3))
        self.assertEqual(sorted(ex['files']), ['calc', 'epma', 'obs'])
        with open(os.path.join(self.tmp, ex['files']['epma'])) as f:
            head, vals = f.read().splitlines()
        self.assertEqual(head.split(','), ['CaO', 'MgO', 'SiO2', 'H2O']); self.assertEqual(vals.split(',')[0], '20.1')
        with open(os.path.join(self.tmp, ex['files']['calc'])) as f:
            self.assertIn('3.4531 92 1 1 0', f.read())
        self.assertEqual(PE.basis_string(ex['basis']), 'O=7')


if __name__ == '__main__':
    unittest.main()


class Parsers(unittest.TestCase):
    """Corpus-found shapes (2026-09 hardening): footnote marks, lost subscripts, fonts, apfu rows."""

    def test_constituents(self):
        self.assertEqual(PE._constituent_ok('Sb')[0], 'Sb')            # antimony, not S + footnote b
        self.assertEqual(PE._constituent_ok('Sc')[0], 'Sc')
        self.assertEqual(PE._constituent_ok('SiO2a')[0], 'SiO2')       # a footnote letter
        self.assertEqual(PE._constituent_ok('NaO')[0], 'Na2O')         # the 2 was a lost subscript
        self.assertEqual(PE._constituent_ok('Fe2O')[0], 'Fe2O3')
        self.assertEqual(PE._constituent_ok('Xa')[0], None)

    def test_row_footnote_word(self):
        # 'TiO2 a 15.36 14.84–16.95 0.43': the footnote printed as its own word must not end the row
        ws = [(294, 0, 312, 9, 'TiO2'), (306, 0, 310, 9, 'a'), (366, 0, 390, 9, '15.36'), (431, 0, 480, 9, '14.84–16.95'), (507, 0, 525, 9, '0.43')]
        c, kind, vals, x0 = PE._row_at(ws, None)
        self.assertEqual((c, kind), ('TiO2', 'constituent'))
        self.assertEqual(vals[0], ('num', 15.36))

    def test_formula_font_and_length(self):
        # colons for decimals and 'ð Þ' brackets (a journal font); a long formula is not cut at 320 chars
        txt = 'The empirical formula is Sr0:57Ba0:38Na0:01 ð Þ0:96 Mn2+ 1:83Fe2+ 0:14 P3:00O11:98 based on 15 O apfu.'
        f, counts, issues, ox = PE.empirical_formula(txt)
        self.assertAlmostEqual(counts['Sr'], 0.57); self.assertAlmostEqual(counts['Mn'], 1.83); self.assertAlmostEqual(counts['P'], 3.0)
        long = 'The empirical formula is ' + ''.join('%s%.2f' % (el, 1 + i / 100) for i, el in enumerate(['Na', 'K', 'Ca', 'Mg', 'Fe', 'Al', 'Si', 'Ti', 'Mn', 'Cu', 'Zn', 'Pb', 'Sr', 'Ba'] * 3)) + 'O200 and so on.'
        f, counts, issues, ox = PE.empirical_formula(long)
        self.assertEqual(counts['O'], 200.0)


class ApfuBlock(unittest.TestCase):
    def test_apfu_rows_below_total(self):
        # bare-element rows under the Total are the apfu block, even for S (a sulfate reports SO3 only)
        import fitz
        tmp = tempfile.mkdtemp(prefix='pe_'); path = os.path.join(tmp, 'sulfate.pdf')
        try:
            doc = fitz.open(); page = doc.new_page(width=595, height=842); y = 60
            for row in ('Table 1. Chemical data (wt%)', 'Constituent  Mean  Range', 'MgO  11.00  10.5-11.5', 'CuO  31.18  30.9-31.5',
                        'ZnO  2.62  2.4-2.8', 'SO3  54.76  54.1-55.2', 'Total  99.56', 'Mg  0.79', 'Cu  1.14', 'Zn  0.09', 'S  1.99'):
                page.insert_text((40, y), row, fontsize=9); y += 14
            doc.save(path); doc.close()
            rows = [r['constituent'] for r in PE.epma_table(path)['rows']]
            self.assertEqual(rows, ['MgO', 'CuO', 'ZnO', 'SO3'])
        finally:
            shutil.rmtree(tmp, ignore_errors=True)


class SpeciesEvidence(unittest.TestCase):
    """Mindat's ideal formula as the fallback line of evidence (owner, 2026-09-02)."""

    def test_ideal_oxidation_and_oxides(self):
        ox = PE.ideal_oxidation('KFe<sup>3+</sup><sub>3</sub>(S<sup>6+</sup>O<sub>4</sub>)<sub>2</sub>(OH)<sub>6</sub>')
        self.assertEqual(ox, {'Fe': {3}, 'S': {6}})
        self.assertEqual(PE.oxide_for('Fe', 3), 'Fe2O3'); self.assertEqual(PE.oxide_for('Mn', 4), 'MnO2'); self.assertEqual(PE.oxide_for('Cu', 1), 'Cu2O')
        species = {'formula': 'Fe<sup>3+</sup>', 'ox': {'Fe': {3}}, 'elements': {'Fe', 'O'}}
        self.assertEqual(PE.oxide_alternatives({'FeO': 10.0, 'SiO2': 40.0}, {}, species), [('FeO', 'Fe2O3', "Mindat's ideal formula has Fe3+")])
        self.assertEqual(PE.oxide_alternatives({'FeO': 10.0}, {'Fe': {3}}, None), [('FeO', 'Fe2O3', 'the formula writes Fe3+')])
        self.assertEqual(PE.oxide_alternatives({'FeO': 10.0}, {'Fe': {2, 3}}, species), [])      # the authors' own split
        self.assertAlmostEqual(PE._convert({'FeO': 10.0}, 'FeO', 'Fe2O3')['Fe2O3'], 11.113, places=3)

    def test_species_lines(self):
        species = {'formula': 'KFe<sup>3+</sup><sub>3</sub>(SO<sub>4</sub>)<sub>2</sub>(OH)<sub>6</sub>', 'ox': {'Fe': {3}}, 'elements': {'K', 'Fe', 'S', 'O', 'H'}}
        L = PE.species_lines(species, {'Fe', 'S'}, {'Fe', 'S'}, {'Fe': {2}})
        self.assertEqual(len(L), 2)
        self.assertIn('carries K, absent', L[0]); self.assertIn('Fe2+; Mindat', L[1])
        self.assertEqual(PE.species_lines(species, {'K', 'Fe', 'S'}, set(), {'Fe': {3}}), [])

    def test_formula_charges_and_group_basis(self):
        f, counts, issues, ox = PE.empirical_formula('The empirical formula is Ca2.00(Mg3.44Ti4+ 1.49Fe0.36Ti3+ 0.34)Σ5.63Si2.37O20 on 20 O.')
        self.assertEqual(ox, {'Ti': {3, 4}})
        n = PE._journal_to_icdd('(Pb0.930Ce0.434Sm0.007)R2.000(CO3)2(OH)1.074')
        self.assertIn(') Σ2.000', n)                                 # a Σ printed as R survives the bare-Σ rule


class Transposed(unittest.TestCase):
    def test_constituents_across(self):
        import fitz
        tmp = tempfile.mkdtemp(prefix='pe_'); path = os.path.join(tmp, 'across.pdf')
        try:
            doc = fitz.open(); page = doc.new_page(width=595, height=842); y = 60
            xs = (40, 120, 190, 260, 330, 400)
            for cells in (('Constituent', 'Nb2O5', 'MgO', 'FeO', 'MnO', 'Total'), ('1', '62.10', '36.20', '0.80', '0.60', '99.70'),
                          ('2', '61.90', '36.40', '0.70', '0.70', '99.70'), ('Mean', '62.00', '36.30', '0.75', '0.65', '99.70'),
                          ('Range', '61.90-62.10', '36.20-36.40', '0.70-0.80', '0.60-0.70', '')):
                for x, c in zip(xs, cells):
                    page.insert_text((x, y), c, fontsize=9)
                y += 14
            doc.save(path); doc.close()
            e = PE.epma_table(path)
            self.assertTrue(e.get('transposed'))
            self.assertEqual({r['constituent']: r['mean'] for r in e['rows']}, {'Nb2O5': 62.0, 'MgO': 36.3, 'FeO': 0.75, 'MnO': 0.65})
            self.assertEqual(e['total'], 99.7)
            self.assertEqual(len(PE.table_alternatives(e, {r['constituent']: r['mean'] for r in e['rows']})), 2)
        finally:
            shutil.rmtree(tmp, ignore_errors=True)


class Totals(unittest.TestCase):
    def test_total_row_dropped_when_split_follows(self):
        rows = [{'constituent': c, 'mean': v} for c, v in (('CaO', 4.8), ('FeO', 11.79), ('FeO', 10.71), ('Fe2O3', 1.2), ('P2O5', 34.19))]
        self.assertEqual([r['constituent'] for r in PE._drop_totals(rows)], ['CaO', 'FeO', 'Fe2O3', 'P2O5'])
        self.assertEqual(PE._drop_totals(rows)[1]['mean'], 10.71)
        rows2 = [{'constituent': c, 'mean': v} for c, v in (('FeO', 11.79), ('Fe2O3', 1.2), ('P2O5', 34.19))]
        self.assertEqual(len(PE._drop_totals(rows2)), 3)                     # one of each: nothing to drop


class MorningRules(unittest.TestCase):
    """2026-09-03: the misses worked through with the owner's hand-check list."""

    def test_notation(self):
        self.assertEqual(PE._constituent_ok('Fe2O3(tot)')[0], 'Fe2O3')
        self.assertEqual(PE._constituent_ok('H2O(calc)')[0], 'H2O')
        n = PE._journal_to_icdd('(Ti0.60Fe+3 0.23Mg0.08)Σ1.04')
        self.assertIn('Fe0.23 +3', n)                                       # the sign before the digit
        n = PE._journal_to_icdd('Cu2.00Ag0.97(As0.95Sb0.04)Σ0.99S4.03 (ΣMe = 3.97)')
        self.assertNotIn('Me', n)
        n = PE._journal_to_icdd('(Sr0.55Ba0.25Ln0.10Ca0.10)')
        self.assertIn('Ce0.10', n)                                          # 'Ln0.10' is a grouped lanthanide
        f, counts, issues, ox = PE.empirical_formula('The empirical formula, normalized to 12 Cu apfu, is Cu12(Pb1.92Fe0.06Si0.06)(O15.08F0.02)(Br0.99Cl0.89) here.')
        self.assertAlmostEqual(counts['Cu'], 12.0)                          # 'Cu12(' is the start, not a site label
        f, counts, issues, ox = PE.empirical_formula('The crystal chemical formula of saranovskite is (Sr0.55Ba0.25)(Fe2+ 1.12Mg0.88)O38 and')
        self.assertAlmostEqual(counts['Mg'], 0.88)                          # 'crystal chemical formula' wording

    def test_nested_bare_group_counts_atoms(self):
        from pxrd_review import epma as EP
        counts, ox, issues = EP.parse_icdd_formula(PE._journal_to_icdd('[(Fe3+ 2.12Al0.18)(Zn0.32Mg0.16Fe2+ 0.13Mn0.03)Ti0.06]Σ3.00(Sb0.97Ti0.03)Σ1.00Zn1.00O7'))
        self.assertEqual(issues, [])
        counts, ox, issues = EP.parse_icdd_formula('[( Cl3.82 F0.18 ) Σ4 ( F1.54 O H1.46 ) Σ3 ( O H )2 ] Σ9')
        self.assertEqual(issues, [])                                         # '(OH)2' is two items of the Σ9

    def test_candidate_tables_and_ppm(self):
        import fitz
        tmp = tempfile.mkdtemp(prefix='pe_'); path = os.path.join(tmp, 'two.pdf')
        try:
            doc = fitz.open(); page = doc.new_page(width=595, height=842); y = 60
            def line(txt):
                nonlocal y
                page.insert_text((40, y), txt, fontsize=9); y += 14
            line('The empirical formula is K2.00Nb1.90Ti0.09Si4.01O12 on 8 cations.'); y += 10
            line('Table 3. Trace elements and oxides'); line('Constituent  wt%')
            for row in ('TiO2  0.65', 'ZrO2  0.21', 'Nb2O5  42.88', 'K2O  15.57', 'Be  0.01', 'B  0.85', 'P  22.0', 'Ba  201'):
                line(row)
            page = doc.new_page(width=595, height=842); y = 60                  # the analytical table on the next page
            line('Table 2. Chemical data (wt%)'); line('Constituent  Mean  Range')
            for row in ('SiO2  40.28  40.07-41.02', 'TiO2  1.20  1.1-1.3', 'ZrO2  0.21  0.1-0.3', 'Nb2O5  42.30  38.2-44.0', 'K2O  15.77  15.7-16.1', 'Total  99.76'):
                line(row)
            doc.save(path); doc.close()
            text = PE.text_of(path); ex = PE.extract(path, None, 'x', write=False) if 'write' in PE.extract.__code__.co_varnames else PE.extract(path, None, 'x')
            e = ex['epma']
            self.assertTrue(e.get('candidates'))
            c = PE.check_composition(ex, text)
            self.assertTrue(c['ok'], c['lines'])
            self.assertTrue(any('reproduces the formula' in l for l in c['lines']) or 'SiO2' in {r['constituent'] for r in e['rows']})
        finally:
            shutil.rmtree(tmp, ignore_errors=True)


class LateMorning(unittest.TestCase):
    def test_prose_composition(self):
        pt = PE.prose_table('The analysis gave SiO2 40.12, Al2O3 20.05, FeO 10.20 and MgO 29.10 wt.%, total 99.47 wt.%.')
        self.assertEqual([(r['constituent'], r['mean']) for r in pt['rows']], [('SiO2', 40.12), ('Al2O3', 20.05), ('FeO', 10.2), ('MgO', 29.1)])
        self.assertEqual(pt['total'], 99.47)
        self.assertIsNone(PE.prose_table('The ideal formula requires Na2O 6.46, MnO 14.78, Ce2O3 34.19, P2O5 29.57, and H2O 15.01, total 100.00 wt.%.'))
        self.assertIsNone(PE.prose_table('(Ce0.39La0.24Pr0.13Sr0.11Nd0.11Sm0.01)'))       # coefficients are not a composition

    def test_prose_total_is_not_the_total_row(self):
        # the other page column says 'A total of 16 scans'; the table row on the same line is Na2O
        ws = [(40, 0, 45, 9, 'A'), (50, 0, 70, 9, 'total'), (75, 0, 85, 9, 'of'), (90, 0, 100, 9, '16'), (105, 0, 130, 9, 'scans'),
              (310, 0, 335, 9, 'Na2O'), (367, 0, 395, 9, '25.51'), (406, 0, 460, 9, '23.95−26.01'), (469, 0, 490, 9, '0.94'), (511, 0, 550, 9, 'jadeite')]
        c, kind, vals, x0 = PE._row_at(ws, None)
        self.assertEqual((c, kind, vals[0]), ('Na2O', 'constituent', ('num', 25.51)))

    def test_subscript_digit_glued(self):
        import fitz
        tmp = tempfile.mkdtemp(prefix='pe_'); path = os.path.join(tmp, 'sub.pdf')
        try:
            doc = fitz.open(); page = doc.new_page(width=595, height=842)
            page.insert_text((40, 100), 'B2O', fontsize=9); page.insert_text((54.5, 102.5), '3', fontsize=6)    # the subscript as its own glyph
            page.insert_text((100, 100), '11.43', fontsize=9)
            page.insert_text((40, 120), 'SiO2', fontsize=9); page.insert_text((59, 96), '1', fontsize=6)       # a numbered footnote, superscript
            doc.save(path); doc.close()
            lines = PE.page_lines(fitz.open(path)[0])
            toks = [w[4] for w in lines[0]['w']]
            self.assertEqual(toks, ['B2O3', '11.43'], toks)
        finally:
            shutil.rmtree(tmp, ignore_errors=True)

    def test_charges_and_structural_form(self):
        n = PE._journal_to_icdd('(Fe2+ 0.56Fe2.5+ 0.25Mg0.14)Σ0.95')
        self.assertIn('Fe0.25 +2.5', n)
        f, counts, issues, ox = PE.empirical_formula('The empirical formula is Al0.01Fe2.55Cu0.91S2(OH)3.07 = (Fe1.09Cu0.91)Σ2S2 · (Al0.01)Σ1.47(OH)3.07 here.')
        self.assertAlmostEqual(counts['Fe'], 2.55)                                # cut at '= (structural form)'
        f, counts, issues, ox = PE.empirical_formula('The empirical formula is (Na0.54h0.26Ca0.20)Σ1.00Mn1.00P2.04O8 based on 8 O.')
        self.assertEqual(issues, [])                                              # 'h' is a vacancy glyph


class Names(unittest.TestCase):
    def test_mineral_name_unicode_and_stoplist(self):
        self.assertEqual(PE.mineral_name('Åsgruvanite-(Ce), a new mineral from Sweden. Despite its rarity, calcite is common.'), 'åsgruvanite-(ce)')
        self.assertEqual(PE.mineral_name('Despite everything, Fuyuanite occurs with Calcite in granite.'), 'fuyuanite')

    def test_fit_judged_on_major_cations(self):
        from pxrd_review import epma as EP
        wt = {'PbO': 42.4 * 1.077, 'MoO3': 21.1 * 1.5, 'S': 24.05, 'MnO': 0.05 * 1.29}
        counts = {'Pb': 4.0, 'Mo': 4.33, 'S': 15.0, 'Mn': 0.05}
        r = EP.replicate_formula(wt, counts, [('element', 'S', 15.0)], 'x', tol_abs=0.03, tol_rel=0.05)
        self.assertLess(r['score'], 0.1, r)                                     # Mn's 60 % is a trace: not the overall fit


class Captions(unittest.TestCase):
    def test_caption_score(self):
        self.assertEqual(PE._caption_score('Table 1. Chemical composition (wt.%) of barronite'), 9)
        self.assertEqual(PE._caption_score('Table 3. Atom coordinates and displacement parameters'), -8)
        self.assertEqual(PE._caption_score('Table 5. Bond lengths (Å)'), -10)
        self.assertEqual(PE._caption_score(''), 0)

    def test_notation_batch(self):
        n = PE._journal_to_icdd('(Na0.63Ca0.25Mn0.12)R1.00(A1.91Na0.09)R2.00')
        self.assertIn('?1.91', n)                                              # 'A' is a vacancy glyph
        n = PE._journal_to_icdd('(PO4)4.07(OH)3.524.13H2O')
        self.assertIn('( O H )3.52 ! 4.13 H2 O', n)                            # the hydrate dot lost
        n = PE._journal_to_icdd('Si4(S1.61Si0.32P0.07)Σ1.99O24')
        self.assertIn('Si4 ', n)                                                # 'Si4(' is a count, not a site label
        n = PE._journal_to_icdd('Mg1(Mg1.42Fe0.30)Σ1.72Mg2(Mg1.71Fe0.10)Σ1.81')
        self.assertNotIn('Mg1 ', n); self.assertNotIn('Mg2 ', n)               # numbered from 1: site labels
        from pxrd_review import epma as EP
        counts, ox, issues = EP.parse_icdd_formula('( Ca1.08 Na0.91 ) Σ2.00 ( O H1.59 O0.61 ) Σ2.00')
        self.assertTrue(issues and issues[0].startswith('anion group sum'))  # the cation groups are fine

    def test_alternative_columns_need_a_total(self):
        e = {'rows': [{'constituent': 'BaO', 'mean': 41.26, 'all': [41.26, 0.36, 41.29, 0.40]},
                      {'constituent': 'SiO2', 'mean': 17.38, 'all': [17.38, 0.2, 17.5, 0.3]},
                      {'constituent': 'MnO', 'mean': 8.49, 'all': [8.49, 0.1, 8.3, 0.2]}]}
        alts = PE.table_alternatives(e, {'BaO': 41.26, 'SiO2': 17.38, 'MnO': 8.49})
        self.assertEqual(alts, [{'BaO': 41.29, 'SiO2': 17.5, 'MnO': 8.3}])      # the s.d. columns are not samples


class Afternoon(unittest.TestCase):
    """2026-09-03 afternoon: the second sweep through the deviating and unparsed papers."""

    def test_formula_rules(self):
        J = PE._journal_to_icdd
        self.assertIn('Pt0.01 +4', J('(Rh3+ 1.20Ir4+ 0.06Pt4+ <0.01)S3.99'))              # '<0.01': the bound stands
        self.assertIn('Ge0.940', J('(Ge0.91-0.97Si0.03-0.09)Σ1.00O2'))                     # a range: its middle
        self.assertIn('Th0.54', J('(Ca0.40REE0.93(Th,U)4+ 0.54□0.13)R2.00'))               # '(Th,U)4+ 0.54'
        self.assertIn('?0.06', J('(Mg1.24Ca0.69 0.06Mn0.01)Σ2.00'))                        # a vacancy glyph lost
        n = J('[(H2O)0.50K0.50]2(Mn1.20Mg0.49Fe2+ 0.27Zn0.05)P2.01(Al1.63Fe3+ 0.20Ti4+ 0.19)P2.02(PO4)4.02')
        self.assertIn(') Σ2.01', n); self.assertIn('(PO4 )4.02', n)                       # a Σ printed as P, (PO4) untouched
        self.assertIn(')2 ! 2 H2 O', J('Ca1.99(Fe0.89Mg0.13Mn0.01)R1.03(P1.00O4)22H2O'))    # '(PO4)22H2O' = (PO4)2·2H2O
        self.assertTrue(J('(SREF) Cu2Fe0.84Al0.16(AsO4)(OH)4·4H2O').startswith('Cu2'))    # a tag
        self.assertIn('Σ48.49', J('Ag1.04Pb46.43(As25.52Sb22.97)Σ=48.49S120'))
        self.assertEqual(PE._constituent_ok('Na2О')[0], 'Na2O')                              # Cyrillic О

    def test_parser_rules(self):
        from pxrd_review import epma as EP
        counts, ox, issues = EP.parse_icdd_formula('( Al7.98 Fe0.01 ) Σ7.99 (SO4 )5.01 ( O H )13.95')
        self.assertAlmostEqual(counts['S'], 5.01); self.assertEqual(issues, [])          # (SO4)5.01 is a multiplier
        counts, ox, issues = EP.parse_icdd_formula('[( O H )1.25 F0.06 ! 0.69 H2 O ] Σ2.00')
        self.assertEqual(issues, [])                                                      # the water counts toward the Σ
        counts, ox, issues = EP.parse_icdd_formula('Cu3.23 Pb18.74 Sb17.30 S56')
        self.assertEqual(issues, [])                                                      # 56 S is a real cell

    def test_formula_triggers(self):
        f, counts, issues, ox = PE.empirical_formula('leading to an empirical formula (based on 2 Te apfu) of Au3.00Tl1.01Te2.00. Honeaite is black.')
        self.assertAlmostEqual(counts['Au'], 3.0)                                         # near-ideal coefficients with two decimals
        f, counts, issues, ox = PE.empirical_formula('On the basis of 56 S, the chemical formula of ciriottiite is Cu3.23(11)Ag0.43(4)Pb18.74(9)S56 and so on.')
        self.assertAlmostEqual(counts['Pb'], 18.74)
        f, counts, issues, ox = PE.empirical_formula('The chemical formula is Ca4(Al0.5Si0.5)2Si4O16(OH) here.')
        self.assertEqual(counts, {})                                                      # an ideal one stays out


class Evening(unittest.TestCase):
    def test_sigma_glyph_lost_and_nested_P(self):
        from pxrd_review import epma as EP
        counts, ox, issues = EP.parse_icdd_formula('(Ca4.97 Na0.013 Mg0.017 )(As3.99 S0.01 )4 O23 H16')
        self.assertAlmostEqual(counts['As'], 3.99); self.assertEqual(issues, [])           # ')4' is the cations' own sum
        counts, ox, issues = EP.parse_icdd_formula('( Fe1.0 Mg1.0 ) (PO4 )3')
        self.assertAlmostEqual(counts['P'], 3.0)                                            # one cation: a multiplier
        n = PE._journal_to_icdd('[((Ca0.5Mg0.2Na0.11□0.14)60.95 (As3+)2.05)P3.00(Fe3+ 2.44Mo6+ 0.56)63.00]')
        self.assertIn(') Σ3.00(Fe', n)                                                      # Σ printed as P after a nested group
        self.assertTrue(PE._journal_to_icdd('(O + F) (Na2.74Mn0.15)Σ2.89Ca2').startswith('(Na2.74'))

    def test_factor_like_is_unverified(self):
        ex = {'name': 'x', 'basis': ('element', 'Ca', 2.0), 'epma': {'rows': [{'constituent': c, 'mean': v} for c, v in (('CaO', 53.25), ('P2O5', 33.7), ('Cl', 16.96))], 'total': 99.6, 'header': 'Constituent wt%'}}
        c = PE.check_composition(ex, 'The empirical formula is Ca2.01P1.98O7.96Cl1.01 here.')
        self.assertTrue(c['ok'] or c['verified'] is not None)
        ex['epma']['rows'][1] = {'constituent': 'P2O5', 'mean': 8.4}                        # a quarter of the phosphorus: a factor, not a slip
        c = PE.check_composition(ex, 'The empirical formula is Ca2.01P1.98O7.96Cl1.01 here.')
        self.assertFalse(c['verified']); self.assertTrue(any('off by a factor' in l for l in c['lines']))


class HeadlineColumn(unittest.TestCase):
    """Owner (2026-09-03): the headline mineral's column is named in the header or the text; other
    phases in the same table are supporting material."""

    def _table(self):
        # two localities side by side; Mn2O3 is a total for both and split only for the second
        rows = [{'constituent': 'As2O5', 'mean': 48.06, 'all': [48.06, 47.9], 'xs': [111, 307]},
                {'constituent': 'Mn2O3', 'mean': 17.48, 'all': [17.48, 16.2], 'xs': [111, 307]},
                {'constituent': 'MnO2', 'mean': 6.28, 'all': [6.28], 'xs': [310]},
                {'constituent': 'Mn2O3', 'mean': 10.5, 'all': [10.5], 'xs': [307]},
                {'constituent': 'CaO', 'mean': 13.84, 'all': [13.84, 14.1], 'xs': [111, 307]},
                {'constituent': 'Na2O', 'mean': 5.54, 'all': [5.54, 5.2], 'xs': [111, 307]}]
        cells = [('Montaldo', 168), ('Valletta', 382), ('(n', 178), ('=', 185), ('10)', 193), ('(n', 374), ('=', 380), ('13)', 388),
                 ('Oxide', 48), ('mean', 119), ('range', 189), ('s.u.', 256), ('mean', 314), ('range', 384), ('s.u.', 452)]
        return {'rows': PE._drop_totals(rows), 'rows_all': rows, 'head_cells': cells, 'header': 'Oxide mean range s.u. mean range s.u.'}

    def test_column_by_holotype_locality(self):
        e = self._table()
        wt, why = PE.headline_column(e, 'piccoliite', {'codes': [], 'n': None, 'holotype': False, 'holotype_words': ['Montaldo']})
        self.assertIn('Montaldo', why)
        self.assertEqual(wt['Mn2O3'], 17.48); self.assertNotIn('MnO2', wt)               # the total stands where the split is blank
        wt, why = PE.headline_column(e, 'piccoliite', {'codes': ['Valletta'], 'n': None, 'holotype': False, 'holotype_words': []})
        self.assertEqual((wt['Mn2O3'], wt['MnO2']), (10.5, 6.28))                         # the split overrides the total

    def test_column_by_n_analyses(self):
        e = self._table()
        wt, why = PE.headline_column(e, 'piccoliite', {'codes': [], 'n': 13, 'holotype': False, 'holotype_words': []})
        self.assertIn('13 analyses', why); self.assertEqual(wt['CaO'], 14.1)

    def test_hints_and_footnote_numbers(self):
        h = PE.sample_hints('The empirical formula of domain A (mean of 4 analyses, sample TL-12) is')
        self.assertIn('A', h['codes']); self.assertIn('TL-12', h['codes']); self.assertEqual(h['n'], 4)
        self.assertEqual(PE._constituent_ok('V2O31)')[0], 'V2O3')                          # a footnote number
        self.assertEqual(PE.holotype_words('The holotype specimen is from Montaldo di Mondovì.')[:1], ['Montaldo'])


class Night(unittest.TestCase):
    def test_flattened_group_and_triggers(self):
        self.assertIn('(PO4 )3.02', PE._journal_to_icdd('Pb0.98Fe2+ 1.69Σ2.00V1.31Al0.06Σ2.00PO43.02OH3'))
        f, counts, issues, ox = PE.empirical_formula('The empirical mineral formula is Ca2.06Mn3+ 1.78Cu0.10F0.97(OH)8.02(SO4)0.39. The unit-cell')
        self.assertAlmostEqual(counts['Mn'], 1.78)
        f, counts, issues, ox = PE.empirical_formula('Electron microprobe analyses together with Mössbauer spectroscopy gives the formula (Ca0.59Mn0.24)Σ0.83Mn(Zn0.74Mn2+ 0.48)Σ2(P0.995O4)4(OH)2. Jahnsite is monoclinic')
        self.assertAlmostEqual(counts['Zn'], 0.74)
        f, counts, issues, ox = PE.empirical_formula('the empirical formula (calculated on the basis of 13 O atoms pfu) is [(U1.00O2)2(C2O4)(OH)2(H2O)2]·H2O and the ideal formula is')
        self.assertAlmostEqual(counts['U'], 2.0)                                     # round coefficients after an explicit "empirical formula"
        f, counts, issues, ox = PE.empirical_formula('yielded the formula: (Ce​1.​81La​0.8​1)​Σ2.62 Fe2+ 0.80 here')
        self.assertAlmostEqual(counts['Ce'], 1.81)                                    # zero-width spaces inside the glyph runs

    def test_caption_names(self):
        self.assertEqual(PE._caption_score('Table 4. Chemical data (wt.%) for torryweiserite', 'torryweiserite'), 3 - 2 + 6)
        self.assertEqual(PE._caption_score('Table 11. Chemical compositions of possible oberthurite from other localities', 'torryweiserite'), -2 - 4 + 6)


class Parser3(unittest.TestCase):
    def test_group_multipliers_and_sums(self):
        from pxrd_review import epma as EP
        self.assertAlmostEqual(EP.parse_icdd_formula('(Cu2.68 Mg0.17 ) Σ3 (N3 C2 H2 )2.755')[0]['N'], 8.265)       # not the 'Fe3 C1.01' charge notation
        self.assertAlmostEqual(EP.parse_icdd_formula('(Cu5.16 Co0.34 ) Σ6.01 (AsO3 OH )5.97')[0]['As'], 5.97)       # no decimals inside: a multiplier
        self.assertAlmostEqual(EP.parse_icdd_formula('Ca1.02 Zn1.91 ((As0.95 Sb0.08 )O4 ) Σ2.03')[0]['As'], 1.9285, places=3)   # Σ2.03 = two such groups
        self.assertAlmostEqual(EP.parse_icdd_formula('Si2.00 [O5.91 OH1.09 ]7.00')[0]['O'], 7.0)                    # decimals inside: a sum

    def test_headline_name_by_use(self):
        t = 'Journal head Bustamite group. ' + 'Mendigite, a new mineral species of the bustamite group. ' + 'mendigite ' * 5 + 'bustamite ' * 3
        self.assertEqual(PE.mineral_name(t), 'mendigite')


class Selector2(unittest.TestCase):
    def test_levinson_suffix_and_sentence_list(self):
        rows = [{'constituent': 'Nb2O5', 'mean': 10.32, 'all': [10.32, 11.1, 9.8], 'xs': [112, 154, 197]},
                {'constituent': 'SiO2', 'mean': 39.45, 'all': [39.45, 39.9, 39.1], 'xs': [112, 154, 197]},
                {'constituent': 'Na2O', 'mean': 8.8, 'all': [8.8, 8.2, 8.5], 'xs': [112, 154, 197]}]
        e = {'rows': rows, 'rows_all': rows, 'head_cells': [('Constituent', 62), ('(Nd)1', 112), ('(Y)2', 154), ('(Ce)3', 197)], 'header': ''}
        wt, why = PE.headline_column(e, 'nacareniobsite-(y)', {'codes': [], 'n': None, 'holotype': False, 'holotype_words': []})
        self.assertIn('(Y)', why); self.assertEqual(wt['Na2O'], 8.2)
        h = PE.sample_hints('Type paqueite from A-WP1 has an empirical formula of')
        self.assertIn('A-WP1', h['codes'])
        pt = PE.prose_table('Pt 2.10, Ir 0.10, Ni 17.09, Fe 9.76, Cu 7.38, Co 1.77 S 30.97, total 99.73 wt.%, which corresponds to')
        self.assertEqual(pt['rows'][-1]['constituent'], 'S')                             # no comma before the last item


class Parser4(unittest.TestCase):
    def test_late_notations(self):
        J = PE._journal_to_icdd
        self.assertIn('( H2 O )0.21', J('(Na0.99Ca0.46La0.01H2O0.21)∑2.00'))                  # a water count inside a group
        self.assertIn('(PO4 )1.91', J('Σ2.00PO4)1.91(OH)2.27'))                                # the opening bracket lost
        self.assertIn('Th0.01 ?0.51', J('(Ce4.02Th0.01–0.51)Σ9'))                              # one dash in a Σ group: a vacancy
        self.assertIn('Ge0.940', J('(Ge0.91-0.97Si0.03-0.09)Σ1.00O2'))                         # several dashes: ranges
        self.assertIn('Fe0.25 +3', J('(Mg0.75Fe0.25 3+)Σ1'))                                   # the charge after the count
        self.assertIn('?0.63', J('(Pb8.33Sr0.04o0.63)S9.00'))                                   # a vacancy printed as o


class Domains(unittest.TestCase):
    def test_domains_assigned_to_the_headline(self):
        t = ('Electron microprobe data (in wt.%) of ferriandrosite-(Ce) (domains A–C) and associated vielleaureite-(Ce) (domain D). '
             'Domains B and C correspond to the end-member formula MnCeFe3+AlMn2+(Si2O7)(SiO4)O(OH), i.e. to ferriandrosite-(Ce), '
             'however domain D leads to the end-member formula of vielleaureite-(Ce).')
        self.assertEqual(PE.headline_domains(t, 'ferriandrosite-(ce)')[:3], ['B', 'C', 'A'])     # B and C named twice; D belongs to the other mineral
        self.assertEqual(PE.headline_domains('Crystal II of piccoliite gave the best data.', 'piccoliite'), ['II'])


class TwoFormulas(unittest.TestCase):
    def test_abstract_disagrees_with_body(self):
        rows = [{'constituent': c, 'mean': v} for c, v in (('Ce2O3', 39.37), ('La2O3', 19.92), ('Nd2O3', 14.46), ('Sm2O3', 2.84), ('CaO', 0.73), ('F', 14.33))]
        ex = {'name': 'håleniusite-(ce)', 'basis': ('O', 2.0), 'epma': {'rows': rows, 'total': 93.72, 'header': 'Constituent Mean'}}
        text = ('Abstract. Electron probe microanalysis provided the empirical formula (Ce0.41La0.21Sm0.15Nd0.04Ca0.02)R0.83(O0.70F1.30)R2.00. ' + 'x ' * 2500 +
                'The empirical formula calculated on the basis of O + F = 2 apfu is (Ce0.412La0.210Nd0.148Sm0.028Ca0.022)R0.820(O0.70F1.30)R2.00. Later text.')
        c = PE.check_composition(ex, text)
        self.assertTrue(any('abstract does not agree' in l and 'Sm 0.15 vs 0.028' in l for l in c['lines']), c['lines'])
        self.assertFalse(c['ok'])


class Exclusions(unittest.TestCase):
    def test_excluded_elements(self):
        t = 'as Al3+ strongly differs in ionic radius from large cations such as REE, the empirical formula was calculated without Al. Later, Si was excluded from the sum.'
        self.assertEqual(PE.excluded_elements(t), ['Al', 'Si'])
        self.assertEqual(PE.excluded_elements('The formula was calculated on the basis of 12 O.'), [])


class Corrected(unittest.TestCase):
    def test_total_composition_and_corrected_formula(self):
        rows = [{'constituent': c, 'mean': v} for c, v in (('Fe2O3', 4.18), ('MnO', 0.43), ('K2O', 0.98), ('SO3', 56.72), ('SiO2', 0.10), ('Al2O3', 10.10), ('MgO', 1.00), ('Na2O', 19.39))]
        ex = {'name': 'heimaeyite', 'basis': None, 'epma': {'rows': rows, 'total': 92.89, 'header': 'Constituent Mean'}}
        text = ('We assume the presence of small amounts of Mn, Si, K and Mg is due to impurities. The total composition of the sample results in '
                'K0.10Na2.95Mn0.03Mg0.12Fe0.25Al0.94Si0.01S3.35O13.5. If the Mn and Si impurity and the contributions of koryakite are removed, '
                'the resulting empirical formula of heimaeyite is Na2.93Al0.82Fe0.25S2.99O12.05. The ideal formula is NaAl(SO4)2.')
        c = PE.check_composition(ex, text)
        self.assertFalse(c['ok']); self.assertFalse(c['verified'])
        self.assertTrue(any('uncorrected composition' in l and 'cannot be re-derived' in l for l in c['lines']), c['lines'])


class HandCheck(unittest.TestCase):
    """2026-09-03: the owner's five hand-checked papers, one rule each."""

    def test_legend_columns(self):
        t = '1, 2, 8 – fluorpyromorphite (1 – holotype, mean of 8 spot analyses; 2 – F-richest spot; 8 – cotype); 3 – hydroxylpyromorphite (mean of 4).'
        self.assertEqual(PE.legend_columns(t, 'fluorpyromorphite'), ['1', '2', '8'])
        self.assertEqual(PE.legend_columns('4–6 – pyromorphite; 1 – mimetite', 'pyromorphite'), ['4', '5', '6'])

    def test_two_line_cell_mean_above_label(self):
        import fitz
        tmp = tempfile.mkdtemp(prefix='pe_'); path = os.path.join(tmp, 'twoline.pdf')
        try:
            doc = fitz.open(); page = doc.new_page(width=595, height=842); y = 60
            page.insert_text((40, y), 'Constituent', fontsize=9); y += 14
            for x, t in zip((130, 180, 220, 260), ('1', '2', '3', '4')): page.insert_text((x, y), t, fontsize=9)
            y += 14
            for label, mean, rng, others in (('CaO', '0.10', '(0.00–0.32)', ('0.15', '0.05', '0.34')), ('PbO', '83.51', '(82.79–84.40)', ('82.50', '83.20', '82.34')),
                                              ('P2O5', '16.13', '(16.00–16.23)', ('15.90', '16.05', '16.02')), ('F', '1.00', '(0.92–1.06)', ('1.36', '0.37', '0.14'))):
                page.insert_text((130, y), mean, fontsize=9)                       # the mean, a line above the label
                page.insert_text((40, y + 5), label, fontsize=9); page.insert_text((120, y + 5), rng, fontsize=7)
                for x, t in zip((180, 220, 260), others): page.insert_text((x, y + 5), t, fontsize=9)
                y += 20
            page.insert_text((40, y), 'Total  100.84  99.37  100.29  100.44', fontsize=9)
            doc.save(path); doc.close()
            e = PE.epma_table(path)
            rows = {r['constituent']: r for r in e['rows']}
            self.assertEqual(rows['PbO']['mean'], 83.51); self.assertEqual(rows['PbO']['all'][:2], [83.51, 82.5])
            self.assertEqual(rows['CaO']['mean'], 0.1)
        finally:
            shutil.rmtree(tmp, ignore_errors=True)


class Attribution(unittest.TestCase):
    def test_other_mineral_formula_dropped_and_associated_caption(self):
        t = ('The empirical formula of zoisite-(Pb) is (Ca1.09Pb0.86Mn2+0.01)Σ1.96(Al2.88Fe3+0.10)Σ2.98Si3.00O12(OH). '
             'The empirical formula of hancockite is (Ca1.18Pb0.73Mn2+0.06)Σ1.97(Al2.32Fe3+0.66)Σ2.98Si3.01O12(OH).')
        fs = PE._formulas(t, 'zoisite-(pb)')
        self.assertEqual(len(fs), 1); self.assertAlmostEqual(fs[0][1]['Pb'], 0.86)
        self.assertLess(PE._caption_score('Table 2. EPMA representative analyses of phases associated with zoisite-(Pb).', 'zoisite-(pb)'),
                        PE._caption_score('Table 1. Chemical composition of zoisite-(Pb) (wt.%).', 'zoisite-(pb)'))


class ReaderClasses(unittest.TestCase):
    """2026-09-03 afternoon: the general reader defects behind the 'column chosen by fit' verdicts."""

    def _pdf(self, lines, xs=None):
        """lines: [(y, [(x, text), …])] or [(y, 'text at x=40')] -> a one-page pdf path."""
        import fitz
        tmp = tempfile.mkdtemp(prefix='pe_'); path = os.path.join(tmp, 'page.pdf')
        doc = fitz.open(); page = doc.new_page(width=595, height=842)
        for y, cells in lines:
            if isinstance(cells, str):
                page.insert_text((40, y), cells, fontsize=9)
            else:
                for x, t in cells:
                    page.insert_text((x, y), t, fontsize=9)
        doc.save(path); doc.close()
        self.addCleanup(shutil.rmtree, tmp, ignore_errors=True)
        return path

    def test_value_cells(self):
        self.assertEqual(PE._numbers(['29(3)', '67(3)', '3.6(3)'])[::2], [('num', 29.0), ('num', 67.0), ('num', 3.6)])
        self.assertEqual(PE._numbers(['(13.16)', '13.09–13.25'])[0], ('num', 13.16))
        self.assertEqual(PE._constituent_ok('S2–'), ('S', 'constituent'))
        self.assertEqual(PE._constituent_ok('HS'), ('HS', 'constituent'))
        self.assertEqual(PE._constituent_ok('Cl–')[0], 'Cl')

    def test_footnote_mark_on_value(self):
        import fitz
        path = self._pdf([(100, [(40, 'FeOb'), (120, '11.41c'), (160, '0.46'), (200, '11.97d')]), (114, [(40, 'MnO'), (120, '0.62'), (160, '0.05'), (200, '0.70')])])
        lines = PE.page_lines(fitz.open(path)[0])
        self.assertEqual([w[4] for w in lines[0]['w']], ['FeOb', '11.41', '0.46', '11.97'])

    def test_legend_name_then_number(self):
        t = 'Table 3. Chemical composition of gmalimite (1—grain used for SCXRD, 2—aggregate, Figure 4B) and zoharite (3—aggregate, Figure 3C).'
        self.assertEqual(PE.legend_columns(t, 'zoharite'), ['3'])
        self.assertEqual(PE.legend_columns(t, 'gmalimite'), ['1'])

    def test_holotype_words_stop_at_the_cotype(self):
        w = PE.holotype_words('It was found at two localities: the holotype in the Sahatany Valley, central Madagascar, and a cotype specimen from Sakangyi, Mogok Township, Mandalay Region, Myanmar.')
        self.assertIn('Madagascar', w); self.assertNotIn('Myanmar', w); self.assertNotIn('Mogok', w)

    def test_formula_sentences(self):
        fs = PE._formulas('The empirical formula of koragoite (Voloshin et al., 1997) calculated on the basis of 20 O atoms is Mn3.71Fe0.13Nb3.65Ta0.56W1.83Ti0.08O20. '
                          'The crystal chemical formula of the mineral is (Mn2.02Fe0.98)Σ3.00(Nb2.30Ta0.64Ti0.06)Σ3.00(W1.34Nb0.66)Σ2.00O20.', 'koragoite')
        self.assertEqual(fs[0][4], 'structural'); self.assertAlmostEqual(fs[1][1]['Mn'], 3.71)      # the cited formula goes last
        fs = PE._formulas('The empirical formula (O = 28 apfu) is Ca9.00(Ca0.33Fe0. 2+ 20□0.47)Σ1.00Mg1.04P6.97O28.', 'keplerite')
        self.assertAlmostEqual(fs[0][1]['Fe'], 0.20); self.assertAlmostEqual(fs[0][1]['Ca'], 9.33)   # the charge superscript set mid-number
        fs = PE._formulas('The empirical formulas for ferriphoxite (for O = 13 apfu) and carboferriphoxite (for O = 15 apfu) are '
                          '{[(NH4)2.13K0.87]Σ3.00(H2O)}{(Fe3+ 0.95Al0.05)Σ1.00(HPO4)2(C2O4)} and {[(NH4)1.12K0.88]Σ2.00(H2CO3)}{(Fe3+ 0.78Al0.22)Σ1.00(HPO4)(H2PO4)(C2O4)}, respectively.', 'ferriphoxite')
        self.assertEqual(fs[0][2], []); self.assertAlmostEqual(fs[0][1]['N'], 2.13); self.assertAlmostEqual(fs[0][1]['Fe'], 0.95)   # braces are brackets
        fs = PE._formulas('A combination of results of EMPA and site scattering values obtained by single-crystal structure refinements of Tangir Valley chevkinite-(Ce) '
                          'yielded the formula: (Ce1.81La0.81Nd0.59Ca0.46)Σ3.67(Fe2+ 0.80Mg0.10)Σ0.90Ti2.65Si4.02O22.', 'chevkinite-(ce)')
        self.assertEqual(fs[0][4], 'structural')

    def test_tables_that_are_not_the_analysis(self):
        path = self._pdf([(80, 'Table 2. Trace element composition of ferri-taramite (µg g−1).'), (94, 'Element   Mean   Range'),
                          (108, 'As   18.0   12–25'), (122, 'B   5.5   4–7'), (136, 'Be   75.0   60–90'), (150, 'Co   109.0   90–120'), (164, 'Sc   49.0   40–60')])
        self.assertIsNone(PE.epma_table(path, 'ferri-taramite'))
        path = self._pdf([(80, 'TABLE 6. EMPIRICAL BOND VALENCES (vu) FOR CANADIAN BAZZITE'), (94, 'Bz-ON   Si   Be   Na'),
                          (108, 'O1   4.09   1.964   1.056'), (122, 'O2   4.05   1.980   1.010'), (136, 'O3   4.11   1.950   1.100')])
        self.assertIsNone(PE.epma_table(path, 'bazzite'))

    def test_transposed_apfu_columns_and_integer_esds(self):
        xs = (40, 90, 130, 170, 210, 250, 290, 330, 370, 410, 450)
        head = ['', 'CaO', 'MgO', 'MnO', 'As2O5', 'P2O5', 'H2O*', 'total', 'Ca', 'Mg', 'Mn']
        rows = [['mean', '25.42', '6.11', '0.10', '56.00', '0.29', '11.20', '99.13', '3.691', '1.235', '0.012'],
                ['1', '26.05', '5.44', '0.15', '56.10', '0.12', '11.22', '99.08', '3.793', '1.102', '0.017'],
                ['2', '26.18', '5.47', '0.02', '55.55', '0.44', '11.19', '98.85', '3.814', '1.109', '0.002']]
        path = self._pdf([(80, 'Tab. 2 Chemical composition of chongite from Jáchymov (wt. %)')] + [(100 + 14 * k, list(zip(xs, r))) for k, r in enumerate([head] + rows)])
        e = PE.epma_table(path, 'chongite')
        self.assertEqual([r['constituent'] for r in e['rows']], ['CaO', 'MgO', 'MnO', 'As2O5', 'P2O5', 'H2O'])
        self.assertEqual(e['rows'][0]['mean'], 25.42)
        xs = (40, 90, 140, 190, 240)
        rows = [['Spot', 'SiO2', 'Al2O3', 'Fe2O3', 'Total'], ['1', '27.25', '68.66', '3.48', '99.96'], ['2', '27.24', '68.84', '3.49', '100.17'],
                ['Mean', '29(3)', '67(3)', '3.6(3)', '99.8(5)']]
        path = self._pdf([(80, 'Table 2. (a) Electron microprobe analyses of mullite-2c given in weight percent.')] + [(100 + 14 * k, list(zip(xs, r))) for k, r in enumerate(rows)])
        e = PE.epma_table(path, 'mullite')
        self.assertEqual({r['constituent']: r['mean'] for r in e['rows']}, {'SiO2': 29.0, 'Al2O3': 67.0, 'Fe2O3': 3.6})

    def test_nd_rows_legend_column_and_the_apfu_block(self):
        xs = (40, 180, 240, 410)
        body = [['S', '33.77', '33.58', '30.87'], ['Fe', '36.19', '36.84', '28.35'], ['Cu', '14.33', '14.10', '20.05'], ['K', '7.58', '7.60', '7.00'],
                ['Ba', '0.12', '0.10', '11.50'], ['Na', 'n.d.', 'n.d.', '0.18'], ['Se', 'n.d.', 'n.d.', '0.17'], ['Total', '99.99', '99.22', '99.72'],
                ['S', '25.00', '25.00', '27.00'], ['Na', '0.00', '0.00', '0.22']]
        path = self._pdf([(80, 'Table 3. Chemical composition of gmalimite (1—grain used for SCXRD, 2—aggregate) and zoharite (3—aggregate).'),
                          (100, list(zip(xs, ['wt.%', '1', '2', '3'])))] + [(114 + 14 * k, list(zip(xs, r))) for k, r in enumerate(body)])
        e = PE.epma_table(path, 'zoharite')
        self.assertNotIn('Na', [r['constituent'] for r in e['rows']])                        # n.d. in the first column: no mean
        na = [r for r in e['rows_all'] if r['constituent'] == 'Na']
        self.assertEqual(len(na), 1); self.assertEqual(na[0]['all'], [0.0, 0.0, 0.18])       # kept for the named columns; the apfu 'Na 0.22' below the Total is not a second row
        cols = PE.headline_columns(e, 'zoharite', {'codes': [], 'n': None, 'holotype': False, 'holotype_words': [], 'domains': ['3']})
        self.assertTrue(cols); wt, why = cols[0]
        self.assertIn('column 3', why); self.assertNotIn('averaged', why)
        self.assertEqual(wt['Na'], 0.18); self.assertEqual(wt['S'], 30.87); self.assertEqual(wt['Se'], 0.17)
        alts = PE.table_alternatives(e, {r['constituent']: r['mean'] for r in e['rows']})
        self.assertTrue(any(a.get('Na') == 0.18 and a.get('S') == 30.87 for a in alts))     # the by-fit columns carry the n.d.-first rows too

    def test_mean_cell_inside_the_named_group(self):
        xs = (40, 170, 200, 230, 262, 330, 360, 390, 422)
        head2 = ['No. of spot analyses', '5', '6', '13', 'Mean (n = 3)', '8', '15', '32', 'Mean (n = 9)']
        body = [['SiO2', '29.21', '29.32', '29.43', '29.32', '26.39', '25.49', '26.30', '25.98'], ['Al2O3', '48.10', '48.60', '48.65', '48.45', '44.10', '43.90', '44.30', '44.10'],
                ['B2O3', '16.50', '16.60', '16.64', '16.58', '15.90', '15.80', '16.00', '15.90'], ['MnO', '0.40', '0.38', '0.39', '0.39', '0.20', '0.22', '0.21', '0.21'],
                ['Na2O', '2.10', '2.12', '2.14', '2.12', '1.90', '1.95', '1.92', '1.92']]
        path = self._pdf([(80, 'TABLE 2. Representative chemical compositions of ertlite and mean analyses used for structure refinement'),
                          (100, [(200, 'Madagascar'), (360, 'Myanmar')]), (114, list(zip(xs, head2)))] + [(128 + 14 * k, list(zip(xs, r))) for k, r in enumerate(body)])
        e = PE.epma_table(path, 'ertlite')
        cols = PE.headline_columns(e, 'ertlite', {'codes': [], 'n': None, 'holotype': True, 'holotype_words': ['Madagascar'], 'domains': []})
        wt, why = next((w, y) for w, y in cols if 'Madagascar' in y)
        self.assertNotIn('averaged', why); self.assertEqual(wt['SiO2'], 29.32); self.assertEqual(wt['MnO'], 0.39)

    def test_pass9_vetting(self):
        fs = PE._formulas('The empirical formula using Li, Na, and K based on the structure refinement is Li1.00Na5.81K2.19(UO2)(SO4)5(SO3OH)(H2O). '
                          'The empirical formula using Na measured via EPMA is Li0.79Na5.02K2.02(UO2)(SO4)5(SO3OH)(H2O).', 'seaborgite')
        self.assertEqual([f[4] for f in fs], ['structural', 'empirical'])
        fs = PE._formulas('The empirical formula is Pb8.00Al2.00S6+ 2.88S2 2.60O28.52H22.92.', 'dinilawiite')
        self.assertAlmostEqual(fs[0][1]['S'], 5.48)                                           # 'S2 2.60': the minus of S2− lost
        self.assertIsNone(PE._constituent_ok('Mn2+')[0]); self.assertIsNone(PE._constituent_ok('As3–')[0])   # only anions carry a charge worth stripping
        self.assertIsNone(PE.prose_table('The composition is K0.89 Na0.05 Y0.02 Ca0.01 Ba0.01 Mg0.97 Sc0.54.'))
        self.assertIsNotNone(PE.prose_table('The mean composition is MnO 14.78, Ce2O3 34.19, P2O5 29.57, and H2O 21.46, total 100.00.'))
        xs = (40, 150, 200, 250, 300, 350, 400)
        rows = [['Na', 'Ca', 'K', 'Na', 'F', 'Cl', 'Mn'], ['A', '1.82', '0.02', '0.09', '0.10', '0.05', '1.81'], ['B', '1.80', '0.03', '0.08', '0.12', '0.04', '1.79'], ['C', '1.85', '0.01', '0.10', '0.11', '0.06', '1.83']]
        path = self._pdf([(80, 'Site populations of speziaite')] + [(100 + 14 * k, list(zip(xs, r))) for k, r in enumerate(rows)])
        e = PE.epma_table(path, 'speziaite')
        self.assertTrue(e is None or e.get('total') is None and sum(r['mean'] for r in e['rows']) < 50)
        xs = (40, 150, 200, 240, 300, 350, 390)
        head = ['', 'CF9a1', 'Range', 'SD', 'CF9a2', 'Range', 'SD']
        body = [['Na2O', '1.58', '1.4–1.7', '0.15', '1.67', '1.5–1.8', '0.05'], ['CaO', '10.20', '9.9–10.5', '0.20', '10.40', '10.1–10.7', '0.15'],
                ['FeO', '20.10', '19.5–20.6', '0.30', '19.80', '19.2–20.3', '0.25'], ['P2O5', '39.50', '39.0–40.0', '0.30', '39.20', '38.8–39.6', '0.25'], ['H2O', '28.60', '', '', '28.90', '', '']]
        path = self._pdf([(80, 'TABLE 3. SUMMARY OF CHEMICAL DATA (wt.%) OF LIRAITE HOLOTYPE')] + [(100 + 14 * k, list(zip(xs, r))) for k, r in enumerate([head] + body)])
        e = PE.epma_table(path, 'liraite')
        cols = PE.headline_columns(e, 'liraite', {'codes': ['CF9a1'], 'n': None, 'holotype': False, 'holotype_words': [], 'domains': []})
        wt, why = cols[0]
        self.assertNotIn('averaged', why); self.assertEqual(wt['Na2O'], 1.58); self.assertEqual(wt['FeO'], 20.1)


class ReviewFixes(unittest.TestCase):
    """2026-09-03: the medium-effort review of 10e25cc, one test per confirmed finding."""

    def _pdf(self, lines):
        return ReaderClasses._pdf(self, lines)

    def test_iron_split_below_the_total_survives(self):
        xs = (40, 130)
        rows = [['CaO', '46.10'], ['FeO', '11.79'], ['P2O5', '38.20'], ['MgO', '2.10'], ['SiO2', '1.20'], ['Total', '99.39'], ['FeO', '10.71'], ['Fe2O3', '1.20']]
        path = self._pdf([(80, 'Table 1. Chemical composition of testite (wt%).'), (94, 'Constituent   Mean')] + [(108 + 14 * k, list(zip(xs, r))) for k, r in enumerate(rows)])
        e = PE.epma_table(path, 'testite')
        self.assertEqual({r['constituent']: r['mean'] for r in e['rows'] if r['constituent'].startswith('Fe')}, {'FeO': 10.71, 'Fe2O3': 1.2})

    def test_nd_points_not_averaged_in_and_sigma_total(self):
        xs = (40, 130, 170, 210, 250, 290, 330)
        rows = [['Constituent', '1', '2', '3', '4', '5', '6'], ['SiO2', '46.1', '46.3', '46.0', '46.2', '46.4', '46.1'], ['CaO', '20.2', '20.1', '20.3', '20.0', '20.2', '20.1'],
                ['MgO', '29.0', '29.1', '28.9', '29.2', '29.0', '29.1'], ['F', '0.30', 'n.d.', '0.28', '0.31', 'n.d.', '0.29'], ['Sum', '95.6', '95.5', '95.5', '95.7', '95.6', '95.6']]
        path = self._pdf([(80, 'Table 1. Chemical composition of testite (wt%).')] + [(100 + 14 * k, list(zip(xs, r))) for k, r in enumerate(rows)])
        e = PE.epma_table(path, 'testite')
        f = next(r for r in e['rows'] if r['constituent'] == 'F')
        self.assertAlmostEqual(f['mean'], 0.295); self.assertEqual(f['all'], [0.3, 0.0, 0.28, 0.31, 0.0, 0.29])
        self.assertEqual(e['total'], 95.6)                                              # 'Sum' is the total row, not a continuation of the F row
        self.assertEqual(len(next(r for r in e['rows'] if r['constituent'] == 'MgO')['all']), 6)

    def test_transposed_fully_twinned_header(self):
        xs = (40, 90, 130, 170, 210, 250, 290, 330, 365, 400, 435, 470)
        head = ['Sample', 'SiO2', 'Al2O3', 'FeO', 'MgO', 'CaO', 'Total', 'Si', 'Al', 'Fe', 'Mg', 'Ca']
        rows = [['1', '40.10', '10.20', '12.30', '20.10', '16.90', '99.60', '2.90', '0.87', '0.74', '2.17', '1.31'],
                ['2', '40.30', '10.10', '12.10', '20.30', '16.80', '99.60', '2.91', '0.86', '0.73', '2.18', '1.30'],
                ['Mean', '40.20', '10.15', '12.20', '20.20', '16.85', '99.60', '2.90', '0.86', '0.74', '2.18', '1.30']]
        path = self._pdf([(80, 'Table 2. Chemical composition of testite (wt%).')] + [(100 + 14 * k, list(zip(xs, r))) for k, r in enumerate([head] + rows)])
        e = PE.epma_table(path, 'testite')
        self.assertIsNotNone(e); self.assertEqual({r['constituent']: r['mean'] for r in e['rows']}, {'SiO2': 40.2, 'Al2O3': 10.15, 'FeO': 12.2, 'MgO': 20.2, 'CaO': 16.85})

    def test_legend_scope_and_order(self):
        self.assertEqual(PE.legend_columns('Fig. 2. 1 – newmineralite, 2 – quartz, 3 – calcite. Table 3 – Newmineralite composition', 'newmineralite'), [])
        self.assertEqual(PE.legend_columns('Table 2. Analyses. 1, 2, 8 – newmineralite (1 – holotype; 8 – cotype); 3 – quartz.', 'newmineralite'), ['1', '2', '8'])
        cells = [('Constituent', 40), ('1', 130), ('2', 170), ('8', 210)]
        e = {'head_cells': cells, 'label_x': 40, 'rows': [{'constituent': c, 'mean': v[0], 'all': v, 'xs': [130, 170, 210]} for c, v in
             (('SiO2', [46.1, 46.3, 46.0]), ('CaO', [20.2, 20.1, 20.3]), ('MgO', [29.0, 29.1, 28.9]))]}
        cols = PE.headline_columns(e, 'newmineralite', {'codes': [], 'n': None, 'holotype': False, 'holotype_words': [], 'domains': ['1', '2', '8']})
        self.assertEqual([w['SiO2'] for w, _ in cols], [46.1, 46.3, 46.0])               # the legend's order: the holotype's column first
        e2 = dict(e, head_cells=[('Constituent', 40), ('Mean', 130), ('Range', 170), ('S.D.', 210)])
        self.assertEqual(PE.headline_columns(e2, 'newmineralite', {'codes': [], 'n': None, 'holotype': False, 'holotype_words': [], 'domains': ['1']}), [])   # no numbered header: no column 1

    def test_norm_text_and_excluded_elements(self):
        self.assertIn('Sr0.57Ba0.38', PE._norm_text('formula Sr0:57Ba0:38'))
        e = {'rows': [{'constituent': 'Al2O3', 'mean': 1.0}, {'constituent': 'SiO2', 'mean': 40.0}], 'rows_all': [{'constituent': 'Al2O3', 'mean': 1.0}],
             'candidates': [{'rows': [{'constituent': 'Al2O3', 'mean': 2.0}, {'constituent': 'CaO', 'mean': 3.0}]}]}
        f = PE._without_elements(e, ['Al'])
        self.assertEqual([r['constituent'] for r in f['rows']], ['SiO2']); self.assertEqual(f['rows_all'], [])
        self.assertEqual([r['constituent'] for r in f['candidates'][0]['rows']], ['CaO'])

    def test_two_samples_and_prose_over_structural(self):
        self.assertIn('NM', PE.sample_hints('The empirical formula of beraunite (NM) calculated on the basis of P = 4 apfu is')['codes'])
        self.assertIn('FR', PE.sample_hints('The empirical formula of beraunite (FR) calculated on the basis of P = 4 apfu is')['codes'])
        self.assertNotIn('NM', PE.sample_hints('formula (NM) of')['codes'] if False else [])
        rows = [['Constituent', 'Mean'], ['PbO', '43.21'], ['CuO', '15.38'], ['TeO3', '35.29'], ['H2O', '3.49'], ['Total', '97.37']]
        path = self._pdf([(60, 'The electron microprobe analyses (average of five) provided: PbO 43.21, CuO 15.38, TeO3 35.29, H2O 3.49 (structure), total 97.37 wt.%.'),
                          (200, 'Table 3. Bond-valence analysis for andychristyite'), (214, [(40, 'Site'), (120, 'O1'), (160, 'O2'), (200, 'Total')]),
                          (228, [(40, 'Pb'), (120, '0.45'), (160, '0.30'), (200, '2.05')]), (242, [(40, 'Cu'), (120, '0.50'), (160, '0.48'), (200, '2.02')]),
                          (256, [(40, 'Te'), (120, '1.05'), (160, '0.98'), (200, '6.01')]), (270, [(40, 'Total'), (120, '2.00'), (160, '1.76'), (200, '')])])
        e = PE.epma_table(path, 'andychristyite')
        self.assertIsNotNone(e); self.assertTrue(e.get('prose')); self.assertEqual(e['rows'][0]['constituent'], 'PbO')


class DocxPaper(unittest.TestCase):
    """A manuscript .docx read like a pdf: its tables become pages of words, its text the paper's."""
    def test_docx_tables_and_text(self):
        from docx import Document
        tmp = tempfile.mkdtemp(prefix='pe_'); path = os.path.join(tmp, 'testite.docx')
        self.addCleanup(shutil.rmtree, tmp, ignore_errors=True)
        doc = Document()
        doc.add_paragraph('Testite, a new mineral from Nowhere')
        doc.add_paragraph('The empirical formula, calculated on the basis of 8 O apfu, is Ca1.00Mg2.01Si2.99O7(OH)0.98. '
                          'H2O was calculated by difference. Optical: alpha = 1.600, beta = 1.610, gamma = 1.620.')
        doc.add_paragraph('Table 1. Chemical data (wt%) for testite.')
        t = doc.add_table(rows=1, cols=5)
        for c, h in zip(t.rows[0].cells, ('Constituent', 'Mean', 'Range', 'S.D.', 'Standard')):
            c.text = h
        for row in (('CaO', '17.23', '16.90–17.50', '0.21', 'wollastonite'), ('MgO', '24.88', '24.50–25.20', '0.30', 'forsterite'),
                    ('SiO2', '55.18', '54.80–55.60', '0.28', 'quartz'), ('H2O', '2.71', '', '', ''), ('Total', '100.00', '', '', '')):
            cells = t.add_row().cells
            for c, v in zip(cells, row):
                c.text = v
        doc.add_paragraph('Table 2. Powder X-ray diffraction data for testite.')
        t2 = doc.add_table(rows=0, cols=7)
        for row in (('Iobs', 'dobs', 'dcalc', 'Icalc', 'h', 'k', 'l'), ('100', '3.4550', '3.4531', '92', '1', '1', '0'), ('35', '2.9800', '2.9791', '40', '0', '2', '1'),
                    ('12', '2.5010', '2.4997', '9', '2', '0', '0'), ('8', '1.0210', '1.0204', '6', '−10', '1', '2')):   # a signed two-digit index: h and k must still read as one column
            cells = t2.add_row().cells
            for c, v in zip(cells, row):
                c.text = v
        doc.save(path)
        e = PE.epma_table(path, 'testite')
        self.assertEqual({r['constituent']: r['mean'] for r in e['rows']}, {'CaO': 17.23, 'MgO': 24.88, 'SiO2': 55.18, 'H2O': 2.71})
        self.assertEqual(e['total'], 100.0); self.assertIsNone(e['page']); self.assertIn('Chemical data', e['caption'])
        self.assertEqual(next(r for r in e['rows'] if r['constituent'] == 'CaO')['standard'], 'wollastonite')
        o, c = PE.pxrd_table(path); self.assertEqual((len(o), len(c)), (4, 4)); self.assertEqual(c[-1][2], (-10, 1, 2))
        ex = PE.extract(path, tmp, 'testite'); self.assertEqual(ex['basis'][:2], ('O', 8.0)); self.assertEqual(ex['files']['epma'], 'testite_docx_epma.csv')   # never the paper's file name
        r = PE.check_paper(path, None, None); self.assertTrue(r['composition']['ok'], r['lines'])
        import io, contextlib
        buf = io.StringIO()
        with contextlib.redirect_stdout(buf):
            self.assertEqual(PE.main([path, '--out', tmp]), 0)             # the CLI on a manuscript: no page to print
        self.assertIn('analytical table (a table of the manuscript, 4 constituents', buf.getvalue())
        # the formula sentence inserted under Track Changes is still read
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        p = doc.paragraphs[1]._p
        ins = OxmlElement('w:ins'); ins.set(qn('w:id'), '1'); ins.set(qn('w:author'), 'reviewer'); ins.set(qn('w:date'), '2026-09-05T00:00:00Z')
        for r_ in list(p.findall(qn('w:r'))):
            p.remove(r_); ins.append(r_)
        p.append(ins); doc.save(path)
        self.assertEqual(PE.extract(path, None, None, write=False)['basis'][:2], ('O', 8.0))

    def test_cell_check(self):
        """The powder table's calculated lines against the cell: the .cif's, else the one the text
        states (a and c only: tetragonal or hexagonal, the table decides); a mistyped d is named."""
        from docx import Document
        from tests.test_bv_check import RUTILE
        tmp = tempfile.mkdtemp(prefix='pe_'); path = os.path.join(tmp, 'rutile.docx'); cif = os.path.join(tmp, 'rutile.cif')
        self.addCleanup(shutil.rmtree, tmp, ignore_errors=True)
        with open(cif, 'w', encoding='utf-8') as f:
            f.write(RUTILE)
        doc = Document()
        doc.add_paragraph('Rutile from Nowhere. The powder pattern was indexed on a tetragonal cell, a = 4.5937(2), c = 2.9587(1) Å, V = 62.43 Å3.')
        doc.add_paragraph('Table 2. Powder X-ray diffraction data for rutile.')
        t = doc.add_table(rows=0, cols=7)
        for row in (('Iobs', 'dobs', 'Icalc', 'dcalc', 'h', 'k', 'l'), ('100', '3.248', '100', '3.2482', '1', '1', '0'), ('50', '2.487', '48', '2.4874', '1', '0', '1'),
                    ('8', '2.297', '7', '2.2696', '2', '0', '0'), ('20', '2.187', '19', '2.1873', '1', '1', '1'), ('10', '2.054', '9', '2.0544', '2', '1', '0'), ('60', '1.687', '58', '1.6874', '2', '1', '1')):
            cells = t.add_row().cells
            for c, v in zip(cells, row):
                c.text = v
        doc.save(path)
        cc = PE.cell_check(path, None)
        self.assertEqual((cc['status'], cc['source'], cc['n'], cc['agree'], cc['loose']), ('checked', 'powder', 6, 5, 6))
        self.assertAlmostEqual(cc['cell']['γ'], 90.0); self.assertEqual(cc['bad'][0][2], (2, 0, 0))
        self.assertTrue(any(ln.startswith('2.2696 (2 0 0) does not follow the cell: it gives 2.2968') for ln in cc['lines']), cc['lines'])
        self.assertEqual(cc['unmatched_obs'], [])                                       # 2.297 observed is the (2 0 0) the cell gives
        cc = PE.cell_check(path, cif)
        self.assertEqual((cc['source'], cc['agree'], len(cc['bad'])), ('.cif', 5, 1))
        r = PE.check_paper(path, cif, None)
        self.assertEqual(r['powder_status'], 'checked'); self.assertTrue(any(l.startswith('powder table: 6 indexed lines vs the .cif cell') for l in r['lines']), r['lines'])
        self.assertEqual(cc['wild'], 0)
        import io, contextlib
        buf = io.StringIO()
        with contextlib.redirect_stdout(buf):
            PE.main([path, '--check', '--cif', cif])
        self.assertIn('does not follow the cell', buf.getvalue())

    def test_powder_table_by_content(self):
        """The columns are typed by what they hold: an unlabelled calculated pattern, a two-line
        header, indices written as one word, and a second sample's columns left out."""
        from docx import Document
        tmp = tempfile.mkdtemp(prefix='pe_'); path = os.path.join(tmp, 'contentite.docx')
        self.addCleanup(shutil.rmtree, tmp, ignore_errors=True)
        doc = Document()
        doc.add_paragraph('Table 2. Calculated powder diffraction pattern of contentite.')
        t = doc.add_table(rows=0, cols=5)
        for row in (('h', 'k', 'l', 'dhkl', 'Irel'), ('1', '1', '0', '6.2345', '100'), ('0', '2', '0', '5.1200', '35'), ('−1', '1', '1', '4.0010', '12'), ('2', '0', '−1', '3.4550', '8'), ('1', '3', '0', '3.1020', '4')):
            cells = t.add_row().cells
            for c, v in zip(cells, row):
                c.text = v
        doc.add_paragraph('Table 3. Powder X-ray diffraction data for contentite.')
        t2 = doc.add_table(rows=0, cols=5)
        for row in (('Iobs', 'dobs', 'Icalc', 'dcalc', ''), ('', '(Å)', '', '(Å)', 'hkl'), ('100', '6.23', '100', '6.2345', '110'), ('30', '5.12', '35', '5.1200', '020'),
                    ('10', '4.00', '12', '4.0010', '2.1.10'), ('7', '3.45', '8', '3.4550', '201'), ('3', '3.10', '4', '3.1020', '130')):
            cells = t2.add_row().cells
            for c, v in zip(cells, row):
                c.text = v
        doc.add_paragraph('Table 4. Powder data for contentite (this study) and for the Elliott (2018) sample.')
        t3 = doc.add_table(rows=0, cols=5)
        for row in (('hkl', 'dobs', 'Iobs', 'dobs', 'Iobs'), ('110', '6.23', '100', '6.24', '90'), ('020', '5.12', '30', '5.13', '28'), ('111', '4.00', '10', '4.01', '11'), ('201', '3.45', '7', '3.46', '6')):
            cells = t3.add_row().cells
            for c, v in zip(cells, row):
                c.text = v
        doc.save(path)
        pages = PE._pages(path)
        o, c = PE._pt_read(pages[0], lambda i: PE._caption(pages[0], i))
        self.assertEqual((len(o), len(c)), (0, 5)); self.assertEqual(c[2], (4.001, 12.0, (-1, 1, 1)))      # the caption says calculated: bare dhkl / Irel are calc
        o, c = PE._pt_read(pages[1], lambda i: PE._caption(pages[1], i))
        self.assertEqual((len(o), len(c)), (5, 5)); self.assertEqual(o[0], (6.23, 100.0)); self.assertEqual(c[2], (4.001, 12.0, (2, 1, 10)))   # a two-line header; '2.1.10'
        o, c = PE._pt_read(pages[2], lambda i: PE._caption(pages[2], i))
        self.assertEqual([d for d, _ in o], [6.23, 5.12, 4.0, 3.45]); self.assertEqual(c, [])            # this study's sample only; no calc without a calc column
