"""Unit tests for pxrd_review.bv_check — bond distances, bond-valence sums, manuscript tables.

    python3 -m unittest tests.test_bv_check -v

Rutile (TiO2, P42/mnm) is the reference structure: Ti–O 1.949 ×4 and 1.980 ×2, and the Ti sum is
known (≈ 4.0 vu with Brese & O'Keeffe 1991). A synthetic manuscript table exercises the checker."""
import os, math, shutil, tempfile, unittest

from pxrd_review import bv_check as B

RUTILE = """data_rutile
_chemical_name_mineral rutile
_cell_length_a 4.5937
_cell_length_b 4.5937
_cell_length_c 2.9587
_cell_angle_alpha 90
_cell_angle_beta 90
_cell_angle_gamma 90
_space_group_name_H-M_alt 'P 42/m n m'
loop_
_space_group_symop_operation_xyz
'x, y, z'
'-x, -y, z'
'-y+1/2, x+1/2, z+1/2'
'y+1/2, -x+1/2, z+1/2'
'-x+1/2, y+1/2, -z+1/2'
'x+1/2, -y+1/2, -z+1/2'
'y, x, -z'
'-y, -x, -z'
'-x, -y, -z'
'x, y, -z'
'y+1/2, -x+1/2, -z+1/2'
'-y+1/2, x+1/2, -z+1/2'
'x+1/2, -y+1/2, z+1/2'
'-x+1/2, y+1/2, z+1/2'
'-y, -x, z'
'y, x, z'
loop_
_atom_site_label
_atom_site_type_symbol
_atom_site_fract_x
_atom_site_fract_y
_atom_site_fract_z
_atom_site_occupancy
Ti1 Ti4+ 0 0 0 1
O1 O2- 0.30479 0.30479 0 1
loop_
_geom_bond_atom_site_label_1
_geom_bond_atom_site_label_2
_geom_bond_distance
_geom_bond_site_symmetry_2
Ti1 O1 1.9485(5) .
Ti1 O1 1.9800(5) 3_555
"""


# A synthetic P1 hydrate (8 Å cube): Ca at the origin with a sulfate-like O1 (2.42 Å) and a water
# OW1 (2.40 Å); O2 and O3 sit 2.80 Å from OW1 at 120° to the Ca–OW1 bond, 120° apart — the two
# acceptors a water molecule wants; O1 is 2.61 Å from OW1 but on the Ca polyhedron (an edge, not
# a hydrogen bond). H1 (0.96 Å from OW1 towards O2) is appended for the H-located tests.
HYDRATE = """data_hydrate
_chemical_name_mineral testhydrate
_cell_length_a 8
_cell_length_b 8
_cell_length_c 8
_cell_angle_alpha 90
_cell_angle_beta 90
_cell_angle_gamma 90
_space_group_name_H-M_alt 'P 1'
loop_
_atom_site_label
_atom_site_type_symbol
_atom_site_fract_x
_atom_site_fract_y
_atom_site_fract_z
Ca1 Ca 0 0 0
O1 O 0.275 0.125 0
OW1 O 0 0.3 0
O2 O 0.30311 0.475 0
O3 O 0.69689 0.475 0
"""
HYDRATE_H = HYDRATE + "H1 H 0.10392 0.36 0\n"


def _write(tmp, name, text):
    p = os.path.join(tmp, name)
    with open(p, 'w', encoding='utf-8') as f:
        f.write(text)
    return p


class Symmetry(unittest.TestCase):
    def test_parse_symop(self):
        rot, tr = B.parse_symop('-y+1/2, x-y, z+0.25')
        self.assertEqual(rot, [[0, -1, 0], [1, -1, 0], [0, 0, 1]])
        self.assertEqual(tr, [0.5, 0, 0.25])
        rot, tr = B.parse_symop('x, y, z')
        self.assertEqual(rot, [[1, 0, 0], [0, 1, 0], [0, 0, 1]])
        with self.assertRaises(ValueError):
            B.parse_symop('x, y')
        with self.assertRaises(ValueError):
            B.parse_symop('x, y, __import__')

    def test_element_and_charge(self):
        for sym, want in (('Fe3+', ('Fe', 3)), ('Bi+3', ('Bi', 3)), ('S-2', ('S', -2)), ('Cl-', ('Cl', -1)),
                          ('O2-', ('O', -2)), ('Al', ('Al', None))):
            self.assertEqual(B._element_of('X1', sym), want, sym)
        self.assertEqual(B._element_of('OH1', '')[0], 'O')
        self.assertEqual(B._element_of('OW2', '')[0], 'O')
        self.assertEqual(B._element_of('LiY', '')[0], 'Li')


class Rutile(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.tmp = tempfile.mkdtemp(prefix='bv_')
        cls.cif = _write(cls.tmp, 'rutile.cif', RUTILE)
        cls.st = B.Structure(cls.cif)

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.tmp, ignore_errors=True)

    def test_sites_and_multiplicity(self):
        st = self.st
        self.assertEqual(st.n_ops, 16)
        ti = next(s for s in st.sites if s.label == 'Ti1'); o = next(s for s in st.sites if s.label == 'O1')
        self.assertEqual((ti.mult, o.mult), (2, 4))
        self.assertEqual([sp.ox for sp in ti.species], [4])
        self.assertAlmostEqual(st.volume, 62.43, places=1)

    def test_distances(self):
        ti = next(s for s in self.st.sites if s.label == 'Ti1')
        nb = [(round(d, 4), n) for o, d, n in self.st.neighbours(ti, 2.5) if o.label == 'O1']
        self.assertEqual(nb, [(1.9485, 4), (1.9801, 2)])

    def test_bvs_and_selfcheck(self):
        P = B.Params(prefer='bo')
        result, anion_sum, cells, hbonds = B.compute(self.st, P)
        self.assertEqual(hbonds, [])                         # no hydrogen anywhere
        site, bonds, bvs, expected, mean_d = result[0]
        self.assertAlmostEqual(bvs, 4.07, places=2)          # Ti4+–O R0 1.815, b 0.37
        self.assertAlmostEqual(anion_sum['O1'], bvs / 2, places=6)  # 2 Ti per 4 O: each O gets half a Ti's sum
        # a half-occupied O: the Ti column sees it half the time, its own row still gets the whole bond
        half = _write(self.tmp, 'half.cif', RUTILE.replace('O1 O2- 0.30479 0.30479 0 1', 'O1 O2- 0.30479 0.30479 0 0.5'))
        r2, a2, c2, _ = B.compute(B.Structure(half), P)
        self.assertAlmostEqual(r2[0][2], bvs / 2, places=6)
        self.assertAlmostEqual(a2['O1'], anion_sum['O1'], places=6)
        self.assertEqual(cells[('O1', 'Ti1')][0][1:], (4, 2))  # 1.9485 ×4 into Ti, ×2 into O
        self.assertIn('consistent', B.geom_self_check(self.st, result))
        text = B.report_text(self.st, P, result, anion_sum, cells)
        self.assertIn("R0 1.815  b 0.370", text)
        self.assertIn('0.70×4↓×2→', text)

    def test_gh_default_and_override(self):
        P = B.Params()                                       # gh
        self.assertEqual(P.get('Ti', 4, 'O', -2)[2], 'bs')
        self.assertEqual(P.get('U', 6, 'O', -2)[2], 'r')     # Burns et al. 1997 for uranyl
        st = B.Structure(self.cif, ox_override={'Ti': 3})
        self.assertEqual(st.cations[0].species[0].ox, 4)     # the .cif's own 'Ti4+' wins over the default
        cif2 = _write(self.tmp, 'rutile2.cif', RUTILE.replace('Ti4+', 'Ti').replace('O2-', 'O'))
        st2 = B.Structure(cif2, ox_override={'Ti': 3})
        self.assertEqual(st2.cations[0].species[0].ox, 3)    # --ox applies when the .cif says nothing

    def test_run_writes_report_and_word(self):
        out = os.path.join(self.tmp, 'out')
        st, result, anion_sum, cells, text = B.run(self.cif, params='bo', word=True, out_dir=out, quiet=True)
        self.assertTrue(os.path.exists(os.path.join(out, 'rutile_bv.txt')))
        from docx import Document
        d = Document(os.path.join(out, 'rutile_bv.docx'))
        self.assertEqual(len(d.tables), 2)
        self.assertEqual(d.tables[1].rows[1].cells[0].text, 'O1')


class HydrogenBonds(unittest.TestCase):
    def test_ferraris_ivaldi_and_labels(self):
        self.assertAlmostEqual(B.fi_valence(2.55), 0.33, places=2)      # the paper's own anchor
        self.assertAlmostEqual(B.fi_valence(2.926), 0.146, places=3)    # szilagyiite O1 ← OW1: 0.15 in the owner's table
        self.assertAlmostEqual(B.fi_valence(2.657), 0.250, places=3)    # O6 ← OW1: 0.25
        for lab, n in (('OH1', 1), ('Oh2', 1), ('OW1', 2), ('Ow3', 2), ('W4', 2), ('Wat1', 2), ('O6H', 1), ('O7W', 2), ('F1/OH1', 1), ('O5', 0), ('OH', 1)):
            self.assertEqual(B.label_h(lab), n, lab)

    def test_reference_names_and_note(self):
        P = B.Params(prefer='bo')
        self.assertEqual(P.short_ref('a'), "Brese and O'Keeffe (1991)")   # BO91 reprints every BA85 value
        self.assertEqual(B.Params(prefer='ba').short_ref('a'), 'Brown and Altermatt (1985)')
        self.assertEqual(P.short_ref('s'), 'García-Rodríguez et al. (2000)')
        self.assertEqual(P.short_ref('bo'), 'Nyman et al. (2010)')       # derived from the file's citation
        P = B.Params()
        P.get('Ti', 4, 'O', -2); P.get('U', 6, 'O', -2); P.get('NH', 1, 'O', -2)
        self.assertEqual(P.note(), 'Bond-valence parameters from Gagné and Hawthorne (2015); U6+–O from Burns et al. (1997); '
                                   'NH4+–O from García-Rodríguez et al. (2000)')
        self.assertEqual(B.Params(u6='params').get('U', 6, 'O', -2)[2], 'bs')

    def test_blind_proposal_from_oo_geometry(self):
        tmp = tempfile.mkdtemp(prefix='bv_')
        try:
            cif = _write(tmp, 'h.cif', HYDRATE)
            st = B.Structure(cif); P = B.Params()
            result, anion_sum, cells, hb = B.compute(st, P)
            pairs = sorted((x.donor.label, x.acceptor.label, round(x.d, 2), round(x.s, 3), x.via) for x in hb)
            s = round(B.fi_valence(2.8), 3)
            self.assertEqual(pairs, [('OW1', 'O2', 2.8, s, 'OO'), ('OW1', 'O3', 2.8, s, 'OO')])   # not O1: a Ca polyhedron edge
            self.assertAlmostEqual(anion_sum['O2'], s, places=3)                 # Σan = cations + accepted
            self.assertAlmostEqual(B.donated(hb)['OW1'], 2 * (1 - B.fi_valence(2.8)), places=3)   # the O–H part, reported apart
            self.assertFalse([n for n in st.notes if 'deficit' in n])            # donors came from the label
            # overrides: one H only; a pair forced across the polyhedron edge
            st = B.Structure(cif)
            _, _, _, hb1 = B.compute(st, P, donors={'OW1': 1})
            self.assertEqual(len(hb1), 1)
            st = B.Structure(cif)
            _, _, _, hb2 = B.compute(st, P, force=[('OW1', 'O1')])
            self.assertIn(('OW1', 'O1'), [(x.donor.label, x.acceptor.label) for x in hb2])
            # no labels: the valence deficit decides, and says so
            cif2 = _write(tmp, 'h2.cif', HYDRATE.replace('OW1 O', 'O4 O'))
            st = B.Structure(cif2)
            _, _, _, hb3 = B.compute(st, P)
            self.assertIn('O4', {x.donor.label for x in hb3})               # (O2/O3 have no cation at all: 'water' too)
            self.assertTrue(any('deficit' in n and 'O4' in n for n in st.notes))
            self.assertEqual(B.compute(B.Structure(cif), P, hbond='none')[3], [])
            text = B.report_text(st, P, *B.compute(st, P)[:3], hbonds=hb3)
            self.assertIn('HYDROGEN BONDS', text); self.assertIn('Ferraris', text); self.assertIn('table note:', text)
        finally:
            shutil.rmtree(tmp, ignore_errors=True)

    def test_with_located_h(self):
        tmp = tempfile.mkdtemp(prefix='bv_')
        try:
            cif = _write(tmp, 'hh.cif', HYDRATE_H)
            st = B.Structure(cif); P = B.Params()
            geo = B.hbond_geometry(st)                       # computed: OW1–H1⋯O2, 180°
            self.assertEqual([(g['labels'], g['ang']) for g in geo], [(('OW1', 'H1', 'O2'), '180')])
            result, anion_sum, cells, hb = B.compute(st, P)  # 'oo': strengths from D⋯A, H not a column
            self.assertEqual([(x.donor.label, x.acceptor.label, x.via) for x in hb], [('OW1', 'O2', 'H')])
            self.assertAlmostEqual(hb[0].s, B.fi_valence(2.8), places=3)
            self.assertNotIn('H1', [r[0].label for r in result])
            result, anion_sum, cells, hb = B.compute(st, P, hbond='h')   # the older convention: H as a cation
            self.assertIn('H1', [r[0].label for r in result])
            got = {(x.donor.label, x.acceptor.label, x.via): x.s for x in hb}
            self.assertIn(('OW1', 'O2', 'H···A'), got)     # (the older mode takes every O within 2.4 Å of the H, O1 included)
            self.assertAlmostEqual(got[('OW1', 'O2', 'H···A')], math.exp((0.990 - 1.84) / 0.59), places=2)   # Brown 2002, H⋯O 1.84 Å
        finally:
            shutil.rmtree(tmp, ignore_errors=True)


class ManuscriptTable(unittest.TestCase):
    def test_bond_and_bvs_tables(self):
        from docx import Document
        tmp = tempfile.mkdtemp(prefix='bv_')
        try:
            cif = _write(tmp, 'rutile.cif', RUTILE)
            doc = Document()
            t = doc.add_table(rows=0, cols=2)
            for a, b in (('Ti1–O1 ×4', '1.949(1)'), ('Ti1–O1 ×2', '1.985(1)'), ('<Ti1–O>', '1.965')):
                r = t.add_row().cells; r[0].text = a; r[1].text = b
            t2 = doc.add_table(rows=0, cols=3)
            for row in (('Atom', 'Ti1', 'Σ'), ('O1', '0.70×4↓×2→, 0.64×2↓', '2.03'), ('Σ', '4.07', '')):
                r = t2.add_row().cells
                for i, x in enumerate(row):
                    r[i].text = x
            path = os.path.join(tmp, 'ms.docx'); doc.save(path)
            st = B.Structure(cif); P = B.Params(prefer='bo')
            result, anion_sum, cells, _ = B.compute(st, P)
            tables = B.read_tables(path)
            bonds = B.check_bond_table(st, result, tables)
            self.assertIn('1 bond distances agree with the .cif, 1 do not', bonds[0])
            self.assertTrue(any('1.985 but the .cif gives 1.980' in x for x in bonds), bonds)
            self.assertTrue(any('<Ti1–O> given as 1.965' in x and 'average 1.961' in x for x in bonds), bonds)
            bvs = B.check_bvs_table(st, result, cells, anion_sum, tables, 'test')
            self.assertIn('1 cells compared, 0 disagree', bvs[0])
            self.assertEqual([x for x in bvs[1:] if 'Σ' in x], [], bvs)
        finally:
            shutil.rmtree(tmp, ignore_errors=True)


if __name__ == '__main__':
    unittest.main()
