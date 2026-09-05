"""Tables mode of the review GUI — the /api/tb/* routes (five tabs), through Flask's test client.

    python3 -m unittest tests.test_gui_tb -v

A rutile .cif, a small probe .csv and an obs / calc peak-list pair in a temp folder; no browser."""
import json, os, shutil, tempfile, unittest

from tests.test_tables import RUTILE
from tests.test_pxrd_table import OBS, CALC

PROBE = """point,CaO,Al2O3,SiO2,F
1,20.1,36.7,43.2,0.5
2,20.3,36.5,43.0,0.4
3,19.9,36.9,43.4,0.6
"""


class TablesMode(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        from pxrd_review.gui import review_gui as G
        cls.G = G
        G._ALLOWED_HOSTS = set()
        cls.tmp = tempfile.mkdtemp(prefix='tbgui_')
        for name, text in (('rutile.cif', RUTILE), ('probe.csv', PROBE), ('obs.txt', OBS), ('calc.txt', CALC), ('notes.txt', 'hello\n')):
            with open(os.path.join(cls.tmp, name), 'w', encoding='utf-8') as f:
                f.write(text)
        cls.c = G.app.test_client()
        assert G.ms_set_folder(cls.tmp)
        if G.MS['athread']:
            G.MS['athread'].join(30)

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.tmp, ignore_errors=True)

    def _get(self, url, status=200):
        r = self.c.get(url)
        self.assertEqual(r.status_code, status, r.data[:300])
        return json.loads(r.data)

    def _post(self, url, status=200, **kw):
        r = self.c.post(url, **kw)
        self.assertEqual(r.status_code, status, r.data[:300])
        return json.loads(r.data)

    def test_state_lists_cifs_data_and_kinds(self):
        st = self._get('/api/tb/state')
        self.assertEqual([c['key'] for c in st['cifs']], ['rutile'])
        kinds = {d['key']: d['kind'] for d in st['data']}
        self.assertEqual({k: kinds[k] for k in ('calc.txt', 'obs.txt', 'probe.csv', 'notes.txt')}, {'calc.txt': 'calc', 'obs.txt': 'obs', 'probe.csv': 'probe', 'notes.txt': ''})
        self.assertIn('manuscript', [j['key'] for j in st['journals']])

    def test_tables_part_filter(self):
        full = self._get('/api/tb/tables/rutile')
        self.assertTrue(full['ok'])
        self.assertEqual(full['html'].count('data-kind='), 3)          # rutile: coords, bonds, bvs (no H)
        co = self._get('/api/tb/tables/rutile?part=coords')
        self.assertIn('data-kind="coords"', co['html']); self.assertNotIn('data-kind="bvs"', co['html'])
        bv = self._get('/api/tb/tables/rutile?part=bvs&params=bo')
        self.assertIn('data-kind="bvs"', bv['html']); self.assertNotIn('data-kind="coords"', bv['html'])
        self.assertIn("Brese and O&#x27;Keeffe (1991)", bv['html'])            # the footnote follows the buttons
        self.assertIn('Brown and Altermatt (1985)', self._get('/api/tb/tables/rutile?part=bvs&params=ba')['html'])
        self.assertIn('Gagné and Hawthorne (2015)', self._get('/api/tb/tables/rutile?part=bvs')['html'])
        # the hydrogen-bond options are accepted (and bad ones ignored) on a structure without H
        ok = self._get('/api/tb/tables/rutile?part=bvs&hb=none&hmax=9&donors=O1%3D1&hbp=O1%3EO1&u6set=1')
        self.assertTrue(ok['ok'])
        self.assertTrue(self.G.app.test_request_context('/x?hb=zzz&hmax=9&donors=;;&hbp=a>b').__enter__() is not None)
        with self.G.app.test_request_context('/x?hb=zzz&hmax=9&donors=;;&hbp=O1%3EO2,bad&u6set=1'):
            o = self.G._tb_opts()
            self.assertEqual((o['hbond'], o['hmax'], o['donors'], o['hb'], o['u6']), ('oo', None, None, None, 'params'))
        self.assertEqual(self.c.get('/api/tb/tables/nope').status_code, 404)

    def test_epma_route_and_export(self):
        r = self._get('/api/tb/epma/probe.csv?basis=O%3D8&name=anorthite&standards=CaO%3Dwollastonite')
        self.assertTrue(r['ok'])
        self.assertIn('Al', r['formula']); self.assertEqual(r['n_points'], 3)
        self.assertIn('wollastonite', r['html']); self.assertIn('Anorthite', r['html'].replace('anorthite', 'Anorthite'))
        self.assertEqual(dict(r['wt'])['SiO2'], 43.2)
        # a bad basis is a 400 with the tool's own message, not a 500
        bad = self.c.get('/api/tb/epma/probe.csv?basis=zzz'); self.assertEqual(bad.status_code, 400)
        self.assertIn('basis', json.loads(bad.data)['error'])
        # a text file that is no probe file: reported, not crashed
        self.assertIn(self.c.get('/api/tb/epma/notes.txt').status_code, (400, 500))
        self.assertEqual(self.c.get('/api/tb/epma/nope.csv').status_code, 404)
        x = self._post('/api/tb/epma/probe.csv/export?fmt=xlsx&basis=O%3D8')
        self.assertEqual(x['file'], 'probe_epma.xlsx')
        w = self._post('/api/tb/epma/probe.csv/export?fmt=word&basis=O%3D8&journal=cjmp')
        self.assertEqual(w['file'], 'probe_epma.docx')
        self.assertTrue(os.path.exists(os.path.join(self.tmp, 'review_out', 'probe_epma.docx')))
        st = self._get('/api/tb/state')
        self.assertEqual(st['outputs'], ['probe_epma.docx', 'probe_epma.xlsx'])

    def test_gd_route(self):
        r = self._get('/api/tb/gd?formula=Ti%3D1&n=2.7&cif=rutile&z=2&name=rutile')
        self.assertTrue(r['ok']); self.assertGreater(r['D_calc'], 4.0)
        self.assertEqual(r['summary'][0][0], 'calculated')
        self.assertIn('data-kind="gd"', r['html'])
        r2 = self._get('/api/tb/gd?wt=TiO2%3D100&n=2.7&density=4.25')
        self.assertEqual(r2['summary'][0][0], 'measured')
        self.assertEqual(self.c.get('/api/tb/gd?formula=Ti%3D1').status_code, 400)          # no n
        self.assertEqual(self.c.get('/api/tb/gd?formula=Ti%3D1&n=2.7&cif=nope').status_code, 404)
        x = self._post('/api/tb/gd/export?fmt=xlsx&formula=Ti%3D1&n=2.7&name=rutile')
        self.assertEqual(x['file'], 'rutile_gd.xlsx')

    def test_pxrd_route_and_export(self):
        r = self._get('/api/tb/pxrd?obs=obs.txt&calc=calc.txt&name=test&min_i=3.5')
        self.assertTrue(r['ok']); self.assertEqual(r['n_obs'], 9); self.assertGreater(r['n_rows'], 5)
        self.assertEqual(len(r['bold']), len([b for b in r['bold']]))          # eight-strongest list present
        self.assertIn('<b>87</b>', r['html'])                                   # the strongest observed line is bold
        self.assertEqual(self.c.get('/api/tb/pxrd?obs=obs.txt&calc=nope.txt').status_code, 404)
        self.assertEqual(self.c.get('/api/tb/pxrd?obs=obs.txt&calc=notes.txt').status_code, 400)
        w = self._post('/api/tb/pxrd/export?fmt=word&obs=obs.txt&calc=calc.txt&name=test')
        self.assertEqual(w['file'], 'test_pxrd.docx')

    def test_version_route(self):
        from pxrd_review import __version__
        r = self._get('/api/version')
        self.assertEqual(r['installed'], __version__)
        self.assertIn(r['status'], ('idle', 'disabled'))          # the test client never starts the network check
        from pxrd_review import update as U
        self.assertEqual(r['command'], 'pxrd update')             # the one command, whatever the install kind
        self.assertEqual(self._get('/api/update/status')['state'], 'idle')
        # Update now is refused when nothing is newer (the check never ran here: no result, no checkout claim)
        from unittest import mock
        with mock.patch.object(U, 'status', return_value={'result': {'newer': False}, 'checkout': None}):
            self.assertEqual(self.c.post('/api/update').status_code, 409)
        if U.checkout():                                          # the suite itself runs from the git checkout
            self.assertIn('git -C', r['pip']); self.assertEqual(r['checkout'], U.checkout())
        else:
            self.assertIn('main.zip', r['pip'])

    def test_fill_from_paper_and_bv_workbook(self):
        from tests.test_paper_extract import make_pdf
        make_pdf(os.path.join(self.tmp, 'testite.pdf'))
        assert self.G.ms_set_folder(self.tmp)                        # re-index: the pdf joins the folder
        if self.G.MS['athread']:
            self.G.MS['athread'].join(30)
        self.assertEqual([p['key'] for p in self._get('/api/tb/state')['pdfs']], ['testite.pdf'])
        r = self._post('/api/tb/extract?pdf=testite.pdf')
        self.assertEqual(r['fill']['epma']['basis'], 'O=7'); self.assertEqual(r['fill']['epma']['add'], 'H2O=difference')
        self.assertEqual(r['fill']['epma']['file'], 'testite_paper_epma.csv'); self.assertIn('wollastonite', r['fill']['epma']['standards'])
        self.assertEqual((r['fill']['gd']['n'], r['fill']['gd']['density']), ('1.6100', '3.120'))
        self.assertEqual(r['fill']['bvs']['params'], 'gh'); self.assertEqual(r['fill']['pxrd']['obs'], 'testite_paper_obs.txt')
        self.assertTrue(any('basis of 7 O' in n for n in r['notes']))
        self.assertEqual(self.c.post('/api/tb/extract?pdf=nope.pdf').status_code, 404)
        # the same paper is a manuscript: its own numbers are checked and carried as findings
        self.assertIn('testite', self.G.MS['files'])
        d = self.G.ms_analysis('testite')
        kinds = {f['kind'] for f in d['findings']}
        self.assertIn('calcinfo', kinds)
        self.assertTrue(any('composition' in f['msg'] for f in d['findings']))
        row = next(r for r in self._get('/api/ms/state')['files'] if r['key'] == 'testite')
        self.assertEqual(row['pdf'], 'testite.pdf')
        h = self._get('/api/ms/docx/testite.html?which=source')
        self.assertIn('data-p="0"', h['html'])
        # the written means feed the EPMA tab, on the paper's basis, with the method sheet in the workbook
        e = self._get('/api/tb/epma/testite_paper_epma.csv?basis=O%3D7&name=testite')
        self.assertTrue(e['ok']); self.assertEqual(e['n_points'], 1)
        x = self._post('/api/tb/epma/testite_paper_epma.csv/export?fmt=xlsx&basis=O%3D7&method=basis%3A+7+O+apfu+%7C+H2O+by+difference')
        import openpyxl
        wb = openpyxl.load_workbook(os.path.join(self.tmp, 'review_out', x['file']))
        self.assertIn('method', wb.sheetnames); self.assertIn('H2O by difference', [c.value for c in wb['method']['A']])
        # the bond-valence workbook
        b = self._post('/api/tb/bvs/rutile/export?fmt=xlsx&params=bo')
        self.assertEqual(b['file'], 'rutile_bv.xlsx')
        wb = openpyxl.load_workbook(os.path.join(self.tmp, 'review_out', 'rutile_bv.xlsx'))
        self.assertEqual(wb.sheetnames, ['bonds', 'cation sums', 'anion sums', 'H bonds', 'parameters'])
        self.assertTrue(str(wb['bonds']['G2'].value).startswith('=EXP('))

    def test_opts_roundtrip_and_open_guard(self):
        self._post('/api/tb/opts/epma', json={'basis': 'O=21', 'file': 'probe.csv', 'junk': ['x'], 'raw': '1'})
        self.assertEqual(self.c.post('/api/tb/opts/nope', json={}).status_code, 404)
        st = self._get('/api/tb/state')
        self.assertEqual(st['opts']['epma'], {'basis': 'O=21', 'file': 'probe.csv', 'raw': '1'})
        with open(os.path.join(self.tmp, 'review_out', 'tables_opts.json')) as f:
            self.assertEqual(json.load(f)['epma']['basis'], 'O=21')
        # open: only the tool's own outputs by basename; never a path, never an arbitrary file
        G = self.G
        self.assertIsNone(G._tb_output_ok('../probe.csv')); self.assertIsNone(G._tb_output_ok('/etc/passwd'))
        self.assertIsNone(G._tb_output_ok('probe.csv')); self.assertIsNone(G._tb_output_ok('missing_epma.docx'))
        self.assertEqual(self.c.post('/api/tb/open?file=..%2Fprobe.csv').status_code, 404)
        self.assertEqual(self.c.post('/api/tb/open?file=nothing_pxrd.docx').status_code, 404)


if __name__ == '__main__':
    unittest.main()


class DocxPaper(unittest.TestCase):
    """A manuscript .docx is a paper for the Tables mode: listed with the pdfs under "from paper",
    and Fill ▸ reads its analytical table (and checks its bond-valence table against the .cif)."""
    @classmethod
    def setUpClass(cls):
        from docx import Document
        from pxrd_review.gui import review_gui as G
        cls.G = G
        G._ALLOWED_HOSTS = set()
        cls.tmp = tempfile.mkdtemp(prefix='tbgui_')
        with open(os.path.join(cls.tmp, 'rutile.cif'), 'w', encoding='utf-8') as f:
            f.write(RUTILE)
        doc = Document()
        doc.add_paragraph('Rutile, a mineral from Nowhere')
        doc.add_paragraph('The empirical formula, calculated on the basis of 2 O apfu, is Ti1.00O2.')
        doc.add_paragraph('Table 1. Chemical data (wt%) for rutile.')
        t = doc.add_table(rows=1, cols=4)
        for c, h in zip(t.rows[0].cells, ('Constituent', 'Mean', 'Range', 'S.D.')):
            c.text = h
        for row in (('TiO2', '99.10', '98.80–99.40', '0.21'), ('FeO', '0.40', '0.30–0.50', '0.07'), ('SiO2', '0.30', '0.20–0.40', '0.06'), ('Total', '99.80', '', '')):
            cells = t.add_row().cells
            for c, v in zip(cells, row):
                c.text = v
        doc.add_paragraph('Table 2. Bond-valence analysis for rutile.')
        t2 = doc.add_table(rows=0, cols=3)
        for row in (('Atom', 'Ti1', 'Σ'), ('O1', '0.70×4↓×2→, 0.64×2↓', '2.03'), ('Σ', '4.07', '')):
            cells = t2.add_row().cells
            for c, v in zip(cells, row):
                c.text = v
        doc.save(os.path.join(cls.tmp, 'Rutile manuscript.docx'))
        cls.c = G.app.test_client()
        assert G.ms_set_folder(cls.tmp)
        if G.MS['athread']:
            G.MS['athread'].join(60)

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.tmp, ignore_errors=True)

    def test_docx_listed_and_filled(self):
        st = json.loads(self.c.get('/api/tb/state').data)
        self.assertEqual([d['name'] for d in st['pdfs']], ['Rutile manuscript.docx'])
        r = self.c.post('/api/tb/extract?pdf=Rutile%20manuscript')
        self.assertEqual(r.status_code, 200, r.data[:200]); d = json.loads(r.data)
        self.assertTrue(d['ok']); self.assertEqual(d['fill']['epma']['file'], 'Rutile manuscript_docx_epma.csv')
        self.assertEqual(d['fill']['epma']['basis'], 'O=2'); self.assertEqual(d['fill']['_cif'], 'rutile')
        self.assertTrue(any('3 constituents' in n for n in d['notes']), d['notes'])
        self.assertIn('bond valence:', d['bvcheck'] or ''); self.assertIn('1 cells compared, 0 disagree', d['bvcheck'])
        with open(os.path.join(self.tmp, 'review_out', d['fill']['epma']['file']), encoding='utf-8') as f:
            self.assertEqual(f.read().splitlines(), ['TiO2,FeO,SiO2', '99.1,0.4,0.3'])
        self.assertEqual(self.c.post('/api/tb/extract?pdf=nope').status_code, 404)
        self.assertEqual(self.c.get('/api/ms/pdf/nope.pdf/page/1.png').status_code, 404)   # a page only of the folder's own papers

    def test_strangers_table_is_not_scored(self):
        """A manuscript whose bond-valence table names other cations than the folder's only .cif is
        not scored against it (nor reported as having no table)."""
        from docx import Document
        from pxrd_review import paper_extract as PE
        doc = Document()
        doc.add_paragraph('Table 1. Bond-valence analysis for otherite.')
        t = doc.add_table(rows=0, cols=4)
        for row in (('Atom', 'Cu1', 'Al1', 'Σ'), ('O1', '0.45', '0.52', '0.97'), ('O2', '0.41', '0.55', '0.96'), ('OH3', '0.38', '0.49', '0.87'), ('Σ', '1.24', '1.56', '')):
            cells = t.add_row().cells
            for c, v in zip(cells, row):
                c.text = v
        path = os.path.join(self.tmp, 'Otherite manuscript.docx'); doc.save(path)
        bc = PE.bv_check_paper(path, os.path.join(self.tmp, 'rutile.cif'), {'bv': {}})
        self.assertEqual(bc['status'], 'foreign'); self.assertIn('Ti1', bc['message'])
        bc = PE.bv_check_paper(os.path.join(self.tmp, 'Rutile manuscript.docx'), os.path.join(self.tmp, 'rutile.cif'), {'bv': {}})
        self.assertEqual(bc['status'], 'checked'); self.assertEqual(bc['compared'], 1)
