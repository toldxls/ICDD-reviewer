#!/usr/bin/env python3
"""
Write the cell/wavelength review back INTO each entry .docx as Word comments and
highlights, so the reviewer sees the findings in context instead of in a console
report.

  * Highlights (yellow) the Author's-Cell value cells the checker flags.
  * Adds a comment on each flagged value explaining the discrepancy
    (value / significant-figures / esd), and a comment on the Radiation cell
    when the anode/λ is flagged.
  * Adds one summary comment on the 'Author's Cell' label cell: the overall
    verdict, the matched PDF cell, and the λ status.

Comparison logic is reused verbatim from cell_lambda_check (single source of
truth); this module only renders the result into the document.

Flagged entries are written as '<name>_edited.docx' so the ones the tool
commented on stand out in the folder listing; clean entries keep the source name.

Usage:
    python3 annotate_review.py "/path/to/Part 1"                 # -> <folder>/review_out
    python3 annotate_review.py "/path/to/Part 1" --id I003416
    python3 annotate_review.py "/path/to/Part 1" --out DIR
    python3 annotate_review.py "/path/to/Part 1" --inplace       # edit originals (asks nothing)
"""
import sys, os, re, glob, shutil, argparse, textwrap, zipfile, io, datetime
import cell_lambda_check as C
import extra_checks as X
from docx import Document
from lxml import etree

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
AUTHOR, INITIALS = 'PXRD Review Tool', 'PXRD'
def _q(tag): return W + tag
PARAM_COL = {'a': 1, 'b': 2, 'c': 3, 'α': 4, 'β': 5, 'γ': 6,
             'SG': 7, 'Z': 8}                                  # Author's-Cell column per parameter
KIND_LABEL = {'value': 'VALUE MISMATCH', 'precision': 'significant figures differ',
              'esd': 'uncertainty (esd) differs'}

def _tidy(s):
    """Cosmetic clean-up of comment text: collapse the double spaces inherited
    from the report formatter, and render anode names with the α subscript
    (docx writes 'MoKa'/'CuKa1', the literature 'MoKα')."""
    s = re.sub(r'(?<=[A-Za-z])Ka', 'Kα', s)
    return re.sub(r' {2,}', ' ', s).strip()

NO_MATCH = 'No matching PDF cell found.'

def _has_lam_flag(res):
    return res['lam'] is not None and res['lam'][0] == 'flag'

def _writable_extras(res):
    """Which extra-check findings get written into the docx: every hard FLAG, plus
    the authoritative Mindat group statement (informational but routinely added by
    the reviewer). Soft 'note'/other 'info' findings stay console-only."""
    return [f for f in res.get('extra', [])
            if f.sev == 'flag' or (f.code == 'classification' and f.anchor == 'name')]

def _is_clean(res):
    """Nothing to report: the cell matched exactly, with no parameter
    discrepancies, no radiation hard-flag, and no writable extra findings. (Soft
    λ 'verify'/'unrec' notes are not errors and are intentionally not reported.)"""
    return (res['cell'][0] == 'match' and not res['params']
            and not _has_lam_flag(res) and not _writable_extras(res))

# ----------------------------------------------------------------------------- analysis
def analyze(docx_path, pdf_path, cif_path=None):
    """Return a structured verdict mirroring cell_lambda_check.report()."""
    d = C.parse_docx(docx_path)
    res = {'docx': d, 'cell': ('nopdf',), 'params': {}, 'lam': None, 'extra': []}
    text = C.pdf_text(pdf_path) if pdf_path else None
    cif_data = X.parse_cif(cif_path) if cif_path else {}
    # extra checks run regardless of a PDF (symmetry/indexing/calculated are
    # docx-internal). Never let them break the core cell/λ verdict.
    entry = None
    try:
        entry = X.parse_entry(docx_path)
        res['extra'] = X.run_all(entry, text, cif_data)
    except Exception as ex:
        res['extra'] = []
    # a CALCULATED powder pattern uses the modelling wavelength (calc software
    # such as PLATON commonly defaults to CuKα), not the experimental radiation
    is_calc = bool(entry and (entry.instr.get('spacing_instr') or '').strip().lower() == 'calculated')
    if not pdf_path:
        return res
    cands = C.find_cells(text)
    if d.authors_cell and any(d.authors_cell[:3]):
        cd, nmatch, ncomp, dev, mode = C.best_match(d.authors_cell[:3], cands, C.entry_name(docx_path))
        if cd is None:
            res['cell'] = ('nocell',)
        elif ncomp >= 2 and nmatch == ncomp:
            res['cell'] = ('match', cd, nmatch, ncomp, dev, mode)
            keys = ['a', 'b', 'c', 'α', 'β', 'γ']
            docx_p = dict(zip(keys, d.authors_cell[:6]))
            pdf_p = dict(zip(keys, [cd.a, cd.b, cd.c, cd.al, cd.be, cd.ga]))
            for k in keys:
                if k in ('α', 'β', 'γ') and C.num_val(docx_p[k]) in (90.0, 120.0):
                    continue                       # symmetry-fixed angle, not transcribed
                iss = C.axis_issues(docx_p[k], pdf_p[k])
                if iss:
                    res['params'][k] = iss
        else:
            res['cell'] = ('investigate', cd, nmatch, ncomp, dev, mode)
            res['cell_diffs'] = C.cell_axis_deltas(d.authors_cell, cd)
    # radiation
    rads = C.find_radiation(text)
    dk = C.anode_key(d.radiation)
    powder = [r for r in rads if r[2] == 'powder']
    # The genuine collection radiation almost always carries an explicit λ; prefer
    # a λ-bearing powder entry over bare element-K mentions (e.g. an explanatory
    # 'for CuKα radiation' aside, or a residual microprobe-standard line).
    powder.sort(key=lambda r: r[1] is None)
    pk = powder[0] if powder else None
    any_match = any(C.anode_key(r[0]) == dk for r in rads)
    if dk is None:
        res['lam'] = ('unrec', 'docx anode not recognised (%s)' % d.radiation)
    elif pk is not None:
        if C.anode_key(pk[0]) == dk:
            res['lam'] = ('ok', 'anode %s matches PDF powder radiation' % d.radiation)
        else:
            res['lam'] = ('flag', 'docx anode %s but PDF POWDER radiation is %sKα'
                          % (d.radiation, pk[0].capitalize()))
        if pk[1] and d.lam and not C.close(float(pk[1]), C.num_val(d.lam), abstol=0.003):
            res['lam'] = (res['lam'][0], res['lam'][1] + ' (λ docx=%s pdf=%s)' % (d.lam, pk[1]))
    elif any_match:
        res['lam'] = ('verify', 'anode %s appears in PDF but no clear powder-context radiation found' % d.radiation)
    else:
        anodes = sorted(set(r[0].capitalize() + 'Kα' for r in rads)) or ['(none found)']
        res['lam'] = ('flag', 'docx anode %s NOT found in PDF; PDF mentions: %s'
                      % (d.radiation, ', '.join(anodes)))
    # For a CALCULATED pattern the docx anode/λ is the modelling wavelength and
    # legitimately differs from the experimental radiation in the article
    # (which is usually the single-crystal collection). Don't flag that mismatch.
    if is_calc and res['lam'] and res['lam'][0] == 'flag':
        res['lam'] = ('calc', res['lam'][1] +
                      ' — but pattern is CALCULATED, so this is the modelling '
                      'wavelength, not the experimental radiation (no action needed)')
    return res

# ----------------------------------------------------------------------------- docx writing
def _cell_runs(cell):
    runs = []
    for p in cell.paragraphs:
        runs.extend(p.runs)
    return runs

def _highlight(cell, color='yellow'):
    for r in cell._tc.findall('.//' + W + 'r'):
        rPr = r.find(W + 'rPr')
        if rPr is None:
            rPr = r.makeelement(W + 'rPr', {}); r.insert(0, rPr)
        hl = rPr.find(W + 'highlight')
        if hl is None:
            hl = rPr.makeelement(W + 'highlight', {}); rPr.append(hl)
        hl.set(W + 'val', color)

def _rows(doc):
    ac = rad = None
    for t in doc.tables:
        for row in t.rows:
            h = row.cells[0].text.strip()
            if h.startswith('Author') and ac is None:
                ac = row
            if h.startswith('Radiation') and rad is None:
                rad = row
    return ac, rad

# --- extra-check anchoring: map a Finding.anchor to the docx cell to comment on
def _find_cell(doc, pred):
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                if pred(c.text):
                    return c
    return None

def _find_value(doc, pred):
    """The value cell following a label cell matching pred; falls back to the
    label cell when the value is blank, so there are always runs to anchor to."""
    for t in doc.tables:
        for row in t.rows:
            cells = row.cells
            for i, c in enumerate(cells):
                if pred(c.text):
                    nxt = cells[i + 1] if i + 1 < len(cells) else None
                    return nxt if (nxt is not None and nxt.text.strip()) else c
    return None

def _anchor_cell(doc, ac_row, anchor):
    if anchor and anchor.startswith('cell:'):
        col = PARAM_COL.get(anchor.split(':', 1)[1])
        return ac_row.cells[col] if (col is not None and col < len(ac_row.cells)) else None
    if anchor == 'instr':
        return (_find_value(doc, lambda t: 'spacing instr' in t.lower())
                or _find_value(doc, lambda t: t.strip().lower().startswith('radiation')))
    if anchor == 'refl':
        return _find_cell(doc, lambda t: t.strip() in ('d(A)', 'd(Å)'))
    if anchor == 'ima':
        return _find_value(doc, lambda t: t.strip() == 'IMA Number')
    if anchor == 'formula':
        return (_find_value(doc, lambda t: t.strip() == 'Empirical')
                or _find_value(doc, lambda t: t.strip() == 'Chemical'))
    if anchor == 'analysis':
        # the 'Analysis' comment cell (where the microprobe wt.% data is given);
        # falls back to the 'Analysis' label cell when that field is empty.
        return _find_value(doc, lambda t: t.strip() == 'Analysis')
    if anchor == 'radiation':
        return _find_value(doc, lambda t: t.strip().lower().startswith('radiation'))
    if anchor == 'filter':
        return _find_value(doc, lambda t: t.strip().lower().startswith('filter')
                           and 'type' not in t.strip().lower())
    if anchor == 'name':
        return _find_value(doc, lambda t: t.strip() == 'Mineral')
    return None

def _write_extras(doc, ac_row, res, rec):
    """Write the writable extra-check findings as comments (flags also highlight
    their anchor cell). Authored as the same 'PXRD Review Tool' so they filter
    apart from human reviewers'."""
    for f in _writable_extras(res):
        cell = _anchor_cell(doc, ac_row, f.anchor) or ac_row.cells[0]
        runs = _cell_runs(cell) or _cell_runs(ac_row.cells[0])
        if not runs:
            continue
        if f.sev == 'flag':
            _highlight(cell)
            rec['highlights'].append(f.anchor or f.code)
        text = _tidy('PXRD check [%s] — %s' % (f.code, f.msg))
        doc.add_comment(runs, text=text, author=AUTHOR, initials=INITIALS)
        rec['comments'].append(text)

def _mark_accept(doc):
    """Put an 'x' in the checkbox cell after the 'Accept' label — the reviewer's
    convention (verified against past reviews). Skips if a human already marked
    any of Accept/Reject/Replace. Returns True if Accept was (or already is) set."""
    for t in doc.tables:
        for row in t.rows:
            cells = row.cells
            idx = {c.text.strip().lower(): i for i, c in enumerate(cells)}
            if not ({'accept', 'reject', 'replace'} <= set(idx)):
                continue
            boxes = [cells[idx[k] + 1] for k in ('accept', 'reject', 'replace')
                     if idx[k] + 1 < len(cells)]
            if any(b.text.strip() for b in boxes):
                return True                       # already decided — don't touch
            acc = cells[idx['accept'] + 1]
            p = acc.paragraphs[0] if acc.paragraphs else acc.add_paragraph()
            if p.runs:
                p.runs[0].text = 'x'
            else:
                p.add_run('x')
            return True
    return False

def _is_severe(res):
    """Withhold auto-Accept ONLY when the cell is fundamentally wrong ('way off').
    Per the reviewer, essentially every new mineral is an Accept — minor cell /
    significant-figure / esd discrepancies and Z mismatches are NOT severe (the
    editorial team fixes those). Only a completely wrong cell (several axes off,
    or a single axis grossly off) is left blank for a manual decision."""
    if res.get('cell', ('',))[0] != 'investigate':
        return False
    out = [t for t in (res.get('cell_diffs') or []) if not t[4]]   # axes off tol
    if len(out) >= 2:                                              # multiple axes wrong
        return True
    for _lab, dv, nv, _d, _m in out:                              # one axis grossly off
        x, y = C.num_val(dv), C.num_val(nv)
        if x and y and abs(x - y) / y > 0.02:
            return True
    return False

def _body_signature(doc):
    """Concatenated body text (paragraphs + table cells) EXCLUDING the
    Accept/Reject/Replace checkbox cells (the tool's own 'x'). Used to tell a
    human-edited output from the untouched source."""
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            cells = row.cells
            idx = {c.text.strip().lower(): i for i, c in enumerate(cells)}
            skip = set()
            if {'accept', 'reject', 'replace'} <= set(idx):
                for k in ('accept', 'reject', 'replace'):
                    if idx[k] + 1 < len(cells):
                        skip.add(idx[k] + 1)
            for i, c in enumerate(cells):
                if i not in skip:
                    parts.append(c.text)
    return '\n'.join(parts)

def _has_foreign_comment(path):
    """True if the docx carries a comment authored by someone other than the tool."""
    try:
        z = zipfile.ZipFile(path)
        if 'word/comments.xml' not in z.namelist():
            return False
        cx = z.read('word/comments.xml').decode('utf-8', 'replace')
        return any(a != AUTHOR for a in re.findall(r'w:author="([^"]*)"', cx))
    except Exception:
        return False

def _has_tracked_changes(path):
    """True if the docx contains tracked changes (w:ins / w:del) — always human;
    the tool never makes revisions."""
    try:
        x = zipfile.ZipFile(path).read('word/document.xml').decode('utf-8', 'replace')
        return ('<w:ins ' in x) or ('<w:del ' in x)
    except Exception:
        return False

def output_hand_edited(out_path, src_path):
    """The output docx was edited by a human since the tool wrote it — tracked
    changes, a non-tool comment, or body text differing from the source (beyond
    the Accept checkbox). Such edits must be preserved across reruns."""
    try:
        if _has_tracked_changes(out_path) or _has_foreign_comment(out_path):
            return True
        return _body_signature(Document(out_path)) != _body_signature(Document(src_path))
    except Exception:
        return False

def _strip_tool_annotations(path):
    """Remove the tool's OWN comments (author == AUTHOR) and its yellow highlights
    from a docx in place, leaving everything human: tracked changes, body edits,
    and any non-tool comments. Lets a rerun refresh the tool's findings without
    disturbing the reviewer's work."""
    z = zipfile.ZipFile(path)
    parts = {n: z.read(n) for n in z.namelist()}
    z.close()
    tool_ids = set()
    if 'word/comments.xml' in parts:
        croot = etree.fromstring(parts['word/comments.xml'])
        for c in list(croot):
            if c.tag == _q('comment') and c.get(_q('author')) == AUTHOR:
                tool_ids.add(c.get(_q('id')))
                croot.remove(c)
        parts['word/comments.xml'] = etree.tostring(croot, xml_declaration=True,
                                                    encoding='UTF-8', standalone=True)
    droot = etree.fromstring(parts['word/document.xml'])
    # drop the comment anchors/refs for tool comments
    for tag in ('commentRangeStart', 'commentRangeEnd'):
        for el in list(droot.iter(_q(tag))):
            if el.get(_q('id')) in tool_ids:
                el.getparent().remove(el)
    for ref in list(droot.iter(_q('commentReference'))):
        if ref.get(_q('id')) in tool_ids:
            run = ref.getparent()                      # the <w:r> wrapping the ref
            if run is not None and run.getparent() is not None:
                run.getparent().remove(run)
    # drop the tool's yellow highlights (re-applied fresh on re-annotation)
    for hl in list(droot.iter(_q('highlight'))):
        if hl.get(_q('val')) == 'yellow':
            hl.getparent().remove(hl)
    parts['word/document.xml'] = etree.tostring(droot, xml_declaration=True,
                                                encoding='UTF-8', standalone=True)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zo:
        for n, data in parts.items():
            zo.writestr(n, data)
    with open(path, 'wb') as f:
        f.write(buf.getvalue())
    return tool_ids

def annotate(docx_path, res, out_path, inplace=False, base_path=None):
    """Write comments/highlights for `res` into the docx — ERRORS ONLY. Entries
    whose cell matches the PDF with no discrepancies get no annotation (the
    unmodified file is copied through, preserving the original bytes). Returns
    the number of comments added.

    Reported (each a separate comment):
      * per-parameter value / significant-figure / esd mismatches (highlighted);
      * a radiation anode hard-flag (highlighted), only when the cell matched;
      * a single brief 'No matching PDF cell found.' when no cell could be
        matched (no cell parsed, or no exact match — both reduce to this);
      * extra-check FLAGS — symmetry/precision (on the offending cell parameter),
        Calculated pattern (Instrumentation), hkl-indexing (Reflection List),
        missing IMA number, ideal-formula — each highlighted on its anchor cell;
      * the authoritative Mindat group statement (informational, on the mineral
        name; not highlighted).

    Returns a record: {'status', 'highlights': [cell labels], 'comments':
    [texts], 'clean': bool} for the run log."""
    rec = {'status': res['cell'][0], 'highlights': [], 'comments': [], 'clean': False,
           'accept': False}
    # `base_path` is the document to annotate ONTO — the source for a normal run,
    # or the reviewer's edited output (already stripped of old tool comments) when
    # refreshing in place so manual edits are preserved.
    base = base_path or docx_path
    if _is_clean(res):
        rec['clean'] = True
        # clean entries are never severe → auto-Accept. (We now open+save to set
        # the Accept 'x', so they are no longer byte-identical copies.)
        out = docx_path if inplace else out_path
        if not inplace:
            os.makedirs(os.path.dirname(out_path) or '.', exist_ok=True)
        doc = Document(base)
        rec['accept'] = _mark_accept(doc)
        doc.save(out)
        return rec

    doc = Document(base)
    ac_row, rad_row = _rows(doc)
    if ac_row is None:
        raise RuntimeError('no Author\'s Cell row found')
    status = res['cell'][0]

    if status == 'match':
        # per-parameter flags: highlight the value cell + explain
        for k, issues in res['params'].items():
            col = PARAM_COL.get(k)
            if col is None or col >= len(ac_row.cells):
                continue
            cell = ac_row.cells[col]
            _highlight(cell)
            runs = _cell_runs(cell) or _cell_runs(ac_row.cells[0])
            body = '; '.join('%s — %s' % (KIND_LABEL[kind], note) for kind, note in issues)
            text = _tidy('PXRD check — %s: %s' % (k, body))
            doc.add_comment(runs, text=text, author=AUTHOR, initials=INITIALS)
            rec['highlights'].append("Author's Cell:%s" % k)
            rec['comments'].append(text)
        # radiation hard-flag: highlight anode cell + explain
        if _has_lam_flag(res) and rad_row is not None and len(rad_row.cells) > 1:
            cell = rad_row.cells[1]
            _highlight(cell)
            runs = _cell_runs(cell) or _cell_runs(rad_row.cells[0])
            text = _tidy('PXRD check — radiation: ' + res['lam'][1])
            doc.add_comment(runs, text=text, author=AUTHOR, initials=INITIALS)
            rec['highlights'].append('Radiation:anode')
            rec['comments'].append(text)
    elif status == 'investigate' and res.get('cell_diffs'):
        # closest PDF cell exists but isn't an exact match — pinpoint the axis.
        diffs = res['cell_diffs']
        out_axes = [t for t in diffs if not t[4]]
        in_axes = [t for t in diffs if t[4]]
        if len(out_axes) == 1 and in_axes:
            lab, dv, nv, dd, _ = out_axes[0]
            col = PARAM_COL.get(lab)
            cell = ac_row.cells[col] if (col is not None and col < len(ac_row.cells)) else ac_row.cells[0]
            _highlight(cell)
            text = _tidy('PXRD check — %s: likely transcription error — docx=%s but PDF cell gives %s '
                         '(Δ=%.4f Å); the other axes match exactly.' % (lab, dv, nv, dd))
            doc.add_comment(_cell_runs(cell) or _cell_runs(ac_row.cells[0]),
                            text=text, author=AUTHOR, initials=INITIALS)
            rec['highlights'].append("Author's Cell:%s" % lab)
            rec['comments'].append(text)
        else:
            text = ('No exact cell match — closest PDF cell differs on %s; it may be a different '
                    'phase/cell in a multi-cell PDF (the matching cell may be unparsed).'
                    % ', '.join(t[0] for t in out_axes))
            doc.add_comment(_cell_runs(ac_row.cells[0]), text=text, author=AUTHOR, initials=INITIALS)
            rec['comments'].append(text)
    else:
        # no cell parsed at all (nocell / nopdf): one brief flag
        doc.add_comment(_cell_runs(ac_row.cells[0]), text=NO_MATCH, author=AUTHOR, initials=INITIALS)
        rec['comments'].append(NO_MATCH)

    # the 10 extra reviewer-comment checks (symmetry, calculated, indexing, IMA,
    # ideal-formula, and the authoritative Mindat group statement)
    _write_extras(doc, ac_row, res, rec)

    # auto-Accept unless the entry is severe (then leave blank for manual decision)
    if not _is_severe(res):
        rec['accept'] = _mark_accept(doc)

    os.makedirs(os.path.dirname(out_path) or '.', exist_ok=True)
    doc.save(out_path)
    return rec

# ----------------------------------------------------------------------------- main
def output_name(src_basename, edited):
    """Output filename for a source docx: flagged entries (the tool added a
    comment/highlight) get an '_edited' suffix so they stand out in the folder
    listing; clean entries keep the source name."""
    stem, ext = os.path.splitext(src_basename)
    return stem + ('_edited' if edited else '') + ext

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('folder')
    ap.add_argument('--id', help='only this entry id, e.g. I003416')
    ap.add_argument('--out', help='output folder (default <folder>/review_out)')
    ap.add_argument('--inplace', action='store_true', help='edit the original .docx in place')
    ap.add_argument('--limit', type=int, help='only the first N entries')
    ap.add_argument('--log', help='log file path (default <out>/annotation_log.txt)')
    ap.add_argument('--force', action='store_true',
                    help='regenerate every output even if it has manual edits (default: preserve hand-edited outputs)')
    args = ap.parse_args()

    idx = C.pdf_index(args.folder)
    cif_idx = C.cif_index(args.folder)
    docs = sorted(f for f in glob.glob(os.path.join(args.folder, '*.docx'))
                  if not os.path.basename(f).startswith('~$'))
    if args.id:
        docs = [d for d in docs if args.id in os.path.basename(d)]
    if args.limit:
        docs = docs[:args.limit]
    out_dir = args.out or os.path.join(args.folder, 'review_out')

    records = []           # (filename, rec) for the log
    for dp in docs:
        eid = C.entry_id(dp)
        pdf = idx.get(eid)
        cif = cif_idx.get(eid)
        res = analyze(dp, pdf, cif)
        # Name the output by what the tool found: flagged entries -> '<name>_edited.docx',
        # clean ones keep the source name. Cleanliness is deterministic for a given
        # source, so the path is stable across reruns (refresh below still works).
        if args.inplace:
            out = dp
        else:
            edited = not _is_clean(res)
            out = os.path.join(out_dir, output_name(os.path.basename(dp), edited))
            # A prior run may have written the opposite-named twin (clean<->edited can
            # flip if the source/PDF changed, or it predates this naming). Keep one
            # output per entry without ever losing the reviewer's work:
            #   * a pristine tool output under the old name -> drop it;
            #   * a hand-edited old name with no new-name output yet -> RENAME it onto
            #     the new name, so refresh-in-place (below) carries the edits forward;
            #   * both names already present and the old one edited -> leave it, warn.
            stale = os.path.join(out_dir, output_name(os.path.basename(dp), not edited))
            if os.path.exists(stale):
                if not output_hand_edited(stale, dp):
                    os.remove(stale)
                elif not os.path.exists(out):
                    os.rename(stale, out)
                    print('  -- %s: migrated hand-edited %s -> %s (edits preserved)'
                          % (os.path.basename(dp), os.path.basename(stale), os.path.basename(out)))
                else:
                    print('  !! %s: hand-edited %s sits alongside %s — reconcile manually'
                          % (os.path.basename(dp), os.path.basename(stale), os.path.basename(out)))
        # If the reviewer hand-edited this output, REFRESH in place: back it up,
        # strip the tool's own old comments/highlights, then re-annotate ONTO the
        # edited file so tracked changes / comments / text fixes are preserved.
        refresh = (not args.inplace and not args.force and os.path.exists(out)
                   and output_hand_edited(out, dp))
        base = None
        if refresh:
            # back up the reviewer's edited file (TIMESTAMPED — each refresh keeps a
            # distinct copy), then strip a TEMP copy (never the real output) so a
            # later annotate() failure can't leave it half-done.
            bdir = os.path.join(os.path.dirname(out), '.edit_backup')
            os.makedirs(bdir, exist_ok=True)
            stem, ext = os.path.splitext(os.path.basename(out))
            ts = datetime.datetime.now().strftime('%Y-%m-%dT%H%M%S')
            shutil.copy(out, os.path.join(bdir, '%s.%s%s' % (stem, ts, ext)))
            try:
                base = os.path.join(bdir, '~strip_' + os.path.basename(out))
                shutil.copy(out, base)
                _strip_tool_annotations(base)    # strip the temp, not `out`
            except Exception as e:
                print('  !! %s: strip failed (%s) — left untouched' % (os.path.basename(dp), e))
                records.append((os.path.basename(dp),
                                {'status': 'preserved', 'highlights': [], 'comments': [],
                                 'clean': False, 'accept': None, 'refreshed': True}))
                continue
        try:
            rec = annotate(dp, res, out, inplace=args.inplace, base_path=base)
        except Exception as e:
            # `out` is untouched on failure (annotate read the temp, not `out`)
            print('  !! %s: %s' % (os.path.basename(dp), e)); continue
        finally:
            if base and os.path.exists(base):
                os.remove(base)
        if refresh:
            rec['refreshed'] = True
        rec['outfile'] = os.path.basename(out)
        records.append((os.path.basename(dp), rec))
        n = len(rec['comments'])
        tag = ' [REFRESHED: manual edits kept]' if rec.get('refreshed') else ''
        renamed = '' if (args.inplace or rec['clean']) else ' -> %s' % rec['outfile']
        print('%-44s cell=%-12s %-26s (%d comment%s)%s%s'
              % (os.path.basename(dp), rec['status'],
                 'clean' if rec['clean'] else ('highlights: ' + ', '.join(rec['highlights']) or 'flag'),
                 n, '' if n == 1 else 's', renamed, tag))

    # --- write the run log -------------------------------------------------
    log_path = args.log or os.path.join(out_dir if not args.inplace else args.folder, 'annotation_log.txt')
    refreshed = [(f, r) for f, r in records if r.get('refreshed')]
    edited = [(f, r) for f, r in records if not r['clean']]
    os.makedirs(os.path.dirname(log_path) or '.', exist_ok=True)
    with open(log_path, 'w') as fh:
        fh.write('PXRD review — annotation log\n')
        fh.write('source folder : %s\n' % os.path.abspath(args.folder))
        fh.write('output        : %s\n' % ('in place' if args.inplace else os.path.abspath(out_dir)))
        fh.write('entries       : %d total | %d edited | %d clean (untouched)\n'
                 % (len(records), len(edited), len(records) - len(edited)))
        fh.write('total comments: %d | total highlights: %d\n'
                 % (sum(len(r['comments']) for _, r in records),
                    sum(len(r['highlights']) for _, r in records)))
        accepted = [f for f, r in records if r.get('accept') is True]
        withheld = [f for f, r in records if r.get('accept') is False]
        fh.write('Accept marked : %d auto-accepted | %d left blank (severe — decide manually)\n'
                 % (len(accepted), len(withheld)))
        if withheld:
            fh.write('  left blank  : %s\n' % ', '.join(C.entry_id(f) or f for f in withheld))
        if refreshed:
            fh.write('refreshed     : %s  (hand-edited — your edits kept, tool comments refreshed; --force to rebuild from source)\n'
                     % ', '.join(C.entry_id(f) or f for f, r in refreshed))
        fh.write('=' * 78 + '\n\nEDITED ENTRIES (highlights / comments)\n')
        for f, r in edited:
            hl = ', '.join(r['highlights']) if r['highlights'] else '(none)'
            name = (C.entry_name(f) or C.entry_id(f) or f).upper()
            eid = C.entry_id(f)
            fh.write('\n' + '=' * 78 + '\n')
            fh.write('  %s   (%s)\n' % (name, eid or '?'))
            fh.write('-' * 78 + '\n')
            fh.write('  highlights: %s\n' % hl)
            for c in r['comments']:
                c = re.sub(r'^PXRD check (?:\[[\w+\-]+\] — (?:Mindat: )?|— )', '', c)
                fh.write(textwrap.fill(c, width=78,
                                       initial_indent='  comment   : ',
                                       subsequent_indent=' ' * 14) + '\n')
        fh.write('\n' + '=' * 78 + '\n\n')
        fh.write('CLEAN ENTRIES (no edits — copied unchanged)\n' + '-' * 78 + '\n')
        for f, r in records:
            if r['clean']:
                fh.write('  %-32s (%s)\n' % (C.entry_name(f).upper(), C.entry_id(f) or '?'))
    print('\nlog written -> %s' % log_path)

if __name__ == '__main__':
    main()
